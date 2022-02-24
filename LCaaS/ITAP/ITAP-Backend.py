#testing edit from different source
from mailmerge import MailMerge
import pythoncom
from datetime import date, datetime
import time
import glob
import os, shutil
from datetime import date
import win32com.client as win32
import copy
import time
from pathlib import Path
import os, json
from math import floor as floor
import docx
import fileinput
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from docx import Document
import shutil
import xlrd
from flask import Flask, request, jsonify
from flask_cors import CORS
import logging, openpyxl
from datetime import datetime, timedelta
import xlwt
from pymongo import MongoClient
import pandas as pd
# pd.options.mode.chained_assignment = None
import subprocess
from time import strftime
from time import gmtime
from openpyxl import load_workbook
import zipfile, io
import schedule
import threading
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from datetime import date, datetime  # import datetime for calculationg date dynamically
import time
import sys
# TestAutomationTool 1
#Testing for demo

app = Flask(__name__)  # Initializing Flask
# iTAP New Version
CORS(app)  # Allow Cross Origin
#second testing
#third testing
client = MongoClient('localhost', 27017)  # Connecting Mongo DB
db = client["TestAutomation2"]  # Creating a Collection

global folderName  # Declare Scenario name as global
global outerFolderName  # Declare application name as global
with open("config.json") as f:
    data_folder = json.load(f)  # Load configuration file

base_location = data_folder['base_location']  # Base location represents the location of iTAP recorded file
base_location_front = base_location.replace("\\", "/")
ui_location = data_folder['ui_location']  # Define location of UI Source code
ui_delay = data_folder['delay']  # Delay between each page if required
ui_minimumdelay = data_folder['minimum_delay']  # Delay between each element if required
download_location = data_folder["download"]  # Download folder location
recorder_path = data_folder["recorder_path"]
chrome_driver_path = data_folder["chrome_driver_path"]
sleep_delay=data_folder["sleep_delay"]

@app.route("/getSavedTestCases")  # Fetch data from db
def getSavedTestCases():
    try:

        connections = db.folderName.find()
        outDict = {}
        finalList = []
        finalDict = {}
        for collection in connections:
            outDict["scenario_Name"] = collection["scenarioName"]
            outDict["application_Name"] = collection["appName"]
            outDict["test_Description"] = collection["test_Description"]
            finalList.append(outDict)
            outDict = {}
        finalDict["data"] = finalList
        finalDict["headers"] = ["application_Name", "scenario_Name", "test_Description", "Actions"]

        return jsonify(finalDict)
    except Exception as e:
        print(e)


@app.route("/getSavedTestRuns")  # view all the test runs
def getSavedTestRuns():  # Fetch all saved multiple scenario
    connections = db.testRun.find()
    dictionary = {}
    finalList = []
    finalDict = {}
    for connection in connections:
        dictionary["test_name"] = connection["test_name"]
        dictionary["test_scenarios_used"] = connection["test_scenarios_used"]
        finalList.append(dictionary)
        dictionary = {}
    print("finalList", finalList)
    finalDict["headers"] = ["test_name", "test_scenarios_used", "Actions"]
    finalDict["data"] = finalList
    return jsonify(finalDict)


@app.route("/getScheduledJobs")  # Fetch data from db
def getScheduledJobs():
    connections = db.scheduledJobs.find()
    outDict = {}
    finalList = []
    finalDict = {}
    for collection in connections:
        outDict["scheduler_name"] = collection["scheduler_name"]
        outDict["test_job_name"] = collection["test_job_name"]
        outDict["scheduled_time"] = collection["scheduled_time"]
        finalList.append(outDict)
        outDict = {}
    finalDict["data"] = finalList
    finalDict["headers"] = ["scheduler_name", "test_job_name", "scheduled_time", "Actions"]
    return jsonify(finalDict)


@app.route("/getAllAppNames")  # Gets called when combine scenario button is clicked
def getAllAppNames():
    appName = []
    collections = db.folderName.find()  # returns all documents in json from folderName collection
    for collection in collections:
        appName.append(collection["appName"])  # fetch only appName from json
    appName = list(dict.fromkeys(appName))  # convert into a list
    return jsonify(appName)  # returns the appName list


@app.route('/uploadRecordedZip', methods=["GET", "POST"])  # import the exported zip
def uploadRecordedZip():
    file = request.get_data()  # get data from UI
    fname = request.args.get("fname")  # get fname from UI
    z = zipfile.ZipFile(io.BytesIO(file))  # zip the file
    if os.path.isdir(base_location + "\\TmpZip") == False:  # if TmpZip not in directory
        os.mkdir(base_location + "\\TmpZip")  # make directory with name TmpZip
    z.extractall(base_location + "\\TmpZip")  # extract files from zip
    populateDb(base_location + "\\TmpZip")  # call the function populateDb using the file location
    return jsonify("Success")  # return success in json format


def populateDb(tmpLoc):  # Used to update db
    with open(tmpLoc + "\\json_file.json", "r") as f:  # Load the json file
        data = json.load(f)
    print(data)

    appName = data["appName"]  # get app name from json
    scenarioName = data["scenarioName"]  # get scenario name from json

    connections = db.folderName.find()  # get entire collection
    print(connections)

    added_updated_count = 0
    for collection in connections:  # if the particular app name and scenarion name matches, run the function
        if appName == collection["appName"]:
            if scenarioName == collection["scenarioName"]:
                print("this falling under this category")
                myquery = {"appName": data["appName"], "scenarioName": data["scenarioName"]}
                newvalues = {"$set": {"test_Description": data["test_Description"], "path": data["path"],
                                      "innerpath": data["innerpath"]}}  # set command to update particular document
                db.folderName.update_many(myquery, newvalues)  # update new data in collection
                added_updated_count += 1
                break
            else:
                dbData = [{"path": data["path"],
                           "innerpath": data["innerpath"],
                           "appName": data["appName"],
                           "scenarioName": data["scenarioName"],
                           "test_Description": data["test_Description"]}]
                db.folderName.insert_many(dbData)  # if doesn't matches update it with old values
                added_updated_count += 1
                break
        else:
            continue
    if added_updated_count == 0:
        dbData = [{"path": data["path"],
                   "innerpath": data["innerpath"],
                   "appName": data["appName"],
                   "scenarioName": data["scenarioName"],
                   "test_Description": data["test_Description"]}]
        db.folderName.insert_many(dbData)
    shutil.copytree(tmpLoc, data[
        "innerpath"])  # if nothing is updated, update a new one and copy the files to particular location
    return "Success"


@app.route("/runSavedTestCase")  # Run already exist data, gets called when user clicked on play button
def runSavedTestCase():
    global base_location
    global outerFolderName
    global folderName
    application_name = request.args.get("application_name")
    scenario_name = request.args.get("scenario_name")
    outerFolderName=application_name
    folderName=scenario_name
    connections = db.folderName.find()
    for connection in connections:
        if connection["appName"] == application_name and connection["scenarioName"] == scenario_name:
            appPath = connection["path"]
            scenarioPath = connection["innerpath"]
            if os.path.isdir(appPath):
                if os.path.isdir(scenarioPath):
                    if os.path.exists(scenarioPath + "\\Batchfiles\\Test_batch.bat"):
                        setPath = scenarioPath + "\\Batchfiles\\Test_batch.bat"
                        print("HI IM COMBINED SCENARIO")
                        subprocess.call(setPath)
                        log_filename = "conf.log"
                        conf_log = conf_log_modification(log_filename)
                        output_json = {"status": "Success", "conf_log": conf_log}
                        print("output log", output_json)

                        return jsonify(output_json)

                    else:
                        setPath = scenarioPath + "\\batch.bat"
                        print("HI IM NORM SCENARIO")
                        subprocess.call(setPath)
                        log_filename = "conf.log"
                        conf_log = conf_log_modification(log_filename)
                        output_json = {"status": "Success", "conf_log": conf_log}
                        print("output log", output_json)

                        return jsonify(output_json)

                else:
                    print("hey here i am")
                    return jsonify("scenarioFail")
            else:
                print("im returning appfail")
                return jsonify("appFail")


def download_scenario(application_name, scenario_name):
    connections = db.folderName.find()  # load entire collection

    for connection in connections:
        if connection["appName"] == application_name and connection[
            "scenarioName"] == scenario_name:  # checks if collection contains the specific application name and scenario name
            appPath = connection["path"]  # get path from db
            scenarioPath = connection["innerpath"]  # get inner path from db
            if os.path.isdir(appPath):  # check whether the directory appPath exists or not
                if os.path.isdir(scenarioPath):  # check whether the directory scenarioPath exists or not
                    spec_file = open(scenarioPath + "\\spec.js")  # opens spec.js file
                    headers = []
                    values = []
                    open_spec = spec_file.read()  # read the file
                    json_format = {}
                    final_json = {}
                    headerList = []
                    final_list = []
                    split_file = open_spec.split("browser.driver.manage().window().maximize();")[1].split(
                        ";")  # split it based on "browser.driver.manage().window().maximize();" and picks the 1st index and then splits it based on ";"
                    for elm in split_file:
                        if (elm.find("element") >= 0):  # hceck if the line contains element
                            if (elm.find("sendKeys") >= 0):  # check if the line contains sendKeys

                                fin_elm = elm.split("xpath")[1].split(".sendKeys")[
                                    0]  # split it based on xpath and sendkeys to get the xpath
                                headers.append(fin_elm)  # append it to headers which is required for excel sheet header
                                # fin_val = elm.split("sendKeys")[1].split("'")[1]
                                if elm.split("sendKeys")[1].__contains__("'"):  # get values using this operation
                                    fin_val = elm.split("sendKeys")[1].split("'")[1]
                                elif elm.split("sendKeys")[1].__contains__(
                                        '"'):  # if  double quotes exists instead of single quotes
                                    fin_val = elm.split("sendKeys")[1].split('"')[1]
                                values.append(fin_val)  # append the value to values list
                    workbook = xlwt.Workbook()  # create a new workbook
                    sheet = workbook.add_sheet(scenario_name)  # add a new sheet to it

                    header_style = xlwt.easyxf('font: bold 1')  # add header style
                    val_style = xlwt.easyxf('font: bold 0')  # add value style

                    for index in range(len(headers)):  # iterate the headers

                        if headers[index] not in headerList:
                            headerList.append(headers[index])  # store all the unique headers in headerList
                        json_format[headers[index]] = values[index]  # store header and value in unique format

                    final_list.append(json_format)  # append json to the final_ist
                    final_list_1 = final_list[::-1]  # reverse the obtained list so that it will be in right order in ui

                    final_json = {"headers": headerList,
                                  "data": final_list_1}  # send it to ui the form of headers and data as key
                    return jsonify(final_json)  # return the obtained json
                else:
                    return "scenarioFail"  # if scenario path doesn't exists
            else:
                return "appFail"  # if app path does't exists

@app.route("/download_spec_scenario")  # Download specific scenario, gets called when user clicks on download button in ui
def download_spec_scenario():
    # folderName = request.args.get("folderName")
    application_name = request.args.get("application_name")  # get application name
    scenario_name = request.args.get("scenario_name")  # get scenario name

    print(application_name, scenario_name, "----app sc name")
    op_download = download_scenario(application_name, scenario_name)
    return op_download


def generateSpecForUploadedJob(uploaded_data,testRunName):
    newlist = []
    scenarioNameList = []
    appNameList = []
    connections = db.testRun.find()
    replace_spec=''
    ct=0

    for connection in connections:
        if connection["test_name"] == testRunName:
            appPath = connection["path"]
            scenario = connection["test_scenarios_used"]
            split_scenario = scenario.split(",")
            for names in split_scenario:
                appNameList.append(names.split("[")[1].split("]")[0])
                scenarioNameList.append(names.split("[")[0])
            upload_spec_file = open(appPath + "\\job_upload_spec.js", "w+")
            upload_batch_file = open(appPath + "\\job_upload_batch.bat", "w+")
            upload_conf_file = open(appPath + "\\job_upload_conf.js", "w+")
    for key,testApps,testScenarios in zip(uploaded_data.keys(),appNameList,scenarioNameList):
        print("here", uploaded_data[key]["data"])
        newlist.clear()
        newlist.append(uploaded_data[key]["data"])
        print("newlist", newlist)
        data_set = 1
        for i in newlist:
            print("listitem", i)

            for j in i:
                point_until_iterated = 0

                print("j is", j)
                filePath = base_location + "\\" + testApps + "\\" + testScenarios
                if db.folderName.find({"innerpath": filePath}).count() > 0:
                    if os.path.isdir(filePath):
                        print("THE FILE PATH", filePath)
                        if os.path.exists(filePath + "\\spec.js"):
                            spec_file = open(filePath + "\\spec.js", "r").read()

                            for index,keys in enumerate(j):
                                # it_block_flag = False
                                print("indexxx",index)
                                length_of_keys=len(j.keys())-1
                                print("lengthh",length_of_keys)
                                print(length_of_keys)
                                pointerFlag = 0
                                text_search = keys
                                value_ = j[keys]
                                value_to_change=str(value_)
                                if text_search.__contains__("('"):
                                    text_to_search=str(text_search).split("('",1)[1].split("')",1)[0]
                                else:
                                    text_to_search=str(text_search)
                                print("key", text_to_search)
                                print("val", value_to_change)
                                newFile=spec_file.split("describe(")[1].rsplit('});', 1)[0].strip()
                                spec_to_be_used=''.join(newFile)
                                final_spec=spec_to_be_used.split(";")
                                length_of_spec=len(final_spec)
                                print(final_spec)
                                print("length_of_spec",length_of_spec)
                                for idx,lines in enumerate(final_spec):
                                    # print(spec_to_be_used)
                                    if idx>point_until_iterated and not lines=='':
                                        print(idx, point_until_iterated)
                                        if lines.__contains__("browser.driver.manage().window().maximize()") :
                                            lines = '''it(\'''' + testScenarios + '''_Uploaded Data ''' + str(
                                                data_set) + '''\', function() {
                                                browser.ignoreSynchronization =true;\n'''+lines+''''''
                                            data_set += 1
                                        if lines.__contains__("browser.takeScreenshot"):
                                            lines='''browser.takeScreenshot().then(function (png) {var dir=\"'''+base_location.replace("\\","\\\\") + "\\\\" + "testRun" + '\\\\' + testRunName +'''\"'''
                                        if lines.__contains__("var fname="):
                                            ct += 1
                                            lines="var fname=\"page"+str(ct)+".png\""
                                        if lines.__contains__(text_to_search) and not lines.__contains__("browser.wait") and not lines.__contains__("clear(") and  lines.__contains__("sendKeys"):
                                            pointerFlag=1
                                            lines="element(by.xpath('"+text_to_search+"')).sendKeys('"+value_to_change+"')"
                                        if text_to_search.__contains__("//select") and lines.__contains__(text_to_search) and lines.__contains__("//select") and lines.__contains__("browser.wait"):
                                            pointerFlag = 1
                                            lines="browser.wait(until.presenceOf(element(by.xpath(\'"+text_to_search+"//option[@value=\""+value_to_change+"\"]\'))), delay, 'Element taking too long to appear in the DOM')"
                                        if text_to_search.__contains__("//select") and lines.__contains__(text_to_search) and lines.__contains__("//select") and not lines.__contains__("browser.wait"):
                                            pointerFlag = 1
                                            lines="element(by.xpath(\'"+text_to_search+"//option[@value=\""+value_to_change+"\"]\')).click()"
                                        # if lines=="})" and not final_spec[idx-1].__contains__("stream1.end") and not final_spec[idx-1].__contains__("send"):
                                        #     print("itsss founddd",lines)
                                        #     lines=''
                                        if not lines=='' and (pointerFlag==0  or int(index)==int(length_of_keys) ) :
                                            replace_spec+=lines+";"+"\n"
                                            # it_block_flag=0
                                        if not lines=='' and pointerFlag==1 and not int(index)==int(length_of_keys):
                                            replace_spec+=lines+";"+"\n"
                                            point_until_iterated = idx
                                            break

                                print("endpoint",point_until_iterated)


    # print(replace_spec)
    write_spec=provideJasmineSpecReporter(testRunName) + replace_spec +"});"
    write_conf=provideConf("job_upload_spec.js",base_location,"testRun",testRunName)
    write_batch="START /B cmd " + appPath + " && protractor " + appPath + "\\job_upload_conf.js >" + appPath + "\\upload_log_file.log"
    upload_spec_file.write(write_spec)
    upload_conf_file.write(write_conf)
    upload_batch_file.write(write_batch)
    upload_spec_file.close()
    upload_conf_file.close()
    upload_batch_file.close()
    # formatted_spec=''
    # with open(appPath+"\\job_upload_spec.js", "r") as read_job_spec:
    #     buf=read_job_spec.readlines()
    # with open(appPath+"\\job_upload_spec.js", "w+") as format_job_spec:
    #     for line in buf:
    #         if :
    #             line = line.split("});")[0]
    #             print(line)
    #         formatted_spec += line
    #     format_job_spec.write(formatted_spec)
    return "success"
@app.route("/uploadExcelRun", methods=['GET', 'POST'])  # upload the test scenario excel file after filling the data
def uploadExcelRun():
    global folderName
    global outerFolderName
    file = request.get_data()
    str_file = file.decode("utf8")
    testRunName = request.args.get("testName")
    print("str_file", str_file)
    uploaded_data = json.loads(str_file)
    print("json_format", uploaded_data)
    generateSpecForUploadedJob(uploaded_data,testRunName)
    outerFolderName="testRun"
    folderName=testRunName
    subprocess.call(base_location+"\\"+outerFolderName+"\\"+testRunName+"\\"+"job_upload_batch.bat")
    moveReport(outerFolderName,testRunName)
    log_filename = "upload_log_file.log"
    conf_log = conf_log_modification(log_filename)
    output_json = {"status": "Success", "conf_log": conf_log}
    print("output log", output_json)
    return jsonify(output_json)


@app.route("/uploadUIRun")  # create a json object to generate excel file in ui
def uploadUIRun(testRunName, sheetLength, flag):  # Name and sheet length as parameter
    replace_spec = ''
    head_list = []
    cell_list = []
    filePath = ''
    # folderName = request.args.get("folderName")
    connections = db.testRun.find()  # get all collections from db
    for connection in connections:
        if connection[
            "test_name"] == testRunName:  # if the name in db same as testRunName then this condition will get executed
            appPath = connection["path"]  # get path of the particular collection
            scenario = connection["test_scenarios_used"]  # get the scenario files
            split_scenario = scenario.split(",")[sheetLength]  # split based on sheetlength
            appNameRun = split_scenario.split("[")[1].split("]")[
                0]  # split the application name which is inside the bracket
            scenarioNameRun = split_scenario.split("[")[0]  # get scenario name which is outside the bracket
            print("scenario Name run", scenarioNameRun)
            collections = db.folderName.find()
            scenarioNamedict = {}
            for collection in collections:
                if collection["appName"] == appNameRun and collection[
                    "scenarioName"] == scenarioNameRun:  # if app name and scenario name matches this condition wil get executed
                    filePath = collection["innerpath"]  # get the path
            if os.path.isdir(filePath):
                if os.path.exists(filePath + "\\spec.js"):  # check if the file exists or not
                    spec_file = open(
                        filePath + "\\spec.js")  # open the particular file
                    path = appPath + "\\sample.xlsx"  # create a new workbook called sample.xlsx
                    new_spec_file = open(
                        appPath + "\\multiple_scenario_file.js",
                        "w+")  # create new file to store multiple scenarios

                    combined_multi_scenario_file = open(appPath + "\\combined_multi_scenario_file.js",
                                                        "a+")  # create new file to store specs
                    # new_conf_log_file = open(appPath + "\\multiple_conf_file.log", 'w+')
                    new_conf_file = open(appPath + "\\multiple_conf_file.js", 'w+')  # create new file for conf
                    combined_conf_file = open(appPath + "\\combined_conf_file.js",
                                              'w+')  # create a temporary file to store data
                    append_file = open(appPath + "\\append_spec", "a+")  # append each data in this spec fie
                    batch_file = open(appPath + "\\multiple_batch_file.bat",
                                      "w+")  # create a separate batch file to run this
                    combined_batch_file = open(appPath + "\\combined_batch_file.bat",
                                               "w+")  # create another separate batc file to run all the test cases

                    # wb = xlrd.open_workbook(path)
                    wb = openpyxl.load_workbook(path)  # load a new workbook
                    sheet = wb.active  # make this sheet active
                    start_row = 1  # initialize row value note: row value should start from 0
                    end_row = sheet.max_row  # shows the maximun non empty rows
                    num_of_columns = sheet.max_column  # shows maximum  non empty columns
                    open_file = spec_file.read()  # read the spec file
                    split_file = open_file[:-3]  # }); not required
                    if (split_file.endswith("//")):  # split each line based on "//"
                        split_file = split_file[:-2] + "});"  # add }); at the end

                    new_list = []

                    if scenarioNameRun in new_list:  # get unique header and value
                        scenarioNamedict[scenarioNameRun] += 1  # check if it is unique
                    else:
                        scenarioNamedict[scenarioNameRun] = 0  # if it is not present initialize it to 0
                    new_list.append(scenarioNameRun)  # append all the unique name in new_list
                    sec_test = "it(\'" + scenarioNameRun + " " + str(
                        scenarioNamedict[scenarioNameRun]) + "\', function() { browser.ignoreSynchronization =true;" + \
                               "browser.driver.manage().window().maximize();" + \
                               split_file.split("browser.driver.manage().window().maximize();")[
                                   1]  # add it condition in our new file

                    for r in range(start_row, end_row):  # iterate each value we get
                        head_list = []
                        cell_list = []
                        for c in range(num_of_columns):  # iterate based on column value
                            cell_obj = sheet.cell(row=r + 1, column=c + 1).value  # get value of each column
                            head_obj = sheet.cell(row=1, column=c + 1).value  # get header of each column
                            head_list.append(head_obj)  # append it to headerlist
                            cell_list.append(cell_obj)  # append column value in cell list
                        for index_value, line in enumerate(sec_test.split(";")):  # get each line based on ";"
                            for idx, header in enumerate(head_list):  # iterate header from header list
                                line = str(line)  # convert any datatype to string
                                if index_value == len(sec_test.split(";")) - 1:  # if reaches the end, break
                                    break
                                if (line.__contains__(header) and not line.__contains__("browser.wait(until")
                                        and not line.__contains__(".clear(") and not line.strip().startswith(
                                            "expect")):  # don;t consider line containing browser.wat or clear or expect
                                    if (type(cell_list[idx]) == int or type(cell_list[
                                                                                idx]) == float):  # if type is int or float, need to convert it to int first, then str
                                        line = line.split("sendKeys")[0] + "sendKeys(" + str(int(cell_list[idx])) + ");"
                                    else:
                                        line = line.split("sendKeys")[0] + "sendKeys('" + str(cell_list[idx]) + "');"
                                    break

                            replace_spec += line + ";"  # add ";" to each line
                    new_spec = provideJasmineSpecReporter(testRunName) + str(replace_spec) + "});"  # create a file structure with required imports and values
                    if os.stat(appPath + "\\combined_multi_scenario_file.js").st_size == 0:  # represents the size of the file in bytes
                        new_combined_file = provideShortSpec(testRunName) + str(replace_spec)
                        combined_multi_scenario_file.write(new_combined_file)  # create a new file
                        # print("Combined_spec_file", new_combined_file)
                    else:
                        new_combined_file = str(replace_spec)
                        combined_multi_scenario_file.write(new_combined_file)  # append with already existing file
                        # print("Combined_spec_file", new_combined_file)
                    outerFolder="testRun"
                    spec_file_name='multiple_scenario_file.js'
                    folderName=testRunName
                    new_conf=provideConf(spec_file_name,base_location,outerFolder,folderName)
                    combined_spec_filename='combined_multi_scenario_file.js'
                    comb_conf = provideConf(combined_spec_filename,base_location,outerFolder,folderName)
                    # START /B cmd "cd base_location\" && protractor base_location\oldFiles\multiple_conf.js > D:\ProtractorWebDriverTest\oldFiles\multiple_conf.log
                    multipleBatchFile = "START /B cmd " + appPath + " && protractor " + appPath + "\\multiple_conf_file.js >" + appPath + "\\upload_log_file.log"
                    comb_batch_file_data = "START /B cmd " + appPath + " && protractor " + appPath + "\\combined_multi_scenario_file.js >" + appPath + "\\upload_log_file.log"

                    new_conf_file.write(new_conf)
                    combined_conf_file.write(comb_conf)

                    new_spec_file.write(new_spec)
                    append_file.write(new_spec)
                    new_conf_file.close()
                    # new_conf_log_file.close()
                    new_spec_file.close()
                    batch_file.write(multipleBatchFile)
                    combined_batch_file.write(comb_batch_file_data)
                    batch_file.close()
                    combined_batch_file.close()
                    combined_multi_scenario_file.close()

                    if flag =="Upload":
                        print("in upload")
                        # subprocess.call(appPath + '\\multiple_batch_file.bat')
                        conf_log = open(appPath + "\\upload_log_file.log").read()
                        logVal = multipletestReportGenerationRun(appPath)
                        time.sleep(1)
                        return logVal
                    elif flag == "Download":
                        print("in download")
                        conf_log = open(appPath + "\\upload_log_file.log").read()
                        logVal = multipletestReportGenerationRun(appPath)
                        time.sleep(1)
                        return logVal

            else:
                return "appFail"

def create_CombinedMultiScenarioFile_SameFile(json_format, testRunName, flag):
    global outerFolderName
    global folderName
    outerFolderName="testRun"
    folderName=testRunName
    sheetLength = 0
    final_val = []
    temp_List = []
    connections = db.testRun.find()
    for connection in connections:
        if connection["test_name"] == testRunName:
            mulipleScenarioPath = connection["path"]
            if os.path.exists(mulipleScenarioPath + "\\job_upload_spec.js"):  # check whether the
                p = os.path.join(mulipleScenarioPath + "\\job_upload_spec.js")
                os.remove(p)

    for k, v in json_format.items():
        v1 = json.dumps(v['data'])
        if type(list(eval(v1))[0]) == "str":
            return jsonify("failure")
        else:
            temp_file = list(eval(v1))
            df = pd.DataFrame.from_dict(temp_file)
        connections = db.testRun.find()
        for connection in connections:
            if connection["test_name"] == testRunName:
                mulipleScenarioPath = connection["path"]
                if os.path.isdir(mulipleScenarioPath):
                    df.to_excel(mulipleScenarioPath + '\\sample.xlsx',
                                index=False)
                    time.sleep(1)
                    append_file = open(mulipleScenarioPath + "\\append_spec.txt", "a+")
                    returnVal = uploadUIRun(testRunName, sheetLength, flag)
                    final_val.append(returnVal)
                    sheetLength += 1
                else:
                    return "appFail"
    final_js = {}
    testResult = [0, 0]
    f_log = ''
    for idx, f_val in enumerate(final_val):
        final_js[str(idx)] = f_val
        testResult = [sum(i) for i in zip(testResult, f_val["testResult"])]
        for log_val in f_val["log"]:
            f_log += log_val
    final_js["testResult"] = testResult
    final_js["log"] = f_log

    comb_multi_file = open(mulipleScenarioPath + "\\combined_multi_scenario_file.js", "a+")
    comb_multi_file.write("});")
    log_filename = "upload_log_file.log"
    conf_log = conf_log_modification(log_filename)
    output_json = {"status": "Success", "conf_log": conf_log}
    final_js["consolidated_log"] = conf_log
    return final_js

@app.route("/multipletestReportGenerationRun")  # to generate report for test job
def multipletestReportGenerationRun(scenarioPath):
    # folderName = request.args.get("folderName")
    print(scenarioPath)
    open_file = open(scenarioPath + "\\upload_log_file.log")  # open log file
    log_file = open(scenarioPath + "\\upload_log_file.log").read()  # read the log file
    read_f = (open_file.readlines())
    passStr = '[32m.[0m'  # pass string in log file
    failStr = '[31mF[0m'  # fail string in log file
    resultList = []
    header = []
    header1 = []
    data = []
    js = {}
    # passStr=".F."
    for wholeStr in read_f:  # iterate each line
        if (wholeStr.__contains__(passStr) or wholeStr.__contains__(
                failStr)):  # ceck whether the line contains pass or fail string
            outputLine = [(wholeStr[i:i + 10]) for i in range(0, len(wholeStr), 10)]
            for result in outputLine:
                if result == passStr:
                    resultList.append('Pass')  # if the line contains pass, append it in pass list
                elif result == failStr:
                    resultList.append('Fail')  # if the line contains fail,append it in fail list

        # if(r.__contains__(failStr)):
    print(resultList)
    file_name = scenarioPath + "\\sample.xlsx"  # Path to your file
    df = pd.read_excel(file_name)
    # df["result"]=resultList[0]
    r, c = df.shape
    df['result'] = ""
    ind = 0
    for row in range(r):
        df["result"][row] = resultList[ind]  # create a key called result and append the result
        ind += 1
    df.to_excel(scenarioPath + "\\new_excel_sheet.xlsx", index=False)  # convert dataframe to excel
    x = df.head()
    time.sleep(1)
    excel_sheet = pd.read_excel(scenarioPath + "\\new_excel_sheet.xlsx")
    js["data"] = excel_sheet.to_dict(orient="records")  # convert pandas values to dictionary
    for colName in excel_sheet.head():
        header.append(colName)  # append header in header list
    # for h in header:
    js["headers"] = header
    totPass = resultList.count('Pass')  # total number of pass case
    totFail = resultList.count('Fail')  # total number of fail case
    testResult = [totPass, totFail]
    js["testResult"] = testResult;
    js["log"] = log_file
    return js


@app.route("/editTestDescription")  # update the edited test description in ui
def editTestDescription():
    appName = request.args.get("appName")  # get application name from UI
    scenarioName = request.args.get("scenarioName")  # get scenario name from UI
    testDescription = request.args.get("testDescription")  # get test Description
    testDescription = testDescription.replace('HASH','#')
    myquery = {"appName": appName, "scenarioName": scenarioName}  # make a json with respective headers
    newvalues = {"$set": {"test_Description": testDescription}}  # db query for updating testDescription field
    db.folderName.update_one(myquery, newvalues)  # update in db
    return jsonify("Success")  # return success in json format


@app.route("/deleteScenario")  # Delete selected scenario
def deleteScenario():
    application_name = request.args.get("application_name")  # get application name from user
    scenario_name = request.args.get("scenario_name")  # get scenario name from user
    my_query = {"appName": application_name,
                "scenarioName": scenario_name}  # put application name and scenario name in a json
    connections = db.folderName.find()  # find collection data in db
    for connection in connections:  # iterate through collection data
        if connection["appName"] == application_name and connection[
            "scenarioName"] == scenario_name:  # if application name and scenario name is present in db
            path = connection["path"]  # assign data
            innerPath = connection["innerpath"]  # assign innerpath's value to the variable
            shutil.rmtree(innerPath, ignore_errors=True)  # delete an entire directory tree
    db.folderName.delete_many(my_query)  # delete all records based on query from db
    return jsonify("success")  # return success


@app.route("/moveImage")  # Move Screenshot from our folder to UI
def moveImage():
    application_name = request.args.get("application_name")  # get application name from UI
    scenario_name = request.args.get("scenario_name")  # get scenario name from UI
    connections = db.folderName.find()  # find collection in db
    fileName = []  # initialize a list
    finalJson = {}  # initialize a dictionary
    destination_folder = ui_location + "\\assets\\images\\" + application_name +"\\"+scenario_name  # path of image
    if os.path.isdir(destination_folder) == False:  # if image not in destination folder
        os.makedirs(destination_folder)  # make a destination folder
    for connection in connections:  # iterate through db collection
        if connection["appName"] == application_name and connection[
            "scenarioName"] == scenario_name:  # if apllication name and scenario name is present in db
            appPath = connection["path"]  # assign value of path to appPath
            scenarioPath = connection["innerpath"]  # assign value of innerpath to scenarioPath
            # destination_folder="D:\\TestAutomationTool\\src\\assets\\images"
            print(appPath, 'appPath')
            print(scenarioPath, 'scenarioPath')
            if os.path.isdir(appPath):  # if appPath is in directory
                # print(appPath, 'appPath1')
                if os.path.isdir(scenarioPath):  # if scenarioPath is available in directory
                    # print(scenarioPath, 'scenarioPath1')
                    for img in os.listdir(scenarioPath):  # iterate through list of directories
                        if img.endswith(".png"):  # if image file has extension .png
                            fileName.append(img)  # append the image to the filename
                            print(img)
                            shutil.copy(scenarioPath + "\\" + img, destination_folder)  # copy the files to UI location
                    finalJson["fileName"] = fileName  # assign the filename to a json
                    finalJson["status"] = "success"  # assign success as a value to the json
                    return jsonify(finalJson)  # return the json
                else:
                    return "scenarioFail"  # return string
            else:
                return "appFail"  # return string


@app.route("/exportScenarioAsZip")  # export the recorded scenario files inthe form of zip
def exportScenarioAsZip():
    app_name = request.args.get("appName")
    scenario_name = request.args.get("scenarioName")
    if os.path.isdir(
            base_location + "\\zip_files\\" + scenario_name == False):  # check whehter the directory present or not
        os.mkdir(base_location + "\\zip_files\\" + scenario_name)  # if not create a directory
    output_filename = base_location + "\\zip_files\\" + scenario_name
    dir_name = base_location + "\\" + app_name + "\\" + scenario_name
    archive_format = "zip"
    collections = db["folderName"].find(
        {"$and": [{"appName": app_name}, {"scenarioName": scenario_name}]})  # get only the particular record
    js_struct = {}
    js_file = open(base_location + "\\" + app_name + "\\" + scenario_name + "\\" + 'json_file.json', "w+")
    for collection in collections:
        print("collection", collection)
        js_struct["appName"] = app_name
        js_struct["scenarioName"] = scenario_name
        js_struct["test_Description"] = collection["test_Description"]
        js_struct["path"] = collection['path']
        js_struct["innerpath"] = collection["innerpath"]
        print(js_struct)
        js_file.write(json.dumps(js_struct))
        js_file.close()
        break
    shutil.make_archive(output_filename, archive_format, dir_name)  # create a zip file using make archieve
    destination_folder = ui_location + "\\assets\\zip_files\\"

    shutil.copy(output_filename + '.zip', destination_folder)  # copy each file, even if it is in use
    # time.sleep(2)
    return {"status": "Success"}, 200


@app.route("/conf_log_modification_show")
def conf_log_modification_show():
    try:


        # Time
        t = time.localtime()
        current_time = time.strftime("%H:%M:%S", t)

        # Date
        now = datetime.now()
        today = date.today()
        d1 = today.strftime("%d/%m/%Y")
        global outerFolderName
        global folderName

        # open_file = open(base_location+"\\" + outerFolderName + "\\" + folderName + "\\" +"conf.log", "r")
        # open_conf_file = open(base_location+"\\" + outerFolderName + "\\" + folderName + "\\" +"new_conf_log.log", "a")
        # buf = open_file.readlines()
        final_report = {}
        report1 = {}
        message = []
        final_list = []
        spec_line = []
        report1['Date'] = d1
        report1["Time"] = current_time
        report1["Application Name"] = outerFolderName
        report1["Scenario Name"] = folderName
        # for lines in buf:
        #     # print("line", lines)
        #     if lines.__contains__("[32m."):
        #         report1['Status'] = "passed"
        #     if lines.__contains__("[31mF"):
        #         report1['Status'] = "failed"
        #
        #     if lines.__contains__("Finished in"):
        #         report1['Time in Seconds'] = lines.split()[2].strip()
        #     if lines.__contains__("instance(s) of WebDriver"):
        #         report1['WebDriver Instance'] = lines.split("-")[1].strip()
        #     if lines.__contains__("31m") and (lines.__contains__("Failed") or lines.__contains__("Expected")):
        #         # print("filaure ", line)
        #         message.append(lines.split("31m")[1].strip())
        #     if lines.__contains__("spec.js"):
        #         num = lines.split(":")[2].strip()
        #         if int(num) > 6:
        #             spec_line.append(num)
        #
        # report1['Error Message'] = message
        # report1['Error lines in spec'] = spec_line
        # report1['Test Description']= testDescription()
        # # open_conf_file.close()
        # modified_report = str(report1).replace(",", ",\n")
        # open_conf_file.write(modified_report + ",\n\n")
        # open_conf_file.close()
        open_conf = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\\" + "new_conf_log.log")
        read_conf = open_conf.read().strip().strip("\n")
        full_dict = []
        for d in read_conf.split("},"):
            if d.strip() != "":
                dct = eval(d + "}")
                print("dct", dct)
                full_dict.append(dct)
        print("full dictionary", full_dict)

        final_report['headers'] = ['Date', 'Time', 'Application Name', 'Scenario Name',
                                   'Error lines in spec', 'Error Message', 'Status', 'Time in Seconds',
                                   'WebDriver Instance', 'Actions']

        # final_list.append(b)
        final_report['data'] = full_dic

        print(final_report)
        return final_report
    except Exception as e:
        return jsonify(e)


@app.route('/multioutputreport', methods=['GET'])
def multioutputreport():  # gets called when download Consolidated Log button is clicked
    try:
        global outerFolderName
        global folderName
        outerFolderName = request.args.get("outerFolderName")
        folderName = request.args.get("folderName")
        inputdata = modify_conf()
        path = base_location + "\\" + outerFolderName + "\\" + folderName
        pathconf = base_location + "\\" + outerFolderName + "\\" + folderName + "\\" + "conf.log"
        destination_folder = ui_location + "\\assets\\doc\\"
        screenshot = path + "\\" + "screenshots"
        # try:
        document = docx.Document()
        heading1 = document.add_heading('Test Report', level=1)
        heading1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        imageno = 1
        errorstepno1 = 3
        errorstepno = errorstepno1
        for indexno in range(len(inputdata['data'])):
            date = document.add_paragraph(inputdata["data"][indexno]["Date"])
            date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            time = document.add_paragraph(inputdata["data"][indexno]["Time"])
            time.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            document.add_paragraph('Application name: ' + inputdata["data"][indexno]["Application Name"])
            document.add_paragraph('Scenario name: ' + inputdata["data"][indexno]["Scenario Name"])
            document.add_paragraph('Test case result : ' + inputdata["data"][indexno]["Status"])
            document.add_paragraph('Time of execution(in secs): ' + inputdata["data"][indexno]["Time in Seconds"])
            Lines = open(pathconf, encoding="utf8").readlines()  # reads file
            total = "0"
            passt = "0"
            failures = "0"
            for line in Lines:
                if line.strip().__contains__("spec,") and line.strip().__contains__("failure"):
                    total = line.split("spec,")[0].strip()
                    failures = line.split("failure")[0].split("spec,")[1].strip()
                    passt = str(eval(total) - eval(failures))

            document.add_paragraph('Total number of tests: ' + total)
            document.add_paragraph('Number of test case passed: ' + passt)
            document.add_paragraph('Number of test case failed: ' + failures)
            listofsteps = inputdata["data"][indexno]["Test Description"].replace("<br>", " \n ", 1).split("<br>")
            listofsteps = [item for item in listofsteps if item != '']
            document.add_paragraph('Number of steps present : ' + str(len(listofsteps) + 1))
            document.add_paragraph(
                'Instances running in webdriver manager: ' + inputdata["data"][indexno]["WebDriver Instance"])
            document.add_heading('Test Description : ', level=3)
            for index, step in enumerate(listofsteps):
                # document.add_paragraph(step[step.find(".")+1:].strip(), style='List Number')
                document.add_paragraph(step.strip())
                if index == errorstepno and inputdata["data"][indexno]["Status"].lower().__contains__("fail"):
                    # em = [escape_ansi(item) for item in inputdata["data"][indexno]["Error Message"]]
                    em = [item.encode('unicode_escape').decode() for item in
                          inputdata["data"][indexno]["Error Message"]]
                    document.add_paragraph(str("\n".join(em)))
                if os.path.isfile(path + '\\page' + str(imageno) + '.png') and index != 0:
                    if not os.path.exists(screenshot):
                        os.mkdir(screenshot)
                    temp = destination_folder + "\\" + inputdata["data"][indexno]["Application Name"] + "\\" + \
                           inputdata["data"][indexno]["Scenario Name"]
                    if not os.path.exists(temp):
                        os.makedirs(temp)
                    shutil.copy(path + '\\page' + str(imageno) + '.png', screenshot)
                    shutil.copy(path + '\\page' + str(imageno) + '.png', temp)
                    destpath = screenshot + '\\page' + str(imageno) + '.png'
                    document.add_picture(destpath, width=Inches(5.0))
                imageno = imageno + 1

        # destination_folder = "C:\\Users\\kn00636678\\Downloads\\"
        # destination_folder = "D:\\iTAP\\iTAP_UI\\src\\assets\\doc\\"
        document.save(destination_folder + "ConsolidatedReport.docx")
        return jsonify("success")
    except Exception as e:
        return jsonify(e)


def modify_conf():
    try:
        # list1=[]
        finaList = []
        open_file = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\\" + "new_conf_log.log", "r")
        buf = open_file.read()
        list1 = buf.split("},")
        del list1[-1]
        outDict = {}
        headers = []
        list2 = []
        for index, line in enumerate(list1):
            if len(list1) > index:
                line += '}'
            print(line)
            dictionary = eval(line)
            list2.append(dictionary)
            dictionary = {}

            for key, value in dictionary.items():
                headers.append(key)
        outDict['headers'] = ['Date', 'Time', 'Application Name', 'Scenario Name', 'Status', 'Time in Seconds',
                              'WebDriver Instance', 'Error Message', 'Error lines in spec', 'Test Description']
        outDict['data'] = list2
        finaList.append(outDict)
        print(outDict)
        return outDict
    except Exception as e:
        return jsonify(e)


@app.route("/scheduledJobList")  # View all the scheduled job list, gets called when browser compatablility test button is clicked
def scheduledJobList():
    cursor = db.testRun.find()
    totList = []
    for collection in cursor:
        totList.append(collection["test_name"])  # returns all the test job files
    return jsonify(totList)


@app.route("/getAllScenarios")  # Get all saved scenario, called when add job button is clicked
def getAllScenarios():
    collections = db.folderName.find()
    string = []
    for collection in collections:
        string.append(collection['scenarioName'] + "[" + collection["appName"] + "]")

    return jsonify(string)
    # return jsonify(["Deposit_flow[XYZ_Bank]","Withdraw_flow[XYZ_Bank]", "OpenAccount_flow[XYZ_Bank]", "CreateOrderFlow[CRST]", "Download_dashboard_charts_Flow[LCaaS]", "MasterInvFlow[LCaaS]"])


@app.route("/runSavedTestRun")  # run already exist test job, called when play button in test job is clicked
def runSavedTestRun():
    global folderName
    global outerFolderName
    testRunName = request.args.get("testName")
    folderName=testRunName
    outerFolderName="testRun"
    subprocess.call(base_location + '\\' + "testRun" + '\\' + testRunName + '\\batch.bat')
    log_filename = "conf.log"
    conf_log = conf_log_modification(log_filename)
    output_json = {"status": "Success", "conf_log": conf_log}
    print("output log", output_json)
    moveReport(outerFolderName, folderName)
    return jsonify(output_json)


def download_testjob(testRunName):
    flag = 0
    collections = db.testRun.find()
    complete_list = []
    complete_json_format = {}
    for collection in collections:
        if collection["test_name"] == testRunName:
            scenarios = collection["test_scenarios_used"]
            path = collection["path"]
            flag = 1
    if flag == 1:
        split_scenario = scenarios.split(",")
        for app_scenario in split_scenario:
            app_name = app_scenario.split("[")[1].split("]")[0]
            scenario_name = app_scenario.split("[")[0]
            spec_file = open(
                base_location + "\\" + app_name.strip() + '\\' + scenario_name.strip() + "\\spec.js")
            headers = []
            values = []
            open_spec = spec_file.read()
            json_format = {}
            final_json = {}
            headerList = []
            final_list = []
            split_file = open_spec.split("browser.driver.manage().window().maximize();")[1].split(";")
            for elm in split_file:
                if (elm.find("element") >= 0):
                    if (elm.find("sendKeys") >= 0):
                        fin_elm = elm.split("xpath")[1].split(".sendKeys")[0]
                        headers.append(fin_elm)
                        # fin_val = elm.split("sendKeys")[1].split("'")[1]
                        if elm.split("sendKeys")[1].__contains__("'"):
                            fin_val = elm.split("sendKeys")[1].split("'")[1]
                        elif elm.split("sendKeys")[1].__contains__('"'):
                            fin_val = elm.split("sendKeys")[1].split('"')[1]
                        values.append(fin_val)
            workbook = xlwt.Workbook()
            sheet = workbook.add_sheet(testRunName)
            # Specifying style
            header_style = xlwt.easyxf('font: bold 1')
            val_style = xlwt.easyxf('font: bold 0')

            for index in range(len(headers)):
                if headers[index] not in headerList:
                    headerList.append(headers[index])
                json_format[headers[index]] = values[index]
            final_list.append(json_format)
            # final_list_1 = final_list[::-1]
            final_json = {"headers": headerList, "data": final_list}
            complete_list.append(final_json)
            final_json = {}
            headerList = []
            final_list_1 = []
            final_list = []

        for index, val in enumerate(complete_list):
            complete_json_format[str(index + 1)] = val
        return complete_json_format
    else:
        return "appFail"


@app.route("/download_run_scenario")  # download selected job
def download_run_scenario():
    # folderName = request.args.get("folderName")
    testRunName = request.args.get("testName")
    op_download_testjob = download_testjob(testRunName)

    return op_download_testjob


@app.route("/generateAndMoveDoc")  # called when Download report button is clicked
def generateAndMoveDoc():
    jobName = request.args.get("testRunName")
    screenshotFlag = request.args.get("screenshotFlag")
    print(screenshotFlag, 'screenshotFlag')
    op_testjob_json = download_testjob(jobName)
    print("op_testjob_json", op_testjob_json)
    generateSpecForUploadedJob(op_testjob_json, jobName)
    create_CombinedMultiScenarioFile_SameFile(op_testjob_json, jobName, "Download")
    filePath = base_location + "\\" + "testRun" + "\\" + jobName
    # multipletestReportGenerationRun(filePath)
    group_jsons = []
    overviewDetailsJson = getOverViewDetailsforTestJobDoc(jobName)
    group_jsons.append(overviewDetailsJson)
    metricDetailJson = getMetricsDetailsforTestJobDoc(jobName)
    group_jsons.append(metricDetailJson)
    environmentDetailJson = getEnvironmentDetailsforTestJobDoc(jobName)
    group_jsons.append(environmentDetailJson)
    screenshotAndLogDetailJson = getLogsAndScreenshotsforTestJobDoc(jobName, screenshotFlag)
    group_jsons.append(screenshotAndLogDetailJson)
    document_detials = combineReturnJsons(group_jsons)
    print('Document Generation Data: ', document_detials)

    # call document generation method to generate

    DocWrite(document_detials)

    destination_folder = ui_location + "\\assets\\doc\\"
    # Frame the string of file name
    fileName = "Detailed Test Report " + jobName + '.docx'
    # Get all data in test job database
    connections = db.testRun.find()
    # Loop through cursor object
    for connection in connections:
        # Find current Job record
        if connection["test_name"] == jobName:
            # If found, check if document exists
            if os.path.exists(connection["path"] + "\\" + fileName):
                # If exists, Copy file to UI folder & Return Success
                print("1088 line",connection["path"] + "\\" + fileName)
                shutil.copy(connection["path"] + "\\" + fileName, destination_folder)
                return jsonify("Success")
            else:
                # If doesn't exist, Return Failure
                return jsonify("File Not Generated")


def getOverViewDetailsforTestJobDoc(jobName):
    '''
    Get section 1 for generating Test Job Document
    :param jobName: test job name
    :return: json object with section 1 data
    '''
    final_json = {}
    # add job name to json
    final_json["test_job_name"] = jobName
    print(jobName)
    # find all records in job table
    collections = db.testRun.find()
    # loop through cursor
    for collection in collections:
        # find current job record
        if collection["test_name"] == jobName:
            # get scenarios used and path and exit loop
            test_scenarios_used = collection["test_scenarios_used"]
            jobPath = collection["path"]
            break
    # split and find scenarios used and add to final json
    scenarios_used = test_scenarios_used.split(',')
    final_json["test_scenarios_used"] = scenarios_used
    final_json["test_description"] = []
    # loop through scenarios used and do the following
    for scenario_used in scenarios_used:
        test_desc = {}
        # get application name and scenario name for each record
        scenario_name = scenario_used.split('[')[0]
        application_name = scenario_used.split('[')[1][:-1]
        test_desc["scenario_name"] = scenario_name
        test_desc["application_name"] = application_name
        # get all scenarios from scenario table
        scenario_collections = db.folderName.find()
        # for every scenario in job, do the following
        for scenario_collection in scenario_collections:
            # find current loop record
            if scenario_collection["scenarioName"] == scenario_name and \
                    scenario_collection["appName"] == application_name:
                # find test description and add to temp json and continue for next record
                test_description = scenario_collection["test_Description"]
                test_desc["test_description"] = test_description
                break
        # add test steps to final json
        final_json["test_description"].append(test_desc)
        test_desc = {}
    # get report location and add to final json
    final_json["report_loc"] = jobPath + "\\new_excel_sheet.xlsx"
    print('Overview: ', final_json)
    # return final json
    return final_json


def getMetricsDetailsforTestJobDoc(jobName):
    '''
  Get Section 2 to generate Test Job Report Document
  :param jobName: test job name
  :return: json object with dataset
  '''
    final_json = {}
    final_json["test_job_name"] = jobName
    # find all records from test job table
    collections = db.testRun.find()
    # loop through the cursor object, do the following
    for collection in collections:
        # find the corresponding job record
        if collection["test_name"] == jobName:
            # if found, find path
            path = collection["path"]
            # find last modified time of conf.log to find when the batch was run last
            mod_time = time.ctime(os.path.getmtime(path + "\\conf.log"))
            # add the last modified time to final json and exit loop
            final_json["last_run"] = "last modified: %s" % time.ctime(os.path.getmtime(path + "\\conf.log"))
            break
    # read config log generated in path
    confLog = open(path + "\\upload_log_file.log").readlines()
    # for every line, do following
    for line in confLog:
        # find Finished time (in sections)
        if line.__contains__("Finished in"):
            time_secs = line.split(" ")[2]
            # add time taken to final json and break loop
            t_min = str(strftime("%H:%M:%S", gmtime(float(line.split(" ")[2])))).split(":")[1]
            t_secs = str(strftime("%H:%M:%S", gmtime(float(line.split(" ")[2])))).split(":")[2]
            final_json["time_elapsed"] = line.split(" ")[2].split(".")[
                                             0] + " seconds " + "[" + t_min + "m," + t_secs + "s]"
            break
    final_json["time_stamp"] = []
    time_stamp = {}
    # get end time from last modified time
    print(final_json["last_run"])
    time_stamp["end_time"] = final_json["last_run"].split(": ")[1]
    print(mod_time)
    last_run_time = mod_time.strip().split(" ")[3].strip()
    # to find start time, subtract with time taken
    # sub_min = str(strftime("%H:%M:%S", gmtime(float(time_secs)))).split(":")[1]
    sub_min = "10"
    # sub_secs = str(strftime("%H:%M:%S", gmtime(float(time_secs)))).split(":")[2]
    sub_secs = "2"
    FMT = '%H:%M:%S'
    print(last_run_time)
    tdelta = datetime.strptime(last_run_time, FMT) - timedelta(minutes=int(sub_min), seconds=int(sub_secs))
    # add start time to temp json
    time_stamp["start_time"] = mod_time.replace(last_run_time, str(tdelta).split(" ")[1])
    # add time stamp to final json
    final_json["time_stamp"].append(time_stamp)
    # initialize pass count and fail count
    passCount = 0
    failCount = 0
    # load workbook to calculate total pass & fail
    wb = load_workbook(path + "\\new_excel_sheet.xlsx")
    # find number of sheets
    tot_sheets = len(wb.worksheets)
    tot_runs = 0
    # for every sheet, do the following
    for sheet in wb.worksheets:
        # find tot rows and tot cols in sheet
        tot_rows = sheet.max_row
        # calculate tot runs with rows count
        tot_runs += tot_rows
        column_count = sheet.max_column
        # for every row, do the following
        for row in sheet.iter_rows(1, tot_rows):
            # find last column value and increase count for pass and fail
            if row[column_count - 1].value == 'Pass':
                passCount += 1
            if row[column_count - 1].value == 'Fail':
                failCount += 1
    # add pass count and fail count to final json
    testResult = {"pass": passCount, "fail": failCount}
    final_json["detailed_result"] = testResult
    # calculate avg
    avg = tot_sheets / (tot_runs - tot_sheets)
    # add avg to final json
    final_json["avgerage_time"] = avg
    print('Metrics: ',final_json)
    # return final json
    return final_json


def getEnvironmentDetailsforTestJobDoc(jobName):
    '''
  Get Section 3 to populate Test Job Document
  :param jobName: test job name
  :return: final json with data
  '''
    final_json = {}
    final_json["test_job_name"] = jobName
    # get all test job records
    collections = db.testRun.find()
    # loop through cursor object
    for collection in collections:
        # find current job name
        if collection["test_name"] == jobName:
            # get path of the job and exit loop
            path = collection["path"]
            break
    # with path information, read the spec file
    jsFile = open(path + "\\" +"spec.js").readlines()
    urls_unedited = []
    # for each line in spec js file, do the following
    for line in jsFile:
        # find all urls used
        if line.__contains__("browser.get("):
            # append the value to list
            urls_unedited.append(line)
    urls = []
    print("urls undedited", urls_unedited)
    # loop through lines, and extract urls
    for url_unedited in urls_unedited:
        urls.append(url_unedited.strip().replace('\n', ''))
    # append urls to final json
    print("urls", urls)
    final_json["urls"] = urls
    # read config file to find browser used
    confFile = open(path + "\\conf.js").readlines()
    final_json["browser"] = []
    # for every line in config file, do the following
    for line in confFile:
        # split the required browser name and save to final json
        if line.__contains__("browserName"):
            final_json["browser"].append(line.split(':')[1].split('\"')[1].split('\"')[0])
            break
    # print('Environment: ',final_json)
    # return final json
    return final_json


def getLogsAndScreenshotsforTestJobDoc(jobName, screenshotFlag):
    '''
  Get section 4 for populating test job document
  :param jobName: Test Job Name, for which information has to be fetched
  :param screenshotFlag: Boolean - True / False to get screenshots on user's details
  :return: json object with data
  '''
    final_json = {}
    final_json["test_job_name"] = jobName
    # Only when user requests for screenshot to be added.
    if screenshotFlag == True or screenshotFlag == 'true':
        final_json["screenshots"] = []
        # Find all records from test job table
        collections = db.testRun.find()
        # loop through the cursor object
        for collection in collections:
            # find current job record
            if collection["test_name"] == jobName:
                # if found, get path and scnearios used
                appPath = collection["path"]
                scenario = collection["test_scenarios_used"]
                split_scenario = scenario.split(",")

                print("split scenario in ss log", split_scenario)
                # for ever scenario used, do the following
                for app_scene in split_scenario:
                    # get application name and scenario name from the string
                    application_name = app_scene.split("[")[1].split("]")[0]
                    scenario_name = app_scene.split("[")[0]
                    # get all scenarios from scenario table
                    scnearios_connections = db.folderName.find()
                    # loop through the cursor object
                    for scneario_connections in scnearios_connections:
                        fileName = []
                        # find corresponding record by matching app name and scenario name
                        if scneario_connections["appName"] == application_name \
                                and scneario_connections["scenarioName"] == scenario_name:
                            # get path and innerpath
                            s_appPath = scneario_connections["path"]
                            scenarioPath = scneario_connections["innerpath"]
                            # check if folder exists
                            if os.path.isdir(s_appPath):
                                # check if inner folder exists
                                if os.path.isdir(scenarioPath):
                                    # find all images (.png) files
                                    for img in os.listdir(scenarioPath):
                                        if img.endswith(".png"):
                                            fileName.append(img)
                                    # save image names with path to final json
                                    list_names = {s_appPath + '\\' + scenario_name: fileName}
                                    final_json["screenshots"].append(list_names)
    else:
        # if user doesn't want screenshots in report
        # find current record in test run table and get its path
        collections = db.testRun.find()
        for collection in collections:
            if collection["test_name"] == jobName:
                appPath = collection["path"]
    # finally get the conf log details and save to final json
    conf_log = open(appPath + "\\conf.log", "r").readlines()
    final_json["log"] = ''.join(conf_log)
    # print(final_json)
    # return final json
    return final_json


def combineReturnJsons(jsonList):
    '''
  Combine list of jsons / dicts to a single json / dict
  :param jsonList: list of json / dict objects
  :return: final json / dict value
  '''
    document_details = {}
    # loop through the list of dict
    for json_dict in jsonList:
        # for ever key in json, do the following
        for k in json_dict:
            # find key's value and append to final json
            document_details[k] = json_dict[k]
    # after processing, return the final json
    return document_details


def DocWrite(TestJobDetails):
    Mergeval = TestJobDetails

    if "screenshots" in Mergeval.keys():
        # Define the templates - assumes they are in the same directory as the code
        template_path = base_location + "\\" + "Detailed Test Report - with SS.docx"
        pythoncom.CoInitialize()
        word = win32.gencache.EnsureDispatch('Word.Application')
        pythoncom.CoInitialize()
        doc = word.Documents.Open(template_path)
        word.Visible = False

        try:
            if os.path.exists(Mergeval['report_loc']):
                doc.InlineShapes.AddOLEObject(FileName=Mergeval['report_loc'], DisplayAsIcon=1,
                                              Range=doc.Paragraphs(47).Range, IconLabel="Individual Data Set Report",
                                              IconFileName="Individual Data Set Report",
                                              LinkToFile=False)

            word.ActiveDocument.SaveAs(
                base_location + "\\" + Mergeval['test_job_name'] + ".docx")
            doc.Close()
        except Exception as e:
            print(e)

    else:
        # Define the templates - assumes they are in the same directory as the code
        template_path = base_location + "\\" + "Detailed Test Report.docx"
        pythoncom.CoInitialize()
        word = win32.gencache.EnsureDispatch('Word.Application')
        pythoncom.CoInitialize()
        doc = word.Documents.Open(template_path)
        word.Visible = False

        try:
            if os.path.exists(Mergeval['report_loc']):
                doc.InlineShapes.AddOLEObject(FileName=Mergeval['report_loc'], DisplayAsIcon=1,
                                              Range=doc.Paragraphs(46).Range, IconLabel="Report",
                                              IconFileName="Individual Data Set Report",
                                              LinkToFile=False)

            word.ActiveDocument.SaveAs(
                base_location + "\\" + Mergeval['test_job_name'] + ".docx")
            doc.Close()
        except Exception as e:
            print(e)

    cleaned_string = ''.join(c for c in Mergeval['log'] if valid_xml_char_ordinal(c))
    test_description = ""

    for item in Mergeval['test_description']:
        print(item['scenario_name'])
        test_description += "\n" + item['scenario_name'] + "\n"
        test_description += item['test_description'].replace("<br>", "\n") + "\n"
    # Show a simple example
    document_1 = MailMerge(base_location + "\\" + Mergeval['test_job_name'] + ".docx")
    print("Fields included in {}: {}".format(template_path,
                                             document_1.get_merge_fields()))
    print("Mergevale", Mergeval)
    Separator = "\n"
    document_1.merge(Date='{:%d-%b-%Y}'.format(date.today()),
                     TestJobName=str(Mergeval['test_job_name']),
                     Scenario=str(Separator.join(Mergeval['test_scenarios_used'])),
                     ScenarioDetails=test_description,
                     LastRun=str(Mergeval['last_run']),
                     StartTime=str(Mergeval['time_stamp'][0]['start_time']),
                     EndTime=Mergeval['time_stamp'][0]['end_time'],
                     TimeElapsed=Mergeval['time_elapsed'],
                     TestCasePassed=str(Mergeval['detailed_result']['pass']),
                     TestCaseFailed=str(Mergeval['detailed_result']['fail']),
                     URL=str(Separator.join(Mergeval['urls'])),
                     Browser=str(Separator.join(Mergeval['browser'])),
                     Logs=cleaned_string
                     )
    if "screenshots" in Mergeval.keys():

        test_desc_mapping_list = []

        for item in Mergeval['screenshots']:
            print("starting from here")
            print("item",item)
            for k in item:
                print("hereee",item[k])
                for v in item[k]:
                    print("v is",v)
                    folderName = k.split("\\")[-1]
                    outerFolderName = k.split("\\")[-2]
                    if os.path.exists(k + "\\" + v):
                        catch_it = testDescMapping(outerFolderName, folderName)
                        print("catchit from testDescMapping",catch_it)
                        print(len(catch_it))
                        test_desc_mapping_list.append(catch_it)
                        break
        print("test_desc_mapping_list",test_desc_mapping_list)

        tDesc = []
        for item in test_desc_mapping_list:
            for ite in item:
                tDesc.append(ite)
        print("tDesc",tDesc)

        document_1.merge_rows('TestStep', tDesc)

        fileName = "Detailed Test Report " + Mergeval['test_job_name'] + '.docx'
        # Get all data in test job database
        connections = db.testRun.find()
        # Loop through cursor object
        for connection in connections:
            # Find current Job record
            if connection["test_name"] == Mergeval['test_job_name']:
                # If found, check if document exists
                document_1.write(connection["path"] + "\\" + fileName)
                doc = Document(connection["path"] + "\\" + fileName)
                # document_1.write(base_location + "\\Detailed Test Report " + Mergeval['test_job_name'] + ".docx")
                # doc = Document(base_location + "\\Detailed Test Report " + Mergeval['test_job_name'] + ".docx")
                tables = doc.tables

                # outerFolderName = Mergeval["test_scenarios_used"][0].split("[")[1].strip().replace("]", "")

                iteration_num = len(tDesc)
                print(iteration_num)
                # iteration_num = iteration_num+ 1
                num = 1
                for catch_it in test_desc_mapping_list:
                    print("catch_it", catch_it)
                    print(catch_it[0])
                    for i in range(len(catch_it)):
                        print(i)
                        folderName = catch_it[0]['ScenarioName']
                        outerFolderName = catch_it[0]['AppName']
                        print(outerFolderName)
                        if num <= iteration_num-1:
                            p = tables[0].rows[num + 1].cells[2].add_paragraph()
                            r = p.add_run()

                            img_loc = base_location + "\\" + outerFolderName + "\\" + folderName + "\\" + 'page' + str(
                                i + 1) + '.png'
                            print("img_loc",img_loc)
                            if os.path.exists(img_loc):
                                r.add_picture(img_loc, width=Inches(3.0), height=Inches(2.0))
                                print("img_loc", img_loc)
                            num += 1
                # doc.save(base_location + "\\Detailed Test Report " + Mergeval['test_job_name'] + ".docx")
                doc.save(connection["path"] + "\\" + fileName)
                break

                # Save the document as example 1
        # document_1.write(base_location + "\\Detailed Test Report " + Mergeval['test_job_name'] + ".docx")
        # doc = Document(base_location + "\\Detailed Test Report " + Mergeval['test_job_name'] + ".docx")
        # tables = doc.tables
        #
        # outerFolderName = Mergeval["test_scenarios_used"][0].split("[")[1].strip().replace("]", "")
        #
        # iteration_num = len(tDesc)
        # print(iteration_num)
        # num = 0
        # for catch_it in test_desc_mapping_list:
        #     for i in range(len(catch_it)):
        #         print(i)
        #         folderName = catch_it[0]['ScenarioName']
        #         if num <= iteration_num:
        #             p = tables[0].rows[num+1].cells[2].add_paragraph()
        #             r = p.add_run()
        #
        #             img_loc = base_location + "\\" + outerFolderName + "\\" + folderName + "\\OldFiles\\" + 'page' + str(
        #                 i + 1) + '.png'
        #             print(img_loc)
        #             if os.path.exists(img_loc):
        #                 r.add_picture(img_loc, width=Inches(3.0), height=Inches(2.0))
        #             num += 1
        # doc.save(base_location + "\\Detailed Test Report " + Mergeval['test_job_name'] + ".docx")
    else:
        fileName = "Detailed Test Report " + Mergeval['test_job_name'] + '.docx'
        # Get all data in test job database
        connections = db.testRun.find()
        # Loop through cursor object
        for connection in connections:
            # Find current Job record
            if connection["test_name"] == Mergeval['test_job_name']:
                # If found, check if document exists
                document_1.write(connection["path"] + "\\" + fileName)
                break
        # document_1.write(base_location + "\\Detailed Test Report " + Mergeval['test_job_name'] + ".docx")

    return "Success"


def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
            0x20 <= codepoint <= 0xD7FF or
            codepoint in (0x9, 0xA, 0xD) or
            0xE000 <= codepoint <= 0xFFFD or
            0x10000 <= codepoint <= 0x10FFFF
    )


def testDescMapping(outerFolderName, folderName):
    open_file = open(base_location + "\\" + outerFolderName + '\\' + folderName + "\\spec.js")

    lines = open_file.readlines()
    full_line = ''
    flag = 0
    # global sentence
    sentence = ''
    for line in lines:
        if line.__contains__("extractModal();"):
            break
        if line.strip().startswith("browser.get(") or line.strip().startswith("browser.takeScreenshot("):
            flag = 1

        if flag == 1:
            full_line += line

    print(full_line)
    sentence_list = []
    number = 0
    for item in full_line.split("browser.takeScreenshot("):
        sentence = ''
        for sline in item.split(";"):

            if sline.__contains__("browser.get("):  # Get Url
                number += 1
                url_line = sline.split("browser.get(")[1].split(")")[0]
                if len(url_line) < 50:
                    sentence += str(number) + ". Open browser and navigate to URL: " + str(
                        url_line) + " <br>"
                else:
                    sentence += str(number) + ". Open browser and navigate to URL: " + "<br>" + str(
                        url_line) + " <br>"

            if sline.__contains__("until.presenceOf") == False and \
                    sline.__contains__("by.xpath(") and \
                    sline.strip().endswith("click()"):  # button click
                if sline.__contains__("@id="):
                    number += 1
                    sentence += str(number) + ". Click on the button with id " + sline.split("@id=")[1].split("]")[
                        0] + "<br>"
                elif sline.__contains__("@name="):
                    number += 1
                    sentence += str(number) + ". Click on the button with name " + sline.split("@name=")[1].split("]")[
                        0] + "<br>"
                elif sline.__contains__(".get("):
                    number += 1
                    x = str(sline.split(".get(")[1].split(")")[0]).strip()
                    if x.endswith("1"):
                        inBetween = str(floor(int(x) / 10)) + "1st"
                    elif x.endswith("2"):
                        inBetween = str(floor(int(x) / 10)) + "2nd"
                    elif x.endswith("3"):

                        inBetween = str(floor(int(x) / 10)) + "3rd"
                    else:
                        inBetween = x + "th"
                    sentence += str(number) + ". Click on " + inBetween + " button in the current page<br>"

            elif sline.__contains__("until.presenceOf") == False and \
                    sline.__contains__("by.xpath(") and \
                    sline.__contains__("click()") == False and \
                    sline.__contains__("sendKeys"):  # Sendkey
                if sline.__contains__("@id="):
                    number += 1
                    if sline.__contains__(".get("):
                        x = str(sline.split(".get(")[1].split(")")[0]).strip()
                        if x.endswith("1"):
                            inBetween = str(floor(int(x) / 10)) + "1st"
                        elif x.endswith("2"):
                            inBetween = str(floor(int(x) / 10)) + "2nd"
                        elif x.endswith("3"):
                            inBetween = str(floor(int(x) / 10)) + "3rd"
                        else:
                            inBetween = x + "th"
                            if sline.lower().__contains__("password"):
                                sentence += str(
                                    number) + ". Enter value " + "******" + " on the" + inBetween + "field with id " + \
                                            sline.split("@id=")[1].split("]")[0] + "<br>"
                            else:
                                sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                                    0] + " on the" + inBetween + "field with id " + sline.split("@id=")[1].split("]")[
                                                0] + "<br>"
                    else:
                        if sline.lower().__contains__("password"):
                            sentence += str(number) + ". Enter value " + "******" + " in the field with id " + \
                                        sline.split("@id=")[1].split("]")[0] + "<br>"
                        else:
                            sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                                0] + " in the field with id " + sline.split("@id=")[1].split("]")[0] + "<br>"
                elif sline.__contains__("@name="):
                    number += 1
                    if sline.__contains__(".get("):
                        x = str(sline.split(".get(")[1].split(")")[0]).strip()
                        if x.endswith("1"):
                            inBetween = str(floor(int(x) / 10)) + "1st"
                        elif x.endswith("2"):
                            inBetween = str(floor(int(x) / 10)) + "2nd"
                        elif x.endswith("3"):
                            inBetween = str(floor(int(x) / 10)) + "3rd"
                        else:
                            inBetween = x + "th"
                            if sline.lower().__contains__("password"):
                                sentence += str(number) + ". Enter value " + "******" + " in the field with name " + \
                                            sline.split("@name=")[1].split("]")[0] + "<br>"
                            else:
                                sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                                    0] + " in the field with name " + sline.split("@name=")[1].split("]")[0] + "<br>"
                    else:
                        if sline.lower().__contains__("password"):
                            sentence += str(number) + ". Enter value " + "*****" + " in the field with name " + \
                                        sline.split("@name=")[1].split("]")[0] + "<br>"
                        else:
                            sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                                0] + " in the field with name " + sline.split("@name=")[1].split("]")[0] + "<br>"
                elif sline.__contains__(".get(") and sline.__contains__("@id=") == False:
                    number += 1
                    x = str(sline.split(".get(")[1].split(")")[0]).strip()
                    if x.endswith("1"):
                        inBetween = str(floor(int(x) / 10)) + "1st"
                    elif x.endswith("2"):
                        inBetween = str(floor(int(x) / 10)) + "2nd"
                    elif x.endswith("3"):
                        inBetween = str(floor(int(x) / 10)) + "3rd"
                    else:
                        inBetween = x + "th"
                    # sentence += "Enter value " + sline.split("sendKeys(")[1].split(")")[0] + " in the field with id " + \
                    #             inBetween+" in the field from the webpage<br>"
                    sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                        0] + " on " + inBetween + " field in the current page<br>"
                elif sline.__contains__("//input"):
                    number += 1
                    sentence += str(number) + ". Entering input as " + sline.split("sendKeys(")[1].split(")")[
                        0] + " in the current page<br>"

        if sentence != "":
            sentence = sentence.replace("<br>", "\n")
            sentence_list.append(sentence.strip())

    myquery = {"appName": outerFolderName, "scenarioName": folderName}
    newvalues = {"$set": {"test_Description": sentence}}
    # db.folderName.update_many(myquery, newvalues)
    print(sentence)
    print(sentence_list)
    print(len(sentence_list))

    final_list = []
    img_no = 0
    for sen in sentence_list:
        sen_dict = {}
        sen_dict["AppName"] = outerFolderName
        sen_dict["ScenarioName"] = folderName
        sen_dict["TestStep"] = sen
        # sen_dict["ScreenShot"] = "Page" + str(img_no+1) + ".png"
        img_no = img_no + 1
        final_list.append(sen_dict)
    print(final_list)
    return final_list

@app.route("/deleteRun")  # delete the specific test job
def deleteRun():
    testRun = request.args.get("testName")
    connections = db.testRun.find()
    myQuery = {"test_name": testRun}
    # "D:\\ProtractorWebDriverTest\\testRun\\" + testRun
    for connection in connections:
        if connection["test_name"] == testRun:
            path = connection["path"]
            shutil.rmtree(path, ignore_errors=True)
    db.testRun.delete_one(myQuery)
    return jsonify("success")


@app.route("/moveRunImage")  # move image from base folder to ui folder
def moveRunImage():
    testRun = request.args.get("testRun")
    fileName = []
    finalJson = {}
    filePath=base_location+"\\"+"testRun"+"\\"+testRun
    if db.testRun.find({"path": filePath}).count() > 0:
        destination_folder = ui_location + "\\assets\\images\\" + testRun
        if os.path.isdir(destination_folder) == False:
            os.makedirs(destination_folder)
        for img in os.listdir(filePath):
            if img.endswith(".png") and img.startswith("page"):
                fileName.append(img)
                shutil.copy(filePath + "\\" + img, destination_folder)
        finalJson["fileName"] = fileName
        finalJson["status"] = "success"
    print(finalJson)
    return jsonify(finalJson)

@app.route("/checkMultiScenarioPath")  # check the scenario path already exist or not
def checkMultiScenarioPath():
    testRunName = request.args.get("testRunName")
    schedulerName = request.args.get("schedulerName")
    if os.path.exists(
            base_location + "\\testRun\\" + testRunName + "\\combined_multi_scenario_file.js"):
        return jsonify({"status": "success"})
    else:
        return jsonify({"status": "failure"})


@app.route("/deleteBatchFile")  # delete the scheduled batch file
def deleteBatchFile():
    schedulerName = request.args.get("schedulerName")
    testRunName = request.args.get("testRunName")
    ScheduleDateTime = request.args.get("ScheduleDateTime")
    my_query = {"test_job_name": testRunName, "scheduled_time": ScheduleDateTime}
    connections = db.scheduledJobs.find()
    for connection in connections:
        if connection["scheduler_name"] == schedulerName and connection["test_job_name"] == testRunName \
                and connection["scheduled_time"] == ScheduleDateTime:
            path = connection["path"]
            os.remove(path + "\\" + schedulerName + "\\scheduleTestRun.bat")
            break
    db.scheduledJobs.delete_many(my_query)
    return jsonify("Success")


@app.route("/checkTestRunExist")  # check all the scenario folders exists in the test run
def checkTestRunExist():  # Check whether folder exists for specific scenario
    value = request.args.get("testName")
    # value="testRun1"
    connections = db.testRun.find()
    flag = 0
    for connection in connections:
        if value == connection["test_name"]:
            # if value == 'testRun1':
            flag = 1
            break

    if flag == 1:
        return jsonify("Fail")
    else:
        return jsonify("Success")


@app.route("/testRunCreation")  # create new test job
def testRunCreation():
    scenarios = request.args.get("scenarios")
    testRunName = request.args.get("testRunName")
    scenario = scenarios
    # .replace(",", ", ")
    if os.path.isdir(base_location + "\\testRun") == False:
        os.mkdir(base_location + "\\testRun")
    os.mkdir(base_location + "\\testRun\\" + testRunName)
    split_scenario = scenarios.split(",")
    dbData = [{
        "test_name": testRunName,
        "test_scenarios_used": scenario,
        "path": base_location + "\\testRun\\" + testRunName
    }]
    db.testRun.insert_many(dbData)
    open(base_location + "\\testRun\\" + testRunName + "\\" + "conf.log", "w+").close()
    log = runSavedTestRun1(testRunName)

    if log == 'Success':
        js = {"status": "success", "log": "Test Job Created.."}
    else:
        js = {"status": "success", "log": log}
    return jsonify(js)


@app.route("/checkSchedulerNameExists")  # check if the scheduled name already exists or not
def checkSchedulerNameExists():
    schedulerName = request.args.get("schedulerName")
    collections = db.scheduledJobs.find()
    flag = False
    if collections.count() == 0:
        return jsonify("Success")
    else:
        for collection in collections:
            if collection['scheduler_name'] != schedulerName:
                flag = True
            else:
                flag = False
                break
        if flag == True:
            return jsonify("Success")
        else:
            return jsonify("Failure")


# @app.route("/ScheduleTestRunWithExcel", methods=['GET', 'POST'])  # upload excel and schedule the job
# def ScheduleTestRunWithExcel():
#     thread = threading.Thread(target=sched)
#     thread.start()
#
#     ScheduleDateTime = request.args.get("ScheduleDateTime")
#     dailyFlag = request.args.get("ScheduleDateTimeFlag")
#     testRunName = request.args.get("testRunName")
#     email = request.args.get("email")
#     schedulerName = request.args.get("schedulerName")
#     file = request.get_data()
#     print('file', file)
#
#     catchit = uploadingExcelInSchedule(schedulerName, file, testRunName)
#
#     print(catchit)
#     with app.app_context():
#
#         ScheduleDate = ScheduleDateTime.split("T")[0]
#         ScheduleTime = ScheduleDateTime.split("T")[1]
#         hours = ScheduleTime.split(":")[0]
#         mins = ScheduleTime.split(":")[1]
#         meridian = "AM"
#         print('hours: ', hours, type(hours), int(hours) < 10)
#         print(ScheduleTime, ScheduleDate, "Date & time")
#         if int(hours) > 12:
#             hours = int(hours) - 12
#             if hours < 10:
#                 hours = '0' + str(hours)
#             meridian = "PM"
#         elif int(hours) < 10 and int(hours) != 0:
#             print(int(hours))
#             hours = '0' + str(hours)
#             meridian = "AM"
#         elif int(hours) == 12:
#             meridian = "PM"
#         ScheduleTime1 = str(hours) + ":" + str(mins)
#         ScheduleDate = ScheduleDate.split("-")[1] + "/" + ScheduleDate.split("-")[2] + "/" + ScheduleDate.split("-")[0]
#         print(ScheduleTime, ScheduleDate)
#
#         connections = db.testRun.find()
#         for connection in connections:
#             if connection["test_name"] == testRunName:
#                 appPath = connection["path"]
#                 appPath = appPath.replace("\\\\", "\\")
#                 if os.path.isdir(appPath + "\\" + schedulerName) == False:
#                     os.mkdir(appPath + "\\" + schedulerName)
#                 if os.path.isdir(appPath + "\\" + schedulerName + "\\combined_conf_file.log") == False:
#                     open(appPath + "\\" + schedulerName + "\\combined_conf_file.log", "w").close()
#                 open(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat", "w").close()
#                 batch_file = open(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat", "w+")
#
#                 if dailyFlag == 'false':
#                     batchScript = f'''
#
#                            START /B cmd "cd {base_location}\\" && protractor {appPath}\\{schedulerName}\\combined_conf_file.js > {appPath}\\{schedulerName}\\combined_conf_file.log
#                        '''
#                     # db_scheduledTime = ScheduleDate + " at " + str(hours) + ":" + str(mins) +" "+ meridian
#                     if int(hours) == 00:
#                         db_scheduledTime = ScheduleDate + " at " + str(int(hours) + 12) + ":" + str(
#                             mins) + " " + meridian
#                     else:
#                         db_scheduledTime = ScheduleDate + " at " + str(hours) + ":" + str(mins) + " " + meridian
#
#                     dbData = [{"scheduler_name": schedulerName,
#                                "test_job_name": testRunName,
#                                "scheduled_time": db_scheduledTime,
#                                "path": appPath}]
#
#                     db.scheduledJobs.insert_many(dbData)
#
#                     batch_file.write(batchScript)
#                     batch_file.close()
#
#                     # job(appPath, schedulerName, ScheduleDate)
#
#                     month, day, year = ScheduleDate.split('/')
#                     born = datetime(int(year), int(month), int(day))
#                     my_day = born.strftime("%A")
#                     print(my_day, ScheduleTime)
#
#                     if my_day == "Monday":
#                         schedule.every().monday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
#                     elif my_day == "Tuesday":
#                         schedule.every().tuesday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
#                     elif my_day == "Wednesday":
#                         schedule.every().wednesday.at(ScheduleTime).do(
#                             lambda: job(appPath, schedulerName, ScheduleDate))
#                     elif my_day == "Thursday":
#                         schedule.every().thursday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
#                     elif my_day == "Friday":
#                         print("entered into this friday")
#
#                         schedule.every().friday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
#                         # schedule.every().friday.at("15:55").do(lambda: job('Hello ', 'world!'))
#                         print("schedule function is called..")
#                     elif my_day == "Saturday":
#                         schedule.every().saturday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
#                     elif my_day == "Sunday":
#                         schedule.every().sunday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
#
#                     # subprocess.call(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")
#                 else:
#
#                     batchScript = f'''
#
#                                      START /B cmd "cd {base_location}\\" && protractor {appPath}\\{schedulerName}\\combined_conf_file.js > {appPath}\\{schedulerName}\\combined_conf_file.log
#
#                          '''
#
#                     # db_scheduledTime = "Everyday at " + str(hours) + ":" + str(mins) + " " + meridian
#                     if int(hours) == 00:
#                         db_scheduledTime = "Everyday at " + str(int(hours) + 12) + ":" + str(mins) + " " + meridian
#                     else:
#                         db_scheduledTime = "Everyday at " + str(hours) + ":" + str(mins) + " " + meridian
#                     dbData = [{"scheduler_name": schedulerName,
#                                "test_job_name": testRunName,
#                                "scheduled_time": db_scheduledTime,
#                                "path": appPath}]
#                     db.scheduledJobs.insert_many(dbData)
#                     batch_file.write(batchScript)
#                     batch_file.close()
#
#                     # daily_job(appPath, schedulerName)
#                     schedule.every().day.at(ScheduleTime).do(lambda: daily_job(appPath, schedulerName))
#                     # subprocess.call(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")
#         send_mail(email, testRunName)
#         return jsonify("Success")


# @app.route("/ScheduleTestRun")  # schedule the test job either daily or specific time
# def ScheduleTestRun():
#     thread = threading.Thread(target=sched)
#     thread.start()
#     with app.app_context():
#         ScheduleDateTime = request.args.get("ScheduleDateTime")
#         dailyFlag = request.args.get("ScheduleDateTimeFlag")
#         testRunName = request.args.get("testRunName")
#         email = request.args.get("email")
#         schedulerName = request.args.get("schedulerName")
#         ScheduleDate = ScheduleDateTime.split("T")[0]
#         ScheduleTime = ScheduleDateTime.split("T")[1]
#         hours = ScheduleTime.split(":")[0]
#         mins = ScheduleTime.split(":")[1]
#         meridian = "AM"
#
#         print('hours: ', hours, type(hours), int(hours) < 10)
#         print(ScheduleTime, ScheduleDate, "Date & time")
#         if int(hours) > 12:
#             hours = int(hours) - 12
#             if hours < 10:
#                 hours = '0' + str(hours)
#             meridian = "PM"
#         elif int(hours) < 10 and int(hours) != 0:
#             print(int(hours))
#             hours = '0' + str(hours)
#             meridian = "AM"
#         elif int(hours) == 12:
#             meridian = "PM"
#         ScheduleTime1 = str(hours) + ":" + str(mins)
#         ScheduleDate = ScheduleDate.split("-")[1] + "/" + ScheduleDate.split("-")[2] + "/" + ScheduleDate.split("-")[0]
#         print(ScheduleTime, ScheduleDate)
#
#         connections = db.testRun.find()
#         for connection in connections:
#             if connection["test_name"] == testRunName:
#                 print("entere")
#                 appPath = connection["path"]
#                 appPath = appPath.replace("\\\\", "\\")
#                 os.mkdir(appPath + "\\" + schedulerName)
#                 open(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat", "w").close()
#                 batch_file = open(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat", "w+")
#
#                 if dailyFlag == 'false':
#                     batchScript = f'''
#                                     START /B cmd cd {base_location}\\ && protractor {appPath}\\conf.js > {appPath}\\conf.log
#                         '''
#                     # db_scheduledTime = ScheduleDate + " at " + str(hours) + ":" + str(mins) +" "+ meridian
#                     if int(hours) == 00:
#                         db_scheduledTime = ScheduleDate + " at " + str(int(hours) + 12) + ":" + str(
#                             mins) + " " + meridian
#                     else:
#                         db_scheduledTime = ScheduleDate + " at " + str(hours) + ":" + str(mins) + " " + meridian
#
#                     dbData = [{"scheduler_name": schedulerName,
#                                "test_job_name": testRunName,
#                                "scheduled_time": db_scheduledTime,
#                                "path": appPath}]
#
#                     db.scheduledJobs.insert_many(dbData)
#
#                     batch_file.write(batchScript)
#                     batch_file.close()
#
#                     # job(appPath, schedulerName, ScheduleDate)
#
#                     month, day, year = (int(i) for i in ScheduleDate.split('/'))
#                     # print(year, month, day)
#                     born = datetime(year, month, day)
#
#                     my_day = born.strftime("%A")
#                     print(my_day)
#
#                     if my_day == "Monday":
#                         schedule.every().monday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
#                     elif my_day == "Tuesday":
#                         schedule.every().tuesday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
#                     elif my_day == "Wednesday":
#                         schedule.every().wednesday.at(ScheduleTime).do(
#                             lambda: job(appPath, schedulerName, ScheduleDate))
#                     elif my_day == "Thursday":
#                         schedule.every().thursday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
#                     elif my_day == "Friday":
#                         schedule.every().friday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
#                     elif my_day == "Saturday":
#                         schedule.every().saturday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
#                     elif my_day == "Sunday":
#                         schedule.every().sunday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
#
#                 else:
#                     batchScript = f'''
#                                     START /B cmd cd {base_location}\\ && protractor {appPath}\\conf.js > {appPath}\\conf.log
#                         '''
#                     # db_scheduledTime = "Everyday at " + str(hours) + ":" + str(mins) + " " + meridian
#                     if int(hours) == 00:
#                         db_scheduledTime = "Everyday at " + str(int(hours) + 12) + ":" + str(mins) + " " + meridian
#                     else:
#                         db_scheduledTime = "Everyday at " + str(hours) + ":" + str(mins) + " " + meridian
#                     dbData = [{"scheduler_name": schedulerName,
#                                "test_job_name": testRunName,
#                                "scheduled_time": db_scheduledTime,
#                                "path": appPath}]
#                     db.scheduledJobs.insert_many(dbData)
#                     batch_file.write(batchScript)
#                     batch_file.close()
#
#                     # daily_job(appPath, schedulerName)
#
#                     print(ScheduleTime, "ScheduleTime")
#
#                     schedule.every().day.at(ScheduleTime).do(lambda: daily_job(appPath, schedulerName))
#
#                     # subprocess.call(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")
#         send_mail(email, testRunName)
#         return jsonify("Success")


@app.route("/checkFileExistForParallel")  # check for the files exists
def checkFileExistForParallel():
    testRunName = request.args.get("testRunName")
    if os.path.exists(base_location + "\\testRun\\" + testRunName + "\\combined_multi_scenario_file.js"):
        print("success")
        return jsonify({"status": "success"})
    else:
        print("failure")
        return jsonify({"status": "failure"})


@app.route("/browserCompatibilityTest")  # check the functionality of each browser
def browserCompatibilityTest():
    jobName = request.args.get("test_run")
    spec_file_name="multiple_scenario_file.js"
    parallel_conf = provideConfBrowserCompatability(spec_file_name)

    collections = db.testRun.find()
    for collection in collections:
        if collection["test_name"] == jobName:
            path = collection["path"]
            break
    open(path + "\\parallel_conf.js", "w").close()
    conf_file = open(path + "\\parallel_conf.js", 'w+')
    conf_file.write(parallel_conf)
    conf_file.close()
    parallel_batch_file = open(path + "\\parallelTestRun.bat", "w+")
    multipleBatchFile = "START /B cmd " + base_location + " && protractor " + path + "\\parallel_conf.js > " + path + "\\parallel_conf_file.log"
    parallel_batch_file.write(multipleBatchFile)
    parallel_batch_file.close()
    os.chdir(path)
    subprocess.call(path + "\\parallelTestRun.bat")
    conf_data = open(path + "\\parallel_conf_file.log").read()
    # print(conf_data)
    data = getParallelJobDetails(jobName, 'parallel_conf_file')
    json_timer = elapseTime(jobName, 'parallel_conf_file')
    data[0].insert(2, json_timer)
    dataList = json_json_json({"data": data[0], "headers": data[1]}, jobName)
    print('T$$T', {"log": conf_data, "data": dataList,
                   "headers": ["Test_Job", "Time_Taken_chrome", "Time_Taken_Firefox", "Result_Chrome",
                               "Result_Firefox"]})
    return jsonify({"log": conf_data, "data": dataList,
                    "headers": ["Test_Job", "Time_Taken_chrome", "Time_Taken_Firefox", "Result_Chrome",
                                "Result_Firefox"]})


@app.route("/browserCompatibilityTestWithExcel", methods=["GET", "POST"])  # check browser compatibility test with excel
def browserCompatibilityTestWithExcel():
    jobName = request.args.get("test_run")
    file = request.get_data()
    print(file)

    catchit = uploadingExcelForBrowserCompatibility(file, jobName)

    print(catchit)
    spec_file_name="multiple_scenario_file.js"
    parallel_conf = provideConfBrowserCompatability(spec_file_name)
    collections = db.testRun.find()
    for collection in collections:
        if collection["test_name"] == jobName:
            path = collection["path"]
            break
    open(path + "\\parallel_new_conf.js", "w").close()
    conf_file = open(path + "\\parallel_new_conf.js", "w")
    conf_file.write(parallel_conf)
    conf_file.close()
    if not os.path.isfile(path + "\\parallelTestRun1.bat"):
        parallel_batch_file = open(path + "\\parallelTestRun1.bat", "w+")
        multipleBatchFile = "START /B cmd " + path + " && protractor " + path + "\\parallel_new_conf.js >" + path + "\\parallel_new_conf_file.log"
        parallel_batch_file.write(multipleBatchFile)
        parallel_batch_file.close()
    subprocess.call(path + "\\parallelTestRun1.bat")
    conf_data = open(path + "\\parallel_new_conf_file.log").read()
    data = getParallelJobDetails(jobName, 'parallel_new_conf_file')
    json_timer = elapseTime(jobName, 'parallel_new_conf_file')
    data[0].insert(2, json_timer)
    dataList = json_json_json({"data": data[0], "headers": data[1]}, jobName)
    print('T$$T', {"log": conf_data, "data": dataList,
                   "headers": ["Test_Job", "Time_Taken_chrome", "Time_Taken_Firefox", "Result_Chrome",
                               "Result_Firefox"]})
    return jsonify({"log": conf_data, "data": dataList,
                    "headers": ["Test_Job", "Time_Taken_chrome", "Time_Taken_Firefox", "Result_Chrome",
                                "Result_Firefox"]})
    # return jsonify({"log": conf_data, "data": data[0], "headers": data[1]})


def uploadingExcelForBrowserCompatibility(file, testRunName):

    global mulipleScenarioPath

    str_file = file.decode("utf-8")

    sheetLength = 0
    json_format = json.loads(str_file)
    print("str_file", str_file)

    # Checking the muliple scenario file exist or not, if exists remove the file
    connections = db.testRun.find()
    for connection in connections:
        if connection["test_name"] == testRunName:
            mulipleScenarioPath = connection["path"]
            if os.path.exists(mulipleScenarioPath + "\\parallel_multi_scenario_file.js"):
                p = os.path.join(mulipleScenarioPath + "\\parallel_multi_scenario_file.js")
                os.remove(p)

    for k, v in json_format.items():
        v1 = json.dumps(v['data'])
        if type(list(eval(v1))[0]) == "str":
            return jsonify("failure")
        else:
            temp_file = list(eval(v1))
            df = pd.DataFrame.from_dict(temp_file)

        connections = db.testRun.find()
        for connection in connections:
            if connection["test_name"] == testRunName:
                mulipleScenarioPath = connection["path"]
                if os.path.isdir(mulipleScenarioPath):
                    df.to_excel(mulipleScenarioPath + '\\sample.xls',
                                index=False)
                    time.sleep(1)
                    append_file = open(mulipleScenarioPath + "\\append_spec.txt", "a+")
                    returnVal = uploadUIWithoutRunForBrowserCompatibility(testRunName, sheetLength)
                    sheetLength += 1
                else:
                    return "appFail"

    comb_multi_file = open(mulipleScenarioPath + "\\parallel_multi_scenario_file.js", "a+")
    comb_multi_file.write("});")
    return jsonify("Success")


def uploadUIWithoutRunForBrowserCompatibility(testRunName, sheetLength):
    replace_spec = ''
    head_list = []
    cell_list = []
    filePath = ''
    scenarioNamedict = {}
    # folderName = request.args.get("folderName")
    connections = db.testRun.find()
    for connection in connections:
        if connection["test_name"] == testRunName:
            appPath = connection["path"]
            scenario = connection["test_scenarios_used"]
            split_scenario = scenario.split(",")[sheetLength]
            appNameRun = split_scenario.split("[")[1].split("]")[0]
            scenarioNameRun = split_scenario.split("[")[0]
            collections = db.folderName.find()
            for collection in collections:
                if collection["appName"] == appNameRun and collection["scenarioName"] == scenarioNameRun:
                    filePath = collection["innerpath"]
            if os.path.isdir(filePath):
                if os.path.exists(filePath + "\\spec.js"):
                    spec_file = open(
                        filePath + "\\spec.js")
                    path = appPath + "\\sample.xlsx"

                    parallel_multi_scenario_file = open(appPath + "\\" + "parallel_multi_scenario_file.js", "a+")
                    parallel_conf_file = open(appPath + "\\parallel_conf_file.js", 'w+')
                    parallel_batch_file = open(appPath + "\\parallel_batch_file.bat", "w+")

                    wb = openpyxl.load_workbook(path)
                    sheet = wb.active
                    start_row = 1
                    end_row = sheet.max_row
                    num_of_columns = sheet.max_column
                    open_file = spec_file.read()
                    split_file = open_file.split("extract")[0]
                    if (split_file.endswith("//")):
                        split_file = split_file[:-2] + "});"
                    new_list = []
                    if scenarioNameRun in new_list:
                        scenarioNamedict[scenarioNameRun] += 1
                    else:
                        scenarioNamedict[scenarioNameRun] = 0
                    new_list.append(scenarioNameRun)
                    sec_test = "it(\'" + scenarioNameRun + " " + str(
                        scenarioNamedict[scenarioNameRun]) + "\', function() {" + \
                               "browser.driver.manage().window().maximize();" + \
                               split_file.split("browser.driver.manage().window().maximize();")[1]
                    print("startrow", start_row, end_row)
                    print("end row", end_row)

                    for r in range(start_row, end_row):
                        head_list = []
                        cell_list = []
                        for c in range(num_of_columns):
                            print("row,c", r, c)
                            cell_obj = sheet.cell(row=r + 1, column=c + 1).value
                            head_obj = sheet.cell(row=1, column=c + 1).value
                            print("cel ", cell_obj, head_obj)
                            head_list.append(head_obj)
                            cell_list.append(cell_obj)
                        for line in sec_test.split(";"):
                            for idx, header in enumerate(head_list):
                                line = str(line)
                                if (line.__contains__(header) and not line.__contains__("browser.wait(until")
                                        and not line.__contains__(".clear(")):
                                    if (type(cell_list[idx]) == int or type(cell_list[idx]) == float):
                                        line = line.split("sendKeys")[0] + "sendKeys(" + str(int(cell_list[idx])) + ");"
                                    else:
                                        line = line.split("sendKeys")[0] + "sendKeys('" + str(cell_list[idx]) + "');"
                                    break
                            replace_spec += line + ";"
                    new_spec = provideJasmineSpecReporter() + str(replace_spec) + "});"

                    if os.stat(appPath + "\\parallel_multi_scenario_file.js").st_size == 0:
                        new_parallel_file = provideJasmineSpecReporter() + str(replace_spec)
                        parallel_multi_scenario_file.write(new_parallel_file)
                        print("parallel_multi_scenario_file", parallel_multi_scenario_file)
                    else:
                        new_parallel_file = str(new_spec)
                        parallel_multi_scenario_file.write(new_parallel_file)
                        print("parallel_multi_scenario_file", parallel_multi_scenario_file)
                    parallel_multi_scenario_file.close()
                    spec_file_name="parallel_multi_scenario_file.js"
                    comb_conf = provideConf(spec_file_name,base_location,outerFolderName,folderName)
                    parallel_conf_file.write(comb_conf)
                    parallel_conf_file.close()

                    parallel_batch_file_data = "START /B cmd " + appPath + " && protractor " + appPath + "\\parallel_multi_scenario_file.js >" + appPath + "\\parallel_conf_file.log"
                    parallel_batch_file.write(parallel_batch_file_data)
                    parallel_batch_file.close()
                    time.sleep(1)
                    return "Success"
            else:
                return "appFail"


@app.route("/editScheduleRunWithExcel", methods=["GET", "POST"])  # edit the scheduled batch file with excel
def editScheduleRunWithExcel():
    ScheduleDateTime = request.args.get("ScheduleDateTime")
    oldScheduleDateTime = request.args.get("oldScheduleDateTime")
    testRunName = request.args.get("testRunName")
    dailyFlag = request.args.get("ScheduleDateTimeFlag")
    email = request.args.get("email")
    schedulerName = request.args.get("schedulerName")
    file = request.get_data()

    catchit = uploadingExcelInSchedule(schedulerName, file, testRunName)

    print(catchit)
    ScheduleDate = ScheduleDateTime.split("T")[0]
    ScheduleTime = ScheduleDateTime.split("T")[1]
    hours = ScheduleTime.split(":")[0]
    mins = ScheduleTime.split(":")[1]
    meridian = "AM"
    if int(hours) > 12:
        hours = int(hours) - 12
        meridian = "PM"
        if hours < 10:
            hours = '0' + str(hours)
    elif int(hours) < 10 and int(hours) != 0:
        hours = str(hours)
    elif int(hours) == 12:
        meridian = "PM"
    ScheduleTime1 = str(hours) + ":" + str(mins)
    ScheduleDate = ScheduleDate.split("-")[1] + "/" + ScheduleDate.split("-")[2] + "/" + ScheduleDate.split("-")[0]
    print(ScheduleTime, ScheduleDate)
    connections = db.scheduledJobs.find()

    for connection in connections:
        if connection["test_job_name"] == testRunName:
            if connection["scheduled_time"] == oldScheduleDateTime:
                appPath = connection["path"]
                print(appPath)
                appPath = appPath.replace("/", "\\").replace("\\\\", "\\")
                open(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat", "w").close()
                deleteTestJob(schedulerName, testRunName, oldScheduleDateTime)
                batch_file = open(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat", "w+")
                if os.path.isdir(appPath + "\\" + schedulerName + "\\combined_conf_file.log") == False:
                    open(appPath + "\\" + schedulerName + "\\combined_conf_file.log", "w").close()
                if dailyFlag == "false":
                    batchScript = f'''
                       START /B cmd "cd {base_location}\\" && protractor {appPath}\\{schedulerName}\\combined_conf_file.js > {appPath}\\{schedulerName}\\combined_conf_file.log
                            '''
                    # db_scheduledTime = ScheduleDate + " at " + str(hours) + ":" + str(mins) + " " + meridian
                    if int(hours) == 00:
                        db_scheduledTime = ScheduleDate + " at " + str(int(hours) + 12) + ":" + str(
                            mins) + " " + meridian
                    else:
                        db_scheduledTime = ScheduleDate + " at " + str(hours) + ":" + str(mins) + " " + meridian

                    dbData = [{"scheduler_name": schedulerName,
                               "test_job_name": testRunName,
                               "scheduled_time": db_scheduledTime,
                               "path": appPath}]

                    db.scheduledJobs.insert_many(dbData)

                    batch_file.write(batchScript)
                    batch_file.close()
                    month, day, year = ScheduleDate.split('/')
                    born = datetime(int(year), int(month), int(day))
                    my_day = born.strftime("%A")
                    print(my_day, ScheduleTime)

                    if my_day == "Monday":
                        schedule.every().monday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Tuesday":
                        schedule.every().tuesday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Wednesday":
                        schedule.every().wednesday.at(ScheduleTime).do(
                            lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Thursday":
                        schedule.every().thursday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Friday":
                        print("entered into this friday")

                        schedule.every().friday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                        # schedule.every().friday.at("15:55").do(lambda: job('Hello ', 'world!'))
                        print("schedule function is called..")
                    elif my_day == "Saturday":
                        schedule.every().saturday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Sunday":
                        schedule.every().sunday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                    # subprocess.call(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")
                else:
                    batchScript = f'''
                               START /B cmd "cd {base_location}\\" && protractor {appPath}\\{schedulerName}\\combined_conf_file.js > {appPath}\\{schedulerName}\\combined_conf_file.log
                        '''
                    # db_scheduledTime = "Everyday at " + str(hours) + ":" + str(mins) + " " + meridian
                    if int(hours) == 00:
                        db_scheduledTime = "Everyday at " + str(int(hours) + 12) + ":" + str(mins) + " " + meridian
                    else:
                        db_scheduledTime = "Everyday at " + str(hours) + ":" + str(mins) + " " + meridian

                    dbData = [{"scheduler_name": schedulerName,
                               "test_job_name": testRunName,
                               "scheduled_time": db_scheduledTime,
                               "path": appPath}]

                    db.scheduledJobs.insert_many(dbData)
                    batch_file.write(batchScript)
                    batch_file.close()
                    schedule.every().day.at(ScheduleTime).do(lambda: daily_job(appPath, schedulerName))
                    # subprocess.call(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")
    send_mail(email, testRunName)
    return jsonify("Success")


def deleteTestJob(schedulerName, testRunName, ScheduleDateTime):  # delete the scheduled batch file
    my_query = {"scheduler_name": schedulerName, "test_job_name": testRunName, "scheduled_time": ScheduleDateTime}
    connections = db.scheduledJobs.find()
    for connection in connections:
        if connection["scheduler_name"] == schedulerName and connection["test_job_name"] == testRunName and connection[
            "scheduled_time"] == ScheduleDateTime:
            path = connection["path"]
            os.remove(path + "\\" + schedulerName + "\\scheduleTestRun.bat")
            break
    db.scheduledJobs.delete_many(my_query)


@app.route("/editScheduleRun")  # edit the scheduled batch file
def editScheduleRun():
    ScheduleDateTime = request.args.get("ScheduleDateTime")
    oldScheduleDateTime = request.args.get("oldScheduleDateTime")
    testRunName = request.args.get("testRunName")
    dailyFlag = request.args.get("ScheduleDateTimeFlag")
    email = request.args.get("email")
    schedulerName = request.args.get("schedulerName")

    ScheduleDate = ScheduleDateTime.split("T")[0]
    ScheduleTime = ScheduleDateTime.split("T")[1]
    hours = ScheduleTime.split(":")[0]
    mins = ScheduleTime.split(":")[1]
    meridian = "AM"
    if int(hours) > 12:
        hours = int(hours) - 12
        meridian = "PM"
        if hours < 10:
            hours = '0' + str(hours)
    elif int(hours) < 10 and int(hours) != 0:
        hours = str(hours)
    elif int(hours) == 12:
        meridian = "PM"
    ScheduleTime1 = str(hours) + ":" + str(mins)
    ScheduleDate = ScheduleDate.split("-")[1] + "/" + ScheduleDate.split("-")[2] + "/" + ScheduleDate.split("-")[0]
    print(ScheduleTime, ScheduleDate)
    connections = db.scheduledJobs.find()

    for connection in connections:
        if connection["test_job_name"] == testRunName:
            if connection["scheduled_time"] == oldScheduleDateTime:
                appPath = connection["path"]
                print(appPath)
                appPath = appPath.replace("/", "\\").replace("\\\\", "\\")
                open(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat", "w").close()
                deleteTestJob(schedulerName, testRunName, oldScheduleDateTime)
                batch_file = open(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat", "w+")

                if dailyFlag == "false":
                    batchScript = f'''
                        START /B cmd "cd {base_location}\\" && protractor {appPath}\\conf.js > {appPath}\\conf.log
                    '''
                    # db_scheduledTime = ScheduleDate + " at " + str(hours) + ":" + str(mins) + " " + meridian
                    if int(hours) == 00:
                        db_scheduledTime = ScheduleDate + " at " + str(int(hours) + 12) + ":" + str(
                            mins) + " " + meridian
                    else:
                        db_scheduledTime = ScheduleDate + " at " + str(hours) + ":" + str(mins) + " " + meridian

                    dbData = [{"scheduler_name": schedulerName,
                               "test_job_name": testRunName,
                               "scheduled_time": db_scheduledTime,
                               "path": appPath}]

                    db.scheduledJobs.insert_many(dbData)

                    batch_file.write(batchScript)
                    batch_file.close()
                    month, day, year = ScheduleDate.split('/')
                    born = datetime(int(year), int(month), int(day))
                    my_day = born.strftime("%A")
                    print(my_day, ScheduleTime)

                    if my_day == "Monday":
                        schedule.every().monday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Tuesday":
                        schedule.every().tuesday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Wednesday":
                        schedule.every().wednesday.at(ScheduleTime).do(
                            lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Thursday":
                        schedule.every().thursday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Friday":
                        print("entered into this friday")

                        schedule.every().friday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                        # schedule.every().friday.at("15:55").do(lambda: job('Hello ', 'world!'))
                        print("schedule function is called..")
                    elif my_day == "Saturday":
                        schedule.every().saturday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Sunday":
                        schedule.every().sunday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                    # subprocess.call(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")
                else:
                    batchScript = f'''
                      START /B cmd "cd {base_location}\\" && protractor {appPath}\\conf.js > {appPath}\\conf.log
                            '''
                    # db_scheduledTime = "Everyday at " + str(hours) + ":" + str(mins) + " " + meridian
                    if int(hours) == 00:
                        db_scheduledTime = "Everyday at " + str(int(hours) + 12) + ":" + str(mins) + " " + meridian
                    else:
                        db_scheduledTime = "Everyday at " + str(hours) + ":" + str(mins) + " " + meridian

                    dbData = [{"scheduler_name": schedulerName,
                               "test_job_name": testRunName,
                               "scheduled_time": db_scheduledTime,
                               "path": appPath}]

                    db.scheduledJobs.insert_many(dbData)
                    batch_file.write(batchScript)
                    batch_file.close()

                    schedule.every().day.at(ScheduleTime).do(lambda: daily_job(appPath, schedulerName))
                    # subprocess.call(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")
    send_mail(email, testRunName)
    return jsonify("Success")


@app.route(
    "/getAllScenariosWithAppName")  # In combine scenarios, while selecting application name, this service gets called
def getAllScenariosWithAppName():
    scenariosList = []  # initialize an empty list.
    appName = request.args.get("appName")  # get application name from UI.
    collections = db.folderName.find()  # find collection data
    for collection in collections:  # iterate through database collection
        if collection["appName"] == appName:  # if application name already in db.
            scenariosList.append(collection["scenarioName"] + '[' + collection[
                "appName"] + ']')  # append scenarioname and appname to scenariolist
    return jsonify(scenariosList)  # return scenariolist


@app.route("/checkFileExist")  # Check whether folder exist inside ProtractorWebdriverTest/outerFolder
def checkFileExist():
    global folderName
    global outerFolderName
    folderName = request.args.get("folderName")
    outerFolderName = request.args.get("outerFolder")

    newFile = []
    if os.path.isdir(base_location + "\\" + outerFolderName) == False:
        os.mkdir(base_location + "\\" + outerFolderName)
    for file in os.listdir(base_location + "\\" + outerFolderName):
        if not file.lower().__contains__("."):
            newFile.append(file.lower())

    if folderName.lower() in newFile:
        return jsonify("Fail")
    else:

        return jsonify("Success")


@app.route("/combineFiles")
def combine_files():
    scenarios = request.args.get("scenarios")
    outerFolderName = request.args.get("appName")
    folderName = request.args.get("testRunName")
    split_scenario = scenarios.split(",")
    scenario_array = []
    dictionary = {}
    array=[]
    writeArray=[]
    point1=''
    point2=''
    flag=0
    scenario_array1 = []
    for uniq in split_scenario:
        application_name = uniq.split("[")[1].split("]")[0].strip()
        array.append(str(application_name))
        scenario_name = uniq.split("[")[0].strip()
        print(scenario_name)
        scenario_array.append(scenario_name)

    print(scenario_array)

    if len(set(array)) == 1:
        # for items in range(1,len(scenario_array)):

        for filename in scenario_array:
            scenario_array1.append(filename)
            print(filename)
            # with fileinput.FileInput(base_location + "\\" + outerFolderName + "\\" + filename + "\\spec.js",
            #                          inplace=True, backup='.bak') as file:
            #     text_to_search = filename
            #     replacement_text = folderName
            #     for line in file:
            #         print(line.replace(text_to_search, replacement_text), end='')
        scenario_array1.append("@")

        for idx, name in enumerate(scenario_array1):
            if not name == "@":
                firstFile = scenario_array1[idx].strip()
                var1 = len(scenario_array1)
                # print(scenario_array[var1-1])
                # secondFile1=scenario_array1[var1-1]

                if not firstFile == "@":
                    secondFile = scenario_array1[idx + 1].strip()

                    firstSpec = open(
                        base_location + "\\" + outerFolderName + "\\" + firstFile + "\\spec.js").readlines()

                    if not secondFile == "@":
                        secondSpec = open(
                            base_location + "\\" + outerFolderName + "\\" + secondFile + "\\spec.js").readlines()

                        for i, j in zip(firstSpec, secondSpec):
                            if i == j and i not in writeArray and not firstSpec is None and i.__contains__("browser.wait") or i.__contains__("element.all") or i.__contains__("element(by.") or i.__contains__("browser.sleep("+str(sleep_delay)+");"):
                                writeArray.append(i)
                            elif not i == j  and i.__contains__("browser.wait") or i.__contains__("element.all") or i.__contains__("element(by.") or i.__contains__("browser.sleep("+str(sleep_delay)+");"):
                                point1 = i
                                # point2 = j
                                break
                        firstSpec = open(
                            base_location + "\\" + outerFolderName + "\\" + firstFile + "\\spec.js").readlines()
                        flag = 0
                        for a in firstSpec:
                            print("point1", point1)
                            if a == point1 and not point1=='':
                                flag = 1
                                writeArray.append("\n")
                            if flag == 1 and not a in writeArray and not firstSpec is None and a.__contains__("browser.wait") or a.__contains__("element.all") or a.__contains__("element(by.") or a.__contains__("browser.sleep("+str(sleep_delay)+");"):

                                writeArray.append(a)

                    if secondFile == "@":

                        openPrev = scenario_array1.index("@") - 1
                        openPrevPrev = scenario_array1.index("@") - 2
                        prevFile = scenario_array1[openPrev]
                        prevPrevFile = scenario_array1[openPrevPrev]
                        prevSpec = open(
                            base_location + "\\" + outerFolderName + "\\" + prevFile + "\\spec.js").readlines()
                        PrevprevSpec = open(
                            base_location + "\\" + outerFolderName + "\\" + prevPrevFile + "\\spec.js").readlines()
                        for line1, line2 in zip(prevSpec, PrevprevSpec):
                            if line1 == line2:
                                continue
                            if not line1==line2 and line1.__contains__("browser.wait") or line1.__contains__("element.all") or line1.__contains__("element(by.") or line1.__contains__("browser.sleep("+str(sleep_delay)+");"):
                                point1 = line2
                                point2 = line1
                                break
                        flag = 0
                        for z in prevSpec:

                            if z == point2 and not point2=='':
                                flag = 1
                                writeArray.append("\n")
                            if flag == 1 and z.__contains__("browser.wait") or z.__contains__("element.all") or z.__contains__("element(by.") or z.__contains__("browser.sleep("+str(sleep_delay)+");"):

                                writeArray.append(z)

            if name == "@":
                break
        CombineSpec(outerFolderName, folderName, writeArray)
        # formatSpec(outerFolderName, folderName)
        CombineBatch(outerFolderName, folderName)
        CombineConf(outerFolderName, folderName)
        combineTestDescription(outerFolderName, folderName, scenario_array)
        return jsonify("Success")
    else:
        return jsonify("Fail")

def formatSpec(outerFolderName,folderName):
    new_spec_file = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\spec.js",
                         "r").readlines()
    insert_line=""
    lengthOfFile=len(new_spec_file)
    print(lengthOfFile)
    with open(base_location + "\\" + outerFolderName + "\\" + folderName + "\\spec.js",
              "w+") as out_file:  # open spec.js in write mode


        for idx,line in enumerate(new_spec_file):
            insert_line += line
            if line.__contains__("var fname=") and not new_spec_file[idx+1].__contains__("stream1.write"):
                insert_line+='''stream1.write(new Buffer(png, 'base64'));
                stream1.end();
                });\n'''
            if line.__contains__("browser.sleep("+str(sleep_delay)+");});") and new_spec_file[idx+1].startswith("});"):
                insert_line=insert_line.replace(line,"browser.sleep("+str(sleep_delay)+");")

        out_file.write(insert_line)
    appendFile = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\spec.js",
                         "a")
    appendFile.write("\n });")

def CombineSpec(outerFolderName,folderName,writeArray):
    if os.path.isdir(base_location + "\\" + outerFolderName + "\\" + folderName) == False:
        os.mkdir(base_location + "\\" + outerFolderName + "\\" + folderName)
    new_spec_file = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\spec.js",
                         "w+")
    for items in writeArray:
        new_spec_file.write(items)


def CombineBatch(outerFolderName,folderName):
    cbn_batch = '''
          START /B cmd "cd ''' + base_location + "\\\" && protractor " + base_location + "\\" + outerFolderName + "\\" + folderName + "\\conf.js >" + base_location + "\\" + outerFolderName + "\\" + folderName + '''\\conf.log
          '''
    if os.path.isdir(base_location + "\\" + outerFolderName + "\\" + folderName + "\\Batchfiles") == False:
        os.mkdir(base_location + "\\" + outerFolderName + "\\" + folderName + "\\Batchfiles")
    f6 = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\\Batchfiles\Test_batch.bat", "w+")
    f6.write(cbn_batch)
    f6.close()
    dbData = [{"path": base_location + "\\" + outerFolderName,
               "innerpath": base_location + "\\" + outerFolderName + "\\" + folderName,
               "appName": outerFolderName,
               "scenarioName": folderName,
               "test_Description": ''}]
    db.folderName.insert_many(dbData)

def CombineConf(outerFolderName,folderName):
    spec_file_name="spec.js"
    conf = provideConf(spec_file_name,base_location,outerFolderName,folderName)
    f5 = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\\conf.js", "w+")
    f5.write(conf)
    f5.close()

def combineTestDescription(outerFolderName, folderName, scenario_array):
    full_line = ''
    for arr in scenario_array:
        full_line += arr + "<br>"
    sentence = "Combined Scenarios of :<br>" + full_line[:-4]
    print("combined sentence", sentence)
    myquery = {"appName": outerFolderName, "scenarioName": folderName}
    newvalues = {"$set": {"test_Description": sentence}}
    db.folderName.update_many(myquery, newvalues)
    return "done"


def format_spec_new(content):
    try:
        print("insid eformat spec")
        file_one = content
        string = ""
        flag = 0
        count_url = 0
        url_line = ''
        temp_val = []
        tempflag = 0
        for f in file_one:
            # print(f)
            if f.__contains__(".isDisplayed()"):
                if tempflag == 1 and f.strip("\t") == temp_val[0].strip("\t"):
                    f = ''
                temp_val.append(f)
                tempflag = 1
            if f.__contains__("browser.get("):
                print(f)
            if f.__contains__("browser.get(") and count_url == 0:
                url_line = f.strip()
                print("url line", url_line)
                count_url += 1
                print("count", count_url)

            if url_line != "" and f.strip().__contains__(url_line) and f.strip() != '':
                print("inside f", f)
                print("count url", count_url)
                if count_url > 1:
                    f = f.replace(url_line, "")
                count_url += 1

            if f.__contains__("it('Test 1', function() {"):
                f = f.replace("it('Test 1', function() {", "\n\t\t\t\tit('Test 1', function() {")
            if f.__contains__("browser.ignoreSynchronization"):
                print("inside")
                f = f.replace("browser.ignoreSynchronization", "\n\t\t\t\t\t\t\t\t\tbrowser.ignoreSynchronization")
                f = f.replace(";", ";\n\t\t\t")
                # f=f
                flag = 1
            if flag == 1:
                # print(flag, "inside")
                f = f.replace(";", ";\n\t\t\t")
                f = "\t" + f
            if f.__contains__("extract()"):
                # print(f)
                flag = 0
            if flag == 0:
                f = f.replace(";", ";\n")
            if f.strip() != '':
                string += f

        whole_string = ''''''
        for s in string.split("\n"):
            if s.strip() != "":
                whole_string += s + "\n"
        return whole_string
    except Exception as e:
        return jsonify(e)


def num_img(outerFolderName, folderName):
    spec_file = open(
        base_location + "\\" + outerFolderName + "\\" + folderName + "\\" + "spec.js").readlines()
    num_line = 0
    whole_line = ''
    for lin in spec_file:
        if lin.__contains__(".png"):
            num_line += 1
            full_line = lin.split("page")[0] + "page" + str(num_line) + "." + lin.split(".")[1]
            print(full_line)
            print(lin)
            whole_line += full_line
        else:
            whole_line += lin
    if whole_line.__contains__("async function extractModal"):
        whole_line = whole_line[:whole_line.index("async function extractModal")] + "});"
    spec_file = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\\" + "spec.js", "w+")
    spec_file.write(whole_line)


@app.route("/checkFolderExist")  # Check folder exist inside ProtractorWebdriverTest
def checkFolderExist():
    # global folderName
    folderName1 = request.args.get("folderName")
    newFile = []
    if os.path.isdir(base_location) == False:
        os.mkdir(base_location)
    for file in os.listdir(base_location):
        if not file.lower().__contains__("."):
            newFile.append(file.lower())

    if folderName1.lower() in newFile:
        return jsonify("Fail")
    else:
        return jsonify("Success")


@app.route("/startExecution")  # Called when user clicks on Record button in ui
def startExecution():
    # Get url from ui
    url = request.args.get("url")
    print("url", url)
    if os.path.isdir(
            base_location + '\\' + outerFolderName + '\\' + folderName) == False:  # Check whether the directory exist or not
        os.mkdir(base_location + '\\' + outerFolderName + '\\' + folderName)
    unpacked_extension_path = recorder_path
    chrome_options = Options()
    chrome_options.add_argument('load-extension={}'.format(unpacked_extension_path))
    driver = webdriver.Chrome(options=chrome_options,
                              executable_path=chrome_driver_path,
                              chrome_options=chrome_options)
    driver.get(url)
    os.remove(download_location)
    return jsonify("Success")

def generateSpec(outerFolderName,folderName):
    input_file = base_location + "\\" + outerFolderName + "\\" + folderName + "\\" + "new_file.txt"
    output_file_spec = base_location + "\\" + outerFolderName + "\\" + folderName + "\\spec.js"
    robocorder_file = open(input_file).readlines()
    robocorder_str = ''.join(robocorder_file)
    splitted_robocorder = robocorder_str.split(",")
    print("splitted", splitted_robocorder)
    flag = 0
    protractor_code = ''

    for idx, line in enumerate(splitted_robocorder):
        if flag == 0:
            print("line 1", line)

            if line.lower().__contains__("click") and not line.lower().__contains__("undefined"):

                if (line.split("//", 1)[1].strip()[-2]) == "\"":
                    protractor_code += "browser.wait(until.presenceOf(element(by.xpath(\'//" + line.split("//", 1)[
                        1].strip() + "\'))), delay, 'Element taking too long to appear in the DOM');\nelement(by.xpath(\'//" + \
                                       line.split("//", 1)[
                                           1].strip() + "\')).click();\nbrowser.sleep(500);\n"
                else:
                    num = ''
                    if line.split("//", 1)[1].strip().endswith("]"):
                        num = str(line.split("//", 1)[1].rsplit("[", 1)[1].rsplit("]", 1)[0])
                        print("num value", num)

                    if line.__contains__("]"):
                        select_id = line.split("//", 1)[1].split("]")[0] + "]"
                    else:
                        select_id = line.split("//", 1)[1]
                    select_id = "".join(select_id.rsplit(")", 1)[0])
                    if num != '' and type(eval(num)) == int:
                        # next_line= line.split("//")[1].strip().rsplit("[",1)[0].replace(")","")
                        select_id = line.split("//", 1)[1].rsplit("[", 1)[0]
                        select_id = "".join(select_id.rsplit(")", 1)[0])
                        # next_line = next_line[:-1]
                        protractor_code += "\nbrowser.wait(until.presenceOf(element.all(by.xpath(\'//" + select_id + "\')).get(" + str(
                            int(
                                num) - 1) + ")), delay, 'Element taking too long to appear in the DOM');\nelement.all(by.xpath(\'//" + select_id + "\')).get(" + str(
                            int(num) - 1) + ").click();\nbrowser.sleep(500);\n"
                    else:
                        num = 0
                        protractor_code += "\nbrowser.wait(until.presenceOf(element.all(by.xpath(\'//" + select_id + "\')).get(" + str(
                            int(
                                num)) + ")), delay, 'Element taking too long to appear in the DOM');\nelement.all(by.xpath(\'//" + select_id + "\')).get(" + str(
                            int(num)) + ").click();\nbrowser.sleep(500);\n"


            elif (line.__contains__("//") and line.__contains__("Open Browser") == False):
                input_line = line.split("//", 1)[1].strip().split()[-1]
                print(input_line)
                next_line_one = line.split("//", 1)[1].strip().rsplit(" ", 1)[0].strip().replace(")", "")
                print(next_line_one, "im next line")
                num = str(line.split("//", 1)[1].split()[0].strip()[-2])
                print("im number", num)
                print("num 1, num", num, type(num), num.isnumeric())

                if num.isnumeric():
                    print("inside line", line)

                    if line.__contains__("]"):
                        select_id = line.split("//", 1)[1].split("]")[0] + "]"
                        if select_id.__contains__(")"):
                            select_fin = select_id.split(")")[0].strip()
                            select_id = str(select_fin)

                        select_ele = line.split("]")[1].strip()
                    else:
                        select_id = line.split("//", 1)[1]
                        select_ele = line.split(")")[1].split("[")[1].split(']')[0]

                    # select_ele = line.split(")")[1].split("[")[1].split(']')[0]

                    if select_id.lower().startswith("select"):
                        protractor_code += "browser.wait(until.presenceOf(element(by.xpath(\'//" + select_id + "//option[@value=\"" + str(
                            select_ele) + "\"]\')).get(" + str(
                            int(
                                num) - 1) + ")), delay, 'Element taking too long to appear in the DOM');\nelement(by.xpath(\'//" + select_id + "//option[@value=\"" + str(
                            select_ele) + "\"]\')).get(" + str(
                            int(num) - 1) + ").click();\nbrowser.sleep(500);\n"
                    else:
                        protractor_code += "\nbrowser.wait(until.presenceOf(element.all(by.xpath(\'//" + select_id + "\')).get(" + str(
                            int(
                                num) - 1) + ")), delay, 'Element taking too long to appear in the DOM');\nelement.all(by.xpath(\'//" + select_id + "\')).get(" + str(
                            int(num) - 1) + ").clear();\nelement.all(by.xpath(\'//" + select_id + "\')).get(" + str(
                            int(num) - 1) + ").sendKeys(\"" + str(
                            select_ele) + "\");\n"
                else:
                    num = 0
                    # next_line = line.split("//")[1].strip().rsplit("[", 1)[0].replace(")","")
                    if line.split("//", 1)[1].count(" ") > 0:
                        next_line = line.split("//", 1)[1].rsplit(" ", 1)[0].strip()
                    else:
                        next_line = line.split("//", 1)[1].rsplit(" ", 1)[0].strip()
                    print("next_line", next_line)
                    if line.__contains__("]"):
                        select_id = line.split("//", 1)[1].split("]")[0] + "]"
                    else:
                        select_id = line.split("//", 1)[1].split()[0]
                    # select_id = line.split("//")[1].split("]")[0] + "]"

                    # select_ele = line.split("]")[1].strip()
                    if line.split("//", 1)[1].count(" ") > 0:
                        select_ele = line.split("//", 1)[1].rsplit(" ", 1)[1].strip()

                    else:
                        select_ele = line.split("//", 1)[1].strip()

                    print("next_line", next_line)
                    # protractor_code += "browser.sleep(5000);\nelement(by.xpath(\'//" + next_line + "\')).clear();\nelement(by.xpath(\'//" + next_line + "\')).sendKeys(\""+str(input_line)+"\");\n"
                    if next_line.lower().startswith("select"):
                        protractor_code += "browser.sleep(" + str(
                            sleep_delay) + ");\nbrowser.wait(until.presenceOf(element(by.xpath(\'//" + select_id + "//option[@value=\"" + str(
                            select_ele) + "\"]\'))), delay, 'Element taking too long to appear in the DOM');\nelement(by.xpath(\'//" + select_id + "//option[@value=\"" + str(
                            select_ele) + "\"]\')).click();\nbrowser.sleep(500);\n"
                    else:
                        protractor_code += "browser.sleep(" + str(
                            sleep_delay) + ");\nbrowser.wait(until.presenceOf(element(by.xpath(\'//" + select_id + "\'))), delay, 'Element taking too long to appear in the DOM');\nelement(by.xpath(\'//" + select_id + "\')).clear();\nelement(by.xpath(\'//" + select_id + "\')).sendKeys(\"" + str(
                            select_ele) + "\");\n"

            elif line.__contains__("Input Text"):
                input_line = line.split("//", 1)[1].strip().split()[-1]
                next_line = line.split("//", 1)[1].strip().rsplit(" ", 1)[0].strip()
                protractor_code += "\nbrowser.wait(until.presenceOf(element(by.xpath(\'//" + next_line + "\'))), delay, 'Element taking too long to appear in the DOM');\nelement(by.xpath(\'//" + next_line + "\')).clear();\nelement(by.xpath(\'//" + next_line + "\')).sendKeys(\"" + str(
                    input_line) + "\");\n"

    protractor_code += "browser.sleep(" + str(sleep_delay) + ");"

    return protractor_code

@app.route("/stopExecution")
def stopExecution():
    try:

        print("Im in")
        url = request.args.get("url")
        url = url.replace("HASH", "#")
        print(url, "url")
        print("SLEEP DELAY",sleep_delay)
        f = open(base_location+"\\" + outerFolderName + "\\" + folderName + "\\" + "new_file.txt", "w+")
        f.close()
        shutil.copy(download_location,
                    base_location+"\\" + outerFolderName + "\\" + folderName + "\\" + "new_file.txt")
        # input_file = base_location + "\\" + outerFolderName + "\\" + folderName + "\\" + "new_file.txt"
        output_file_spec = base_location + "\\" + outerFolderName + "\\" + folderName + "\\spec.js"
        output_file_conf = base_location + "\\" + outerFolderName + "\\" + folderName + "\\conf.js"
        f = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\\batch.bat", "w+")
        new_spec_file = open(output_file_spec, "w+")
        protractor_code=generateSpec(outerFolderName,folderName)
        spec_file = provideSpec(url,outerFolderName,folderName)
        spec_file_name="spec.js"
        conf_file = provideConf(spec_file_name,base_location,outerFolderName,folderName)
        new_conf_file = open(output_file_conf, "w+")
        new_conf_file.write(conf_file)
        new_conf_file.close()
        # protractor_code+=outputValidation(protractor_code)
        final_line = spec_file + "\n" + protractor_code + "});\n});"
        print(final_line)
        # print(final_line)
        new_spec_file.write(final_line)
        new_spec_file.close()
        finalDict1={}
        finalDict1=assertPage()

        print("finalDict1",finalDict1)

        mapScreenshotWithSpec(outerFolderName,folderName)
        script1 = '''
          START /B cmd "cd ''' + base_location + '\\' + outerFolderName + '\\' + folderName + '''\\" && protractor ''' + base_location + '\\' + outerFolderName + '\\' + folderName + '''\\conf.js >''' + base_location + "\\" + outerFolderName + '\\' + folderName + '''\\\conf.log
          '''
        f.write(script1)  # write the command in batchfiles
        f.close()  # close the file


        description = testDescription()
        dbData = [{"path": base_location + "\\" + outerFolderName,
                   "innerpath": base_location + "\\" + outerFolderName + "\\" + folderName,
                   "appName": outerFolderName,
                   "scenarioName": folderName,
                   "test_Description":description}]  # db structure update
        db.folderName.insert_many(dbData)  # insert db data into db


        return finalDict1
    except Exception as e:
        return "Failure"

def moveReport(outerFolder,folder):
    try:
        destination_folder = ui_location + "\\assets\\doc\\" + outerFolder + "\\" + folder
        if os.path.isdir(destination_folder) == False:
            os.makedirs(destination_folder)
        search_folder = base_location + "\\" + outerFolder + "\\" + folder
        for files in os.listdir(search_folder):
            if files.endswith(".html"):
                shutil.copy(search_folder + "\\" + files, destination_folder)
        return "success"
    except Exception as e:
        return jsonify(e)

@app.route("/rerunRecording")
def rerunRecording():
    try:
        global base_location
        global outerFolderName
        global folderName
        outerFolderName = request.args.get("applicationname")
        folderName = request.args.get("scenarioname")
        url = request.args.get("url")
        url = url.replace("HASH", "#")
        print(base_location,outerFolderName,folderName)
        #assertion()
        check_conditions(url)
        subprocess.call(base_location + '\\' + outerFolderName + '\\' + folderName + '\\batch.bat')
        moveReport(outerFolderName,folderName)
        log_filename = "conf.log"
        conf_log = conf_log_modification(log_filename)
        output_json = {"status": "Success", "conf_log": conf_log}
        print("output log", output_json)
        # exportScenarioAsZip()
        # moveImage()
        return jsonify(output_json)
    except Exception as e:
        return jsonify(e)

def mapScreenshotWithSpec(outerFolderName,folderName):
    try:
        for filename in glob.glob(os.path.join(base_location + '\\' + outerFolderName + '\\' + folderName, '*.png')):
            if filename.split("\\")[-1].startswith("page") and filename.split("\\")[-1].endswith(".png"):
                os.remove(filename)
        counter=1
        insert_line = ''
        file_one = open(base_location + '\\' + outerFolderName + '\\' + folderName + "\\spec.js", "r")
        buf = file_one.readlines()
        file_two = open(base_location + '\\' + outerFolderName + '\\' + folderName + "\\spec.js", "w+")
        for line in buf:
            new_line = '''browser.takeScreenshot().then(function (png) {
                        var dir="''' + base_location.replace("\\","\\\\") + "\\\\" + outerFolderName + '\\\\' + folderName + '''";
                        var fname="page''' + str(counter) + ".png\"; " + \
                       '''var stream1 = fs.createWriteStream(path.join(dir, fname));
                        stream1.write(new Buffer(png, 'base64'));
                        stream1.end();});\n'''
            if line.__contains__(".click()") or line.__contains__("sendKeys("):
                insert_line += line.replace(line, new_line)
                counter += 1
            insert_line += line
        file_two.write(insert_line)
    except Exception as e:
        return jsonify(e)

@app.route("/assertion")
def assertion():# to perform assertion logic
    try:
        global base_location
        global folderName
        global outerFolderName
        newfolderName=''
        find_folder = db.folderName.find().sort("_id", -1).limit(1)
        for data in find_folder:
            newfolderName = data["scenarioName"]
        # print("lasstttt", folderName)
        url = request.args.get("url")
        url = url.replace("HASH", "#")
        outerFolderName=request.args.get("applicationname")
        # folderName = request.args.get("scenarioname")
        folder_path=base_location+"\\"+outerFolderName+"\\"+newfolderName+"\\"
        xpath1 = request.args.get("xpath_list").split(",")
        index_list=request.args.get("index_list").split(",")
        original_value1 = request.args.get("value_list").split(",")
        asserted_value1 = request.args.get("assert_list").split(",")
        condition_list=request.args.get("condition_list").split(",")
        stepnumber_list=request.args.get("stepNumber").split(",")
        index1=[]
        for i in index_list:
            if i=='':
                index1.append("-")
            else:
                index1.append(i)
        # shutil.copy(download_location,
        #             base_location + "\\" + outerFolderName + "\\" + newfolderName + "\\" + "new_file.txt")
        new_spec_file=open(folder_path+"spec.js", "w")
        readnew = open(folder_path+"new_file.txt", "r").read()
        opennew = open(folder_path+"new_file1.txt", "w")
        input_file=folder_path+"new_file1.txt"
        refer=[]
        step1=''
        buf=readnew.split(",")

        for idx, line in enumerate(buf):
            refer.append(line)
            for xpath, step, assertval, orignalval, index, cond in zip(xpath1, stepnumber_list, asserted_value1,
                                                                       original_value1,
                                                                     index1, condition_list):
                step1=step
                if step.__contains__("/"):
                    step=step.split("/")[0].strip()
                if idx == int(step) - 1:

                    if step1.__contains__("/") and not index == '-':

                        refer.append( "indexBased" + '\t' + xpath + '\t' + index + '\t' + cond + "\t" + assertval )


                    elif step1.__contains__("/") and index == '-':

                        refer.append("nonIndexBased"+ '\t'  + xpath + '\t' + cond + "\t" + assertval)


                    elif not step1 == '-' and not step1.__contains__("/") and orignalval == '-':

                        refer.append("existingXpath"+ '\t'  + xpath + '\t' + cond + '\t' + assertval)

                    elif not step1 == '-' and not step1.__contains__("/") and not orignalval == '-':

                        refer.append("XpathWithVal"+ '\t'  + xpath + '\t' + ' ?separate? ' + '\t' + orignalval + '\t' + cond + '\t' + assertval)

        print(refer)

        for i in refer:
            opennew.write(i + ",")
        opennew.close()
        robocorder_file = open(input_file).readlines()
        robocorder_str = ''.join(robocorder_file)
        splitted_robocorder = robocorder_str.split(",")
        flag = 0
        protractor_code = ''
        print(splitted_robocorder)
        for idx, line in enumerate(splitted_robocorder):
            if flag == 0:
                print("line 1", line)

                if line.lower().__contains__("click") and not line.lower().__contains__("undefined") and not line.__contains__("nonIndexBased") and not line.__contains__("indexBased") and not line.__contains__("existingXpath") and not line.__contains__("XpathWithVal"):

                    if (line.split("//",1)[1].strip()[-2]) == "\"":
                        protractor_code += '''browser.wait(until.presenceOf(element(by.xpath(\'//''' + line.split("//",1)[1].strip() + '''\'))), delay, 'Element taking too long to appear in the DOM');\n
                        element.all(by.xpath(\'//''' + line.split("//",1)[1].strip() + '''\')).getText().then(function (text) {
                        if( typeof text ==='object')
                        {
                  Get_Text.set(\'//''' + line.split("//",1)[1].strip() + '''\', text);
                  }
                  else
                  {
                  Get_Text.set(\'//''' + line.split("//",1)[1].strip() + '''\', text[0]);
                  }  
                        });
                        element.all(by.xpath(\'//''' + line.split("//",1)[1].strip() + '''\')).getAttribute('value').then(function (text) {
                        if( typeof text ==='object')
                        {
                  Get_Attribute.set(\'//''' + line.split("//",1)[1].strip() + '''\', text);
                  }
                  else
                  {
                  Get_Attribute.set(\'//''' + line.split("//",1)[1].strip() + '''\', text[0]);
                  }  
                        });
                        element(by.xpath(\'//''' + line.split("//",1)[1].strip() + '''\')).click();\n'''
                    else:
                        num = ''
                        if line.split("//",1)[1].strip().endswith("]"):
                            num = str(line.split("//",1)[1].rsplit("[", 1)[1].rsplit("]", 1)[0])
                            print("num value", num)

                        if line.__contains__("]"):
                            select_id = line.split("//",1)[1].split("]")[0] + "]"
                        else:
                            select_id = line.split("//",1)[1]
                        select_id = "".join(select_id.rsplit(")", 1)[0])
                        if num != '' and type(eval(num)) == int:
                            # next_line= line.split("//")[1].strip().rsplit("[",1)[0].replace(")","")
                            select_id = line.split("//",1)[1].rsplit("[", 1)[0]
                            select_id = "".join(select_id.rsplit(")", 1)[0])
                            # next_line = next_line[:-1]
                            protractor_code += '''\nbrowser.wait(until.presenceOf(element.all(by.xpath(\'//''' + select_id + '''\')).get(''' + str(int(num) - 1) + ''')), delay, 'Element taking too long to appear in the DOM');\n
                                                element.all(by.xpath(\'//''' + select_id + '''\')).get(''' + str(int(num) - 1) + ''').getText().then(function (text) {
                                                if( typeof text ==='object')
                                                {
                                                Get_Text.set(\'//''' + select_id + '''[''' + str(int(num) - 1) + ''']'''+'''\', text);
                                                }
                                                else
                                                {
                                                Get_Text.set(\'//''' + select_id + '''[''' + str(int(num) - 1) + ''']'''+'''\', text[0]);
                                                }  
                                                });
                                                element.all(by.xpath(\'//''' + select_id + '''\')).get(''' + str(int(num) - 1) + ''').getAttribute('value').then(function (text) {
                                                if( typeof text ==='object')
                                                {
                                                Get_Attribute.set(\'//''' + select_id + '''[''' + str(int(num) - 1) + ''']'''+'''\', text);
                                                }
                                               else
                                                {
                                                Get_Attribute.set(\'//''' + select_id + '''[''' + str(int(num) - 1) + ''']'''+'''\', text[0]);
                                                }  
                                                });
                                                element.all(by.xpath(\'//''' + select_id + '''\')).get(''' + str(int(num) - 1) + ''').click();\n'''
                        else:
                            num = 0
                            protractor_code += '''\nbrowser.wait(until.presenceOf(element.all(by.xpath(\'//''' + select_id + '''\')).get(''' + str(int(num)) + ''')), delay, 'Element taking too long to appear in the DOM');\n
                                            element.all(by.xpath(\'//''' + select_id + '''\')).get(''' + str(int(num)) + ''').getText().then(function (text) {
                                                if( typeof text ==='object')
                                               {
                                                Get_Text.set(\'//''' + select_id + '''[''' + str(int(num)) + ''']'''+'''\', text);
                                                }
                                                else
                                                {
                                                Get_Text.set(\'//''' + select_id + '''[''' + str(int(num)) + ''']'''+'''\', text[0]);
                                                }  
                                                });
                                                element.all(by.xpath(\'//''' + select_id + '''\')).get(''' + str(int(num)) + ''').getAttribute('value').then(function (text) {
                                                if( typeof text ==='object')
                                                {
                                                Get_Attribute.set(\'//''' + select_id + '''[''' + str(int(num)) + ''']'''+'''\', text);
                                                }
                                                else
                                                {
                                                Get_Attribute.set(\'//''' + select_id + '''[''' + str(int(num)) + ''']'''+'''\', text[0]);
                                                }  
                                                });
                                            element.all(by.xpath(\'//''' + select_id + '''\')).get(''' + str(int(num)) + ''').click();'''


                elif (line.__contains__("//") and line.__contains__("Open Browser") == False) and not line.__contains__("nonIndexBased") and not line.__contains__("indexBased")  and not line.__contains__("existingXpath") and not line.__contains__("XpathWithVal"):
                    input_line = line.split("//",1)[1].strip().split()[-1]
                    print(input_line)
                    next_line_one = line.split("//",1)[1].strip().rsplit(" ", 1)[0].strip().replace(")", "")
                    print(next_line_one,"im next line")
                    num = str(line.split("//",1)[1].split()[0].strip()[-2])
                    print("im number",num)
                    print("num 1, num", num, type(num), num.isnumeric())

                    if num.isnumeric():
                        print("inside line", line)

                        if line.__contains__("]"):
                            select_id = line.split("//",1)[1].split("]")[0] + "]"
                            if select_id.__contains__(")"):
                                select_fin=select_id.split(")")[0].strip()
                                select_id=str(select_fin)

                            select_ele=line.split("]")[1].strip()
                        else:
                            select_id = line.split("//",1)[1]
                            select_ele = line.split(")")[1].split("[")[1].split(']')[0]


                        #select_ele = line.split(")")[1].split("[")[1].split(']')[0]

                        if select_id.lower().startswith("select"):
                            protractor_code += '''browser.wait(until.presenceOf(element(by.xpath(\'//''' + select_id + '''//option[@value=\"''' + str(select_ele) + '''\"]\')).get(''' + str(int(num) - 1) + ''')), delay, 'Element taking too long to appear in the DOM');
                                            element.all(by.xpath(\'//''' + select_id + '''//option[@value=\"''' + str(select_ele) + '''\"]\')).get(''' + str(int(num) - 1) + ''').getText().then(function (text) {
                                                if( typeof text ==='object')
                                                {
                                                Get_Text.set(\'//''' + select_id + '''//option[@value=\"''' + str(select_ele) + '''\"]\', text);
                                                }
                                                else
                                                {
                                                Get_Text.set(\'//''' + select_id + '''//option[@value=\"''' + str(select_ele) + '''\"]\', text[0]);
                                                }  
                                                });
                                             element.all(by.xpath(\'//''' + select_id + '''//option[@value=\"''' + str(select_ele) + '''\"]\')).get(''' + str(int(num) - 1) + ''').getAttribute('value').then(function (text) {
                                                if( typeof text ==='object')
                                                {
                                                Get_Attribute.set(\'//''' + select_id + '''//option[@value=\"''' + str(select_ele) + '''\"]\', text);
                                                }
                                                else
                                                {
                                                Get_Attribute.set(\'//''' + select_id + '''//option[@value=\"''' + str(select_ele) + '''\"]\', text[0]);
                                                }  
                                                });
                                            element(by.xpath(\'//''' + select_id + '''//option[@value=\"''' + str(select_ele) + '''\"]\')).get(''' + str(int(num) - 1) + ''').click();'''
                        else:
                            protractor_code += '''\nbrowser.wait(until.presenceOf(element.all(by.xpath(\'//''' + select_id + '''\')).get(''' + str(int(num) - 1) + ''')), delay, 'Element taking too long to appear in the DOM');
                            element.all(by.xpath(\'//''' + select_id + '''\')).get(''' + str(int(num) - 1) + ''').clear();
                            element.all(by.xpath(\'//''' + select_id + '''\')).get(''' + str(int(num) - 1) + ''').sendKeys(\"''' + str(select_ele) + '''\");
                            element.all(by.xpath(\'//''' + select_id + '''\')).get(''' + str(int(num) - 1) + ''').getText().then(function (text) {
                            if( typeof text ==='object')
                            {
                            Get_Text.set(\'//''' + select_id + '''[''' + str(int(num) -1) + ''']'''+'''\', text);
                            }
                            else
                            {
                            Get_Text.set(\'//''' + select_id + '''[''' + str(int(num) -1) + ''']'''+'''\', text[0]);
                            }  
                            });
                            element.all(by.xpath(\'//''' + select_id + '''\')).get(''' + str(int(num) - 1) + ''').getAttribute('value').then(function (text) {
                            if( typeof text ==='object')
                            {
                            Get_Attribute.set(\'//''' + select_id + '''[''' + str(int(num) -1) + ''']'''+'''\', text);
                            }
                            else
                            {
                            Get_Attribute.set(\'//''' + select_id + '''[''' + str(int(num) -1) + ''']'''+'''\', text[0]);
                            }  
                            });'''
                    else:
                        num = 0
                        # next_line = line.split("//")[1].strip().rsplit("[", 1)[0].replace(")","")
                        if line.split("//",1)[1].count(" ") > 0:
                            next_line = line.split("//",1)[1].rsplit(" ", 1)[0].strip()
                        else:
                            next_line = line.split("//",1)[1].rsplit(" ", 1)[0].strip()
                        print("next_line", next_line)
                        if line.__contains__("]"):
                            select_id = line.split("//",1)[1].split("]")[0] + "]"
                        else:
                            select_id = line.split("//",1)[1].split()[0]

                        if line.split("//",1)[1].count(" ") > 0:
                            select_ele = line.split("//",1)[1].rsplit(" ", 1)[1].strip()

                        else:
                            select_ele = line.split("//",1)[1].strip()

                        print("next_line", next_line)

                        if next_line.lower().startswith("select"):
                            protractor_code += '''browser.sleep('''+str(sleep_delay)+''');\nbrowser.wait(until.presenceOf(element(by.xpath(\'//''' + select_id + '''//option[@value=\"''' + str(select_ele) + '''\"]\'))), delay, 'Element taking too long to appear in the DOM');
                                element.all(by.xpath(\'//''' + select_id + '''//option[@value=\"''' + str(select_ele) + '''\"]\')).getText().then(function (text) {
                            if( typeof text ==='object')
                            {
                            Get_Text.set(\'//''' + select_id + '''//option[@value=\"''' + str(select_ele) + '''\"]\', text);
                            }
                            else
                            {
                            Get_Text.set(\'//''' + select_id + '''//option[@value=\"''' + str(select_ele) + '''\"]\', text[0]);
                            }  
                            });
                            element.all(by.xpath(\'//''' + select_id + '''//option[@value=\"''' + str(select_ele) + '''\"]\')).getAttribute('value').then(function (text) {
                            if( typeof text ==='object')
                            {
                            Get_Attribute.set(\'//''' + select_id + '''//option[@value=\"''' + str(select_ele) + '''\"]\', text);
                            }
                            else
                            {
                            Get_Attribute.set(\'//''' + select_id + '''//option[@value=\"''' + str(select_ele) + '''\"]\', text[0]);
                            }  
                            });
                                element(by.xpath(\'//''' + select_id + '''//option[@value=\"''' + str(select_ele) + '''\"]\')).click();'''
                        else:
                            protractor_code += '''browser.sleep('''+str(sleep_delay)+''');
                                        browser.wait(until.presenceOf(element(by.xpath(\'//''' + select_id + '''\'))), delay, 'Element taking too long to appear in the DOM');
                                        element(by.xpath(\'//''' + select_id + '''\')).clear();
                                        element(by.xpath(\'//''' + select_id + '''\')).sendKeys(\"'''+ str(select_ele) + '''\");
                                        element.all(by.xpath(\'//''' + select_id + '''\')).getText().then(function (text) {
                                        if( typeof text ==='object')
                                        {
                                        Get_Text.set(\'//''' + select_id + '''\', text);
                                        }
                                        else
                                        {
                                        Get_Text.set(\'//''' + select_id + '''\', text[0]);
                                        }  
                                        });
                                        element.all(by.xpath(\'//''' + select_id + '''\')).getAttribute('value').then(function (text) {
                                        if( typeof text ==='object')
                                        {
                                        Get_Attribute.set(\'//''' + select_id + '''\', text);
                                        }
                                        else
                                        {
                                        Get_Attribute.set(\'//''' + select_id + '''\', text[0]);
                                        }  
                                        });'''

                elif line.__contains__("Input Text")and not line.__contains__("nonIndexBased") and not line.__contains__("indexBased")  and not line.__contains__("existingXpath") and not line.__contains__("XpathWithVal"):
                    input_line = line.split("//",1)[1].strip().split()[-1]
                    next_line = line.split("//",1)[1].strip().rsplit(" ", 1)[0].strip()
                    protractor_code += "\nbrowser.wait(until.presenceOf(element(by.xpath(\'//" + next_line + "\'))), delay, 'Element taking too long to appear in the DOM');\nelement(by.xpath(\'//" + next_line + "\')).clear();\nelement(by.xpath(\'//" + next_line + "\')).sendKeys(\"" + str(
                        input_line) + "\");\n"

                elif line.__contains__('nonIndexBased') and not line.__contains__('indexBased')  and not line.__contains__("existingXpath") and not line.__contains__("XpathWithVal"):


                    assertxpathindex = "-"
                    assertvalueindex = "-"
                    assertcondition = fetchCondition(line)

                    assertxpath = line.split('nonIndexBased')[1].split(assertcondition)[0].strip()
                    print(line)
                    xpath = assertxpath
                    assertvalue=line.split(assertcondition)[1].strip()
                    value = assertvalue

                    if assertcondition=="Presence":
                        method="isPresent()"
                    if assertxpath.__contains__("//") and assertxpath.__contains__("[") and (assertxpath.split("[")[1].split("]")[0]).isnumeric():
                        assertxpathindex = assertxpath.split("[")[1].split("]")[0]

                        print(assertxpathindex)
                    if xpath.__contains__("xpath=("):
                        assertxpath = xpath.split("xpath=(")[1].split(")")[0]
                        print(assertxpath)

                    if assertvalue.__contains__("[") and (assertvalue.split("[")[1].split("]")[0]).isnumeric():
                        assertvalueindex = assertvalue.split("[")[1].split("]")[0]
                        print(assertvalueindex)
                    if value.__contains__("xpath="):
                        assertvalue = value.split("xpath=(")[1].split(")")[0]
                        print(assertvalue)
                    if assertxpath.__contains__("//input"):
                        leftmethod="getAttribute('value')"
                    else:
                        leftmethod="getText()"
                    if assertvalue.__contains__("//input"):
                        rightmethod="getAttribute('value')"
                    else:
                        rightmethod="getText()"

                    # CheckWith = "Color"
                    CheckWith = ""

                    if CheckWith == "Color":
                        leftmethod = "getCssValue('color')"
                        rightmethod = "getCssValue('color')"
                    elif CheckWith ==  "Border Color":
                        leftmethod = "getCssValue('border-color')"
                        rightmethod = "getCssValue('border-color')"
                    elif CheckWith == "Background Color":
                        leftmethod = "getCssValue('background-color')"
                        rightmethod = "getCssValue('background-color')"

                    protractor_code+=assertionSyntax(assertxpath,assertxpathindex,assertvalue,assertvalueindex,assertcondition,leftmethod,rightmethod)
                elif line.__contains__('indexBased') and not line.__contains__('nonIndexBased')  and not line.__contains__("existingXpath") and not line.__contains__("XpathWithVal"):

                    assertcondition = fetchCondition(line)
                    assign = line.split('indexBased')[1].split(assertcondition)[0]
                    assertxpathindex = assign.split()[-1]
                    assertxpath = assign.split(assertxpathindex)[0].strip()
                    assertvalueindex = "-"

                    assertvalue = line.split(assertcondition)[1].strip()
                    value = assertvalue

                    if assertcondition=="Presence":
                        method="isPresent()"
                    if assertvalue.__contains__("[") and (assertvalue.split("[")[1].split("]")[0]).isnumeric():
                        assertvalueindex = assertvalue.split("[")[1].split("]")[0]
                        print(assertvalueindex)
                    if value.__contains__("xpath="):
                        assertvalue = value.split("xpath=(")[1].split(")")[0]
                        print(assertvalue)
                    if assertxpath.__contains__("//input"):
                        leftmethod="getAttribute('value')"
                    else:
                        leftmethod="getText()"
                    if assertvalue.__contains__("//input"):
                        rightmethod="getAttribute('value')"
                    else:
                        rightmethod="getText()"
                    # CheckWith = "Color"
                    CheckWith = ""

                    if CheckWith == "Color":
                        leftmethod = "getCssValue('color')"
                        rightmethod = "getCssValue('color')"
                    elif CheckWith == "Border Color":
                        leftmethod = "getCssValue('border-color')"
                        rightmethod = "getCssValue('border-color')"
                    elif CheckWith == "Background Color":
                        leftmethod = "getCssValue('background-color')"
                        rightmethod = "getCssValue('background-color')"

                    protractor_code += assertionSyntax(assertxpath, assertxpathindex, assertvalue, assertvalueindex,
                                                       assertcondition, leftmethod,rightmethod)

                elif line.__contains__("existingXpath") and not line.__contains__("nonIndexBased") and not line.__contains__("indexBased") and not line.__contains__("XpathWithVal"):

                    leftmethod=''
                    rightmethod=''
                    assertxpathindex = "-"
                    assertvalueindex = "-"
                    assertcondition = fetchCondition(line)
                    assertxpath = line.split('existingXpath')[1].split(assertcondition)[0].strip()
                    xpath = assertxpath
                    assertvalue = line.split(assertcondition)[1].strip()
                    value = assertvalue

                    if assertcondition=="Presence":
                        method="isPresent()"
                    if assertxpath.__contains__("//") and assertxpath.__contains__("[") and (assertxpath.split("[")[1].split("]")[0]).isnumeric():
                        assertxpathindex = assertxpath.split("[")[1].split("]")[0]
                        print(assertxpathindex)
                    if xpath.__contains__("xpath=("):
                        assertxpath = xpath.split("xpath=(")[1].split(")")[0]
                        print(assertxpath)

                    if assertvalue.__contains__("[") and (assertvalue.split("[")[1].split("]")[0]).isnumeric():
                        assertvalueindex = assertvalue.split("[")[1].split("]")[0]
                        print(assertvalueindex)
                    if value.__contains__("xpath="):
                        assertvalue = value.split("xpath=(")[1].split(")")[0]
                        print(assertvalue)
                    if assertxpath.__contains__("//input"):
                        leftmethod = "getAttribute('value')"
                        print("LEFT",leftmethod)
                    else:

                        leftmethod = "getText()"
                        print("LEFT", leftmethod)
                    if assertvalue.__contains__("//input"):

                        rightmethod = "getAttribute('value')"
                        print("RIGHT", rightmethod)
                    else:

                        rightmethod = "getText()"
                        print("RIGHT", rightmethod)

                    # CheckWith = "Color"
                    CheckWith = ""
                    if CheckWith == "Color":
                        leftmethod = "getCssValue('color')"
                        rightmethod = "getCssValue('color')"
                    elif CheckWith == "Border Color":
                        leftmethod = "getCssValue('border-color')"
                        rightmethod = "getCssValue('border-color')"
                    elif CheckWith == "Background Color":
                        leftmethod = "getCssValue('background-color')"
                        rightmethod = "getCssValue('background-color')"
                    print("PARAM",assertxpath, assertxpathindex, assertvalue, assertvalueindex,
                                                       assertcondition, leftmethod,rightmethod)
                    protractor_code += assertionSyntax(assertxpath, assertxpathindex, assertvalue, assertvalueindex,
                                                       assertcondition, leftmethod,rightmethod)

                    # refer.append("XpathWithVal" + '\t' + xpath + '\t' + orignalval + '\t' + cond + '\t' + assertval)

                elif line.__contains__("XpathWithVal") and not line.__contains__("nonIndexBased") and not line.__contains__("indexBased") and not line.__contains__("existingXpath"):
                    assertxpath = line.split('XpathWithVal')[1].split("?separate?")[0].strip()
                    xpath = assertxpath
                    assertxpathindex = "-"
                    assertvalueindex = "-"
                    assertcondition = fetchCondition(line)
                    # assertvalue = line.split()[-1].strip()
                    assertvalue = line.split(assertcondition)[1].strip()
                    value = assertvalue

                    if assertxpath.__contains__("//") and assertxpath.__contains__("[") and (assertxpath.split("[")[1].split("]")[0]).isnumeric():
                        assertxpathindex = assertxpath.split("[")[1].split("]")[0]
                        print(assertxpathindex)
                    if xpath.__contains__("xpath=("):
                        assertxpath = xpath.split("xpath=(")[1].split(")")[0]
                        print(assertxpath)

                    if assertvalue.__contains__("[") and (assertvalue.split("[")[1].split("]")[0]).isnumeric():
                        assertvalueindex = assertvalue.split("[")[1].split("]")[0]
                        print(assertvalueindex)
                    if value.__contains__("xpath="):
                        assertvalue = value.split("xpath=(")[1].split(")")[0]
                        print(assertvalue)
                    if assertxpath.__contains__("//input"):
                        leftmethod = "getAttribute('value')"
                    else:
                        leftmethod = "getText()"
                    if assertvalue.__contains__("//input"):
                        rightmethod = "getAttribute('value')"
                    else:
                        rightmethod = "getText()"

                    # CheckWith = "Color"
                    CheckWith = ""

                    if CheckWith == "Color":
                        leftmethod = "getCssValue('color')"
                        rightmethod = "getCssValue('color')"
                    elif CheckWith == "Border Color":
                        leftmethod = "getCssValue('border-color')"
                        rightmethod = "getCssValue('border-color')"
                    elif CheckWith == "Background Color":
                        leftmethod = "getCssValue('background-color')"
                        rightmethod = "getCssValue('background-color')"

                    protractor_code += assertionSyntax(assertxpath, assertxpathindex, assertvalue, assertvalueindex,
                                                       assertcondition, leftmethod,rightmethod)


        protractor_code += "browser.sleep("+str(sleep_delay)+");"

        spec_file = provideSpec(url,outerFolderName,newfolderName)
        final_line = spec_file + "\n" + protractor_code + "});\n});"
        new_spec_file.write(final_line)
        new_spec_file.close()
        mapScreenshotWithSpec(outerFolderName,newfolderName)
        subprocess.call(base_location + '\\' + outerFolderName + '\\' + newfolderName + '\\batch.bat')
        moveReport(outerFolderName,newfolderName)
        log_filename="conf.log"
        conf_log = conf_log_modification(log_filename)
        output_json = {"status": "Success", "conf_log": conf_log}
        print("output log", output_json)

        return jsonify(output_json)
        # return "success"

    except Exception as e:
        return jsonify(e)


def fetchCondition(line):
    assertcondition = ""
    if line.__contains__("Equal"):
        assertcondition="Equal"
    if line.__contains__("NotEqual"):
        assertcondition="NotEqual"
    if line.__contains__("Contains"):
        assertcondition="Contains"
    if line.__contains__("Presence"):
        assertcondition="Presence"
    return assertcondition


def assertionSyntax(assertXpath, assertXpathIndex, assertvalueXpath, assertvalueXpathIndex, condition, leftmethod,rightmethod):
    syntax = ''
    assign1 = ''
    assign2 = ''
    find_in_dict=''
    # text='text'
    if not assertvalueXpathIndex == "-":
        assign1 = '''element.all(by.xpath(`''' + assertvalueXpath + '''`)).get(''' + assertvalueXpathIndex + ''')'''
        find_in_dict=assertvalueXpath+"["+assertvalueXpathIndex+"]"
    if assertvalueXpathIndex == "-":
        assign1 = '''element.all(by.xpath(`''' + assertvalueXpath + '''`))'''
        find_in_dict=assertvalueXpath
    if not assertXpathIndex == "-":
        assign2 = '''element.all(by.xpath(`''' + assertXpath + '''`)).get(''' + assertXpathIndex + ''')'''
    if assertXpathIndex == '-':
        assign2 = '''element.all(by.xpath(`''' + assertXpath + '''`))'''
    # if method=="getAttribute('value')":
    #     text='text[0]'
    if assertvalueXpath.startswith("//"):

        if condition == "Equal":
            syntax = '''browser.sleep('''+str(sleep_delay)+''');
                       ''' + assign1 + '''.isPresent().then(function (text) {
                           console.log("presence");
                           var assert='';
                           if (text == true){
                               console.log('xpath in same page');
                               
                                ''' + assign1 + '''.''' + rightmethod + '''.then(function (text){
                                    if( typeof text ==='object')
                                        {
                                        assert=text[0];
                                        }
                                        else
                                        {
                                        assert=text;
                                        }  
                                    console.log("assert",assert);});
                                }
                           else{
                                if(\"''' + rightmethod + '''\"=="getText()"){
                               assert=Get_Text.get(`''' + find_in_dict + '''`)  }
                               if(\"''' + rightmethod + '''\"=="getAttribute('value')"){
                               assert=Get_Attribute.get(`''' + find_in_dict + '''`)  }
                               if( typeof assert ==='object')
                                {
                                assert=assert[0];
                                }
                                else
                                {
                                assert=assert;
                                }
                                }
                             ''' + assign2 + '''.''' + leftmethod + '''.then(function (text) {
                           var assert1='';
                           if( typeof text ==='object')
                            {
                            assert1=text[0];
                            }
                            else
                            {
                            assert1=text;
                            }
                        console.log(assert);
                        if (assert1 == assert){
                            console.log('Assertion Passed: Expected', assert1, `for the Xpath''' + assertXpath + ''' contain`, assert ,'. :Assertion Passed');
                        }else{
                            console.log('Assertion Failed: Expected', assert1, `for the Xpath''' + assertXpath + ''' But has`, assert , '. :Assertion Failed');
                            browser.quit();
                        }
                            });

                       });  
     
                '''

        if condition == "Contains":
            syntax ='''browser.sleep('''+str(sleep_delay)+''');
                                   ''' + assign1 + '''.isPresent().then(function (text) {
                                       console.log("presence");
                                       var assert='';
                                       if (text == true){
                                           console.log('xpath in same page');

                                            ''' + assign1 + '''.''' + rightmethod + '''.then(function (text){
                                                if( typeof text ==='object')
                                                    {
                                                    assert=text[0];
                                                    }
                                                    else
                                                    {
                                                    assert=text;
                                                    }  
                                                console.log("assert",assert);});
                                            }
                                       else{
                                            if(\"''' + rightmethod + '''\"=="getText()"){
                                           assert=Get_Text.get(`''' + find_in_dict + '''`)  }
                                           if(\"''' + rightmethod + '''\"=="getAttribute('value')"){
                                           assert=Get_Attribute.get(`''' + find_in_dict + '''`)  }
                                           if( typeof assert ==='object')
                                            {
                                            assert=assert[0];
                                            }
                                            else
                                            {
                                            assert=assert;
                                            }
                                            }
                                         ''' + assign2 + '''.''' + leftmethod + '''.then(function (text) {
                                       var assert1='';
                                       if( typeof text ==='object')
                                        {
                                        assert1=text[0];
                                        }
                                        else
                                        {
                                        assert1=text;
                                        }
                                    console.log(assert);
                                    if (assert1.indexOf(assert)!== -1){
                                   console.log('Assertion Passed: Expected', assert1, `for the Xpath''' + assertXpath + ''' contain `, assert, '. :Assertion Passed');
                                    }else{
                                   console.log('Assertion Failed: Expected', assert1, `for the Xpath''' + assertXpath + ''' But has`, assert , '. :Assertion Failed');
                                   browser.quit();
                                    }
                                        });

                                   });  

                            '''

        if condition == "NotEqual":
            syntax = '''browser.sleep('''+str(sleep_delay)+''');
                       ''' + assign1 + '''.isPresent().then(function (text) {
                           console.log("presence");
                           var assert='';
                           if (text == true){
                               console.log('xpath in same page');
                               
                                ''' + assign1 + '''.''' + rightmethod + '''.then(function (text){
                                    if( typeof text ==='object')
                                        {
                                        assert=text[0];
                                        }
                                        else
                                        {
                                        assert=text;
                                        }  
                                    console.log("assert",assert);});
                                }
                           else{
                                if(\"''' + rightmethod + '''\"=="getText()"){
                               assert=Get_Text.get(`''' + find_in_dict + '''`)  }
                               if(\"''' + rightmethod + '''\"=="getAttribute('value')"){
                               assert=Get_Attribute.get(`''' + find_in_dict + '''`)  }
                               if( typeof assert ==='object')
                                {
                                assert=assert[0];
                                }
                                else
                                {
                                assert=assert;
                                }
                                }
                             ''' + assign2 + '''.''' + leftmethod + '''.then(function (text) {
                           var assert1='';
                           if( typeof text ==='object')
                            {
                            assert1=text[0];
                            }
                            else
                            {
                            assert1=text;
                            }
                        console.log(assert);
                        if (assert1 == assert){
                                        console.log('Assertion Failed: Expected', assert1, `for the Xpath''' + assertXpath + ''' and has same value`, assert,'. :Assertion Failed');
                                        browser.quit();

                                    }else{
                                        console.log('Assertion Passed: Expected', assert1, `for the Xpath''' + assertXpath + ''' contain `, assert,'. :Assertion Passed');

                                    }
                            });

                       });  
     
                '''

        if condition == "Presence":
            syntax = '''browser.sleep('''+str(sleep_delay)+''');
                       ''' + assign2 + '''.isPresent().then(function (text) {
                           console.log("presence");
                           if (text == true){
                               console.log('Assertion Passed: Expected element xpath ''' + assertXpath + ''' is present in current page. :Assertion Passed');


                           }else{
                               console.log(`Assertion Failed: Expected element xpath''' + assertXpath + ''' is not present in current page. :Assertion Failed`);
                               browser.quit();
                           }

                       });

                   '''

    if not assertvalueXpath.startswith("//"):
        if condition == "Equal":
            syntax ='''browser.sleep('''+str(sleep_delay)+''');
                    var assert='';
                    ''' + assign2 + '''.''' + leftmethod + '''.then(function (text) {
                    if( typeof text ==='object')
                            {
                            assert=text[0];
                            }
                            else
                            {
                            assert=text;
                            }
                        if (assert == \"''' + assertvalueXpath + '''\"){
                            console.log('Assertion Passed: Expected', assert, `for the Xpath''' + assertXpath + ''' contain ''' + assertvalueXpath + '''. :Assertion Passed` );

                        }else{
                            console.log('Assertion Failed: Expected', assert, `for the Xpath''' + assertXpath + ''' But has ''' + assertvalueXpath + '''. :Assertion Failed`);
                            browser.quit();
                        }

                    });

                '''
        if condition == "Contains":
            syntax = '''browser.sleep('''+str(sleep_delay)+''');
                   var assert='';
                    ''' + assign2 + '''.''' + leftmethod + '''.then(function (text) {
                    if( typeof text ==='object')
                            {
                            assert=text[0];
                            }
                            else
                            {
                            assert=text;
                            }
                        if (assert.indexOf( \"''' + assertvalueXpath + '''\")!== -1){
                            console.log('Assertion Passed: Expected', assert, `for the Xpath''' + assertXpath + ''' contain ''' + assertvalueXpath + '''. :Assertion Passed` );

                        }else{
                            console.log('Assertion Failed: Expected', assert, `for the Xpath''' + assertXpath + ''' But has ''' + assertvalueXpath + '''. :Assertion Failed`);
                            browser.quit();
                        }

                    });

                '''

        if condition == "NotEqual":
            syntax = '''browser.sleep('''+str(sleep_delay)+''');
                var assert='';
                    ''' + assign2 + '''.''' + leftmethod + '''.then(function (text) {
                    if( typeof text ==='object')
                            {
                            assert=text[0];
                            }
                            else
                            {
                            assert=text;
                            }
                    if (assert == \"''' + assertvalueXpath + '''\"){
                        console.log('Assertion Failed: Expected', assert, `for the Xpath''' + assertXpath + ''' and has same value ''' + assertvalueXpath + '''. :Assertion Failed`);
                        browser.quit();
                    }else{
                        console.log('Assertion Passed: Expected', assert, `for the Xpath''' + assertXpath + ''' contain ''' + assertvalueXpath + '''. :Assertion Passed` );

                    }

                });

            '''
        # if condition == "Presence":
        #     syntax = '''
        #             ''' + assign2 + '''.''' + method + '''.then(function (text) {
        #                 if (text == true){
        #                     console.log("assertion passed for  ", text);
        #
        #                 }else{
        #                     console.log('Expected element xpath'''+ assertXpath + '''is not present in current page, So Failed in assertion');
        #                     browser.quit();
        #                 }
        #
        #             });
        #
        #         '''

    return syntax


def check_conditions(url):
    global outerFolderName
    global folderName
    try:
        xpath1 = request.args.get("xpath_list").split(",")  # get the xpaths from user and split using comma
        original_value1 = request.args.get("value_list").split(
            ",")  # get the original values from user and split using comma
        asserted_value1 = request.args.get("assert_list").split(
            ",")  # get the asserted values from user and split using comma
        flag=0
        print("values", xpath1, original_value1, asserted_value1)
        # readSpec=open(base_location + "\\" + outerFolderName + "\\" + folderName + "\\spec.js","r").readlines()
        xpathFile=base_location+"\\"+outerFolderName+"\\"+folderName+"\\new_file.txt"

        for i in xpath1:
            if not i=='':
                flag=1
                break
            else:
                break
        if flag==1:


            count = db.folderName.count_documents(
                {"$and": [{"appName": outerFolderName}, {"scenarioName": {'$regex': '^' + folderName}}]})
            # occur = count
            occurr = str(count)
            newDir = base_location + "\\" + outerFolderName + "\\" + folderName + "_data" + occurr
            folderName=folderName + "_data" + occurr
            if not os.path.exists(newDir):
                os.makedirs(newDir)
            shutil.copy(xpathFile, newDir)
            newSpec = open(newDir + "\\spec.js", "w+")
            newConf=open(newDir+"\\conf.js","w+")
            spec_file_name = "spec.js"
            conf_file = provideConf(spec_file_name,base_location,outerFolderName,folderName)
            newConf.write(conf_file)
            f = open(newDir+ "\\batch.bat", "w+")
            script1 = '''
                      START /B cmd "cd ''' + newDir + '''\\" && protractor ''' + newDir + '''\\conf.js >''' + newDir + '''\\\conf.log
                      '''
            f.write(script1)
            xpathFile1=base_location+"\\"+outerFolderName+"\\"+folderName+"\\new_file.txt"

            for (xpath, original_value, asserted_val) in zip(xpath1, original_value1,
                                                             asserted_value1):  # iterate through xpath,condition,original value and asserted value list
                original_value = original_value.strip()  # strip spaces
                asserted_val = asserted_val.strip()  # strip spaces
                xpath = xpath.strip()  # strip spaces
                xpath = xpath.replace("HASH", "#")
                if not xpath == '':
                    with fileinput.FileInput(xpathFile1, inplace=True) as file:
                        text_to_search= original_value
                        replacement_text= asserted_val
                        for line1 in file:
                            print(line1.replace(text_to_search, replacement_text), end='')
            spec_file = provideSpec(url,outerFolderName,folderName)
            protractor_code=generateSpec(outerFolderName,folderName)
            final_line = spec_file + "\n" + protractor_code + "});\n});"
            newSpec.write(final_line)
            newSpec.close()
            mapScreenshotWithSpec(outerFolderName, folderName)
            description=testDescription()
            print("DESC", description)
            dbData = [{"path": base_location + "\\" + outerFolderName,
                       "innerpath": base_location + "\\" + outerFolderName + "\\" + folderName,
                       "appName": outerFolderName,
                       "scenarioName": folderName,
                       "test_Description": description}]  # db structure update
            db.folderName.insert_many(dbData)
            if not description=='':
                return "success"
        else:
            return "success"
    except Exception as e:
        return jsonify(e)


def conf_log_fun(path):
    try:
        f_open = open(path + "\\conf.log")  # open conf.log
        content = f_open.read()  # read the content in conf.log
        list1 = content.split("\n")  # split content using new line
        jsonformat = {"output": content}  # store the content in jsonformat
        print(jsonformat)
        return content  # return content of conf.log
    except Exception as e:
        return jsonify(e)  # return the exception occurred


def outputValidation(final_line):
    try:
        popup = request.args.get("popup")
        # popup= "Webpage Element"
        name_id = request.args.get("nameId")
        # name_id= 'id'
        output_string = request.args.get("outputString")
        # output_string = "output"
        #
        idVal = request.args.get("idVal")
        # idVal= "Ok"
        print("popup")

        if (popup != None or popup != "other" or name_id != None or output_string != None or idVal != None):
            print("popup 1", popup, name_id, output_string, "idval", idVal)
            if (popup == "browser-pop-up" and idVal != None):
                # script = open("F:\iTap-Backend\Id_test.txt",'a')
                final_line += '''browser.switchTo().alert().then((alrt) => {
                            alrt.getText().then(function(txt)
                               {expect(txt.toLowerCase()).toContain("''' + idVal + '''".toLowerCase())});alrt.accept();});'''

            elif (
                    name_id != None and idVal != None and output_string != None and name_id != "" and idVal != "" and output_string != ""):
                # script = open("F:\iTap-Backend\Id_test.txt", 'a')
                final_line += (
                        "element(by.xpath(\"//*[@" + name_id + "='" + output_string + "']\")).getText().then(function(txt)"
                                                                                      "{expect(txt.toLowerCase()).toContain(\'" + idVal + "\'.toLowerCase());});")
            print("final line", final_line)
            return final_line
    except Exception as e:
        return jsonify(e)

@app.route("/assertionDataTable")
def assertionDataTable():
    """
    1. Find the last entry in db and get the scenario name, to perform assertion on that scenario.
    2. The above step is done to perform assertion on the latest set of data that's performed rerun.
    3. Get the necessary values for the table, from new_file.txt which contains recorded xpaths.
    :return: Data to be displayed on the assertion table.
    """
    try:
        newfolderName = ''
        find_folder = db.folderName.find().sort("_id", -1).limit(1) # sort db in descending order and limit to 1, to get last record.
        for data in find_folder:
            newfolderName = data["scenarioName"]
        f_open = open(base_location + "\\" + outerFolderName + "\\" + newfolderName + "\\new_file.txt","r")
        content = f_open.read()
        xpath = []
        original_value = []
        line_numers = []
        split_file = content.split(",")
        for idx, line in enumerate(split_file):
            if line.split()[-1].__contains__("//"):
                xpath.append(line.split()[-1])
                original_value.append("-")
                line_numers.append(idx + 1)
            if line.split()[-1].__contains__("//") == False and line.__contains__("Open Browser") == False:
                xpathvar = line.split()[2].strip()
                val = line.split(xpathvar[-1])[-1].strip()
                if xpathvar.startswith(("//select")):
                    xpath.append(xpathvar + '//option[@value="' + val + '"]')
                    original_value.append("-")
                else:
                    xpath.append(xpathvar)
                    original_value.append(val)
                line_numers.append(idx + 1)

        outDict = {}
        finalList = []
        finalDict = {}
        assertConditionDropDown = [
            "Equal",
            "NotEqual",
            "Contains",
            "Presence"
        ]
        xpathdropdown=[]
        xpathdropdown.append("")
        for x_path in xpath:
            xpathdropdown.append(x_path)
        for (path, value, line_numer) in zip(xpath, original_value,line_numers):
            outDict["xpath"] = path
            outDict["value"] = value
            outDict["index"] = ""
            outDict["operation"] = assertConditionDropDown
            outDict["Assertion Value"] = Remove(xpathdropdown) # remove duplicates from list
            outDict["stepNumber"] = line_numer
            finalList.append(outDict)
            outDict = {}
        finalDict['headers'] = ["checkbox", "stepNumber", "xpath","index", "value","operation", "Assertion Value","Add/Remove"]
        finalDict["data"] = finalList
        print("final dict", finalDict)
        return finalDict
    except Exception as e:
        return jsonify(e)

def assertPage():
    '''
    1.Read new_file.txt and get all data to be displayed.
    2.Insert data generated for table, into db.
    :return: Data to be displayed on the alternate data table.
    '''
    try:
        shutil.copy(download_location,
                    base_location + "\\" + outerFolderName + "\\" + folderName + "\\" + "new_file.txt")
        f_open = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\\new_file.txt","r")
        content = f_open.read()
        xpath = []
        original_value = []
        line_numers = []
        split_file = content.split(",")

        for idx, line in enumerate(split_file):
            if line.split()[-1].__contains__("//"):
                xpath.append(line.split()[-1])
                original_value.append("-")
                line_numers.append(idx + 1)

            if line.split()[-1].__contains__("//") == False and line.__contains__("Open Browser") == False:
                xpath.append(line.split()[2])
                original_value.append(line.split(xpath[-1])[-1].strip())
                line_numers.append(idx + 1)
        outDict = {}
        finalList = []
        finalDict = {}
        for (path, value, line_numer) in zip(xpath, original_value,
                                             line_numers):
            outDict["xpath"] = path
            outDict["value"] = value
            outDict["Alternative Value"] = ""
            outDict["stepNumber"] = line_numer
            finalList.append(outDict)
            outDict = {}

        finalDict['headers'] = ["checkbox", "stepNumber", "xpath", "value", "Alternative Value","Add/Remove"]
        finalDict["data"] = finalList
        outputList = finalList

        try:
            db.assertTable.delete_many({})
        except Exception as e:
            print('Error while deleting missing components report:' + str(e))
        db.assertTable.insert_many(outputList)
        for l in finalList:
            if "_id" in l.keys():
                del l["_id"]
        print("final list", finalList)
        print("final dict", finalDict)
        return finalDict
    except Exception as e:
        return jsonify(e)

@app.route("/testDescription")  # Test case description
def testDescription():
    global outerFolderName
    global folderName
    open_file = open(base_location + "\\" + outerFolderName + '\\' + folderName + "\\spec.js")
    lines = open_file.readlines()
    full_line = ''
    flag = 0
    sentence = ''
    for line in lines:
        if line.__contains__("extractModal();"):
            break
        if line.strip().startswith("browser.get("):
            flag = 1
        if flag == 1:
            full_line += line

    number = 0
    for sline in full_line.split(";"):
        if sline.__contains__("browser.get("):  # Get Url
            number += 1
            url_line = sline.split("browser.get(")[1].split(")")[0]
            if len(url_line) < 50:
                sentence += str(number) + ". Open browser and navigate to URL: " + str(
                    url_line) + " <br>"
            else:
                sentence += str(number) + ". Open browser and navigate to URL: " + "<br>" + str(
                    url_line) + " <br>"
        # if sline.__contains__("browser.sleep("):
        #   sentence+="Browser will wait for "+sline.split("browser.sleep(")[1].split(")")[0]+"MilliSeconds<br>"
        if sline.__contains__("until.presenceOf") == False and \
                sline.__contains__("by.xpath(") and \
                sline.strip().endswith("click()"):  # button click
            if sline.__contains__("@id="):
                number += 1
                sentence += str(number) + ". Click on the button with id " + sline.split("@id=")[1].split("]")[
                    0] + "<br>"
            elif sline.__contains__("@name="):
                number += 1
                sentence += str(number) + ". Click on the button with name " + sline.split("@name=")[1].split("]")[
                    0] + "<br>"
            elif sline.__contains__(".get("):
                number += 1
                x = str(sline.split(".get(")[1].split(")")[0]).strip()
                if x.endswith("1"):
                    inBetween = str(floor(int(x) / 10)) + "1st"
                elif x.endswith("2"):
                    inBetween = str(floor(int(x) / 10)) + "2nd"
                elif x.endswith("3"):

                    inBetween = str(floor(int(x) / 10)) + "3rd"
                else:
                    inBetween = x + "th"
                sentence += str(number) + ". Click on " + inBetween + " button in the current page<br>"

        elif sline.__contains__("until.presenceOf") == False and \
                sline.__contains__("by.xpath(") and \
                sline.__contains__("click()") == False and \
                sline.__contains__("sendKeys"):  # Sendkey
            if sline.__contains__("@id="):
                number += 1
                if sline.__contains__(".get("):
                    x = str(sline.split(".get(")[1].split(")")[0]).strip()
                    if x.endswith("1"):
                        inBetween = str(floor(int(x) / 10)) + "1st"
                    elif x.endswith("2"):
                        inBetween = str(floor(int(x) / 10)) + "2nd"
                    elif x.endswith("3"):
                        inBetween = str(floor(int(x) / 10)) + "3rd"
                    else:
                        inBetween = x + "th"
                        if sline.lower().__contains__("password"):
                            sentence += str(
                                number) + ". Enter value " + "******" + " on the" + inBetween + "field with id " + \
                                        sline.split("@id=")[1].split("]")[0] + "<br>"
                        else:
                            sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                                0] + " on the" + inBetween + "field with id " + sline.split("@id=")[1].split("]")[
                                            0] + "<br>"
                else:
                    if sline.lower().__contains__("password"):
                        sentence += str(number) + ". Enter value " + "******" + " in the field with id " + \
                                    sline.split("@id=")[1].split("]")[0] + "<br>"
                    else:
                        sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                            0] + " in the field with id " + sline.split("@id=")[1].split("]")[0] + "<br>"
            elif sline.__contains__("@name="):
                number += 1
                if sline.__contains__(".get("):
                    x = str(sline.split(".get(")[1].split(")")[0]).strip()
                    if x.endswith("1"):
                        inBetween = str(floor(int(x) / 10)) + "1st"
                    elif x.endswith("2"):
                        inBetween = str(floor(int(x) / 10)) + "2nd"
                    elif x.endswith("3"):
                        inBetween = str(floor(int(x) / 10)) + "3rd"
                    else:
                        inBetween = x + "th"
                        if sline.lower().__contains__("password"):
                            sentence += str(number) + ". Enter value " + "******" + " in the field with name " + \
                                        sline.split("@name=")[1].split("]")[0] + "<br>"
                        else:
                            sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                                0] + " in the field with name " + sline.split("@name=")[1].split("]")[0] + "<br>"
                else:
                    if sline.lower().__contains__("password"):
                        sentence += str(number) + ". Enter value " + "*****" + " in the field with name " + \
                                    sline.split("@name=")[1].split("]")[0] + "<br>"
                    else:
                        sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                            0] + " in the field with name " + sline.split("@name=")[1].split("]")[0] + "<br>"
            elif sline.__contains__(".get(") and sline.__contains__("@id=") == False:
                number += 1
                x = str(sline.split(".get(")[1].split(")")[0]).strip()
                if x.endswith("1"):
                    inBetween = str(floor(int(x) / 10)) + "1st"
                elif x.endswith("2"):
                    inBetween = str(floor(int(x) / 10)) + "2nd"
                elif x.endswith("3"):
                    inBetween = str(floor(int(x) / 10)) + "3rd"
                else:
                    inBetween = x + "th"
                # sentence += "Enter value " + sline.split("sendKeys(")[1].split(")")[0] + " in the field with id " + \
                #             inBetween+" in the field from the webpage<br>"
                sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                    0] + " on " + inBetween + " field in the current page<br>"
            elif sline.__contains__("//input"):
                number +=1
                sentence += str(number)+". Entering input as " + sline.split("sendKeys(")[1].split(")")[0] + " in the current page<br>"

    # new_file.write(sentence)
    print(sentence)
    # connections = db.folderName.find()
    # myquery = {"appName": outerFolderName, "scenarioName": folderName}
    # newvalues = {"$set": {"test_Description": sentence}}
    # db.folderName.update_many(myquery, newvalues)
    return sentence



@app.route("/conf_log_modification")
def conf_log_modification(log_filename):
    """
    1.Read the log file that contains the test results.
    2.Get all details - Test case status, time of execution, date, scenario name, error line number, error message.
    3.assign all values to respective headers in a json and write the json into new_conf_log.log file.
    :param log_filename: The log filename that contains the results of a test run.
    :return: Necessary data from the log file to be displayed in UI -- Data for consolidated log table.
    """
    try:
        assert_flag = False
        t = time.localtime()  # get the time and assign it to variable t
        current_time = time.strftime("%H:%M:%S", t)  # formatting time in hour minute second format

        # Date
        now = datetime.now()
        today = date.today()  # get the date and store in today variable
        d1 = today.strftime("%d/%m/%Y")  # format the date into day month year
        global outerFolderName  # use outerFolderName from outside this scope
        global folderName  # use folderName from outside this scope
        # spec_file = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\\" + "spec.js", "r").readlines()
        # xpath_file = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\\" + "new_file.txt", "r").readlines()
        open_file = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\\" + log_filename,
                         "r")  # open conf.log in read mode
        open_conf_file = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\\" + "new_conf_log.log",
                              "a")  # open new_conf.log in append mode
        buf = open_file.readlines()  # read file
        final_report = {}  # initialize an empty dictionary
        report1 = {}  # initialize an empty dictionary
        message = []  # initialize an empty list
        final_list = []
        spec_line = []  # initialize an empty list
        report1['Date'] = d1  # assign date
        report1["Time"] = current_time  # assign time
        report1["Application Name"] = outerFolderName  # assign outer folder name
        report1["Scenario Name"] = folderName  # assign folder name
        count = 0  # initialize count to 0
        count1 = 0  # initialize count1 to 0
        count2 = 0  # initialize count2 to 0
        assert_message = ''
        multi_line = False
        log_index = 0
        for lines in buf:  # iterate through lines in conf.log file
            # print("line", lines)
            if lines.__contains__("Assertion Passed:") and lines.__contains__(":Assertion Passed"):
                assert_flag = True

                log_index += 1

                assert_message += str(log_index) + ". " + lines.replace(":Assertion Passed", '') + "\n"

            elif lines.__contains__("Assertion Passed:"):
                multi_line = True
                assert_flag = True
                log_index += 1

                assert_message += str(log_index) + ". " + lines
                # assert_message+=lines.replace(":Assertion Passed",'')+"\n"
            elif lines.__contains__(":Assertion Passed"):
                multi_line = False
                assert_message += lines.replace(":Assertion Passed", '') + "\n"
            elif lines.__contains__("Assertion Failed:") and lines.__contains__(":Assertion Failed"):
                assert_flag = True
                log_index += 1
                assert_message += str(log_index) + ". " + lines.replace(":Assertion Failed", '') + "\n"
            elif lines.__contains__("Assertion Failed:"):
                assert_flag = True
                multi_line = True
                log_index += 1
                assert_message += str(log_index) + ". " + lines
            elif lines.__contains__(":Assertion Failed"):
                multi_line = False
                assert_message += lines.replace(":Assertion Failed", '') + "\n"
            elif multi_line:
                assert_message += lines
            if lines.__contains__("[32m."):  # if line contains 32m ,test case passed
                report1['Status'] = "passed"  # assign passed to status in dictionary
            if lines.__contains__("[31mF"):  # if line contains 31mF ,test case passed
                report1['Status'] = "failed"  # assign failed to status in dictionary
            count += lines.count("[32m.")  # count the number of lines containing fail string (32m)
            count += lines.count("[31mF")  # count the number of lines containing pass string (31mF)
            count1 += lines.count("Finished in")  # count the number of lines containing 'Finished in' string
            count2 += lines.count(
                "instance(s) of WebDriver")  # count the number of lines containing Instances of webdriver string
            if lines.__contains__("Finished in"):  # if lines contain finishe in split and take the time
                ftime = lines.split()[2].strip()  # split using space and take second index element
                time1 = ftime.split(".")[0]  # split using dot and take zeroth index element
                time2 = ftime.split(".")[1]  # split using dot and take first index element
                report1[
                    'Time in Seconds'] = time1 + "s" + time2 + "ms"  # add finished in time to the dictionary in format i.e 20s156ms seconds and milliseconds
            if lines.__contains__(
                    "instance(s) of WebDriver"):  # if line contains webdriver instance get the number of instances running
                report1['WebDriver Instance'] = lines.split("-")[
                    1].strip()  # split using hyphen and get the first index element
            if lines.__contains__("31m") and lines.__contains__(
                    "Failed:"):  # if line contains 31m and failed get the fail message
                a = lines.split("Failed:")[1].strip()  # split using failed string and take the first index element
                if a.__contains__(":"):  # if splitted and taken elem contains colon
                    # message.clear()
                    err_message = a.split(":")[0].strip() + "\n"  # split using colon and append the message
                else:  # if splitted and taken elem not contains colon
                    # message.clear()
                    err_message = a.strip() + "\n"  # append the message
                if not err_message.__contains__("This driver instance does not have a valid session ID"):
                    message.append(err_message)
            if lines.__contains__("spec.js"):  # if line contains spec.js split and take the error line numbers
                num = lines.split(":")[2].strip()  # split using colon
                if int(num) > 9:  # if number obtained is greater than 6
                    spec_line.append(num)  # append number
        if count == 0:  # if conf.log is empty
            report1['Status'] = "failed"
        if count1 == 0:  # if conf.log is empty
            report1['Time in Seconds'] = "-"
        if count2 == 0:  # if conf.log is empty
            report1['WebDriver Instance'] = "-"

        if assert_flag:
            message.append(assert_message)

        report1['Error Message'] = message  # store message in dictionary
        report1['Error lines in spec'] = spec_line  # store error line numbers in dictionary
        report1[
            'Test Description'] = testDescription()  # call testDescription and get a formatted description of user actions .
        # # modified_report = str(report1).replace(",",
        #                                        ",\n")  # replace comma with new line in report1 json , convert json to string and store it in modified_report variable
        modified_report = str(report1)
        print("mod_rep", modified_report)
        open_conf_file.write(
            modified_report + ",\n\n")  # writen the modified report into new_conf_log.log adding \n\n for formatting
        open_conf_file.close()  # close new_conf_log.log file
        open_conf = open(
            base_location + "\\" + outerFolderName + "\\" + folderName + "\\" + "new_conf_log.log")  # open new_conf_log.log file
        read_conf = open_conf.read().strip().strip("\n")  # read and strip new lines
        full_dict = []  # initialize an empty list

        for d in read_conf.split("},"):  # split
            print("color", d)
            if d.strip() != "":  # if element is not empty
                dct = eval(d + "}")  # evaluate the string and concat }
                print("dct", dct)
                full_dict.append(dct)  # append contents of new_conf_log.log to full_dict dictionary
            # elif d.__contains__("rgba"):
            #     dct = d + "}"
            #     full_dict.append(dct)
        print("full dictionary", full_dict)

        final_report['headers'] = ['Date', 'Time', 'Application Name', 'Scenario Name',
                                   'Error lines in spec', 'Error Message', 'Status', 'Time in Seconds',
                                   'WebDriver Instance', 'Actions']  # assign necessary headers

        # final_list.append(b)
        final_report['data'] = full_dict  # assign full_dict contaioning contents of new_conf_log.log to a dictionary

        print(final_report)
        return final_report  # return json
    except Exception as e:
        return jsonify(e)  # return exception in json format

@app.route("/insertNewRow")
def insertNewRow():
    """
    1.Get the values from the table and the index where the new row has to be inserted.
    2.Insert empty values to the position in list, where a new row has to be added.
    3.Specify the necessary dropdowns for the new empty row that has to be inserted.
    4.Return the data that has to be displayed in the alternate value table, after clicking on add new row.
    :return: Data for alternate value table, after clicking on add new empty row.
    """
    try:
        index = request.args.get("index")  # get index
        xpath_list_row = request.args.get("xpath_list_row").split(",")  # get xpath and split using comma
        value_list_row = request.args.get("value_list_row").split(",")  # get value and split using comma
        assert_list_row = request.args.get("assert_list_row").split(",")  # get asserted value and split using comma
        stepNumber = request.args.get("step_number").split(",")  # get stepNumber and split using comma
        xpath = []  # initialize a list for xpath
        assert_val = []  # initialize a list for asserted value
        original_value = []  # initialize a list for original value
        step_number_list = []  # initialize a list for step numbers
        index_value = 0
        lastindex = len(xpath_list_row) + 1
        newindex = int(index) + 1
        # iterate through stepNumber,xpath,condition,value,asserted value lists.

        for step_number, xpath_list, value_list, assert_list in zip(stepNumber, xpath_list_row,
                                                                                value_list_row,
                                                                               assert_list_row):
            index_value += 1
            xpath_list = xpath_list.replace("HASH", "#")
            if newindex == lastindex and index_value+1 == lastindex:
                step_number_list.extend([step_number,""])  # append stepnumber
                xpath.extend([xpath_list,"xpath name"])  # append the xpath
                original_value.extend([value_list,"value name"])  # append the value
                assert_val.extend([assert_list,""])
                break

            if newindex == index_value:  # check the condition
                print("came in")
                xpath.append("xpath name")  # append xpath name to xpath list
                original_value.append("value name")  # append value name to original value list
                assert_val.append("")  # append empty value since it has to be entered by user
                step_number_list.append("")  # append empty since this is a new step , not in test description
            step_number_list.append(step_number)  # append stepnumber
            xpath.append(xpath_list)  # append the xpath
            original_value.append(value_list)  # append the value
            assert_val.append(assert_list)  # append assert_list to assert_val list


        print(xpath, original_value, assert_val)
        outDict = {}
        finalList = []
        finalDict = {}
        for (step_number, path, value, assert_value) in zip(step_number_list, xpath, original_value,
                                                                        assert_val):  # iterate through step number, xpath, original value, asserted value and condition all at once
            outDict["xpath"] = path  # assign xpath to dictionary
            outDict["value"] = value  # assign original value to dictionary
            outDict["Alternative Value"] = assert_value  # assign asserted value to dictionary
            outDict["stepNumber"] = step_number  # assign step number to dictionary
            finalList.append(outDict)  # append the json to finalList after each iteration
            outDict = {}
        finalDict['headers'] = ["checkbox", "stepNumber", "xpath", "value", "Alternative Value",
                                "Add/Remove"]  # assign headers as value to finalDict
        finalDict["data"] = finalList  # assign the finaList to key - data
        outputList = finalList  # assign finalList to outputList
        try:
            db.assertTable.delete_many({})  # delete empty
        except Exception as e:
            print('Error while deleting missing components report:' + str(e))  # print exception

        db.assertTable.insert_many(outputList)  # insert the outputList into db
        for l in finalList:  # iterate through finalList
            if "_id" in l.keys():  # if _id is a key in finaList
                del l["_id"]  # delete the key value pair
        print("final list", finalList)
        print("final dict", finalDict)
        return finalDict  # return finalDict json
    except Exception as e:
        return jsonify(e)  # return exception in json format


@app.route("/deleteRow")
def deleteRow():
    """

    :return:
    """
    try:

        index = int(request.args.get("index"))
        xpath_list_row = request.args.get("xpath_list_row").split(",")
        value_list_row = request.args.get("value_list_row").split(",")
        assert_list_row = request.args.get("assert_list_row").split(",")
        stepNumber = request.args.get("step_number").split(",")


        dupliacte_xpath_list = copy.deepcopy(xpath_list_row)
        dupliacte_value_list = copy.deepcopy(value_list_row)
        dupliacte_assert_list = copy.deepcopy(assert_list_row)
        duplicate_step_number = copy.deepcopy(stepNumber)

        xpath = []
        assert_val = []
        original_value = []
        step_number_list = []

        index_value = 0
        lastindex=len(dupliacte_xpath_list)
        print(lastindex)
        for step_number, xpath_list, value_list, assert_list in zip(duplicate_step_number,
                                                                               dupliacte_xpath_list,
                                                                               dupliacte_value_list,
                                                                               dupliacte_assert_list):
            index_value += 1
            xpath_list=xpath_list.replace("HASH","#")
            if index==lastindex and index_value==lastindex:
                del xpath_list_row[index-1]
                del value_list_row[index-1]
                del assert_list_row[index-1]
                del stepNumber[index-1]
            elif ((index) == index_value) and index_value<lastindex:
                del xpath_list_row[index]
                del value_list_row[index]
                del assert_list_row[index]
                del stepNumber[index]
            else:
                step_number_list.append(step_number)
                xpath.append(xpath_list)
                original_value.append(value_list)
                assert_val.append(assert_list)

        print(xpath, original_value, assert_val)
        outDict = {}
        finalList = []
        finalDict = {}
        for (step_number, path, value, assert_value) in zip(step_number_list, xpath, original_value,
                                                                        assert_val):
            outDict["xpath"] = path
            outDict["value"] = value
            outDict["Alternative Value"] = assert_value
            outDict["stepNumber"] = step_number
            finalList.append(outDict)
            outDict = {}

        finalDict['headers'] = ["checkbox", "stepNumber", "xpath", "value", "Alternative Value", "Add/Remove"]
        finalDict["data"] = finalList
        outputList = finalList
        try:
            db.assertTable.delete_many({})
        except Exception as e:
            print('Error while deleting missing components report:' + str(e))
        try:
            db.assertTable.insert_many(outputList)
        except Exception as e:
            print("Empty list cannot be inserted:"+ str(e))
        for l in finalList:
            if "_id" in l.keys():
                del l["_id"]
        print("final list", finalList)
        print("final dict", finalDict)
        return finalDict
    except Exception as e:
        return jsonify(e)

def Remove(duplicate):
    final_list = []
    for list_iter in duplicate:
        if list_iter not in final_list:
            final_list.append(list_iter)
    return final_list
@app.route("/insertAssertionRow")  # When user clicks on Add new row in Assert Table
def insertAssertionRow():
    try:

        index = request.args.get("index")  # get index
        xpath_list_row = request.args.get("xpath_list_row").split(",")  # get xpath and split using comma
        value_list_row = request.args.get("value_list_row").split(",")  # get value and split using comma
        assert_list_row = request.args.get("assert_list_row").split(",")  # get asserted value and split using comma
        stepNumber = request.args.get("step_number").split(",")  # get stepNumber and split using comma
        index_list= request.args.get("index_list").split(",")
        operation_list = request.args.get("condition_list").split(",")
        xpath = []  # initialize a list for xpath
        assert_val = []  # initialize a list for asserted value
        original_value = []  # initialize a list for original value
        step_number_list = []  # initialize a list for step numbers
        xpath_index_list = []
        operation=[]
        assertionDropDown=[
            "Equal",
            "NotEqual",
            "Contains",
            "Presence"
        ]
        newdrop2=assertionDropDown
        index_value = 0
        lastindex = len(xpath_list_row) + 1
        newindex = int(index) + 1
        xpathdrop=[]
        newlist=[]
        # iterate through stepNumber,xpath,condition,value,asserted value lists.
        xpathdrop.append("")
        for xpaths in xpath_list_row:
            xpathdrop.append(xpaths)
            print("ROWWWW",xpathdrop)
        newdrop = xpathdrop

        assertdrop1=assertionDropDown

        xpathdropdown = copy.deepcopy(xpathdrop)
        assertDrpDownValue = copy.deepcopy(assertionDropDown)
        xpathdropdown1=xpathdropdown
        xpathdrop1=xpathdropdown
        assertDrpDownValue1 = assertDrpDownValue

        print("xpathdrop1",xpathdrop1)

        for idx, (step_number, xpath_list, value_list, assert_list, xpath_index,action) in enumerate(zip(stepNumber, xpath_list_row,
                                                                                value_list_row,
                                                                               assert_list_row,index_list,operation_list)):
            xpathdropdown = copy.deepcopy(xpathdrop)
            assertDrpDownValue = copy.deepcopy(assertionDropDown)
            index_value += 1
            print("ASSETED LIST VAL ",assert_list)
            xpath_list = xpath_list.replace("HASH", "#")
            print("indexesss",newindex,lastindex,index_value+1)
            if newindex == lastindex and index_value + 1 == lastindex:
                step_number_list.extend([step_number,""])  # append stepnumber
                xpath.extend([xpath_list,"xpath name"])  # append the xpath
                original_value.extend([value_list,"value name"])  # append the value
                print("assert_list",assert_list)
                print("assertval", assert_val)
                if assert_list=='':
                    assert_val.append(xpathdrop1)
                    assert_val.extend([xpathdrop1])
                if not assert_list=='':
                    xpathdropdown1.remove(assert_list)  # remove condition from list.
                    xpathdropdown1.insert(0, assert_list)
                    assert_val.append(xpathdropdown1)
                    print("sadsd",xpathdropdown1)
                    assert_val.extend([newdrop])
                print("HERE",assert_val)
                xpath_index_list.extend([xpath_index,"xpath index"])
                if action=='Equal':
                    operation.append(assertdrop1)
                    operation.extend([assertdrop1])
                if not action=='Equal':
                    assertDrpDownValue1.remove(action)  # remove condition from list.
                    assertDrpDownValue1.insert(0, action)
                    operation.append(assertDrpDownValue1)
                    operation.extend([newdrop2])
                print("THERE",operation)
                break

            if newindex == index_value:  # check the condition
                print("came in")
                xpath.append("xpath name")  # append xpath name to xpath list
                original_value.append("value name")  # append value name to original value list
                assert_val.append(newdrop)  # append empty value since it has to be entered by user
                step_number_list.append("")  # append empty since this is a new step , not in test description
                xpath_index_list.append("xpath index")
                operation.append(newdrop2)
            step_number_list.append(step_number)  # append stepnumber
            xpath.append(xpath_list)  # append the xpath
            original_value.append(value_list)  # append the value
            #assert_val.append(assert_list)  # append assert_list to assert_val list
            xpath_index_list.append(xpath_index)
            if action in assertDrpDownValue:  # if condition available in assertDrpDownValue list
                # print("insdie", cond_list)
                assertionDropDown.remove(action)  # remove condition from list.
                assertionDropDown.insert(0, action)  # insert condition in first position
                print("assert drop down", assertionDropDown)
            operation.append(assertionDropDown)  # append dropdown list to cond list
            assertionDropDown = assertDrpDownValue
            if assert_list in xpathdropdown:
                xpathdrop.remove(assert_list)  # remove condition from list.
                xpathdrop.insert(0, assert_list)  # insert condition in first position
                print("ASSSEERTTEDDD", xpathdrop)
            elif not assert_list in xpathdropdown:
                xpathdrop.insert(0, assert_list)
            assert_val.append(xpathdrop)
            xpathdrop=xpathdropdown
            # assert_val = list(set(assert_val))

        outDict = {}
        finalList = []
        finalDict = {}
        # assertValue=Remove(assert_val)

        for (step_number, path, value, assert_value, action,xpathindex) in zip(step_number_list, xpath, original_value,
                                                                        assert_val,operation,xpath_index_list):  # iterate through step number, xpath, original value, asserted value and condition all at once
            outDict["xpath"] = path  # assign xpath to dictionary
            outDict["value"] = value  # assign original value to dictionary
            outDict["Assertion Value"] = Remove(assert_value)  # assign asserted value to dictionary
            outDict["operation"] = action  # assign condition to dictionary
            outDict["stepNumber"] = step_number  # assign step number to dictionary
            outDict["index"] = xpathindex
            finalList.append(outDict)  # append the json to finalList after each iteration
            outDict = {}
        # print(finalList)
        finalDict['headers'] = ["checkbox", "stepNumber", "xpath", "index", "value","operation", "Assertion Value",
                                "Add/Remove"]  # assign headers as value to finalDict
        finalDict["data"] = finalList  # assign the finaList to key - data
        outputList = finalList  # assign finalList to outputList
        try:
            db.assertTable.delete_many({})  # delete empty
        except Exception as e:
            print('Error while deleting missing components report:' + str(e))  # print exception
        # from bson.json_util import dumps


        # def parse_json(data):
        #     return json.loads(dumps(data))

        db.assertTable.insert_many(outputList)  # insert the outputList into db
        # finalList = parse_json(finalList)
        for l in finalList:  # iterate through finalList
            if "_id" in l.keys():  # if _id is a key in finaList
                del l["_id"]  # delete the key value pair
        print("final dict", finalDict)

        return finalDict  # return finalDict json
    except Exception as e:
        return jsonify(e)  # return exception in json format


@app.route("/deleteAssertionRow")  # To delete a row in assert table
def deleteAssertionRow():
    try:

        # index = request.args.get("index")
        index = int(request.args.get("index"))
        xpath_list_row = request.args.get("xpath_list_row").split(",")
        value_list_row = request.args.get("value_list_row").split(",")
        assert_list_row = request.args.get("assert_list_row").split(",")
        stepNumber = request.args.get("step_number").split(",")
        index_list = request.args.get("index_list").split(",")
        operation_list = request.args.get("condition_list").split(",")


        dupliacte_xpath_list = copy.deepcopy(xpath_list_row)
        dupliacte_value_list = copy.deepcopy(value_list_row)
        dupliacte_assert_list = copy.deepcopy(assert_list_row)
        duplicate_step_number = copy.deepcopy(stepNumber)
        duplicate_index_list1 = copy.deepcopy(index_list)
        duplicate_operation_list1 = copy.deepcopy(operation_list)

        xpath = []
        assert_val = []
        original_value = []
        step_number_list = []
        xpath_index_list = []
        operation = []
        assertionDropDown = [
            "Equal",
            "NotEqual",
            "Contains",
            "Presence"
        ]

        index_value = 0
        lastindex=len(dupliacte_xpath_list)
        print(lastindex)

        xpathdrop=[]
        xpathdrop.append("")
        for xpaths in dupliacte_xpath_list:
            xpathdrop.append(xpaths)

        for step_number, xpath_list, xpathindex,action, value_list, assert_list in zip(duplicate_step_number,
                                                                               dupliacte_xpath_list,
                                                                            duplicate_index_list1,
                                                                               duplicate_operation_list1,
                                                                               dupliacte_value_list,
                                                                               dupliacte_assert_list):
            assertDrpDownValue = copy.deepcopy(assertionDropDown)
            xpathdropdown = copy.deepcopy(xpathdrop)
            index_value += 1
            print(index,"index")
            print(index_value, "index value")
            print(lastindex, "lastindex")

            xpath_list=xpath_list.replace("HASH","#")
            if index==lastindex and index_value==lastindex:
                del xpath_list_row[index-1]
                del value_list_row[index-1]
                del assert_list_row[index-1]
                del stepNumber[index-1]
                del index_list[index-1]
                del operation_list[index-1]
            elif ((index) == index_value) and index_value<lastindex:
                del xpath_list_row[index]
                del value_list_row[index]
                del assert_list_row[index]
                del stepNumber[index]
                del index_list[index]
                del operation_list[index]
            else:
                step_number_list.append(step_number)
                xpath.append(xpath_list)
                original_value.append(value_list)
                xpath_index_list.append(xpathindex)
                if action in assertDrpDownValue:  # if condition available in assertDrpDownValue list
                    # print("insdie", cond_list)
                    assertionDropDown.remove(action)  # remove condition from list.
                    assertionDropDown.insert(0, action)  # insert condition in first position
                    print("assert drop down", assertionDropDown)
                operation.append(assertionDropDown)
                assertionDropDown = assertDrpDownValue
                if assert_list in xpathdropdown:
                    xpathdrop.remove(assert_list)  # remove condition from list.
                    xpathdrop.insert(0, assert_list)  # insert condition in first position
                    print("ASSSEERTTEDDD", xpathdrop)
                elif not assert_list in xpathdropdown:
                    xpathdrop.insert(0, assert_list)
                assert_val.append(xpathdrop)
                xpathdrop = xpathdropdown
        print(xpath, original_value, assert_val)
        outDict = {}
        finalList = []
        finalDict = {}
        for (step_number, path, xindex,action,value, assert_value) in zip(step_number_list, xpath,xpath_index_list,operation, original_value,
                                                                        assert_val):
            outDict["xpath"] = path
            outDict["value"] = value
            outDict["Assertion Value"] = Remove(assert_value)
            outDict["stepNumber"] = step_number
            outDict["index"] = xindex
            outDict["operation"] = action
            finalList.append(outDict)
            outDict = {}
        # print(finalList)
        finalDict['headers'] = ["checkbox", "stepNumber", "xpath","index", "value","operation", "Assertion Value", "Add/Remove"]
        finalDict["data"] = finalList
        outputList = finalList
        try:
            db.assertTable.delete_many({})
        except Exception as e:
            print('Error while deleting missing components report:' + str(e))
        try:
            db.assertTable.insert_many(outputList)
        except Exception as e:
            print("Empty list cannot be inserted:"+ str(e))
        for l in finalList:
            if "_id" in l.keys():
                del l["_id"]
        print("final list", finalList)
        print("final dict", finalDict)
        return finalDict
    except Exception as e:
        return jsonify(e)

@app.route('/uidata', methods=['GET'])  # Called When user clicks on View report button
def uidata():
    try:

        inputdata = modify_conf()
        path = base_location + "\\" + outerFolderName + "\\" + folderName
        pathconf = base_location + "\\" + outerFolderName + "\\" + folderName + "\\" + "conf.log"
        indexno = int(request.args.get("indexValue")) - 1
        try:
            output = []
            output.append("<h1 style='text-align:center;'>Test Report</h1>")
            output.append("<p style='text-align:right;'>" + inputdata['data'][indexno]['Date'] + "</p>")
            output.append("<p style='text-align:right;'>" + inputdata['data'][indexno]['Time'] + "</p>")
            output.append("<p><b>Application name:</b> " + inputdata["data"][indexno]["Application Name"] + "</p>")
            output.append("<p><b>Scenario name:</b> " + inputdata["data"][indexno]["Scenario Name"] + "</b></p>")
            output.append("<p><b>Test case result:</b> " + inputdata["data"][indexno]["Status"] + "</p>")
            output.append(
                "<p><b>Time of execution(in secs):</b> " + inputdata["data"][indexno]["Time in Seconds"] + "</p>")
            Lines = open(pathconf, encoding="utf8").readlines()  # reads file
            total = "0"
            passt = "0"
            failures = "0"
            for line in Lines:
                if line.strip().__contains__("spec,") and line.strip().__contains__("failure"):
                    total = line.split("spec,")[0].strip()
                    failures = line.split("failure")[0].split("spec,")[1].strip()
                    passt = str(eval(total) - eval(failures))

            print("total,", total, failures, passt)
            output.append("<p><b>Total number of tests: </b>" + total + "</p>")
            output.append("<p><b>Number of test case passed: </b>" + passt + "</p>")
            output.append("<p><b>Number of test case failed: </b>" + failures + "</p>")

            output.append("<p><b>Instances running in webdriver manager:</b> " + inputdata["data"][indexno][
                "WebDriver Instance"] + "</p>")
            output.append("<h3><b>Test Description :</b> </h3>")
            listofsteps = inputdata["data"][indexno]["Test Description"].replace("<br>", " \n ", 1).split("<br>")
            listofsteps = [item.replace(" \n ", "<br") for item in listofsteps if item != '']
            output.append("<ol>")
            output.append("<p><b>Number of steps present: </b>" + str(len(listofsteps) + 1) + "</p>")
            for index, step in enumerate(listofsteps):
                if index == len(listofsteps) - 1 and inputdata["data"][indexno]["Status"].lower().__contains__("fail"):
                    output.append(
                        "<div style='color:red;'><li>" + str(step[step.find(".") + 1:].strip()) + "</li></div>")
                    output.append("<div style='color:red;'><p>" + str(
                        "\n".join(inputdata["data"][indexno]["Error Message"])) + "</p></div>")
                else:
                    output.append("<div style='color:green;'><li>" + step[step.find(".") + 1:].strip() + "</li></div>")

                destination_folder = ui_location + "\\assets\\images\\"

                if os.path.isfile(path + '\\page' + str(index) + '.png'):
                    print("image exists in recorded scenarios")
                    # temp = r'D:\ITAP_UI\src\assets\images\LCaaS_Test\test1'
                    temp = destination_folder + "\\" + inputdata["data"][indexno]["Application Name"] + "\\" + \
                           inputdata["data"][indexno]["Scenario Name"]
                    print("temp is", temp)
                    if not os.path.exists(temp):
                        os.makedirs(temp)
                    shutil.copy(path + '\\page' + str(index) + '.png', temp)
                    print(path + '\\page' + str(index) + '.png' + "hi" + temp)
                    destpath = temp + '\\page' + str(index) + '.png'
                    output.append(
                        '<img src=' + "../assets/images/" + inputdata["data"][indexno]["Application Name"] + "/" + \
                        inputdata["data"][indexno]["Scenario Name"] + "/page" + str(
                            index) + '.png' + ' width=500 height=300>')
            output.append("</ol>")
            print("this is output", output)
            output = "<br>".join(output)
            print("otput ", output)
            return jsonify(output)
            # return render_template(output)
        except Exception as e:

            print(e, )
            exc_type, exc_obj, exc_tb = sys.exc_info()
            print(exc_tb.tb_lineno)
            return 'FAIL' + str(e)

    except Exception as e:
        return jsonify(e)


@app.route('/outputreport', methods=['GET'])  # Gets called when user clicks on Download report button
def outputreport():
    try:
        inputdata = modify_conf()
        print(inputdata)
        path = base_location + "\\" + outerFolderName + "\\" + folderName
        print("out", outerFolderName, "fol", folderName)
        pathconf = base_location + "\\" + outerFolderName + "\\" + folderName + "\\" + "conf.log"
        indexno = int(request.args.get("indexValue")) - 1
        destination_folder = ui_location + "\\assets\\doc\\"

        try:
            document = docx.Document()
            heading1 = document.add_heading('Test Report', level=1)
            heading1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date = document.add_paragraph(inputdata["data"][indexno]["Date"])
            date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            time = document.add_paragraph(inputdata["data"][indexno]["Time"])
            time.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            document.add_paragraph('Application name: ' + inputdata["data"][indexno]["Application Name"])
            document.add_paragraph('Scenario name: ' + inputdata["data"][indexno]["Scenario Name"])
            document.add_paragraph('Test case result : ' + inputdata["data"][indexno]["Status"])
            document.add_paragraph('Time of execution(in secs): ' + inputdata["data"][indexno]["Time in Seconds"])

            Lines = open(pathconf, encoding="utf8").readlines()  # reads file
            total = "0"
            passt = "0"
            failures = "0"
            for line in Lines:
                if line.strip().__contains__("spec,") and line.strip().__contains__("failure"):
                    total = line.split("spec,")[0].strip()
                    failures = line.split("failure")[0].split("spec,")[1].strip()
                    passt = str(eval(total) - eval(failures))
            document.add_paragraph('Total number of tests: ' + total)
            document.add_paragraph('Number of test case passed: ' + passt)
            document.add_paragraph('Number of test case failed: ' + failures)
            document.add_paragraph(
                'Instances running in webdriver manager: ' + inputdata["data"][indexno]["WebDriver Instance"])
            document.add_heading('Test Description : ', level=3)
            listofsteps = inputdata["data"][indexno]["Test Description"].replace("<br>", "\n", 1).split("<br>")
            listofsteps = [item for item in listofsteps if item != '']
            for index, step in enumerate(listofsteps):
                document.add_paragraph(step[step.find(".") + 1:].strip(), style='List Number')
                if index == len(listofsteps) - 1 and inputdata["data"][indexno]["Status"].lower().__contains__("fail"):
                    # document.add_paragraph(str("\n".join(inputdata["data"][indexno]["Error Message"])))
                    em = [item.encode('unicode_escape').decode() for item in
                          inputdata["data"][indexno]["Error Message"]]
                    document.add_paragraph(str("\n".join(em)))
                    if os.path.isfile(path + '\\page' + str(index) + '.png'):
                        screenshot = base_location + "\\" + inputdata["data"][indexno]["Application Name"] + "\\" + \
                                     inputdata["data"][indexno]["Scenario Name"] + "\\" + "screenshots"
                        if not os.path.exists(screenshot):
                            os.mkdir(screenshot)
                        temp = destination_folder + "\\" + inputdata["data"][indexno]["Application Name"] + "\\" + \
                               inputdata["data"][indexno]["Scenario Name"]
                        if not os.path.exists(temp):
                            os.makedirs(temp)

                        shutil.copy(path + '\\page' + str(index) + '.png', temp)
                        shutil.copy(path + '\\page' + str(index) + '.png', screenshot)
                        destpath = screenshot + '\\page' + str(index) + '.png'
                        document.add_picture(destpath, width=Inches(5.0))

                document.save(destination_folder + "Report.docx")

            return jsonify("Success")
        except Exception as e:
            return 'FAIL' + str(e)
    except Exception as e:
        return jsonify(e)


@app.route('/getspecfile', methods=['GET'])
def getspecfile():
    try:
        pathspec = base_location + "\\" + outerFolderName + "\\" + folderName + "\\spec.js"  # assign spec.js path to pathspec
        Lines = open(pathspec, encoding="utf8").readlines()  # reads file
        output = {}  # initialize an empty json
        output["headers"] = ["lineno", "specdata"]  # assign headers
        output["data"] = []
        for index, line in enumerate(Lines):  # iterate through lines in spec file and index
            output["data"].append(
                {"lineno": index + 1, "specdata": line})  # append data line number and spec line to output json

        return output  # return json
    except Exception as e:
        return jsonify(e)  # return exception in json format


@app.route("/runSavedTestRun1")  # Run multiple job files
def runSavedTestRun1(testRunName):
    # testRunName = request.args.get("testName")
    ct=0
    testCount = 0  # assign 0 to testCount
    requiredLines = ''  # assign empty string
    it_block=''
    fulString = provideShortSpec(testRunName) # multiline string protractor code
    # scenarios="Deposit_Flow[XYZ_BANK], Bank_flow[XYZ_Bank]"
    collections = db.testRun.find()  # get collection in db
    for collection in collections:  # iterate through collection in db
        if collection["test_name"] == testRunName:  # if testcase name in db
            scenarios = collection["test_scenarios_used"]  # assign value of test_scenarios_used to scenarios
    split_scenario = scenarios.split(",")  # split scenarios using comma
    if os.path.isdir(base_location + "\\testRun") == False:  # if directory not available
        os.mkdir(base_location + "\\testRun")  # make directory
    if os.path.isdir(base_location + "\\testRun\\" + testRunName) == False:  # if directory not available
        os.mkdir(base_location + "\\testRun\\" + testRunName)  # make directory

    scenario_file = open(base_location + "\\testRun\\" + testRunName + "\\" + "spec.js",
                         "w+")  # open file in write mode
    conf_file = open(base_location + "\\testRun\\" + testRunName + "\\" + "conf.js" ,
                     "w+")  # open file in write mode
    spec_file_name = "spec.js"
    outerFolder="testRun"
    get_conf_file = provideConf(spec_file_name,base_location,outerFolder,testRunName)
    conf_file.write(get_conf_file)  # write configuration content in conf_file
    conf_file.close()  # close file
    batch_file = open(base_location + "\\testRun\\" + testRunName + "\\" + "batch.bat",
                      "w+")  # open batch file in write mode
    script1 = '''START /B cmd "cd ''' + base_location + '''\\" && protractor ''' + base_location + '''\\''' + "testRun" + '\\' + testRunName + '''\\conf.js > ''' + base_location + '''\\''' + "testRun" + '\\' + testRunName + '''\\conf.log'''  # commands required for batch file
    for scene_app in split_scenario:  # iterate through list of test scenarios
        testCount += 1  # increement testcount

        flag = 0  # initialize flag to 0
        scenario_name = scene_app.split("[")[
            0].strip()  # split element in split_scenarios using open bracket and taken first index element
        application_name = scene_app.split("[")[1].split("]")[
            0].strip()  # split using brackets and get application name
        path_var = base_location + "\\" + application_name + "\\" + scenario_name  # assign path using obtained scenario and apllication name
        if os.path.isdir(path_var) == False:  # if path_var path not available
            return "Fail"  # return string fail
        else:  # if path_var path available
            specFile = open(path_var + "\\spec.js").read()  # open spec.js file and read
            newFile=specFile.rsplit("});",1)[0].strip()
            finalFile=''.join(newFile)
            print(finalFile)
            for lines in finalFile.split(";"):  # iterate through spec.js file
                if flag == 1 and not lines=='' and not lines.__contains__("browser.takeScreenshot") and not lines.__contains__("var fname="):
                    requiredLines += lines+";"  # add line to requiredLines variable

                if lines.__contains__("it(\'") and lines.__contains__("browser.ignoreSynchronization =true"):  # if line starts with 'it'
                    flag = 1
                    it_block=lines.split("it(\'")[1].split("\'",1)[0].strip()
                    print("itttttt",it_block)
                    requiredLines += "\nit(\'"+it_block+"_testcase"+str(testCount)+"\' , function() {browser.ignoreSynchronization =true;"  # replace testCount

                if lines.__contains__("browser.takeScreenshot"):
                    requiredLines += '''\nbrowser.takeScreenshot().then(function (png) {var dir=\"''' + base_location.replace("\\",
                                                                                                                   "\\\\") + "\\\\" + "testRun" + '\\\\' + testRunName + '''\";'''
                if lines.__contains__("var fname="):
                    ct += 1
                    requiredLines += "\nvar fname=\"page" + str(ct) + ".png\";"

    fulString += requiredLines + "});"  # concat }); to lines obtained through conditions and assign it to fulString variable
    scenario_file.write(fulString)  # write fulString containing the lines , to scenario_file file
    scenario_file.close()  # close file
    batch_file.write(script1)  # write script1 to batch file
    batch_file.close()  # close batch file

    return "Success"  # return string


def sched():  # schedule test job check
    while True:  # if test jobs have been scheduled
        schedule.run_pending()  # calls all jobs that are scheduled to run
        time.sleep(5)  # delays 5 seconds
        print("Checking schedule")


def daily_job(appPath, schedulerName):  # call when scheduled test job gets executed
    print("job function")
    subprocess.call(
        appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")  # run external batch file from this python code


def job(appPath, schedulerName, scheduledDate):  # called when scheduled for specific date
    mon, day, year = scheduledDate.split("/")  # split scheduledDate using /
    print(mon, day, year)

    datestring = str(year) + "-" + str(mon) + "-" + str(day)  # date formatting
    print(scheduledDate, "job function")
    print(date.today(), "joob function")
    print(datestring, "jobbb function")

    if date.today() == date(int(year), int(mon), int(day)):  # if today's date is equal to the format
        print("job function")
        subprocess.call(
            appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")  # run external batch file from this python code


# def uploadingExcelInSchedule(schedulerName, file, testRunName):
#     # folderName = request.args.get("folderName")
#     import pandas as pd
#     global mulipleScenarioPath
#
#     final_val = []
#
#     str_file = file.decode("utf-8")
#     # temp_file = list(eval(str_file))
#     temp_List = []
#     sheetLength = 0
#     json_format = json.loads(str_file)
#     print("str_file", str_file)
#
#     # Checking the muliple scenario file exist or not, if exists remove the file
#     connections = db.testRun.find()
#     for connection in connections:
#         if connection["test_name"] == testRunName:
#             mulipleScenarioPath = connection["path"]
#             if os.path.isdir(mulipleScenarioPath + "\\" + schedulerName) == False:
#                 os.mkdir(mulipleScenarioPath + "\\" + schedulerName)
#             if os.path.exists(mulipleScenarioPath + "\\" + schedulerName + "\\combined_multi_scenario_file.js"):
#                 p = os.path.join(mulipleScenarioPath + "\\" + schedulerName + "\\combined_multi_scenario_file.js")
#                 os.remove(p)
#
#     for k, v in json_format.items():
#         v1 = json.dumps(v['data'])
#         if type(list(eval(v1))[0]) == "str":
#             return jsonify("failure")
#         else:
#             temp_file = list(eval(v1))
#             df = pd.DataFrame.from_dict(temp_file)
#
#         connections = db.testRun.find()
#         for connection in connections:
#             if connection["test_name"] == testRunName:
#                 mulipleScenarioPath = connection["path"]
#                 if os.path.isdir(mulipleScenarioPath):
#                     df.to_excel(mulipleScenarioPath + "\\" + schedulerName + '\\sample.xls',
#                                 index=False)
#                     time.sleep(1)
#                     append_file = open(mulipleScenarioPath + "\\" + schedulerName + "\\append_spec.txt", "a+")
#                     returnVal = uploadUIWithoutRun(schedulerName, testRunName, sheetLength)
#                     # final_val.append(returnVal)
#                     sheetLength += 1
#                 else:
#                     return "appFail"
#
#     comb_multi_file = open(mulipleScenarioPath + "\\" + schedulerName + "\\combined_multi_scenario_file.js", "a+")
#
#     comb_multi_file.write("});")
#     comb_multi_file.close()
#     return jsonify("Success")  # create a json object to generate excel inui


def send_mail(MailID, testRunName):
    try:
        pythoncom.CoInitialize()
        outlook = win32.Dispatch('outlook.application')
        mail = outlook.CreateItem(0)
        mail.To = MailID

        mail.Subject = testRunName + ' Test Report'
        mail.Body = 'Message body'
        cont = """\
    <html>
     <head></head>
     <body>
       <p>Dear iTAP User,<br><br>
          Test Job Name : """ + testRunName + """<br><br>
          Date and Time : """ + datetime.now().strftime("%d/%m/%Y %H:%M:%S") + """<br><br>
          Attached the report of above Test Job, Please find attached document.<br><br>

          Note : Please do not reply to this mail, This is an auto-generated e-mail.<br><br>
          Regards,<br>
          iTAP Team
       </p>
     </body>
    </html>
    """
        mail.HTMLBody = cont  # this field is optional
        try:
            file_one = base_location + "\\testRun\\" + testRunName + "\\new_excel_sheet.xlsx"
            file_two = base_location + "\\testRun\\" + testRunName + "\\iTAPTestDataReport-" + testRunName + ".xlsx"
            os.rename(file_one, file_two)

            attachmentPath = file_two
            mail.Attachments.Add(attachmentPath)
            mail.Send()
            os.rename(file_two, file_one)
            return "Success"
        except:
            return "No File"
    except:
        return "Outlook Missing"  # to send the test job in email


def uploadUIWithoutRun(schedulerName, testRunName, sheetLength):
    replace_spec = ''  # assign empty string
    head_list = []
    cell_list = []
    filePath = ''  # assign empty string
    # folderName = request.args.get("folderName")
    connections = db.testRun.find()  # get collection in db
    for connection in connections:  # iterate through collection in db
        if connection["test_name"] == testRunName:  # if testrunname present in db
            appPath = connection["path"]  # assign value of path to appPath
            scenario = connection["test_scenarios_used"]  # assign value of test scenario used to scenario
            split_scenario = scenario.split(",")[
                sheetLength]  # split scenario using comma and take zero index element i.e. sheetLength=0
            appNameRun = split_scenario.split("[")[1].split("]")[0]  # split split_scenario using brackets
            scenarioNameRun = split_scenario.split("[")[0]  # split split_scenario using open bracket
            collections = db.folderName.find()  # get collections in db
            for collection in collections:  # iterate through collection in db
                if collection["appName"] == appNameRun and collection[
                    "scenarioName"] == scenarioNameRun:  # if appname and scenario name is in db
                    filePath = collection["innerpath"]  # assign value of innerpath to filepath
            if os.path.isdir(appPath + "\\" + schedulerName) == False:  # if following path not present
                os.mkdir(appPath + "\\" + schedulerName)  # make a directory
            if os.path.isdir(filePath):  # if filepath available
                if os.path.exists(filePath + "\\spec.js"):  # if spec.js available in the path
                    spec_file = open(
                        filePath + "\\spec.js")  # open spec.js file
                    path = appPath + "\\" + schedulerName + "\\sample.xls"  # assign excelsheet path
                    new_spec_file = open(
                        appPath + "\\" + schedulerName + "\\multiple_scenario_file.js",
                        "w+")  # open multiple_scenario_file.js in write mode

                    combined_multi_scenario_file = open(
                        appPath + "\\" + schedulerName + "\\combined_multi_scenario_file.js",
                        "a+")  # open combined_multi_scenario_file.js in append mode
                    new_conf_file = open(appPath + "\\" + schedulerName + "\\multiple_conf_file.js",
                                         'w+')  # open multiple_conf_file.js.js in write mode
                    combined_conf_file = open(appPath + "\\" + schedulerName + "\\combined_conf_file.js",
                                              'w+')  # open combined_conf_file.js in write mode
                    append_file = open(appPath + "\\" + schedulerName + "\\append_spec",
                                       "a+")  # open append_spec in append mode
                    batch_file = open(appPath + "\\" + schedulerName + "\\multiple_batch_file.bat",
                                      "w+")  # open multiple_batch_file.bat in write mode
                    combined_batch_file = open(appPath + "\\" + schedulerName + "\\combined_batch_file.bat",
                                               "w+")  # open combined_batch_file.bat in write mode

                    wb = xlrd.open_workbook(path)  # open excelsheet using path i.e ..\\..\\sample.xlsx
                    sheet = wb.sheet_by_index(0)  # opening sheet by index
                    start_row = 1  # starting row
                    end_row = sheet.nrows  # number of rows present in sheet
                    num_of_columns = sheet.ncols  # number of columns present in sheet
                    open_file = spec_file.read()  # read spec_file
                    split_file = open_file.split("extract")[
                        0]  # in spec_file split using extract and take the element in zero index
                    # if (split_file.endswith("//")):  # if zero index element ends with //
                    #     split_file = split_file[
                    #                  :-2] + "});"  # take everything except last two elements and add }); to it.
                    sec_test = "it('Test 1', function() {" +" browser.ignoreSynchronization =true;  "+ "\n"+\
                               "browser.driver.manage().window().maximize();" + \
                               split_file.split("browser.driver.manage().window().maximize();")[1]  # split

                    for r in range(start_row, end_row):  # iterate in excelsheet
                        head_list = []  # initialize an empty list
                        cell_list = []  # initialize an empty list
                        for c in range(num_of_columns):  # iterate through columns in excelsheet
                            cell_obj = sheet.cell_value(r, c)  # read data from a cell r-row c-column
                            head_obj = sheet.cell_value(0, c)  # read data from a cell row-0 c-column
                            head_list.append(head_obj)  # append data
                            cell_list.append(cell_obj)  # append data
                        for line in sec_test.split(";"):  # split and iterate
                            for idx, header in enumerate(
                                    head_list):  # iterate through header and have record of index value idx
                                line = str(line)  # type conversion
                                if (line.__contains__(header) and not line.__contains__("browser.wait(until")
                                        and not line.__contains__(
                                            ".clear(")):  # if line contains header and browser wait until and not contains .clear
                                    if (type(cell_list[idx]) == int or type(
                                            cell_list[
                                                idx]) == float):  # if data taken from excelsheet has index type int or float
                                        line = line.split("sendKeys")[0] + "sendKeys(" + str(
                                            int(cell_list[idx])) + ")"  # split line using sendkeys and add the string
                                    else:
                                        line = line.split("sendKeys")[0] + "sendKeys('" + str(
                                            cell_list[idx]) + "')"  # split line using sendkeys and add the string
                                    break  # exit loop
                            replace_spec += line + ";"  # add semicolon for each line

                    new_spec = provideShortSpec(schedulerName) + str(replace_spec) #+ "});"  # add lines added with semicolon

                    if os.stat(
                            appPath + "\\" + schedulerName + "\\combined_multi_scenario_file.js").st_size == 0:  # os.stats--get status of specified path, st_size--size of file in bytes
                        new_combined_file = provideShortSpec(schedulerName) + str(replace_spec)
                        combined_multi_scenario_file.write(new_combined_file)  # write file
                        print("Combined_spec_file", combined_multi_scenario_file)
                    else:  # if file size in bytes isn't 0
                        new_combined_file = str(replace_spec)
                        combined_multi_scenario_file.write(new_combined_file)  # write file
                        print("Combined_spec_file", combined_multi_scenario_file)
                    spec_file_name='multiple_scenario_file.js'
                    outerFolder="testRun"
                    new_conf = provideConf(spec_file_name,base_location,outerFolder,folderName)
                    comb_specfilename='combined_multi_scenario_file.js'
                    comb_conf = provideConf(comb_specfilename,base_location,outerFolder,folderName)
                    multipleBatchFile = "START /B cmd " + appPath + " && protractor " + appPath + "\\" + schedulerName + "\\multiple_conf_file.js >" + appPath + "\\" + schedulerName + "\\upload_log_file.log"  # path for batch file
                    comb_batch_file_data = "START /B cmd " + appPath + " && protractor " + appPath + "\\" + schedulerName + "\\combined_multi_scenario_file.js >" + appPath + "\\" + schedulerName + "\\upload_log_file.log"  # path for batch file

                    new_conf_file.write(new_conf)  # write configuration file
                    combined_conf_file.write(comb_conf)

                    new_spec_file.write(new_spec)
                    append_file.write(new_spec)
                    new_conf_file.close()  # close configuration file
                    new_spec_file.close()
                    batch_file.write(multipleBatchFile)
                    combined_batch_file.write(comb_batch_file_data)  # write batch file commands
                    batch_file.close()
                    combined_batch_file.close()  # close batch file
                    combined_multi_scenario_file.close()

                    time.sleep(1)  # suspends execution for given number of seconds

                    return "Success"  # return success string

            else:
                return "appFail"  # return appFail string


def getParallelJobDetails(jobName, file):
    with app.app_context():
        final_json = {}  # initialize an dictionary
        connections = db.testRun.find()  # get collections in db
        for connection in connections:  # iterate through collection
            if connection["test_name"] == jobName:  # if jobname present in db
                path = connection["path"]  # assign value of path to path variable
                break  # exit loop
        parallelConf = open(path + "\\" + file + ".log").readlines()  # open file in read mode
        time_taken_seconds = []  # initialize an empty list
        passStr = '[32m.[0m'  # assign pass string
        failStr = '[31mF[0m'  # assign fail string
        resultList = []  # initialize an empty list
        for line in parallelConf:  # iterate through the lines in file
            if line.__contains__("Finished in"):  # if line contains finished in
                time_taken_seconds.append(line.split(" ")[4])  # split and append time to the list
            if line.__contains__(passStr) or line.__contains__(failStr):  # if line contains pass string or fail string
                line = line.split(" ")[2]  # split line with space and take second element
                outputLine = [(line[i:i + 10]) for i in
                              range(0, len(line), 10)]  # iterate through the line and make a list of words
                for result in outputLine:  # iterate through outputLine
                    if result == passStr:  # if element is pass string
                        resultList.append('Pass')  # append pass in list
                    elif result == failStr:  # if element is fail string
                        resultList.append('Fail')  # append fail in list
        final_json["time_taken"] = []
        final_json["time_taken"].append({"chrome": time_taken_seconds[0], "firefox": time_taken_seconds[
            1]})  # append time to json from time_taken_seconds list
        final_json["total_specs"] = []
        chromeList, FirefoxList = split_list(
            resultList)  # split resultList contatining pass/fail and assign to variables
        final_json["total_specs"].append({"chrome": chromeList, "firefox": FirefoxList})  # append value to key in json
        final_json["total_pass_results"] = []
        final_json["total_pass_results"].append(
            {"chrome": chromeList.count('Pass'), "firefox": FirefoxList.count('Pass')})  # append value to key in json
        final_json["total_fail_results"] = []
        final_json["total_fail_results"].append(
            {"chrome": chromeList.count('Fail'), "firefox": FirefoxList.count('Fail')})  # append value to key in json
        print(final_json)
        data_header_list = json_to_json(final_json)  # json_to_json convert a json file into ui specified file
        return data_header_list  # return list


def split_list(a_list):
    half = len(a_list) // 2
    return a_list[:half], a_list[half:]  # split list into two


def json_to_json(given_json):  # convert a json file into ui specified file
    given_json = str(given_json)
    given_json = given_json.replace("'", "\"")
    print(type(given_json), ' ', given_json)
    json_value = json.loads(given_json)
    headers = ["browser_compatibility_test", "chrome", "firefox"]
    data = []
    for val in json_value.keys():
        x = {}
        x = json_value[val][0]
        if val == "time_taken":
            val = str(val).replace("time_taken", "Time Taken (In seconds)")
        if val == "total_specs":
            val = str(val).replace("total_specs", "Total Script Results")
        if val == "total_pass_results":
            val = str(val).replace("total_pass_results", "Total Scripts Passed")
        if val == "total_fail_results":
            val = str(val).replace("total_fail_results", "Total Scripts Failed")
        x["browser_compatibility_test"] = val
        final_json = x
        data.append(final_json)
    lst = [0] * 2
    lst[1] = headers
    lst[0] = data
    print(lst)
    return lst


def elapseTime(jobName, fileName):
    collections = db.testRun.find({"test_name": jobName})  # find data in collection containing the jobName
    if collections.count() > 0:  # if collection present
        for collection in collections:  # iterate through collection in db
            chrome_list = []  # initialize an empty list
            firefox_list = []  # initialize an empty list
            json_format = {}  # initialize an empty dictionary
            path = collection["path"]  # assign path value from db
            f = open(path + "\\" + fileName + ".log", "r")  # open file in read mode
            for line in f:  # iterate through lines in file
                # print(line)
                if line.__contains__("secs)") or line.__contains__("min)") or line.__contains__(
                        "mins)"):  # if line contains either of secs),min),mins)
                    # print(line)
                    if line.__contains__("chrome"):  # if line contains chrome
                        chrome_list.append(line.split("(")[1].split(")")[0])  # split and append to chrome_list
                    if line.__contains__("firefox"):  # if line contains firefox
                        firefox_list.append(line.split("(")[1].split(")")[0])  # split and append to firefox_list

            json_format["browser_compatibility_test"] = "Time Taken Individual"  # json formatting the string
            chrometimer = ''  # assign empty string
            print(chrome_list, 'chrome_list')
            for chrome in chrome_list:  # iterate through chrome_list
                print(chrome, 'chrome')
                chrometimer += '[' + chrome + '],'  # add brackets to the value
            firefoxtimer = ''  # assign empty string
            for firefox in firefox_list:  # iterate through firefox_list
                firefoxtimer += '[' + firefox + '],'  # add brackets to the value
            json_format["chrome"] = chrometimer[:-1]  # assign all elements except last, to the json
            json_format["firefox"] = firefoxtimer[:-1]  # assign all elements except last, to the json
            print(json_format)
            return json_format  # return the json
    else:
        return "No db found"  # return string


def json_json_json(str_file, jobName):
    # str_file=open("D:\ProtractorWebDriverTest\output_json.js","r").read()
    # jobName="something"
    print(str_file)
    str_file = str(str_file)
    str_file = str_file.replace("'", "\"")
    print(str_file)
    json_line = json.loads(str_file)
    Test_Job = json.loads(str(json_line["data"][2]).replace("'", "\""))
    num_of_scripts_chrome = len(Test_Job["chrome"].split(","))
    num_of_scripts_firefox = len(Test_Job["firefox"].split(","))
    chrome_scripts = str(Test_Job["chrome"]).replace("[", "").replace("]", "")
    firefox_script = str(Test_Job["firefox"]).replace("[", "").replace("]", "")
    # print("chrome scripts",chrome_scripts, firefox_script)
    tot_value = max(num_of_scripts_chrome, num_of_scripts_firefox)
    test_job_list = []
    ful_count = count_func_name(jobName)
    count_dict = {}
    arr_val = []
    for val in ful_count:
        if val.split()[0] in arr_val:
            count_dict[val.split()[0]] += 1
        else:
            count_dict[val.split()[0]] = 1

        arr_val.append(val.split()[0])
        test_job_list.append(val.split()[0] + "_DataSet_" + str(count_dict[val.split()[0]]))

    # print(test_job_list)
    result = json.loads(str(json_line["data"][1]).replace("'", "\""))
    newresult1 = json.loads(str(json_line["data"][3]).replace("'", "\""))
    newresult2 = json.loads(str(json_line["data"][4]).replace("'", "\""))
    total_chrome = (newresult1["chrome"], newresult2["chrome"])
    total_firefox = (newresult1["firefox"], newresult2["firefox"])
    # print("total chrome", total_chrome, total_firefox)
    resultchrome = result["chrome"].append(str(total_chrome))
    resultfirefox = (result["firefox"]).append(str(total_firefox))
    tot_time = json.loads(str(json_line["data"][0]).replace("'", "\""))
    tot_time_chrome = float(tot_time["chrome"])
    tot_time_firefox = float(tot_time["firefox"])
    new_tot_time_chrome = str(int(tot_time_chrome // 60)) + "mins" + str(int(tot_time_chrome % 60)) + "secs"
    new_tot_time_firefox = str(int(tot_time_firefox // 60)) + "mins" + str(int(tot_time_firefox % 60)) + "secs"
    print(new_tot_time_chrome, new_tot_time_firefox)
    # print("result chrome",result["chrome"], result["firefox"])
    total_time_chrome = chrome_scripts.split(",")
    total_time_firefox = firefox_script.split(",")
    total_time_chrome.append(new_tot_time_chrome)
    total_time_firefox.append(new_tot_time_firefox)
    print("total time chrome", tot_time_chrome, tot_time_firefox)
    json_format = {}
    test_job_list.append("Total")
    json_format["Test_Job"] = test_job_list
    json_format["Time_Taken_chrome"] = total_time_chrome
    json_format["Time_Taken_Firefox"] = total_time_firefox
    json_format["Result_Chrome"] = result["chrome"]
    json_format["Result_Firefox"] = result["firefox"]
    print(json.dumps(json_format, indent=4))


    new_json_format = {}
    full_list = []
    counter = 0
    for index_val in range(len(json_format["Test_Job"])):
        print(index_val)
        for key, value in json_format.items():
            print("key", key)
            new_json_format[key] = value[index_val]
        full_list.append(new_json_format)
        new_json_format = {}

    print(full_list)
    return full_list


def count_func_name(jobName):
    full_array = []  # initialize an empty list
    collections = db.testRun.find({"test_name": jobName})  # find the records with the jobName in db.
    for collection in collections:  # iterate through the collection in db.
        path = collection["path"]  # assign value of path from db
        f = open(
            path + "\parallel_multi_scenario_file.js").readlines()  # open parallel multi scenario file and readlines
        for line in f:  # iterate through lines in file
            if line.__contains__("it('"):  # if line contains 'it('
                # print(line)
                # print(line.split("it('")[1].split("'")[0])
                full_array.append(line.split("it('")[1].split("'")[0])  # append in the list
        print(full_array)
        # full_array=['Deposit_Flow 0', 'Deposit_Flow 0', 'Withdraw_flow 0', 'Withdraw_flow 0']
        return full_array  # return list


@app.route("/getCombinedTestDescription")  # to display the test description for multiple scenarios
def getCombinedTestDescription():
    appName = request.args.get("appName")  # get application name from UI
    scenarioName = request.args.get("scenarioName")  # get scenarioname from UI
    value = modifiedtestDescription(appName,
                                    scenarioName)  # call modifiedtestDescription with application name and scenario name
    if value != '':  # if value is not equal to empty
        json_format = {
            "status": "success",  # assign status as success
            "test_Description": value  # assign the value to test description key
        }
    else:
        json_format = {
            "status": "fail",  # assign status as fail
            "test_Description": value  # assign the value to test description
        }
    return jsonify(json_format)  # return json format dictionary


def modifiedtestDescription(outerFolderName, folderName):

    if os.path.exists(base_location + "\\" + outerFolderName + '\\' + folderName + "\spec.js"):
        open_file = open(base_location + "\\" + outerFolderName + '\\' + folderName + "\spec.js")
        # new_file = open(r"base_location\\" + outerFolderName + '\\' + folderName + "\oldFiles\TestDescription.html",
        #                 "w+")
        lines = open_file.readlines()
        full_line = ''
        flag = 0
        # global sentence
        sentence = 'Combined Scenarios Test Steps :<br>'
        for line in lines:
            if line.__contains__("extractModal();"):
                break
            if line.strip().startswith("browser.get("):
                flag = 1
            if flag == 1:
                full_line += line

        number = 0
        for sline in full_line.split(";"):
            if sline.__contains__("browser.get("):  # Get Url
                number += 1
                url_line = sline.split("browser.get(")[1].split(")")[0]
                if len(url_line) < 50:
                    sentence += str(number) + ". Open browser and navigate to URL: " + str(
                        url_line) + " <br>"
                else:
                    sentence += str(number) + ". Open browser and navigate to URL: " + "<br>" + str(
                        url_line) + " <br>"
            # if sline.__contains__("browser.sleep("):
            #   sentence+="Browser will wait for "+sline.split("browser.sleep(")[1].split(")")[0]+"MilliSeconds<br>"
            if sline.__contains__("until.presenceOf") == False and \
                    sline.__contains__("by.xpath(") and \
                    sline.strip().endswith("click()"):  # button click
                if sline.__contains__("@id="):
                    number += 1
                    sentence += str(number) + ". Click on the button with id " + sline.split("@id=")[1].split("]")[
                        0] + "<br>"
                elif sline.__contains__("@name="):
                    number += 1
                    sentence += str(number) + ". Click on the button with name " + sline.split("@name=")[1].split("]")[
                        0] + "<br>"
                elif sline.__contains__(".get("):
                    number += 1
                    x = str(sline.split(".get(")[1].split(")")[0]).strip()
                    if x.endswith("1"):
                        inBetween = str(floor(int(x) / 10)) + "1st"
                    elif x.endswith("2"):
                        inBetween = str(floor(int(x) / 10)) + "2nd"
                    elif x.endswith("3"):

                        inBetween = str(floor(int(x) / 10)) + "3rd"
                    else:
                        inBetween = x + "th"
                    sentence += str(number) + ". Click on " + inBetween + " button in the current page<br>"

            elif sline.__contains__("until.presenceOf") == False and \
                    sline.__contains__("by.xpath(") and \
                    sline.__contains__("click()") == False and \
                    sline.__contains__("sendKeys"):  # Sendkey
                if sline.__contains__("@id="):
                    number += 1
                    if sline.__contains__(".get("):
                        x = str(sline.split(".get(")[1].split(")")[0]).strip()
                        if x.endswith("1"):
                            inBetween = str(floor(int(x) / 10)) + "1st"
                        elif x.endswith("2"):
                            inBetween = str(floor(int(x) / 10)) + "2nd"
                        elif x.endswith("3"):
                            inBetween = str(floor(int(x) / 10)) + "3rd"
                        else:
                            inBetween = x + "th"
                            if sline.lower().__contains__("password"):
                                sentence += str(
                                    number) + ". Enter value " + "******" + " on the" + inBetween + "field with id " + \
                                            sline.split("@id=")[1].split("]")[0] + "<br>"
                            else:
                                sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                                    0] + " on the" + inBetween + "field with id " + sline.split("@id=")[1].split("]")[
                                                0] + "<br>"
                    else:
                        if sline.lower().__contains__("password"):
                            sentence += str(number) + ". Enter value " + "******" + " in the field with id " + \
                                        sline.split("@id=")[1].split("]")[0] + "<br>"
                        else:
                            sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                                0] + " in the field with id " + sline.split("@id=")[1].split("]")[0] + "<br>"
                elif sline.__contains__("@name="):
                    number += 1
                    if sline.__contains__(".get("):
                        x = str(sline.split(".get(")[1].split(")")[0]).strip()
                        if x.endswith("1"):
                            inBetween = str(floor(int(x) / 10)) + "1st"
                        elif x.endswith("2"):
                            inBetween = str(floor(int(x) / 10)) + "2nd"
                        elif x.endswith("3"):
                            inBetween = str(floor(int(x) / 10)) + "3rd"
                        else:
                            inBetween = x + "th"
                            if sline.lower().__contains__("password"):
                                sentence += str(number) + ". Enter value " + "******" + " in the field with name " + \
                                            sline.split("@name=")[1].split("]")[0] + "<br>"
                            else:
                                sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                                    0] + " in the field with name " + sline.split("@name=")[1].split("]")[0] + "<br>"
                    else:
                        if sline.lower().__contains__("password"):
                            sentence += str(number) + ". Enter value " + "*****" + " in the field with name " + \
                                        sline.split("@name=")[1].split("]")[0] + "<br>"
                        else:
                            sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                                0] + " in the field with name " + sline.split("@name=")[1].split("]")[0] + "<br>"
                elif sline.__contains__(".get(") and sline.__contains__("@id=") == False:
                    number += 1
                    x = str(sline.split(".get(")[1].split(")")[0]).strip()
                    if x.endswith("1"):
                        inBetween = str(floor(int(x) / 10)) + "1st"
                    elif x.endswith("2"):
                        inBetween = str(floor(int(x) / 10)) + "2nd"
                    elif x.endswith("3"):
                        inBetween = str(floor(int(x) / 10)) + "3rd"
                    else:
                        inBetween = x + "th"
                    # sentence += "Enter value " + sline.split("sendKeys(")[1].split(")")[0] + " in the field with id " + \
                    #             inBetween+" in the field from the webpage<br>"
                    sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                        0] + " on " + inBetween + " field in the current page<br>"

    else:
        sentence = ''
    sentence = sentence.replace('<br>', '\n')
    return sentence


# writing a function to find between,before and after values
def between(value, a, b):
    # Find and validate before-part.
    pos_a = value.find(a)
    if pos_a == -1: return ""
    # Find and validate after part.
    pos_b = value.rfind(b)
    if pos_b == -1: return ""
    # Return middle part.
    adjusted_pos_a = pos_a + len(a)
    if adjusted_pos_a >= pos_b: return ""
    return value[adjusted_pos_a:pos_b]


@app.route("/uploadExcelScenario", methods=['GET', 'POST'])  # Upload specific sceanrio
def uploadExcelScenario():
    # folderName = request.args.get("folderName")

    file = request.get_data()
    print("file", file)
    str_file = file.decode("utf-8")
    # temp_file = list(eval(str_file))
    print("str file", str_file)
    temp_List = []
    if type(list(eval(str_file))[0]) == str:
        return jsonify("failure")
    else:
        temp_file = list(eval(str_file))
        df = pd.DataFrame.from_dict(temp_file)
    application_name = request.args.get("application_name")
    scenario_name = request.args.get("scenario_name")
    connections = db.folderName.find()
    for connection in connections:
        if connection["appName"] == application_name and connection["scenarioName"] == scenario_name:
            appPath = connection["path"]
            scenarioPath = connection["innerpath"]
            if os.path.isdir(appPath):
                if os.path.isdir(scenarioPath):
                    df.to_excel(scenarioPath + '\sample1.xlsx',
                                index=False)
                    time.sleep(1)


                else:
                    return "scenarioFail"
            else:
                return "appFail"

    returnVal = uploadUI(application_name, scenario_name)
    return returnVal

@app.route("/uploadUI")  # Generate excel sheet while uploading
def uploadUI(application_name, scenario_name):

    replace_spec = ''

    connections = db.folderName.find()
    for connection in connections:
        if connection["appName"] == application_name and connection["scenarioName"] == scenario_name:
            appPath = connection["path"]
            scenarioPath = connection["innerpath"]
            if os.path.isdir(appPath):
                if os.path.isdir(scenarioPath):
                    if os.path.exists(scenarioPath + "\\spec.js"):
                        spec_file = open(
                            scenarioPath + "\\spec.js")
                        path = scenarioPath + "\\sample1.xlsx"
                        new_spec_file = open(
                            scenarioPath + "\\multiple_scenario_file.js",
                            "w+")
                        new_conf_file = open(scenarioPath + "\\multiple_conf_file.js", 'w+')
                        if os.path.isdir(scenarioPath + "\\Batchfiles") == False:
                            os.mkdir(scenarioPath + "\\Batchfiles")
                        batch_file = open(scenarioPath + "\\Batchfiles\\multiple_batch_file.bat", "w+")
                        wb = load_workbook(path)
                        sheet = wb.worksheets[0]
                        start_row = 1
                        end_row = sheet.max_row
                        num_of_columns = sheet.max_column
                        open_file = spec_file.read()
                        split_file = open_file.split("extract")[0]
                        if (split_file.endswith("//")):
                            split_file = split_file[:-2] + "});"
                        sec_test = "it('Test 1', function() {" + \
                                   "browser.driver.manage().window().maximize();" + \
                                   split_file.split("browser.driver.manage().window().maximize();")[1]

                        for r in range(start_row, end_row):
                            head_list = []
                            cell_list = []
                            for c in range(num_of_columns):
                                cell_obj = sheet.cell(r + 1, c + 1).value
                                head_obj = sheet.cell(1, c + 1).value
                                head_list.append(head_obj)
                                cell_list.append(cell_obj)

                            print("headlist cell list", head_list, cell_list)
                            for line in sec_test.split(";"):
                                for idx, header in enumerate(head_list):
                                    if (
                                            line.__contains__(header) and not line.__contains__(
                                        "browser.wait(until") and not line.__contains__(
                                        ".clear(")):
                                        if (type(cell_list[idx]) == int or type(cell_list[idx]) == float):
                                            line = line.split("sendKeys")[0] + "sendKeys(" + str(
                                                int(cell_list[idx])) + ");"
                                        else:
                                            line = line.split("sendKeys")[0] + "sendKeys('" + str(
                                                cell_list[idx]) + "');"
                                        break
                                replace_spec += line + ";"
                        new_spec =provideShortSpec(application_name) + str(replace_spec)
                        spec_file_name='multiple_scenario_file.js'
                        outerFolder="testRun"
                        new_conf = provideConf(spec_file_name,base_location,outerFolder,folderName)
                        multipleBatchFile = "START /B cmd " + scenarioPath + " && protractor " + scenarioPath + "\\multiple_conf_file.js >" + scenarioPath + "\\multiple_conf_file.log"
                        new_conf_file.write(new_conf)
                        new_spec_file.write(new_spec[:-4])
                        new_conf_file.close()
                        new_spec_file.close()
                        batch_file.write(multipleBatchFile)
                        batch_file.close()
                        subprocess.call(scenarioPath + '\\Batchfiles\\multiple_batch_file.bat')

                        conf_log = open(scenarioPath + "\\multiple_conf_file.log").read()
                        logVal = multipletestReportGeneration(scenarioPath)
                        time.sleep(1)

                        return jsonify(logVal)

                else:
                    return "scenarioFail"
            else:
                return "appFail"


@app.route("/updateSpecFromUI")
def updateSpecFromUI():
    global outerFolderName
    global folderName
    updatedSpec = request.args.get("updatedspec").split("/n,")
    print(type(updatedSpec))
    print(updatedSpec)
    fileOpened = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\\" + "spec.js", "w")
    for i in updatedSpec:
        if i.__contains__("/n"):
            i=str(i).replace("/n"," ")
        i=str(i).replace('HASH','#')
        i.strip()
        fileOpened.write(i+"\n")
    output_json = {"status": "Success"}
    return jsonify(output_json)


@app.route("/multipletestReportGeneration")  # Run multiple test case
def multipletestReportGeneration(scenarioPath):
    # folderName = request.args.get("folderName")
    open_file = open(scenarioPath + "\\multiple_conf_file.log")
    log_file = open(scenarioPath + "\\multiple_conf_file.log").read()
    read_f = (open_file.readlines())
    passStr = '[32m.[0m'
    failStr = '[31mF[0m'
    resultList = []
    header = []
    header1 = []
    data = []
    js = {}
    # passStr=".F."
    for wholeStr in read_f:
        if (wholeStr.__contains__(passStr) or wholeStr.__contains__(failStr)):
            outputLine = [(wholeStr[i:i + 10]) for i in range(0, len(wholeStr), 10)]
            for result in outputLine:
                if result == passStr:
                    resultList.append('Pass')
                elif result == failStr:
                    resultList.append('Fail')

        # if(r.__contains__(failStr)):

    file_name = scenarioPath + "\\sample1.xlsx"  # Path to your file
    df = pd.read_excel(file_name)
    # df["result"]=resultList[0]
    r, c = df.shape
    df['result'] = ""
    ind = 0
    for row in range(r):
        df["result"][row] = resultList[ind]
        ind += 1
    df.to_excel(scenarioPath + "\\new_excel_sheet.xlsx", index=False)
    x = df.head()
    time.sleep(1)
    excel_sheet = pd.read_excel(scenarioPath + "\\new_excel_sheet.xlsx")
    js["data"] = excel_sheet.to_dict(orient="records")
    for colName in excel_sheet.head():
        header.append(colName)
    # for h in header:
    js["headers"] = header
    totPass = resultList.count('Pass')
    totFail = resultList.count('Fail')
    testResult = [totPass, totFail]
    js["testResult"] = testResult
    js["log"] = log_file
    return js


@app.route("/ScheduleTestRunWithExcel", methods=['GET', 'POST'])  # upload excel and schedule the job
def ScheduleTestRunWithExcel():
    thread = threading.Thread(target=sched)
    thread.start()

    ScheduleDateTime = request.args.get("ScheduleDateTime")
    dailyFlag = request.args.get("ScheduleDateTimeFlag")
    testRunName = request.args.get("testRunName")
    email = request.args.get("email")
    schedulerName = request.args.get("schedulerName")
    file = request.get_data()
    print('file', file)

    catchit = uploadingExcelInSchedule(schedulerName, file, testRunName)

    print(catchit)
    with app.app_context():

        ScheduleDate = ScheduleDateTime.split("T")[0]
        ScheduleTime = ScheduleDateTime.split("T")[1]
        hours = ScheduleTime.split(":")[0]
        mins = ScheduleTime.split(":")[1]
        meridian = "AM"
        print('hours: ', hours, type(hours), int(hours) < 10)
        print(ScheduleTime, ScheduleDate, "Date & time")
        if int(hours) > 12:
            hours = int(hours) - 12
            if hours < 10:
                hours = '0' + str(hours)
            meridian = "PM"
        elif int(hours) < 10 and int(hours) != 0:
            print(int(hours))
            hours = '0' + str(hours)
            meridian = "AM"
        elif int(hours) == 12:
            meridian = "PM"
        ScheduleTime1 = str(hours) + ":" + str(mins)
        ScheduleDate = ScheduleDate.split("-")[1] + "/" + ScheduleDate.split("-")[2] + "/" + ScheduleDate.split("-")[0]
        print(ScheduleTime, ScheduleDate)

        connections = db.testRun.find()
        for connection in connections:
            if connection["test_name"] == testRunName:
                appPath = connection["path"]
                appPath = appPath.replace("\\\\", "\\")
                if os.path.isdir(appPath + "\\" + schedulerName) == False:
                    os.mkdir(appPath + "\\" + schedulerName)
                if os.path.isdir(appPath + "\\" + schedulerName + "\\combined_conf_file.log") == False:
                    open(appPath + "\\" + schedulerName + "\\combined_conf_file.log", "w").close()
                open(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat", "w").close()
                batch_file = open(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat", "w+")

                if dailyFlag == 'false':
                    batchScript = f'''

                           START /B cmd "cd {base_location}\\" && protractor {appPath}\\{schedulerName}\\combined_conf_file.js > {appPath}\\{schedulerName}\\combined_conf_file.log
                       '''
                    # db_scheduledTime = ScheduleDate + " at " + str(hours) + ":" + str(mins) +" "+ meridian
                    if int(hours) == 00:
                        db_scheduledTime = ScheduleDate + " at " + str(int(hours) + 12) + ":" + str(
                            mins) + " " + meridian
                    else:
                        db_scheduledTime = ScheduleDate + " at " + str(hours) + ":" + str(mins) + " " + meridian

                    dbData = [{"scheduler_name": schedulerName,
                               "test_job_name": testRunName,
                               "scheduled_time": db_scheduledTime,
                               "path": appPath}]

                    db.scheduledJobs.insert_many(dbData)

                    batch_file.write(batchScript)
                    batch_file.close()

                    # job(appPath, schedulerName, ScheduleDate)

                    month, day, year = ScheduleDate.split('/')
                    born = datetime(int(year), int(month), int(day))
                    my_day = born.strftime("%A")
                    print(my_day, ScheduleTime)

                    if my_day == "Monday":
                        schedule.every().monday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Tuesday":
                        schedule.every().tuesday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Wednesday":
                        schedule.every().wednesday.at(ScheduleTime).do(
                            lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Thursday":
                        schedule.every().thursday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Friday":
                        print("entered into this friday")

                        schedule.every().friday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                        # schedule.every().friday.at("15:55").do(lambda: job('Hello ', 'world!'))
                        print("schedule function is called..")
                    elif my_day == "Saturday":
                        schedule.every().saturday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Sunday":
                        schedule.every().sunday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))

                    # subprocess.call(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")
                else:

                    batchScript = f'''

                                     START /B cmd "cd {base_location}\\" && protractor {appPath}\\{schedulerName}\\combined_conf_file.js > {appPath}\\{schedulerName}\\combined_conf_file.log

                         '''

                    # db_scheduledTime = "Everyday at " + str(hours) + ":" + str(mins) + " " + meridian
                    if int(hours) == 00:
                        db_scheduledTime = "Everyday at " + str(int(hours) + 12) + ":" + str(mins) + " " + meridian
                    else:
                        db_scheduledTime = "Everyday at " + str(hours) + ":" + str(mins) + " " + meridian
                    dbData = [{"scheduler_name": schedulerName,
                               "test_job_name": testRunName,
                               "scheduled_time": db_scheduledTime,
                               "path": appPath}]
                    db.scheduledJobs.insert_many(dbData)
                    batch_file.write(batchScript)
                    batch_file.close()

                    # daily_job(appPath, schedulerName)
                    schedule.every().day.at(ScheduleTime).do(lambda: daily_job(appPath, schedulerName))
                    # subprocess.call(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")
        send_mail(email, testRunName)
        return jsonify("Success")

@app.route("/ScheduleTestRun")  # schedule the test job either daily or specific time
def ScheduleTestRun():
    thread = threading.Thread(target=sched)
    thread.start()
    with app.app_context():
        ScheduleDateTime = request.args.get("ScheduleDateTime")
        dailyFlag = request.args.get("ScheduleDateTimeFlag")
        testRunName = request.args.get("testRunName")
        email = request.args.get("email")
        schedulerName = request.args.get("schedulerName")
        ScheduleDate = ScheduleDateTime.split("T")[0]
        ScheduleTime = ScheduleDateTime.split("T")[1]
        hours = ScheduleTime.split(":")[0]
        mins = ScheduleTime.split(":")[1]
        meridian = "AM"

        print('hours: ', hours, type(hours), int(hours) < 10)
        print(ScheduleTime, ScheduleDate, "Date & time")
        if int(hours) > 12:
            hours = int(hours) - 12
            if hours < 10:
                hours = '0' + str(hours)
            meridian = "PM"
        elif int(hours) < 10 and int(hours) != 0:
            print(int(hours))
            hours = '0' + str(hours)
            meridian = "AM"
        elif int(hours) == 12:
            meridian = "PM"
        ScheduleTime1 = str(hours) + ":" + str(mins)
        ScheduleDate = ScheduleDate.split("-")[1] + "/" + ScheduleDate.split("-")[2] + "/" + ScheduleDate.split("-")[0]
        print(ScheduleTime, ScheduleDate)

        connections = db.testRun.find()
        for connection in connections:
            if connection["test_name"] == testRunName:
                print("entere")
                appPath = connection["path"]
                appPath = appPath.replace("\\\\", "\\")
                os.mkdir(appPath + "\\" + schedulerName)
                open(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat", "w").close()
                batch_file = open(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat", "w+")

                if dailyFlag == 'false':
                    batchScript = f'''
                                    START /B cmd cd {base_location}\\ && protractor {appPath}\\conf.js > {appPath}\\conf.log
                        '''
                    # db_scheduledTime = ScheduleDate + " at " + str(hours) + ":" + str(mins) +" "+ meridian
                    if int(hours) == 00:
                        db_scheduledTime = ScheduleDate + " at " + str(int(hours) + 12) + ":" + str(
                            mins) + " " + meridian
                    else:
                        db_scheduledTime = ScheduleDate + " at " + str(hours) + ":" + str(mins) + " " + meridian

                    dbData = [{"scheduler_name": schedulerName,
                               "test_job_name": testRunName,
                               "scheduled_time": db_scheduledTime,
                               "path": appPath}]

                    db.scheduledJobs.insert_many(dbData)

                    batch_file.write(batchScript)
                    batch_file.close()

                    # job(appPath, schedulerName, ScheduleDate)

                    month, day, year = (int(i) for i in ScheduleDate.split('/'))
                    # print(year, month, day)
                    born = datetime(year, month, day)

                    my_day = born.strftime("%A")
                    print(my_day)

                    if my_day == "Monday":
                        schedule.every().monday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Tuesday":
                        schedule.every().tuesday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Wednesday":
                        schedule.every().wednesday.at(ScheduleTime).do(
                            lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Thursday":
                        schedule.every().thursday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Friday":
                        schedule.every().friday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Saturday":
                        schedule.every().saturday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
                    elif my_day == "Sunday":
                        schedule.every().sunday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))

                else:
                    batchScript = f'''
                                    START /B cmd cd {base_location}\\ && protractor {appPath}\\conf.js > {appPath}\\conf.log
                        '''
                    # db_scheduledTime = "Everyday at " + str(hours) + ":" + str(mins) + " " + meridian
                    if int(hours) == 00:
                        db_scheduledTime = "Everyday at " + str(int(hours) + 12) + ":" + str(mins) + " " + meridian
                    else:
                        db_scheduledTime = "Everyday at " + str(hours) + ":" + str(mins) + " " + meridian
                    dbData = [{"scheduler_name": schedulerName,
                               "test_job_name": testRunName,
                               "scheduled_time": db_scheduledTime,
                               "path": appPath}]
                    db.scheduledJobs.insert_many(dbData)
                    batch_file.write(batchScript)
                    batch_file.close()

                    # daily_job(appPath, schedulerName)

                    print(ScheduleTime, "ScheduleTime")

                    schedule.every().day.at(ScheduleTime).do(lambda: daily_job(appPath, schedulerName))

                    # subprocess.call(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")
        send_mail(email, testRunName)
        return jsonify("Success")


def uploadingExcelInSchedule(schedulerName, file, testRunName):
    # folderName = request.args.get("folderName")

    global mulipleScenarioPath

    final_val = []

    str_file = file.decode("utf-8")
    # temp_file = list(eval(str_file))
    temp_List = []
    sheetLength = 0
    json_format = json.loads(str_file)
    print("str_file", str_file)

    # Checking the muliple scenario file exist or not, if exists remove the file
    connections = db.testRun.find()
    for connection in connections:
        if connection["test_name"] == testRunName:
            mulipleScenarioPath = connection["path"]
            if os.path.isdir(mulipleScenarioPath + "\\" + schedulerName) == False:
                os.mkdir(mulipleScenarioPath + "\\" + schedulerName)
            if os.path.exists(mulipleScenarioPath + "\\" + schedulerName + "\\combined_multi_scenario_file.js"):
                p = os.path.join(mulipleScenarioPath + "\\" + schedulerName + "\\combined_multi_scenario_file.js")
                os.remove(p)

    for k, v in json_format.items():
        v1 = json.dumps(v['data'])
        if type(list(eval(v1))[0]) == "str":
            return jsonify("failure")
        else:
            temp_file = list(eval(v1))
            df = pd.DataFrame.from_dict(temp_file)

        connections = db.testRun.find()
        for connection in connections:
            if connection["test_name"] == testRunName:
                mulipleScenarioPath = connection["path"]
                if os.path.isdir(mulipleScenarioPath):
                    df.to_excel(mulipleScenarioPath + "\\" + schedulerName + '\\sample.xls',
                                index=False)
                    time.sleep(1)
                    append_file = open(mulipleScenarioPath + "\\" + schedulerName + "\\append_spec.txt", "a+")
                    returnVal = uploadUIWithoutRun(schedulerName, testRunName, sheetLength)
                    # final_val.append(returnVal)
                    sheetLength += 1
                else:
                    return "appFail"

    # comb_multi_file = open(mulipleScenarioPath + "\\" + schedulerName + "\\job_upload_spec.js", "a+")
    #
    # comb_multi_file.write("});")
    # comb_multi_file.close()
    return jsonify("Success")  # create a json object to generate excel inui

def provideShortSpec(jobName):
    spec='''const jsdom = require("jsdom");
                         const fs = require('fs');
                         const path = require('path');
                         const Get_Text = new Map();
                        const Get_Attribute = new Map();
                         describe(\''''+jobName+'''\', function() { \n'''
    return spec

def provideSpec(url, outerfolderName, folderName):
    spec = '''
              const jsdom = require("jsdom");
              const fs = require('fs');
              const path = require('path');
              const Get_Text = new Map();
              const Get_Attribute = new Map();
              describe(\'''' + outerfolderName + '''\', function() {
                    it(\'''' + folderName + '''\', function() {
                        browser.ignoreSynchronization =true;
                        browser.driver.manage().window().maximize();
                        browser.get(\"''' + url + '''\");
                        browser.waitForAngularEnabled(false);
                        browser.sleep(10);
                        var until = protractor.ExpectedConditions;
                        var delay = 5000;
                        browser.sleep('''+str(sleep_delay)+''');
                        function hexToRgbA(hex){
                                            var c;
                                            if(/^#([A-Fa-f0-9]{3}){1,2}$/.test(hex)){
                                                c= hex.substring(1).split('');
                                                if(c.length== 3){
                                                    c= [c[0], c[0], c[1], c[1], c[2], c[2]];
                                                }
                                                c= '0x'+c.join('');
                                                return 'rgba('+[(c>>16)&255, (c>>8)&255, c&255].join(', ')+', 1)';
                                            }
                                            throw new Error('Bad Hex');
                                        }
        '''
    return spec


def provideJasmineSpecReporter(jobName):
    spec='''const jsdom = require("jsdom");
                         const fs = require('fs');
                         const path = require('path');
                         const Get_Text = new Map();
                        const Get_Attribute = new Map();
                         const {SpecReporter} = require("jasmine-spec-reporter");
                        jasmine.getEnv().addReporter(new SpecReporter({spec: {displayStacktrace: true, displayDuration: true}}));
                         describe(\''''+jobName+'''\', function() { \n'''
    return spec

def provideConf(spec_file_name,base_location,outerFolderName,folderName):
    print(spec_file_name,"here")
    conf = '''
             exports.config = {
            framework: 'jasmine',
            seleniumAddress: 'http://localhost:4444/wd/hub',
            specs: [\'''' + spec_file_name + '''\'],
            multiCapabilities: [
            {
                "browserName":"chrome",
                "shardTestFiles":true,
                "chromeOptions":{
                    "args":["--disable-gpu","-disable-dev-shm-usage","--no-sandbox","-disable-popup-blocking","--start-maximized","--disable-web-security","--allow-running-insecure-content","--disable-infobars"]
                }
            }
        ],
        useAllAngular2AppRoots: true,
      onPrepare: function(){
	//Getting XML report    
			var jasmineReporters = require('jasmine-reporters');
			jasmine.getEnv().addReporter(new jasmineReporters.JUnitXmlReporter({
			   consolidateAll: true,
			   filePrefix: 'guitest-xmloutput',
			   savePath: \'''' + base_location.replace("\\",
                                                       "\\\\") + '''\\\\''' + outerFolderName + '''\\\\''' + folderName + '''\'
			}));
	//Getting screenshots  
		var fs = require('fs-extra');
		fs.emptyDir(`''' + base_location.replace("\\", "\\\\") + '''\\\\''' + outerFolderName + '''\\\\''' + folderName + '''\\\\screenshots`, function (err) {
				 console.log(err);
			 });
			 jasmine.getEnv().addReporter({
				 specDone: function(result) {
					 if (result.status == 'failed') {
                         browser.takeScreenshot().then(function (png) {
                        var stream = fs.createWriteStream(`''' + base_location.replace("\\",
                                                                                       "\\\\") + '''\\\\''' + outerFolderName + '''\\\\''' + folderName + '''\\\\screenshots\\\\`+ result.fullName+ '.png');
                        stream.write(new Buffer.from(png, 'base64'));
                             stream.end();
                         });
						
					 }
				 }
			 });
			 }, 
		 onComplete: function() {
			 //Getting HTML reportvar browserName, browserVersion;
				var browserName = "chrome";
				var platform = "iTAP"
				var HTMLReport = require('protractor-html-reporter-2');
				testConfig = {
					reportTitle: `Test Report for'''+outerFolderName+'''-'''+folderName+'''`,
					outputPath: \'''' + base_location.replace("\\",
                                                              "\\\\") + '''\\\\''' + outerFolderName + '''\\\\''' + folderName + '''\',
					outputFilename: 'ProtractorTestReport',
					//screenshotPath: `''' + base_location.replace("\\",
                                                                 "\\\\") + '''\\\\''' + outerFolderName + '''\\\\''' + folderName + '''\\\\screenshots`,
					testBrowser: browserName,
					modifiedSuiteName: false,
					//screenshotsOnlyOnFailure: true,
					testPlatform: platform
				};
				new HTMLReport().from(`''' + base_location.replace("\\",
                                                                 "\\\\") + '''\\\\''' + outerFolderName + '''\\\\''' + folderName + '''\\\\guitest-xmloutput.xml`, testConfig);
			
		},
       getPageTimeout: 20000,
                    allScriptsTimeout: 3000000,
                    jasmineNodeOpts: {defaultTimeoutInterval: 3000000},
                    defaultTimeoutInterval: 3000000
      }
        '''
    return conf

def provideConfBrowserCompatability(spec_file_name):
    conf= '''
    exports.config={
    framework:"jasmine",
	seleniumAddress:"http://localhost:4444/wd/hub",
    specs:[\'''' + spec_file_name + '''\'],
	multiCapabilities: [
		{
			"browserName":"chrome",
			"shardTestFiles":true,
			"chromeOptions":{
				"args":["--disable-gpu","-disable-dev-shm-usage","--no-sandbox","-disable-popup-blocking","--start-maximized","--disable-web-security","--allow-running-insecure-content","--disable-infobars"]
		    }	
		},
		{
			"browserName":"firefox",
			"shardTestFiles":true,
			"moz:firefoxOptions":{
				"args":["--disable-gpu","-disable-dev-shm-usage","--no-sandbox","-disable-popup-blocking","--start-maximized","--disable-web-security","--allow-running-insecure-content","--disable-infobars"],
			    'binary' : "C:\\\Program Files (x86)\\\Mozilla Firefox\\firefox.exe"
			},
			'firefoxPath' : "C:\\\Program Files (x86)\\\Mozilla Firefox\\firefox.exe",


		}
	],
	getPageTimeout:1000000,
	allScriptsTimeout:3000000,
	jasmineNodeOpts:{defaultTimeoutInterval:3000000},
	defaultTimeoutInterval:3000000
    };'''
    return conf

if __name__ == "__main__":
    app.run(host='0.0.0.0', port=5010, debug=True)
    
