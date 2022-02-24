from __future__ import print_function
from mailmerge import MailMerge
import pythoncom
from datetime import date
import os, win32com
import win32com.client as win32
import copy
import json
import numpy as np
import time
import os, json
import math
from math import floor as floor
from docx import Document
from docx.shared import Inches
import shutil
import xlrd
from flask import Flask, request, jsonify
from flask_cors import CORS
import logging.config
import logging
from datetime import datetime, timedelta
import xlwt
from pymongo import MongoClient
# import excel2json
from collections import OrderedDict
# import simplejson as json
from zipfile import ZipFile
import openpyxl
import pandas
import subprocess
import pandas as pd
from urllib.parse import unquote
from time import strftime
from time import gmtime
from openpyxl import load_workbook
import winapps
import requests, zipfile, io
import schedule
import threading
#newones
import re
import win32api

#TestAutomationTool
app = Flask(__name__)
CORS(app)

client = MongoClient('localhost', 27017)
db = client["MDB"]

global folderName
global outerFolderName
with open("config.json") as f:
  data_folder = json.load(f)

base_location=data_folder['base_location']
base_location_front= base_location.replace("\\","/")
ui_location=data_folder['ui_location']
ui_delay = data_folder['delay']
ui_minimumdelay = data_folder['minimum_delay']

@app.route("/html_extractor")  # Triggers the batch file which will further trigger protractor confApr9.js command
def html_extractor():
  # creates conf spec and batch files subprocess triggers batch 
  # folderName = request.args.get("folderName")
  if os.path.isdir(base_location) == False:
    os.mkdir(base_location)
  if os.path.isdir(base_location+'\\' + outerFolderName + '\\' + folderName) == False:
    os.mkdir(base_location+'\\' + outerFolderName + '\\' + folderName)
  confApr = '''
  exports.config = {
    framework: 'jasmine',
    seleniumAddress: 'http://localhost:4444/wd/hub',
    specs: ['specApr9.js'],
    commonCapabilities: {
     'browserName': 'chrome',
     'chromeOptions': {
			'args': ['--start-maximized']
		}
   },
   onPrepare: function() {
      return browser.takeScreenshot();
    },
   getPageTimeout: 20000,
	allScriptsTimeout: 3000000,
	jasmineNodeOpts: {defaultTimeoutInterval: 3000000},
	defaultTimeoutInterval: 3000000
  }
  '''
  conf = '''
    exports.config = {
      framework: 'jasmine',
      seleniumAddress: 'http://localhost:4444/wd/hub',
      specs: ['spec.js'],
      commonCapabilities: {
       'browserName': 'chrome',
       'chromeOptions': {
                                                  'args': [ '--disable-gpu','-disable-dev-shm-usage','--no-sandbox','-disable-popup-blocking','--start-maximized','--disable-web-security','--allow-running-insecure-content','--disable-infobars']
                                  }
     },
     onPrepare: function() {
        return browser.takeScreenshot();
      },
     getPageTimeout: 20000,
                  allScriptsTimeout: 3000000,
                  jasmineNodeOpts: {defaultTimeoutInterval: 3000000},
                  defaultTimeoutInterval: 3000000
    }

    '''
  specApr = '''
const jsdom = require("jsdom");
const fileUrl = \''''+base_location_front+'''//''' + outerFolderName + '/' + folderName + '''/inputfile.json';
const fileUrl1 = \''''+base_location_front+'''//''' + outerFolderName + '/' + folderName + '''/inputfile1.json';
const fs = require('fs');
const urlListJSON = require(\''''+base_location_front+'''//''' + outerFolderName + '/' + folderName + '''/urls.json');
var ipJson = require(fileUrl);

describe('Extract HTML elements', function() {
    it('should extract all elements from all webpage(s)', function() {
        if(urlListJSON.length>=1) {
            for(var i=0;i<urlListJSON.length;i++)
            {
                var urlList = [];
                var inputList = [];
				var buttonList = [];
                var dropdownList = [];
                var textareaList = [];
                var linkList = [];
                var tot_divlist = [];
                var divList = [];
                var divTagIdx = [];
                var tot_divtagList = [];
				var spanList = [];
                var inputJson = {};
                var aTagIdx = [];
                var anchorList = [];
                var tot_List = [];
				var complete_List=[];
				var tagName_1=[];

                var tot_anchorList = [];
                var tot_linkList = [];
                var url = '';
				var divCounter=0;
				var divAccList = [];
                var orderOfIpElems = {};
				var un_handle=["html", "head", "style", "title", "meta", "option","link", "br", "body", "script","input",
					"textarea", "a", "button", "span", "select","app-root","app-sidebar", "section","div",  "nav","td", "tr", "app-router", "app-footer",
					"app-header", "router-outlet", "base", "app-login","app-layout", "form", "footer", "header", "li","dd","dt","dl" ];
                var un_app="app-";
                var un_ng="ng-";
                if(urlListJSON[i]!="")
                    url = urlListJSON[i];
                else {
                    console.log('invalid url')
                    break;
                }
                browser.ignoreSynchronization = true;
				browser.driver.manage().window().maximize();
                browser.get(url);

				browser.sleep(3000);
                browser.getTitle().then(function(title){
                    console.log('Title of Application:',title);
                });

                browser.getCurrentUrl().then(function(currUrl){
                    urlList.push(currUrl);
                    inputJson["url"] = urlList;
                });

				element.all(by.xpath('//*')).getTagName().then(function(outerHTML){
					outerHTML.forEach(function (arrayValue, idx) {
						if(tagName_1.indexOf(arrayValue)<0 && un_handle.indexOf(arrayValue)<0 && arrayValue.indexOf(un_app)<0 && arrayValue.indexOf(un_ng)<0){
							tagName_1.push(arrayValue);
							console.log("tagname",tagName_1);
							element.all(by.xpath('//'+arrayValue)).getAttribute("outerHTML").then(function(outerHTML){
								outerHTML.forEach(function (arrayItem) {
									complete_List.push(arrayItem);
									/*const dom = new jsdom.JSDOM(arrayItem);
									var tmp = dom.window.document.querySelector(arrayValue).innerHTML;
									console.log(tmp, 'innerHTML');
									console.log(arrayValue, 'arrayValue');*/
								});
								inputJson[arrayValue] = complete_List;
								complete_List=[];
							});
							element.all(by.xpath('//'+arrayValue)).getLocation().then(function(pos){
								orderOfIpElems[arrayValue]=[];
								for(var i=0;i<pos.length;i++){
									orderOfIpElems[arrayValue].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
								}
								console.log('Elements extracted..');
							});
						}
					});
				});

                element.all(by.xpath('//a')).getAttribute("outerHTML").then(function(outerHTML){
                    console.log("inside anchor");
                    outerHTML.forEach(function (arrayItem, idx) {
                        const dom = new jsdom.JSDOM(arrayItem);
                        tot_List.push(arrayItem)
                        if((dom.window.document.querySelector("a").getAttribute("data-toggle")) == "collapse"
						&& (dom.window.document.querySelector("a").getAttribute("aria-expanded")) == "false"){
                            aTagIdx.push(idx);
                            anchorList.push(idx);
                            tot_anchorList.push(arrayItem)
                        } else {
                            linkList.push(idx);
                            tot_linkList.push(arrayItem)
                        }
                    });
                    var remaining = tot_List.filter(item => tot_anchorList.indexOf(item) < 0);
                    if(aTagIdx.length!=0)
                    {   aTagIdx.reverse();
                        aTagIdx.forEach(function(idx){
                            element(by.xpath("//a["+(idx+1)+"]")).click();
                        })
                    }
                });

                element.all(by.xpath('//div[contains(@data-toggle,"collapse")]')).getAttribute("outerHTML").then(function(outerHTML){
                  outerHTML.forEach(function(arrayItem, idx){
          
                      const dom = new jsdom.JSDOM(arrayItem)
                      var div = dom.window.document.querySelector("div");
                      divTagIdx.push(idx);
          
                  });
                  if(divTagIdx.length!=0)
                  {
                      console.log("div Tag Idx b4", divTagIdx);
                      divTagIdx.reverse();
                      console.log("div Tag Idx after", divTagIdx);
                      console.log("Click reversed");
                      divTagIdx.forEach(function(idx){
                          element.all(by.xpath('//div[contains(@data-toggle,"collapse")]')).get(idx).click();
                          console.log("click happening", idx);
                          browser.sleep(1000);
                      });
                  }
              });


                element.all(by.xpath('//input')).getAttribute("outerHTML").then(function(outerHTML){
                    console.log("inside input");
                    outerHTML.forEach(function (arrayItem) {
                        inputList.push(arrayItem);
                    });
					inputJson["linklist"] = tot_List;
                    inputJson["input"] = inputList;
                });

				element.all(by.xpath('//span')).getAttribute("outerHTML").then(function(outerHTML){
				            console.log("inside span");
                    outerHTML.forEach(function (arrayItem) {
                        spanList.push(arrayItem);
                    });
					inputJson["span"] = spanList;
                });

				element.all(by.xpath('//button')).getAttribute("outerHTML").then(function(outerHTML){
				            console.log("inside button");
                    outerHTML.forEach(function (arrayItem) {
                        buttonList.push(arrayItem);
                    });
                    inputJson["button"] = buttonList;
                });

                element.all(by.xpath('//select')).getAttribute("outerHTML").then(function(outerHTML){
                    console.log("inside select");
                    outerHTML.forEach(function (arrayItem) {
                        dropdownList.push(arrayItem);
                    });
                    inputJson["select"] = dropdownList;
                });
                element.all(by.xpath('//textarea')).getAttribute("outerHTML").then(function(outerHTML){
                    console.log("inside textarea");
                    outerHTML.forEach(function (arrayItem) {
                        textareaList.push(arrayItem);
                    });
                    inputJson["textarea"] = textareaList;
                });

				element.all(by.xpath('//input')).getLocation().then(function(pos){
				console.log("inside input getloc");
					orderOfIpElems['input']=[];
                    for(var i=0;i<pos.length;i++){

                        orderOfIpElems['input'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
                    }
                })
                element.all(by.xpath('//a')).getLocation().then(function(pos){
                console.log("inside anchor getloc");
					orderOfIpElems['linklist']=[]
                    for(var i=0;i<pos.length;i++){
                        orderOfIpElems['linklist'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
                    }
                });
				element.all(by.xpath('//span')).getLocation().then(function(pos){
				console.log("inside span getloc");
					orderOfIpElems['span']=[];
                    for(var i=0;i<pos.length;i++){
                        orderOfIpElems['span'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
                    }
                });
                element.all(by.xpath('//select')).getLocation().then(function(pos){
					orderOfIpElems['select']=[];
                    for(var i=0;i<pos.length;i++){
                        orderOfIpElems['select'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
                    }
                });
				element.all(by.xpath('//button')).getLocation().then(function(pos){
				console.log("inside button getloc");
					orderOfIpElems['button']=[];
                    for(var i=0;i<pos.length;i++){
                        orderOfIpElems['button'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
                    }
                });

                element.all(by.xpath('//textarea')).getLocation().then(function(pos){
                console.log("inside textarea getloc");
					orderOfIpElems['textarea']=[];
                    for(var i=0;i<pos.length;i++){
                        orderOfIpElems['textarea'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
                    }
                    inputJson['order'] = [orderOfIpElems]
                    ipJson.push(inputJson);
                    // console.log('current Push json:', inputJson);
                    // console.log('File input json:',ipJson);
                    var data = JSON.stringify(ipJson);
                    fs.writeFileSync(fileUrl, data.replace(/\\n/g, '').replace(/\\t/g, ''));
                    fs.writeFileSync(fileUrl1, data.replace(/\\n/g, '').replace(/\\t/g, ''));
                    console.log('Elements extracted..')
                    urlList = [];
                    inputList = [];
                    dropdownList = [];
					buttonList = [];
                    textareaList = [];
                    linkList = [];
                    inputJson = {};
                    aTagIdx = [];
					spanList=[];
                    anchorList = [];
                    url = '';
                    orderOfIpElems = {};
                });
            }
        } else {
            console.log("Please enter URL");
        }
    });
});

  '''
  if os.path.isfile(base_location+'\\' + outerFolderName + '\\' + folderName + "\\specApr9.js") == False:
    specFile = open(base_location+'\\' + outerFolderName + '\\' + folderName + "\\specApr9.js", "w+")
    specFile.write(specApr)
    specFile.close()
  if os.path.isfile(base_location+'\\' + outerFolderName + '\\' + folderName + "\\confApr9.js") == False:
    confFile = open(base_location+'\\' + outerFolderName + '\\' + folderName + "\\confApr9.js", "w+")
    confFile.write(confApr)
    confFile.close()
  if os.path.isdir(base_location+'\\' + outerFolderName + '\\' + folderName + '\\Batchfiles') == False:
    os.mkdir(base_location+'\\' + outerFolderName + '\\' + folderName + '\\Batchfiles')

  # if os.path.isfile('base_location\\'+folderName+"\\Batchfiles\\batch.bat")==False:
  f = open(base_location+'\\' + outerFolderName + '\\' + folderName + "\\Batchfiles\\batch.bat", "w+")
  script1 = '''
  START /B cmd "cd '''+ base_location+'''\\" && protractor '''+ base_location+'''\\''' + outerFolderName + '\\' + folderName + '''\\confApr9.js >'''+ base_location+'''\\''' + outerFolderName + '\\' + folderName + '''\\\confApr9.log
  '''
  f.write(script1)
  f.close()

  subprocess.call(base_location+'\\' + outerFolderName + '\\' + folderName + '\\Batchfiles\\batch.bat')
  return jsonify({"status": "success"})

@app.route("/test_run")  # Triggers the batch file which will further trigger protractor conf.js command
def test_run():   # creation of batch file and triggers conf file
  # folderName=request.args.get("folderName")
  if os.path.isdir(base_location+"\\" + outerFolderName + '\\' + folderName + "\OldFiles") == False:
    os.mkdir(base_location+"\\" + outerFolderName + '\\' + folderName + "\OldFiles")

  if os.path.isdir(base_location+'\\' + outerFolderName + '\\' + folderName + '\Batchfiles') == False:
    os.mkdir(base_location+'\\' + outerFolderName + '\\' + folderName + '\Batchfiles')
  script = '''
  START /B cmd "cd '''+ base_location+'''\\''' + outerFolderName + '\\' + folderName + '''\OldFiles\\" && protractor '''+ base_location+'''\\''' + outerFolderName + '\\' + folderName + '''\OldFiles\conf.js >'''+ base_location+'''\\''' + outerFolderName + '\\' + folderName + '''\OldFiles\conf.log
  '''
  f = open(base_location+'\\' + outerFolderName + '\\' + folderName + '\Batchfiles\Test_batch.bat', "w+")
  f.write(script)
  f.close()
  subprocess.call(base_location+'\\' + outerFolderName + '\\' + folderName + '\Batchfiles\Test_batch.bat')
  return jsonify("success")

@app.route("/handle_elem")  # Triggers the batch file which will further trigger protractor conf.js command
def handle_elem():
  # folderName=request.args.get("folderName")
  if os.path.isdir(base_location+'\\' + outerFolderName + '\\' + folderName + '\Batchfiles'):
    os.mkdir(base_location+'\\' + outerFolderName + '\\' + folderName + '\Batchfiles')

  f = open(base_location+'\\' + outerFolderName + '\\' + folderName + '\Batchfiles\Handle_batch.bat', "w+")
  script = '''
  START /B cmd "cd '''+ base_location+'''\" && protractor '''+ base_location+'''\\''' + outerFolderName + '\\' + folderName + '''\confJul17.js > '''+ base_location+'''\\''' + outerFolderName + '\\' + folderName + '''\confJul17.log
  '''
  f.write(script)
  f.close()
  subprocess.call(base_location+'\\Batchfiles\\Handle_batch.bat')
  # f = open(r"base_location\requried_3.txt")
  # x = f.read().split(",")
  return jsonify("Success")
  # return "success"

@app.route("/downloadTestCase")  # Download spec file
def downloadTestCase():
  # folderName = request.args.get("folderName")
  date_time = datetime.now().strftime("%d-%m-%Y %H:%M:%S")
  conf_data = open(base_location+"\\" + outerFolderName + '\\' + folderName + "\OldFiles\conf.js").read()
  conf_data = conf_data.replace('spec.js', "spec" + date_time + ".js")
  spec_data = open(base_location+"\\" + outerFolderName + '\\' + folderName + "\OldFiles\spec.js").read()
  json_format = {
    "conf_fname": "conf" + date_time + ".js",
    "spec_fname": "spec" + date_time + ".js",
    "conf": conf_data,
    "spec": spec_data
  }
  return jsonify(json_format)

@app.route("/extractHTML")  # Insert url in urls.json file and creates inputfile.json
def extractHTML():
  global outerFolderName
  dbData = [{"path": base_location+"\\" + outerFolderName,
             "innerpath": base_location+"\\" + outerFolderName + "\\" + folderName,
             "appName": outerFolderName,
             "scenarioName": folderName,
             "test_Description": ''}]
  db.folderName.insert_many(dbData)
  global counter
  counter = 0
  global final_line
  final_line = ''
  global whole_script
  whole_script = ''
  global modal_count

  modal_count = 0
  urls = request.args.get('url').replace("HASH", "#")
  # folderName=request.args.get("folderName")
  # folderName="CRST_Application"
  url_list = []
  if ',' in urls:
    url_list = urls.split(',')
  else:
    url_list.insert(0, urls)
  # callNpm()
  if os.path.isdir(base_location+"/" + outerFolderName + '/' + folderName) == False:
    os.mkdir(base_location+"/" + outerFolderName + '/' + folderName)

  if os.path.isdir(base_location+"/" + outerFolderName + '/' + folderName + "/OldFiles") == False:
    os.mkdir(base_location+"/" + outerFolderName + '/' + folderName + "/OldFiles")

  with open(base_location+"\\" + outerFolderName + '\\' + folderName + "\\urls.json", "w") as outfile:
    outfile.write(json.dumps(url_list, indent=4))
  with open(base_location+"\\" + outerFolderName + '\\' + folderName + "\\inputfile.json", "w") as outfile:
    outfile.write(json.dumps([], indent=4))

  with open(base_location+"\\" + outerFolderName + '\\' + folderName + "\\inputfile1.json",
            "w") as outfile:
    outfile.write(json.dumps([], indent=4))

  if os.path.isfile(
    base_location+"\\" + outerFolderName + '\\' + folderName + "\OldFiles\spec.js") == True:
    open(base_location+"\\" + outerFolderName + '\\' + folderName + "\OldFiles\spec.js", "w").close()
  # open(base_location+"\\OldFiles\\spec.js", "w").close()
  if os.path.isfile(
    base_location+"\\" + outerFolderName + '\\' + folderName + "\OldFiles\multiple_spec_file.js") == True:
    open(base_location+"\\" + outerFolderName + '\\' + folderName + "\OldFiles\multiple_spec_file.js",
         "w").close()
  # open(base_location+"\\OldFiles\\multiple_spec_file.js", "w").close()
  if os.path.isfile(
    base_location+"\\" + outerFolderName + '\\' + folderName + "\\uniqure_output_file.js") == True:
    open(base_location+"\\" + outerFolderName + '\\' + folderName + "uniqure_output_file.txt", "w").close()
  # open(base_location+"\\uniqure_output_file.txt", "w").close()
  return jsonify({"status": "success"})

t = 0

@app.route("/log_files")  # Create confApr9.log file
def log_files():
  # folderName = request.args.get("folderName")
  loc = base_location+"\\" + outerFolderName + '\\' + folderName + "confApr9.log"
  f = open(loc, "r")
  c = os.path.getmtime(loc)
  global t
  if t != c:
    t = c
    return jsonify(f.read())

tt = 0

@app.route("/test_log_files")  # Create conf.log file
def test_log_files():
  # folderName = request.args.get("folderName")
  loc = base_location+"\\" + outerFolderName + '\\' + folderName + "\OldFiles\conf.log"
  f = open(loc, "r")
  c = os.path.getmtime(loc)
  global tt
  if tt != c:
    tt = c
    return jsonify(f.read())

@app.route("/new_spec_file")
def new_spec_file():
  tags = request.args.get("tagName")
  # tags="div,p"
  if (tags == ""):
    tags = "textarea"
  split_tag = tags.split(",")
  line1 = ''
  line2 = ''
  spec_line = ''
  full_code = ''
  ord_line = ''
  order_line = ''
  open_spec = open(base_location+"\\specApr9.js").readlines()
  new_spec = open(base_location+"\\specJul23.js", "w+")
  for tag in split_tag:
    line1 += "        var " + tag + "List=[]" + '''
        '''
    line2 += ",\"" + tag + "\":[]"
    spec_line += '''        element.all(by.xpath('//''' + tag + '''\')).getAttribute("outerHTML").then(function(outerHTML){
                      outerHTML.forEach(function (arrayItem) {
                          ''' + tag + '''List.push(arrayItem);
                      });
                    inputJson["''' + tag + '''"] =''' + tag + '''List;
                });
        '''
    order_line += '''        element.all(by.xpath('//''' + tag + '''\')).getLocation().then(function(pos){
                    for(var i=0;i<pos.length;i++){

                        orderOfIpElems[\'''' + tag + '''\'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y)
                        + ")"); } })

    '''

  for line in open_spec:
    if line.strip().__contains__("urlList = [];"):
      full_code += "        " + line1 + '''
''' + line
    elif line.strip().__contains__("element.all(by.xpath('//a')).getAttribute"):
      full_code += "        " + spec_line + '''
      ''' + line

    elif line.strip().__contains__("element.all(by.xpath('//input')).getLocation()"):
      full_code += "         " + order_line + '''
      ''' + line
    elif line.__contains__("var orderOfIpElems = {\"input\":[], \"select\":[], \"button\":[]"):
      full_code += " " + line.split("};")[0] + line2 + "};"
    else:
      full_code += line

  new_spec.write(full_code)
  clr = open(base_location+"\\requried_3.txt", "w+")
  return jsonify("success")

@app.route("/new_spec_file_1")
def new_spec_file_1():
  tags = request.args.get("tagName")
  # tags="div,p"
  if (tags == ""):
    tags = "textarea"
  split_tag = tags.split(",")
  line1 = ''
  line2 = ''
  spec_line = ''
  full_code = ''
  ord_line = ''
  order_line = ''
  open_spec = open(base_location+"\\oldFiles\\spec.js").readlines()
  new_spec = open(base_location+"\\oldFiles\\updated_spec.js", "w+")
  for tag in split_tag:
    line1 += "        var " + tag + "List=[]" + '''
        '''
    line2 += ",\"" + tag + "\":[]"
    spec_line += '''        element.all(by.xpath('//''' + tag + '''\')).getAttribute("outerHTML").then(function(outerHTML){
                      outerHTML.forEach(function (arrayItem) {
                          ''' + tag + '''List.push(arrayItem);
                      });
                    inputJson["''' + tag + '''"] =''' + tag + '''List;
                });
        '''
    order_line += '''        element.all(by.xpath('//''' + tag + '''\')).getLocation().then(function(pos){
                    for(var i=0;i<pos.length;i++){

                        orderOfIpElems[\'''' + tag + '''\'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y)
                        + ")"); } })

    '''

  for line in open_spec:
    if line.strip().__contains__("var urlList = [];"):
      full_code += "        " + line1 + '''
''' + line
    elif line.strip().__contains__("urlList = [];"):
      full_code += "        " + line1.replace("var", "") + '''
  ''' + line
    elif line.strip().__contains__("element.all(by.xpath('//a')).getAttribute"):
      full_code += "        " + spec_line + '''
      ''' + line

    elif line.strip().__contains__("element.all(by.xpath('//input')).getLocation()"):
      full_code += "         " + order_line + '''
      ''' + line
    elif line.__contains__("var orderOfIpElems = { \""  "orderOfIpElems = { \"") or \
      line.__contains__("orderOfIpElems = { \""):
      full_code += " " + line.split("};")[0] + line2 + "};"
    else:
      full_code += line

  new_spec.write(full_code)
  clr = open(base_location+"\\requried_3.txt", "w+")
  return jsonify("success")

@app.route("/new_spec_file_1_modal")
def new_spec_file_1_modal():
  tags = request.args.get("tagName")
  # tags="div,p"
  if (tags == ""):
    tags = "textarea"
  split_tag = tags.split(",")
  line1 = ''
  line2 = ''
  spec_line = ''
  full_code = ''
  ord_line = ''
  order_line = ''
  open_spec = open(base_location+"\\oldFiles\\spec.js").readlines()
  new_spec = open(base_location+"\\oldFiles\\updated_spec.js", "w+")
  for tag in split_tag:
    line1 += "        var " + tag + "List=[]" + '''
        '''
    line2 += ",\"" + tag + "\":[]"
    spec_line += '''        element.all(by.xpath(\"//div[contains(@style, 'display: block')]//''' + tag + '''\")).getAttribute("outerHTML").then(function(outerHTML){
                      outerHTML.forEach(function (arrayItem) {
                          ''' + tag + '''List.push(arrayItem);
                      });
                    inputJson["''' + tag + '''"] =''' + tag + '''List;
                });
        '''
    order_line += '''        element.all(by.xpath(\"//div[contains(@style, 'display: block')]//''' + tag + '''\")).getLocation().then(function(pos){
                    for(var i=0;i<pos.length;i++){

                        orderOfIpElems[\'''' + tag + '''\'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y)
                        + ")"); } })

    '''

  for line in open_spec:
    if line.strip().__contains__("var urlList = [];"):
      full_code += "        " + line1 + '''
''' + line
    elif line.strip().__contains__("urlList = [];"):
      full_code += "        " + line1.replace("var", "") + '''
  ''' + line
    elif line.strip().__contains__(
      "element.all(by.xpath(\"//div[contains(@style, 'display: block')]//a\")).getAttribute") or \
      line.strip().__contains__(
        "element.all(By.xpath(\"//div[contains(@style, 'display: block')]//a\")).getAttribute("):
      full_code += "        " + spec_line + '''
      ''' + line

    elif line.strip().__contains__(
      "element.all(by.xpath(\"//div[contains(@style, 'display: block')]//input\")).getLocation()") or \
      line.strip().__contains__(
        "element.all(By.xpath(\"//div[contains(@style, 'display: block')]//input\")).getLocation()"):
      full_code += "         " + order_line + '''
      ''' + line
    elif line.__contains__("var orderOfIpElems = { \""  "orderOfIpElems = { \"") or \
      line.__contains__("orderOfIpElems = { \""):
      full_code += " " + line.split("};")[0] + line2 + "};"
    else:
      full_code += line

  new_spec.write(full_code)
  clr = open(base_location+"\\requried_3.txt", "w+")
  return jsonify("success")

@app.route("/control_flow_one")
def control_flow_one():
  f1 = open(base_location+"\\" + outerFolderName + '\\' + folderName + "\oldFiles\\updatespec.js", "r")
  # new_f=open(r"base_location\oldFiles\spec28.js", "r")
  # new_f=open(r"base_location\ExtractScript.txt","r")
  file = open(base_location+"\\" + outerFolderName + '\\' + folderName + "\oldFiles\\fileJul28.js", "w+")
  x = f1.readlines()
  x1 = ''
  for l in x:
    if (l.__contains__("extract();") or l.__contains__("extractModal();") and l.strip() != ''):
      if (l.strip().startswith("//") == False):
        l = l.replace("extract();", "//extract();")
        l = l.replace("extractModal();", "//extractModal();")
      if (l.__contains__("extractModal();")):
        l += '''
                                                  extractUnhandled();'''
      x1 += l
    else:
      x1 += l
  # y=new_f.read()
  file.write(x1)
  subprocess.call(
    base_location+'\\' + outerFolderName + '\\' + folderName + '\\Batchfiles\\unhandled_batch.bat')
  return jsonify("success")

@app.route("/control_flow_one_modal")
def control_flow_one_modal():
  f1 = open(base_location+"\\" + outerFolderName + '\\' + folderName + "\\oldFiles\\updatespec.js", "r")
  # new_f=open(r"base_location\oldFiles\spec28.js", "r")
  # new_f=open(r"base_location\ExtractScript.txt","r")
  file = open(base_location+"\\" + outerFolderName + '\\' + folderName + "\\oldFiles\\fileJul28.js", "w+")
  x = f1.readlines()
  x1 = ''
  for l in x:
    if (l.__contains__("extract();") or l.__contains__("extractModal();") and l.strip() != ''):
      if (l.strip().startswith("//") == False):
        l = l.replace("extract();", "//extract();")
        l = l.replace("extractModal();", "//extractModal();")
      if (l.__contains__("extractModal();")):
        l += '''
                                                  extractUnhandledModal();'''
      x1 += l
    else:
      x1 += l
  # y=new_f.read()
  file.write(x1)
  subprocess.call(
    base_location+'\\' + outerFolderName + '\\' + folderName + '\\Batchfiles\\unhandled_batch.bat')
  return jsonify("success")

@app.route("/download_spec")
def download_spec():
  # folderName = request.args.get("folderName")
  spec_file = open(base_location+"\\" + outerFolderName + '\\' + folderName + "\\oldFiles\\spec.js")
  headers = []
  values = []
  open_spec = spec_file.read()
  json_format = {}
  final_json = {}
  headerList = []
  final_list = []
  split_file = open_spec.split("browser.driver.manage().window().maximize();")[1].split("extract()")[0].split(";")
  for elm in split_file:
    if (elm.find("element") >= 0):
      if (elm.find("sendKeys") >= 0):
        fin_elm = elm.split("xpath")[1].split(".sendKeys")[0]
        headers.append(fin_elm)
        if elm.split("sendKeys")[1].__contains__("'"):
          fin_val = elm.split("sendKeys")[1].split("'")[1]
        elif elm.split("sendKeys")[1].__contains__('"'):
          fin_val = elm.split("sendKeys")[1].split('"')[1]
        values.append(fin_val)
      # elif(elm.find("click()")>=0):
      #   fin_elm = elm.split("xpath")[1].split(".click()")[0]
      #   headers.append(fin_elm)
      #   fin_val = "click()"
      #   values.append(fin_val)
  workbook = xlwt.Workbook()
  sheet = workbook.add_sheet("Sheet Name")
  # Specifying style
  header_style = xlwt.easyxf('font: bold 1')
  val_style = xlwt.easyxf('font: bold 0')
  # Specifying column

  for index in range(len(headers)):
    if headers[index] not in headerList:
      headerList.append(headers[index])
    json_format[headers[index]] = values[index]
  final_list.append(json_format)
  final_list_1 = final_list[::-1]
  # sheet.write(0, index, headers[index], header_style)
  # sheet.write(1, index, values[index], val_style)
  # workbook.save("sample_2.xls")
  # wb = xlrd.open_workbook(r'base_location\oldFiles\sample_2.xls')
  # # sh = wb.sheet_by_index(0)
  # # # List to hold dictionaries
  # # data_list = []
  # # # Iterate through each row in worksheet and fetch values into dict
  # # for rownum in range(1, sh.nrows):
  # #   # data = OrderedDict()
  # #     # Added line
  # # # Serialize the list of dicts to JSON
  # # j = json.dumps(data_list)
  # json_format=excel2json.convert_from_file(r'base_location\oldFiles\sample_2.xls')
  # excel_data_df = pandas.read_excel(r'sample_2.xls')
  # json_str = excel_data_df.to_json()
  # # json_str=dict(json_str)
  # # json_str=pandas.DataFrame.to_json('sample_2.xls')
  # # excel2json.convert_from_file('sample_2.xls')
  # import collections as co
  # import pandas as pd
  # df = pd.read_excel('sample_2.xls')
  # df = df.where(pd.notnull(df), None)
  # od = co.OrderedDict((k.strip().encode('utf8'), v.strip().encode('utf8'))
  #                     for (k, v) in df.values)
  # import pandas as pd
  # df = pd.read_excel("sample_2.xls", index_col=0)
  # df = df.where(pd.notnull(df), None)
  # # for key in headers:
  final_json = {"headers": headerList, "data": final_list_1}
  return jsonify(final_json)

@app.route("/upload_spec")
def upload_spec():
  replace_spec = ''
  head_list = []
  cell_list = []
  # folderName = request.args.get("folderName")
  spec_file = open(base_location+"\\" + outerFolderName + '\\' + folderName + "\\oldFiles\\spec.js")
  path = base_location+"\\" + outerFolderName + '\\' + folderName + "\\oldFiles\\sample1.xlsx"
  new_spec_file = open(
    base_location+"\\" + outerFolderName + '\\' + folderName + "\\oldFiles\\multiple_spec_file.js", "w+")
  # To open the workbook
  # workbook object is created
  wb = xlrd.open_workbook(path)
  sheet = wb.sheet_by_index(0)
  # start_row=request.args.get("start_row")
  # end_row=request.args.get("end_row")
  start_row = 1
  end_row = sheet.nrows

  # Get workbook active sheet object
  # from the active attribute

  num_of_columns = sheet.ncols

  # Cell objects also have row, column,
  # and coordinate attributes that provide
  # location information for the cell.

  # Note: The first row or
  # column integer is 1, not 0.

  # Cell object is created by using
  # sheet object's cell() method.
  open_file = spec_file.read()
  split_file = open_file.split("extract")[0]
  if (split_file.endswith("//")):
    split_file = split_file[:-2] + "});"
  sec_test = "it('Test 1', function() {" + \
             "browser.driver.manage().window().maximize();" + \
             split_file.split("browser.driver.manage().window().maximize();")[1]

  for r in range(start_row, end_row):
    head_list = []
    cell_list = []
    for c in range(num_of_columns):
      cell_obj = sheet.cell_value(r, c)
      head_obj = sheet.cell_value(0, c)
      head_list.append(head_obj)
      cell_list.append(cell_obj)
    for line in sec_test.split(";"):
      for idx, header in enumerate(head_list):
        if (line.__contains__(header) and not line.__contains__("browser.wait(until") and not line.__contains__(".clear(")):
          if (type(cell_list[idx]) == int or type(cell_list[idx]) == float):
            line = line.split("sendKeys")[0] + "sendKeys(" + str(int(cell_list[idx])) + ");"
          else:
            line = line.split("sendKeys")[0] + "sendKeys('" + str(cell_list[idx]) + "');"
          break
      replace_spec += line + ";"
  new_spec = '''const jsdom = require("jsdom");
            const fs = require('fs');
            const path = require('path');
            describe('Protractor Demo App', function() { ''' + str(replace_spec) + "});"
  new_spec_file.write(new_spec)
  time.sleep(1)
  return "Success"

@app.route("/trigger_multiple_spec")
def trigger_multiple_spec():
  # folderName = request.args.get("folderName")
  if os.path.isfile(
    base_location+'\\' + outerFolderName + '\\' + folderName + '\Batchfiles\multiple_batch.bat') == False:
    f = open(base_location+'\\' + outerFolderName + '\\' + folderName + '\Batchfiles\multiple_batch.bat',
             "w+")
    script = '''
    START /B cmd "cd '''+base_location+'''\" && protractor '''+base_location+'''\\''' + '\\' + outerFolderName + '\\' + folderName + '''\oldFiles\multiple_conf.js > '''+ base_location+'''\\oldFiles\\multiple_conf.log
    '''
    f.write(script)
    f.close()
  subprocess.call(
    base_location+'\\' + outerFolderName + '\\' + folderName + '\Batchfiles\multiple_batch.bat')
  return jsonify("Success")

@app.route("/uploadExcel", methods=['GET', 'POST'])
def uploadExcel():
  # folderName = request.args.get("folderName")
  import pandas as pd
  file = request.get_data()
  str_file = file.decode("utf-8")
  # temp_file = list(eval(str_file))
  temp_List = []
  if type(list(eval(str_file))[0]) == str:
    return jsonify("failure")
  else:
    temp_file = list(eval(str_file))
    df = pd.DataFrame.from_dict(temp_file)
  df.to_excel(base_location+'\\' + outerFolderName + '\\' + '\sample1.xlsx', index=False)
  time.sleep(1)
  upload_spec()
  return jsonify("success")

counter = 0

def select_value(line):  # Split the element based on id or name and returns it. if both are not present returns NA
  modified_line = ''
  if ((line.find(" id=") >= 0 and line.find("name=") == -1) or (line.find(" name=") >= 0 and line.find("id=") >= 0)):
    modified_line = str(line).split(" id=")[1].split("\"")[1]
  elif line.find(" name=") == -1 and line.find(" id=") == -1:
    modified_line = "NA"
  elif (line.find(" name=") >= 0):
    modified_line = str(line).split("name=")[1].split("\"")[1]
  return modified_line

def input_func(json_line, x1, xVar):  # returns the line by adding placeholder to it
  if (str(json_line).find(" type=") >= 0):
    if (str(json_line).split(" type=")[1].split("\"")[1] == "text"
      or str(json_line).split(" type=")[1].split("\"")[1] == "email"
      or str(json_line).split(" type=")[1].split("\"")[1] == "submit"
      or str(json_line).split(" type=")[1].split("\"")[1] == "number"
      or str(json_line).split(" type=")[1].split("\"")[1] == "button"
      or str(json_line).split(" type=")[1].split("\"")[1] == "password"
      or str(json_line).split(" type=")[1].split("\"")[1] == "checkbox"
      or str(json_line).split(" type=")[1].split("\"")[1] == "month"
      or str(json_line).split(" type=")[1].split("\"")[1] == "search"
      or str(json_line).split(" type=")[1].split("\"")[1] == "date"
      or str(json_line).split(" type=")[1].split("\"")[1] == "range"
      or str(json_line).split(" type=")[1].split("\"")[1] == "tel"
      or str(json_line).split(" type=")[1].split("\"")[1] == "reset"
      or str(json_line).split(" type=")[1].split("\"")[1] == "week"
      or str(json_line).split(" type=")[1].split("\"")[1] == "time"):
      x2 = json_line[:str(json_line).index(">") - 1] + json_line[
        str(json_line).index(">") - 1] + " " + "placeholder=\'" + x1 + "\'>" + xVar
      if (str(json_line).split("type=")[1].split("\"")[1]) == "text":
        if (x2.find(" value=") >= 0):
          valSplit = x2.split(" value=")
          # print(valSplit)
          x2 = valSplit[0] + " value='' " + valSplit[1].split('"')[2] + ">"
          # print("x2",x2)
      if (str(json_line).split("type=")[1].split("\"")[1]) == "checkbox":
        # print(json_line)
        if (x2 != ""):
          x2 = x2.replace("custom-control-input", "")
      return x2

    elif str(json_line).split(" type=")[1].split("\"")[1] == "radio":
      json_line = str(json_line).replace("checked=\"\"", "")
      if (str(json_line).find("value=") >= 0):
        x1 = str(json_line).split("value=")[1].split("\"")[1]
      elif (str(json_line).find("id=") >= 0 and str(json_line).find("value=") == -1):
        x1 = str(json_line).split("id=")[1].split("\"")[1]
      x2 = json_line + "<span>" + x1 + "</span>" + xVar
      return x2
  elif (str(json_line).find(" type=") == -1):
    x2 = json_line[:str(json_line).index(">") - 1] + json_line[
      str(json_line).index(">") - 1] + " " + "placeholder=\'" + x1 + "\'>" + xVar
    return x2

def link_list(anc_list):  # if href is present it will remove the href and return it
  append_href = ''
  for select_for in anc_list.split():
    if (select_for.find("href=") >= 0):
      # href_split = select_for.split(" href=")[1].split('"')[2]  # Remove href and append
      # append_href += " " + href_split
      print("select for 1", select_for)
      if select_for[select_for.find("href="):].__contains__("='"):
        href_split = select_for.strip().split("href=")[1]  # Remove href and append
        if href_split.__contains__("'"):
          href_split = select_for.strip().split("href=")[1].split("'")[2]  # Remove href and append
      else:
        href_split = select_for.strip().split("href=")[1]  # Remove href and append
        print("inside one", href_split)
        if href_split.count('"') > 1:
          href_split = select_for.strip().split("href=")[1].split('"')[2]
      print("inside href")
      append_href += " " + href_split
    elif (select_for.find(" ng-reflect-router-link=") >= 0):
      href_split = select_for.split(" ng-reflect-router-link=")[1].split('"')[2]  # Remove href and append
      append_href += " " + href_split
    elif (select_for.find(" routerlink=") >= 0):
      href_split = select_for.split(" routerlink=")[1].split('"')[2]  # Remove href and append
      append_href += " " + href_split
    elif (select_for.find(" data-toggle=") >= 0):
      print("inside data toggle", select_for)
      href_split = select_for.split(" data-toggle=")[1].split('"')[2]  # Remove href and append
      append_href += " " + href_split
    elif (select_for.find(" data-target=") >= 0):
      href_split = select_for.split(" data-target=")[1].split('"')[2]  # Remove href and append
      append_href += " " + href_split
    elif (select_for.find("modal-backdrop") >= 0):
      href_split = select_for.replace("modal-backdrop","")  # Remove href and append
      append_href += " " + href_split
    else:
      append_href += " " + select_for
  # print("append href",append_href)
  return append_href

def add_order(line):
  if line.find("<option") >= 0:
    # print("count select", line.count("<select"))
    select_count = line.count("<select")
    opt = -1
    sel_opt = -1
    res = ''
    for scount in range(select_count):
      sel_line = line.find("<select", sel_opt + 1)
      sel_opt = sel_line
      # print("select option",sel_opt)
      opt = line.find("<option", sel_line)
      # print("option", opt)
      add_string = "<option value=''></option>"
      line = line[: opt] + add_string + line[opt:]
      # print("re", line)
    # print("response", line)
    return line
  else:
    return line

def replace_func(x):
  # print(x)
  u = x.replace("(", '').replace(")", '')
  y = u.split(",")
  return int(y[1])

def replace_func1(x):
  u = x.replace("(", '').replace(")", '')
  y = u.split(",")
  z = [int(y[0]), int(y[1])]
  return z

def replace_func2(x):
  x = str(x)
  u = x.replace("[", '(').replace("]", ')').replace(" ", "")
  return u

def test1():    # remove duplicates in json
  # checks the duplicates element in the list of dicts 



  # f = json.load(open(r"base_location\\" + outerFolderName + '\\' + folderName + "\inputfile.json"))
  # input_file = open(r"base_location\crst_test\create_order\inputfile.json")
  # output_file = open(r"base_location\crst_test\create_order\inputfile1.json", "w+")
  # json_val = json.load(input_file)
  # # print(json_val)
  # # # json_file=json.load(input_file, output_file, indent=4)
  # # print(json.dumps(json_val, indent=4))
  # json.dump(json_val, output_file, indent=4)
  # output_file.close()
  # f = json.dump(open(r"base_location\\" + outerFolderName + '\\' + folderName + "\inputfile1.json"))
  # # print(f)
  # for key in f[counter].keys():
  #   print(key)
  f = json.load(open(base_location+"\\" + outerFolderName + '\\' + folderName + "\inputfile1.json"))
  # f1=json.load(open(base_location+"\\" + outerFolderName + '\\' + folderName + "\inputfile1.json"))
  ff1 = open(base_location+"\\" + outerFolderName + '\\' + folderName + "\ipJson.json", "w+")
  print("counter", counter)
  list_items = copy.deepcopy(list(f[counter].items()))
  for first_key, first_value in f[counter].items():
    for second_key, second_value in f[counter].items():
      # if( first_key!= second_key and first_key!="url" and second_key!="url" and first_key!="order" and second_key!="order" and second_key!="button" ):
      #   print("1", first_key)
      #   print("2", second_key)
      #   # print("val 1", f[counter][first_key])
      #   # print("val 2", f[counter][second_key])
      #   for value_list_1 in f[counter][first_key]:
      #     for value_list_2 in f[counter][second_key]:
      #       # print("val1", value_list_1)
      #       # print("avl2", value_list_2)
      #       if value_list_1.__contains__(value_list_2):
      #         ind=f[counter][second_key].index(value_list_2)
      #         f[counter][second_key].remove(value_list_2)
      #         f[counter]["order"][0][second_key].remove(f[counter]["order"][0][second_key][ind])
      # print("keyval",second_key,second_value)
      if (
        first_key != second_key and first_key != "url" and
        second_key != "url" and first_key != "order" and
        second_key != "order" ):
        already_exist = []
        # print("1", first_key)
        # print("2", second_key)
        # print("val 1", f[counter][first_key])
        # print("val 2", f[counter][second_key])
        copy_list_one = copy.deepcopy(f[counter][first_key])
        copy_list_two = copy.deepcopy(f[counter][second_key])
        for value_list_1 in copy_list_one:
          for value_list_2 in copy_list_two:
            # print("val1", value_list_1)
            # print("avl2", value_list_2)

            if value_list_1.__contains__(value_list_2) and value_list_2 not in already_exist:
              already_exist.append(value_list_2)
              ind = f[counter][second_key].index(value_list_2)
              f[counter][second_key][ind] = ''
              f[counter]["order"][0][second_key][ind] = ''
              # if value_list_1.strip().lower().__contains__("register") or value_list_2.strip().lower().__contains__("register"):
              # print("pop logic", ind,first_key, second_key, first_value, second_value)
              # print("end")
              # f[counter][second_key].remove(value_list_2)
              # f[counter]["order"][0][second_key].remove(f[counter]["order"][0][second_key][ind])
            # if value_list_2.__contains__(value_list_1):
            #   ind = f[counter][first_key].index(value_list_1)
            #   f[counter][first_key].remove(value_list_1)
            #   f[counter]["order"][0][first_key].remove(f[counter]["order"][0][first_key][ind])
            #   break

        # if f[counter][first_key]

  # json.dump(f, ff1)
  # ff1.writelines(f)
  json.dump(f, ff1, indent=4)
  return "Success"

def findDuplicateIndex(lst, item):
  return [i for i, x in enumerate(lst) if x == item]

def identifyHideElement(append_href, tagname):
  if append_href.__contains__(" id="):
    id_splittedvalue = append_href.split("id=\"",1)
    hideaddedvalue = id_splittedvalue[0] + "id=\"hide$_" + id_splittedvalue[1]
    return hideaddedvalue
  else:
      splittedvalue = append_href.replace(">", " id=\"hide$\">",1)
      return splittedvalue

@app.route("/TestService")
def TestService():  #setting position of elements 
  #sorting 2d arrays and according to that sorted arrays setting the elements and if conditions for various types of elements like text area etc
 
  test1()
  if os.path.isdir(base_location+"\\" + outerFolderName + '\\' + folderName) == False:
    os.mkdir(base_location+"\\" + outerFolderName + '\\' + folderName)
  data_out = json.load(open(base_location+"\\" + outerFolderName + '\\' + folderName + "\\ipJson.json"))
  global counter
  # counter = 2
  data = data_out[counter]
  # print("full data", data)
  order = data["order"][0]
  new_list = []
  sep_list = []
  dictionary = {}
  # unhandled_elements=request.args.get("unhandled_items") #All handled keys + user selected keys
  # unhandled_list=unhandled_elements.split(",")
  
  dup_json = {}
  for k, v in order.items():
    if '(0,0)' in v:
      dup_json[k] = findDuplicateIndex(v, '(0,0)')
      
  #dup_json = {k:findDuplicateIndex(v, '(0,0)') for k, v in order.items() if '(0,0)' in v}    
      
  xVar = ''
  temp_variable = 0
  for var in order.keys():
    input1 = order[var]  # Add if condition here to check the key is present in required key list
    dictionary[var] = 0
    new_list.extend(input1)
    for val in order[var]:
      if val != '':
        sep_list.append(str(var) + "$" + str(val))
     
  new_list1 = []
  for val in new_list:
    if val != "":
     new_list1.append(replace_func1(val))
  
  yList = sorted(new_list1, key=lambda x: (x[1]))
  temp_list = []
  temp_idx = []
  str_out = ''
  next_index = 0
  for idx, value in enumerate(yList, next_index):
    next_index = 1
    while idx + next_index < len(yList):
      next_value = yList[idx + next_index]
      if value[1] == next_value[1]:
        if idx not in temp_idx:
          temp_list.append(value)
          temp_idx.append(idx)
        if idx + next_index not in temp_idx:
          temp_list.append(next_value)
          temp_idx.append(next_index + idx)
      else:
        break
      next_index += 1
    xTempList = sorted(temp_list, key=lambda x: (x[0]))
    for id1, idx1 in enumerate(temp_idx):
      yList[idx1] = xTempList[id1]
    temp_idx = []
    temp_list = []
  final_list = []
  for val1 in yList:
    final_list.append(replace_func2(val1))
  print('final_list', sep_list)
  fin_list = [0] * (len(final_list))
  inc_val = 0
  increase_position = 1  # to add position if the elements are in same line
  # print("separated list",sep_list)
  print(dup_json, 'dupJSON')
  zero_List = []
  for pos, ar in enumerate(sep_list):# taking the array index and array value
    zeroKey = False
    if ar != "$":
      y_val = replace_func(ar.split("$")[1])
      rem_log = ar.split("$")[1]
      init_pos = 1
      pos_flag = 1


      # print("rtype of rem logic", type(rem_log), rem_log)
      # replace_func(sep_list[pos-1])==y_val
      # =>xFlag=1
      # =>else, xFlag=0
      # if(y_val!=0):
      # if pos+1<len(sep_list) and replace_func(sep_list[pos+1])==y_val :
      #   xVar="&nbsp;&nbsp;&nbsp;&nbsp;"
      # else:
      #   xVar="<br/><br/>"
      if (str(y_val).isnumeric()):  # split it based on the -
        # if str(y_val).strip() == str(0):
        #   try:# If the splitted element is 0, removing it from json as it is hidden element
        #     dictionary[ar.split("$")[0]] -= 1
        #     print("remove logic", data[ar.split("$")[0]][dictionary[ar.split("$")[0]]])
        #     data[ar.split("$")[0]].remove(data[ar.split("$")[0]][dictionary[ar.split("$")[0]]])
        #   except Exception as e:
        #     print("fail", e)
        #     pass
        print("rem log", rem_log, rem_log.strip(), type(rem_log))
        if rem_log != (0, 0) or rem_log == (0, 0) or rem_log == '(0,0)' or rem_log != "(0,0)":
          print("inside rem", rem_log)

          key = ar.split("$")[0]  # fetching the key# input$(251,31), label$(51,23)
          # if key in unhandled_list:
          if ar.split("$")[1] == (0, 0) or ar.split("$")[1] == '(0,0)':
            position = final_list.index(str(ar.split("$")[1]), temp_variable)
            fIndex = final_list.index(str(ar.split("$")[1]), temp_variable)
            print(temp_variable, ', temp_variable', position)
            temp_variable += 1
            # print('0,0 val ->', dup_json[key][0], dup_json[key], key)
            # position = dup_json[key][0]
            # fIndex = dup_json[key][0]
            # dup_json[key].pop(0)
            # print('0,0 val ->', dup_json[key], key)
          else:
            position = final_list.index(str(ar.split("$")[1]))
            fIndex = final_list.index(str(ar.split("$")[1]))
          if (fIndex + 1 < len(final_list) and replace_func(final_list[fIndex + 1]) == replace_func(final_list[fIndex])):
            xVar = "&nbsp;&nbsp;&nbsp;&nbsp;"
          else:
            xVar = "<br/><br/>"
          # if(replace_func(final_list.index(str(ar.split("-")[1]))==replace_func(final_list.index(str(ar.split("-")[1]))+1)
          #<span >Welcome</span> => <span id=hide><Welcome</span>
          #<input id="name"/>=> <input id="hide_name" />

          if (key == "textarea"):
            zeroKey = False
            if rem_log == (0, 0):
              zeroKey = True

              append_href = link_list(data[ar.split("$")[0]][dictionary[ar.split("$")[0]]])
            else:
              append_href = link_list(data[ar.split("$")[0]][dictionary[ar.split("$")[0]]])
            if append_href != "":
              json_line = append_href
              # json_line = json_line.replace("none", "block")
              # json_line = json_line.replace("hidden", "visible")
              x1 = select_value(json_line)  # returns id or name or na
              x2 = json_line[:str(json_line).index(">") - 1] + json_line[
                str(json_line).index(
                  ">") - 1] + " " + "placeholder=\'" + x1 + "\'></textarea>" + xVar  # adding placeholder to the element
              if fin_list[position] != 0:
                while (pos_flag != 0):
                  if fin_list[position + init_pos] == 0:
                    if zeroKey == True:
                      x2 = identifyHideElement(x2, "textarea")
                      x2=add_order(x2)
                      zero_List.append(x2.replace("form-control", ""))
                    else:
                      x2 = add_order(x2)
                      fin_list[position + init_pos] = x2.replace("form-control", "")
                    pos_flag = 0
                    break
                  else:
                    init_pos += 1

                # fin_list.insert(position,x2.replace("form-control", ""))
              else:
                if zeroKey == True:
                  btnstr = identifyHideElement(btnstr, "button")
                  x2 = add_order(btnstr)
                  zero_List.append(x2.replace("form-control", ""))
                else:
                  x2 = add_order(x2)
                  fin_list[position] = x2.replace("form-control", "")
          elif (key == "linklist"):
            zeroKey = False
            try:
              if  ar.split('$')[1] == '(0,0)':
                zeroKey = True
                inp_ind = dup_json['linklist'][0]
                dup_json[key].pop(0)
              else:
                inp_ind = order[ar.split("$")[0]].index(ar.split("$")[1])
              # print("anchor ", inp_ind)
              anc_list = data[ar.split("$")[0]][dictionary[ar.split("$")[0]]]
              anc_list = anc_list.replace("z-index", "")
              append_href = link_list(anc_list)  # removes href to the element
              append_href = append_href.replace("checked", "")
              # append_href = append_href.replace("none", "block")
              # append_href = append_href.replace("hidden", "visible")
              # json_line = json_line.replace(" ", "").replace("display:none", "display:inline")
              if fin_list[position] != 0:
                # fin_list.insert(position, add_order(append_href) + xVar)
                while (pos_flag != 0):
                  if fin_list[position + init_pos] == 0:
                    if zeroKey == True:
                      append_href = identifyHideElement(append_href, "linklist")
                      zero_List.append(add_order(append_href) + xVar)
                    else:
                      fin_list[position + init_pos] = add_order(append_href) + xVar
                    pos_flag = 0
                    break
                  else:
                    init_pos += 1
              else:
                if zeroKey == True:
                  zero_List.append(add_order(append_href) + xVar)
                else:
                  fin_list[position] = add_order(append_href) + xVar
              # fin_list[position] = add_order(append_href) + xVar
            except Exception as e:
              print(" anchor Exception e", e)
              pass
          elif (str(key).strip() == "button"):
            zeroKey = False
            if  ar.split('$')[1] == '(0,0)':
              zeroKey = True
              inp_ind = dup_json['button'][0]
              dup_json[key].pop(0)
            else:
              inp_ind = order[ar.split("$")[0]].index(ar.split("$")[1])
            append_href = link_list(str(data[ar.split("$")[0]][inp_ind]).replace('\n', '').replace('\t',
                                                                                                   ''))
            append_href = append_href.replace("checked", "")
            # append_href = append_href.replace("hidden", "visible")
            # append_href = append_href.replace("none", "block")
            btnstr = append_href
            if fin_list[position] != 0:
              # fin_list.insert(position, btnstr + xVar)
              while (pos_flag != 0):
                if fin_list[position + init_pos] == 0:
                  if zeroKey == True:
                    btnstr=add_order(btnstr)
                    btnstr = identifyHideElement(btnstr, "button")
                    zero_List.append(btnstr + xVar)
                  else:
                    btnstr = add_order(btnstr)
                    fin_list[position + init_pos] = btnstr + xVar
                  pos_flag = 0
                  break
                else:
                  init_pos += 1
            else:
              if zeroKey == True:
                btnstr = add_order(btnstr)
                zero_List.append(btnstr + xVar)
              else:
                btnstr = add_order(btnstr)
                fin_list[position] = btnstr + xVar
            # fin_list[position] = btnstr + xVar
          elif (key == "span"):
            zeroKey = False
            try:
              # inp_ind = order[ar.split("$")[0]].index(ar.split("$")[1])
              # inp_ind = order[ar.split("$")[0]][dictionary[ar.split("$")[0]]]
              # append_href = link_list(data[ar.split("$")[0]][inp_ind] + xVar)
              if  ar.split("$")[1] == '(0,0)':
                append_href = link_list(data[ar.split("$")[0]][
                                          dictionary[ar.split("$")[0]]] + xVar)
                append_href = append_href.replace("checked", "")
                append_href = identifyHideElement(append_href, "span")
                # append_href = append_href.replace("hidden", "visible")
                # append_href = append_href.replace("none", "block")
                zero_List.append(append_href)
              else:
                append_href = link_list(data[ar.split("$")[0]][
                                          dictionary[ar.split("$")[0]]] + xVar)
                append_href = append_href.replace("checked", "")
                append_href = append_href.replace("hidden", "visible")
                append_href = append_href.replace("none", "block")

                if fin_list[position] != 0:
                  # fin_list.insert(position, add_order(append_href))
                  while (pos_flag != 0):
                    if fin_list[position + init_pos] == 0:
                      fin_list[position + init_pos] = add_order(append_href)
                      pos_flag = 0
                      break
                    else:
                      init_pos += 1
                else:
                  fin_list[position] = add_order(append_href)
              # fin_list[position] = add_order(append_href)
            except Exception as e:
              print("span exception", e)
          elif (key == "select"):
            zeroKey = False
            if  ar.split('$')[1] == '(0,0)':
              zeroKey = True
              inp_ind = dup_json['select'][0]
              dup_json[key].pop(0)
            else:
              inp_ind = order[ar.split("$")[0]].index(ar.split("$")[1])
            append_href = link_list(data[ar.split("$")[0]][inp_ind])
            append_href = append_href.replace("checked", "")
            # append_href = append_href.replace("hidden", "visible")
            # append_href = append_href.replace("none", "block")
            add_option = add_order(append_href)
            json_line = add_option
            return_line = select_value(json_line)  # split and return id or name or na
            span_line = data[ar.split("$")[0]][
                          dictionary[ar.split("$")[0]]] + xVar
            if fin_list[position] != 0:
              # fin_list.insert(position, json_line + xVar)
              while (pos_flag != 0):
                if fin_list[position + init_pos] == 0:
                  if zeroKey == True:
                    json_line = identifyHideElement(json_line, "select")
                    zero_List.append(json_line + xVar)
                  else:
                    fin_list[position + init_pos] = json_line + xVar
                  pos_flag = 0
                  break
                else:
                  init_pos += 1
            else:
              if zeroKey == True:
                json_line = add_order(json_line)
                zero_List.append(json_line + xVar)
              else:
                json_line = add_order(json_line)
                fin_list[position] = json_line + xVar
            # fin_list[position] = json_line + xVar
          elif (key == "input"):
            # inp_ind=order[ar.split("$")[0]].index(ar.split("$")[1])
            zeroKey = False
            if  ar.split('$')[1] == '(0,0)':
              zeroKey = True
              inp_ind = dup_json['input'][0]
              dup_json[key].pop(0)
            else:
              inp_ind = order[ar.split("$")[0]].index(ar.split("$")[1])
            print("requ, index", inp_ind)

            print("input", data[ar.split("$")[0]])
            if rem_log == (0, 0):
              append_href = link_list(data[ar.split("$")[0]][inp_ind + inc_val])
              print("inc position", inc_val)
              inc_val += 1
            else:
              append_href = link_list(data[ar.split("$")[0]][inp_ind])
            # append_href = append_href.replace("hidden", "visible")
            # append_href = append_href.replace("none", "block")
            print("append href", append_href)
            # print("input", append_href, dictionary[ar.split("$")[0]])
            # try:
            json_line = append_href
            # json_line = json_line.replace("hidden", "visible")
            # json_line = json_line.replace("none", "block")
            x1 = select_value(json_line)  # split and return id or name or na
            x2 = input_func(json_line, x1, xVar)  # split the type and return by adding placeholder
            print("x2", position, x2, type(x2))
            if (x2 != "" and x2 != None):
              # try:
              if fin_list[position] != 0:
                # fin_list.insert(position, x2.replace("form-control ", ""))
                while (pos_flag != 0):
                  if fin_list[position + init_pos] == 0:
                    if zeroKey == True:
                      x2 = identifyHideElement(x2, "input")
                      x2 = add_order(x2)
                      zero_List.append(x2.replace("form-control", ""))
                    else:
                      x2 = add_order(x2)
                      fin_list[position + init_pos] = x2.replace("form-control", "")
                    pos_flag = 0
                    break
                  else:
                    init_pos += 1
              else:
                if zeroKey == True:
                  x2 = add_order(x2)
                  zero_List.append(x2.replace("form-control", ""))
                else:
                  x2 = add_order(x2)
                  fin_list[position] = x2.replace("form-control ", "")
                # fin_list[position] = x2.replace("form-control ", "")

                # except Exception as e:
                #   print("1 input type")
                #   pass
            # except Exception as e:
            #   print("input exception")
            #   pass
          else:
            zeroKey = False
            if  ar.split("$")[1] == '(0,0)':
              zeroKey = True
              append_href = link_list(data[ar.split("$")[0]][dictionary[ar.split("$")[0]]])
              json_line = add_order(append_href)
              json_line = identifyHideElement(json_line, ">")

              # json_line = json_line.replace("none", "block")
              # json_line = json_line.replace("hidden", "visible")
              zero_List.append(json_line)
            else:
              append_href = link_list(data[ar.split("$")[0]][dictionary[ar.split("$")[0]]])
              json_line = add_order(append_href)

              # json_line = json_line.replace("none", "block")
              # json_line = json_line.replace("hidden", "visible")

              if json_line.__contains__("checked"):
                json_line = json_line.replace("checked", "")
              # tag=json_line.split()[0]
              # x1 = select_value(json_line)  # split and return id or name or na
              # x2 = json_line[:str(json_line).index(">") - 1] + json_line[
              #   str(json_line).index(
              #     ">") - 1] + " " + "placeholder=\'" + x1 + "\'></"+tag+">" + xVar  # adding placeholder to the element
              x2 = json_line + xVar
              if fin_list[position] != 0:
                # while()
                # fin_list.insert(position, x2.replace("form-control ", ""))
                while (pos_flag != 0):
                  if fin_list[position + init_pos] == 0:
                    fin_list[position + init_pos] = x2.replace("form-control", "")
                    pos_flag = 0
                    break
                  else:
                    init_pos += 1
              else:
                fin_list[position] = x2.replace("form-control ", "")
              # fin_list[position] = x2.replace("form-control", "")

          dictionary[ar.split("$")[0]] += 1  # add the value of key to one
          # inc_val+=1
    else:
      pass
  counter += 1
  print(zero_List, 'zl')
  str_out += '<div id="hidden-elem" style="border-style: dashed;border-color: grey;border-radius:8px;"><br><h4>Hidden Elements:</h4><br>'
  for val in zero_List:
    str_out += val.replace("checked","")
  if len(zero_List) == 0:
    str_out += "<h6>There are no hidden elements...!!!"
  str_out += '</div>'
  for outer_for in fin_list:
    if (outer_for != 0 and outer_for != None):
      str_out += outer_for
  print(dup_json, 'dupJSON')
  # print(str_out)
  print("final list", fin_list)
  str_out = str_out.replace("position", "").replace("z-index", "").replace("modal-backdrop","")
  return jsonify('<br><br>' + str_out)

@app.route("/loadTestCase")#
def loadTestCase():
  # folderName = request.args.get("folderName")
  logger = logging.getLogger('myapp')  # Initializing logger
  hdlr = logging.FileHandler(
    base_location+'\\' + outerFolderName + '\\' + folderName + '\oldFiles\load_test.log')  # Logger file path
  formatter = logging.Formatter('%(asctime)s %(levelname)s %(message)s')  # Adding time stamp and formatting the logger
  hdlr.setFormatter(formatter)  # Set the formatter to the file
  logger.addHandler(hdlr)
  logger.setLevel(logging.INFO)  # Setting the level of logger

  elements = request.args.get("elements").replace("HASH", "#")  # Elements from UI
  values = request.args.get("values").replace("HASH", "#")  # Values from UI
  # elements='''<input type="text" id="sameText" placeholder="sameText">,
  # <input type="text" id="adit" placeholder="adit">,
  # <input type="text" id="sameText" placeholder="sameText">,
  # <input type="text" id="sameText" placeholder="sameText"> '''
  # values="sameText$0,adit,sameText$1,sameText$2"
  elements = unquote(elements)
  values = unquote(values)

  # elements = "<label id=\"id1\"> Hi Hello!<label>,<label name=\"name1\"> Hello!<label>,<div class=\"class1\">name</div>"
  # values = "click,click,Yourname"
  elements1 = ''
  # url="None"
  url = request.args.get("url").replace("HASH", "#")  # Url from UI
  url = unquote(url)

  # url = "nothing"
  # flow_completed="no"
  flow_completed = request.args.get("flag")
  elements = elements.replace(",click,", ",")

  s_list = elements.split(",")
  t_list = elements.split(",")
  # modal_click=request.args.get("modalFlag")
  modal_click = "no"
  rem_flag = 0
  # for s1 in s_list:
  #   rem_flag = 0
  #   for s2 in s_list:
  #     if (s1 != s2):
  #       if (str(s1).find(str(s2)) >= 0):
  #         t_list.remove(s2)
  for u in t_list:
    if (len(t_list) - 1 == t_list.index(u)):
      if (u.endswith(">")):
        u = u[:-1]
      elements1 += u
    else:
      elements1 += u + ","
  split_element = str(elements1).split(">,")
  split_values = str(values).split(",")
  for num_idx, val1 in enumerate(split_values):
    if val1.__contains__("$"):
      print("inside", val1)
      if val1.split("$")[1] == "-1":
        val1 = val1[:-3]
        print(val1)
        split_values[num_idx] = val1
  output_file = open(
    base_location+"\\" + outerFolderName + '\\' + folderName + "\\uniqure_output_file.txt", "w+")
  output = []
  unique_name = {}
  unique_id = {}
  glob_dict = {}
  unique_name1 = []
  unique_id1 = []
  unique_span = []
  input_count = 1
  text_area = 1
  inp_pos = 0
  select = 0
  sp_val = 0
  final = ''
  fin_out = ''
  rf = 0
  output_val = ''
  sp_value = ''
  flag = 0
  anc_id1 = []
  anc_name1 = []
  anc_id = {}
  anc_name = {}
  f_span = 0
  span_count = 0
  for val1 in split_element:
    if val1.__contains__("<"):
      elm = val1.strip().split()[0].split("<")[1]
      if elm.__contains__(">"):
        elm = elm.split(">")[0]
      glob_dict[elm] = 0
  for index, val in enumerate(split_element):
    if (val.find(",") >= 0):
      clk = val.split(',')

      if (clk[0] == "click"):
        val1 = clk[1]
        split_element[split_element.index(val)] = val1
        val = val1
    # val = val.strip()
    # elm = val.strip().split()[0].split("<")[1]
    # glob_dict[elm] = 1
    #<input type="text" value="male"/>
    if (str(val).find(" id=") == -1 and str(val).find(" name=") == -1
      and (str(val).find("<a") == -1 or str(val).find("</a") == -1) and str(val).find("<span") == -1
      and str(val).find("<button") == -1 and str(val).find("<select") == -1 and str(val).find("<input") >= 0):
      if modal_click == "yes":

        # inp_val_id = split_values[index]
        output.append(
          "element.all(by.xpath(\"// div[contains( @ style, 'display: block')]//input\")).get(" + str(
            input_count - 1) + ").sendKeys('" + split_values[
            split_element.index(val)] + "');")

      else:
        print("inside")
        output.append(
          "element.all(by.xpath(\"//input\")).get(" + str(input_count - 1) + ").sendKeys('" + split_values[
            split_element.index(val)] + "');")
      input_count += 1
      #<button id='button name='nbitton>button</button>
    elif (str(val).strip().find("<img") >= 0 and str(val).strip().find("<button") < 0 and str(val).strip().find(
            "<input") < 0 and
          str(val).strip().find("<span") < 0 and str(val).strip().find("<a") < 0):
      if str(val).__contains__("src="):
        print("inside image")
        src_val = str(val).split("src=")[1].split('"')[1]
        output.append(
          "element(by.xpath(\"//img[@src=\'" + src_val + "\']\")).click();")
    elif (str(val).strip().find("<button") >= 0 and str(val).strip().find("</button") >= 0):
      print("inside button")
      mod_line = select_value(val)
      but_val_id = split_values[index]
      if mod_line.split(">")[0].__contains__(" id=") == False and mod_line.split(">")[0].__contains__(
        " name=") == False:
        mod_line = "NA"
      if (val.split(">")[0].__contains__(" id=") and val.split(">")[0].__contains__("name=") == False) or (
        val.split(">")[0].__contains__(" name=") and val.split(">")[0].__contains__("id=")):
        if modal_click == "yes":
          mod_line = val.split(">")[0].split("id=")[1].split('"')[1].split('"')[0]
          if but_val_id.__contains__("$"):
            output.append(
              "element.all(by.xpath(\"// div[contains( @ style, 'display: block')]//button[@id=\'" + mod_line + "\']\")).get(" + str(
                but_val_id.split("$")[1]) + ").click();"
            )
          else:
            output.append(
              "element(by.xpath(\"// div[contains( @ style, 'display: block')]//button[@id=\'" + mod_line + "\']\")).click();"
            )
        else:
          mod_line = val.split(">")[0].split("id=")[1].split('"')[1].split('"')[0]
          if but_val_id.__contains__("$"):
            output.append(
              "element.all(by.xpath(\"//button[@id=\'" + mod_line + "\']\")).get(" + but_val_id.split("$")[
                1] + ").click();"
            )
          else:
            output.append(
              "element(by.xpath(\"//button[@id=\'" + mod_line + "\']\")).click();"
            )
      elif (val.split(">")[0].__contains__(" name=") and val.split(">")[0].__contains__("id") == False):
        if modal_click == "yes":
          mod_line = val.split(">")[0].split("name=")[1].split('"')[1].split('"')[0]
          if but_val_id.__contains__("$"):
            output.append(
              "element.all(by.xpath(\"// div[contains( @ style, 'display: block')]//button[@name=\'" + mod_line + "\']\")).get(" +
              but_val_id.split("$")[1] + ").click();"
            )
          else:
            output.append(
              "element(by.xpath(\"// div[contains( @ style, 'display: block')]//button[@name=\'" + mod_line + "\']\")).click();"
            )
        else:
          mod_line = val.split(">")[0].split("name=")[1].split('"')[1].split('"')[0]
          if but_val_id.__contains__("$"):
            output.append(
              "element.all(by.xpath(\"//button[@name=\'" + mod_line + "\']\")).get(" + str(
                but_val_id.split("$")[1]) + ").click();"
            )
          else:
            output.append(
              "element(by.xpath(\"//button[@name=\'" + mod_line + "\']\")).click();"
            )


      elif mod_line == "NA":
        file1 = open(base_location+"\\" + outerFolderName + '\\' + folderName + "\\inputfile.json")
        data_out = json.load(file1)
        find_val = str(val).strip().strip('"')
        if (find_val.endswith(">") == False):
          find_val = str(val).strip().strip('"') + ">"
        ind = 0
        for index1, butVal in enumerate(data_out[counter - 1]["button"]):
          butVal = butVal.replace("\n", "").replace("\t", "")
          butVal = link_list(butVal).replace(" ", "")
          find_val=find_val.replace("id=hide$","").replace("id=hide$","")
          if butVal == find_val.replace(" ", "").replace("wasClicked", ""):
            ind = index1
            break
        # ind = data_out[counter - 1]["button"].index(find_val)
        if data_out[counter - 1]["url"][0].__contains__("modal"):
          output.append(
            "element.all(by.xpath(\"//div[contains(@style, 'display: block')]//button\")).get(" + str(
              ind) + ").click(); browser.sleep(2000);"
          )
        else:
          output.append(
            "element.all(by.xpath(\"//button\")).get(" + str(ind) + ").click(); browser.sleep(2000);")
    elif (str(val).strip().find("<a") >= 0 and str(val).strip().find("</a") >= 0):
      anc_val_id = split_values[index]
      if (val.find(" id=") >= 0 and val.find(" name") == -1):
        x = val.split(" id=")[1].strip().split()[0].strip('"')
      # if x not in anc_id1:
        anc_id1.append(x)
        anc_id[x] = 2
        if modal_click == "yes":
          if anc_val_id.__contains__("$"):
            output.append(
              "element(by.xpath(\"//div[contains(@style, 'display: block')]//a[@id=\'" + x + "\']\")).get(" + str(
                anc_val_id.split("$")[1]) + ").click();")
          else:
            output.append(
              "element(by.xpath(\"//div[contains(@style, 'display: block')]//a[@id=\'" + x + "\']\")).click();")
        else:
          if anc_val_id.__contains__("$"):
            output.append(
              "element(by.xpath(\"//a[@id=\'" + x + "\']\")).get(" + str(anc_val_id.split("$")[1]) + ").click();")
          else:
            output.append(
              "element(by.xpath(\"//a[@id=\'" + x + "\']\")).click();")

      elif (val.find(" id=") == -1 and (val.find(" name=") >= 0 or (val.find(" id=") >= 0))):
        x = val.split(" name=")[1].strip().split()[0].strip('"')
      # if x not in anc_name1:
        anc_name1.append(x)
        anc_name[x] = 2
        if modal_click == "yes":
          if anc_val_id.__contains__("$"):
            output.append(
              "element.all(by.xpath(\"//div[contains(@style, 'display: block')]//input[@name=\'" + x + "\']\")).get(" + str(
                anc_val_id.split("$")[1]) + ").click();")
          else:
            output.append(
              "element(by.xpath(\"//div[contains(@style, 'display: block')]//input[@name=\'" + x + "\']\")).click();")
        else:
          output.append(
            "element(by.xpath(\"//input[@name=\'" + x + "\']\")).click();")

      else:
        x = ''
        for i in val.split('"'):
          if i.__contains__(">") and i.__contains__("<"):
            c = i.count(">")
            for j in range(1, c + 1):
              k = i.split(">", j)[j].split("<", j)[0]
              if k.strip() != '' and k != ' ':
                x = k
                break
        # if flag == 1:
        if x != '':
          if anc_val_id.__contains__("$"):
            output.append("element(by.partialLinkText(\'" + str(x).strip() + "\')).get(" + str(
              anc_val_id.split("$")[1]) + ").click();")
          else:
            output.append("element(by.partialLinkText(\'" + str(x).strip() + "\')).click();")
        else:
          js_file = json.load(
            open(base_location+"\\" + outerFolderName + '\\' + folderName + "\inputfile.json"))
          anc_button = js_file[counter]["linklist"]
          # print("list of linklinst", anc_button)
          pos = anc_button.index(val)
          print("position", pos)

        # output.append("element(by.partialLinkText(\'" + str(x).strip() + "\')).click();")
        # else:
        #   output.append("element(by.linkText(\'" + str(x1) + "\')).click();")
    elif ((str(val).find(" id=") >= 0 and str(val).find(" name=") == -1) or (
      str(val).find(" id=") >= 0 and str(val).find(" name=") >= 0)
          and (str(val).find("<a") == -1 or str(val).find("</a>") == -1) and str(val).find("<span") == -1 and str(
        val).strip().find("<button") == -1):
      if (str(val).find("<input") >= 0):
        flag = 1
        inp_pos += 1
        if (len(val.split("id=")) >= 3):
          x = val.split("id=")[1].strip().split()[0].strip('"')
        # if x not in unique_id1:
          unique_id1.append(x)
          unique_id[x] = 2
          inp_val = split_values[index]
          if inp_val.__contains__("$"):
            output.append("element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + str(
                inp_val.split("$")[1]) + ").clear();")
            output.append(
              "element.all(by.xpath(\"//input[@id=\'" + x + "\']\")).get(" + str(
                inp_val.split("$")[1]) + ").sendKeys('" + inp_val.split("$")[0] + "');")
          else:
            output.append("element(by.xpath(\"//input[@id='" + x + "']\")).clear();")
            output.append(
              "element(by.xpath(\"//input[@id=\'" + x + "\']\")).sendKeys('" + split_values[
                split_element.index(val)] + "');")

        else:
          x = str(val.split("id=")[1]).split('"')[1].strip()
          if (val.find("type") == -1):
            xy = "text"
          else:
            xy = str(val.split(" type=")[1]).split('"')[1]
          y = str(val.split(" id=")[0]).split('"')
          if (xy == "radio"):
            # rad = split_values[split_element.index(val)]
            rad = split_values[index]
            print(str(val.split("id=\"")[1].split('"')[0]))
            if (val.find("value=") >= 0 and val.find("id=") < 0):
              if rad.__contains__("$"):
                rad_ind = rad.split("$")[1]
                output.append('element.all(by.xpath(\'//input[@value="' +
                              str(rad).strip() + '"]\')).get(' + rad_ind + ').click();')
              else:
                output.append('element(by.xpath(\'//input[@value="' +
                              str(rad).strip() + '"]\')).click();')
            elif (val.find("id=") >= 0 and val.find("value=") == -1):
              if rad.__contains__("$"):
                rad_ind = rad.split("$")[1]
                output.append('element.all(by.xpath(\'//input[@id="' +
                              str(val.split("id=\"")[1].split('"')[0]) + '"]\')).get(' + rad_ind + ').click();')
              else:
                output.append('element(by.xpath(\'//input[@id="' +
                              str(val.split("id=\"")[1].split('"')[0]) + '"]\')).click();')
            elif (val.find("id=") >= 0 and val.find("value=") >= 0):
              if rad.__contains__("$"):
                rad_ind = rad.split("$")[1]
                output.append('element.all(by.xpath(\'//input[@id="' +
                              str(val.split("id=\"")[1].split('"')[0]) + '"]\')).get(' + rad_ind + ').click();')
              else:
                output.append('element(by.xpath(\'//input[@id="' +
                              str(val.split("id=\"")[1].split('"')[0]) + '"]\')).click();')
            # elif ((val.find("id=") >= -1 and val.find("name=") >= 0 ) and val.find("value=") == -1):
            #   output.append('element(by.xpath(\'//input[@name="' +
            #                 str(rad).strip() + '"]\')).click();')
            input_count += 1
          elif (xy == "text"):
            txt = split_values[index]
            print("inside")
            print(val)
          # if x not in unique_id1:
            unique_id1.append(x)
            unique_id[x] = 2
            if txt.__contains__("$"):
              # ind_txt=txt.split("$")[1]
              output.append("element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + str(
                  txt.split("$")[1]) + ").clear();")
              output.append(
                "element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + str(
                  txt.split("$")[1]) + ").sendKeys(\"" + txt.split("$")[0] + "\");")
            else:
              output.append("element(by.xpath(\"//input[@id='" + x + "']\")).clear();")
              output.append(
                "element(by.xpath(\"//input[@id='" + x + "']\")).sendKeys(\"" + split_values[
                  split_element.index(val)] + "\");")

          elif (xy == "time"):
            time = split_values[index]
          # if x not in unique_id1:
            unique_id1.append(x)
            unique_id[x] = 2
            if time.__contains__("$"):
              output.append("element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + str(
                  time.split("$")[1]) + ").clear();")
              output.append(
                "element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + str(
                  time.split("$")[1]) + ").sendKeys(\"" + split_values[
                  index] + "\");")
            else:
              output.append("element(by.xpath(\"//input[@id='" + x + "']\")).clear();")
              output.append(
                "element(by.xpath(\"//input[@id='" + x + "']\")).sendKeys(\"" + split_values[
                  split_element.index(val)] + "\");")

          elif (xy == "tel"):
            tel = split_values[index]
          # if x not in unique_id1:
            unique_id1.append(x)
            unique_id[x] = 2
            if tel.__contains__("$"):
              output.append("element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + str(
                  tel.split("$")[1]) + ").clear();")
              output.append(
                "element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + str(
                  tel.split("$")[1]) + ").sendKeys(\"" + tel.split("$")[0] + "\");")
            else:
              output.append("element(by.xpath(\"//input[@id='" + x + "']\")).clear();")
              output.append(
                "element(by.xpath(\"//input[@id='" + x + "']\")).sendKeys(\"" + split_values[
                  split_element.index(val)] + "\");")

          elif (xy == "week"):
            week = split_values[index]
          # if x not in unique_id1:
            unique_id1.append(x)
            unique_id[x] = 2
            if week.__contains__("$"):
              output.append("element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + str(
                  week.split("$")[1]) + ").clear();")
              output.append(
                "element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + str(
                  week.split("$")[1]) + ").sendKeys(\"" + week.split("$")[0] + "\");")
            else:
              output.append("element(by.xpath(\"//input[@id='" + x + "']\")).clear();")
              output.append(
                "element(by.xpath(\"//input[@id='" + x + "']\")).sendKeys(\"" + split_values[
                  split_element.index(val)] + "\");")

          elif (xy == "range"):
            range_val = split_values[index]
          # if x not in unique_id1:
            unique_id1.append(x)
            unique_id[x] = 2

            if range_val.__contains__("$"):
              output.append("element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + str(
                  range_val.split("$")[1]) + ").clear();")
              output.append(
                "element.all(by.xpath('//input[@id=\"" + x +
                "\"]')).get(" + str(
                  range_val.split("$")[1]) + ").sendKeys(protractor.Key.SHIFT, protractor.Key.HOME);")
              for ran in range(int(range_val.split("$")[0])):
                output.append(
                  "element.all(by.xpath('//input[@id=\"" + x + "\"]')).get(" + range_val.split("$")[
                    1] + ").sendKeys(protractor.Key.UP);")
            else:
              output.append("element(by.xpath(\"//input[@id='" + x + "']\")).clear();")
              output.append(
                "element(by.xpath('//input[@id=\"" + x +
                "\"]')).sendKeys(protractor.Key.SHIFT, protractor.Key.HOME);")
              for ran in range(int(range_val.split("$")[0])):
                output.append(
                  "element(by.xpath('//input[@id=\"" + x + "\"]')).sendKeys(protractor.Key.UP);")


          elif (xy == "search"):
            search_val = split_values[index]
          # if x not in unique_id1:
            unique_id1.append(x)
            unique_id[x] = 2
            if search_val.__contains__("$"):
              output.append("element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + search_val.split("$")[1] + ").clear();")
              output.append(
                "element(by.xpath(\"//input[@id='" + x + "']\")).get(" + search_val.split("$")[1] + ").sendKeys(\"" +
                search_val.split("$")[0] + "\");")
            else:
              output.append("element(by.xpath(\"//input[@id='" + x + "']\")).clear();")
              output.append(
                "element(by.xpath(\"//input[@id='" + x + "']\")).sendKeys(\"" + split_values[
                  split_element.index(val)] + "\");")

          elif (xy == "date"):
            date_val = split_values[index]
          # if x not in unique_id1:
            unique_id1.append(x)
            unique_id[x] = 2
            if date_val.__contains__("$"):
              output.append("element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + date_val.split("$")[
                  1] + ").clear();")
              output.append(
                "element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + date_val.split("$")[
                  1] + ").sendKeys(\"" + date_val.split("$")[0] + "\");")
            else:
              output.append("element(by.xpath(\"//input[@id='" + x + "']\")).clear();")
              output.append(
                "element(by.xpath(\"//input[@id='" + x + "']\")).sendKeys(\"" + split_values[
                  split_element.index(val)] + "\");")

          elif (xy == "month"):
            month_val = split_values[index]
          # if x not in unique_id1:
            unique_id1.append(x)
            unique_id[x] = 2
            month_split = (split_values[index]).split()
            if month_val.__contains__("$"):
              output.append("element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + str(
                  month_val.split("$")[1]) + ").clear();")
              output.append(
                "element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + str(
                  month_val.split("$")[1]) + ").sendKeys(\"" + month_split[0] + "\");"
                                                                                "element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + str(
                  month_val.split("$")[1]) + ").sendKeys(protractor.Key.TAB);"
                                             "element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + str(
                  month_val.split("$")[1]) + ").sendKeys(\"" + month_split[1] + "\");")
            else:
              output.append("element(by.xpath(\"//input[@id='" + x + "']\")).clear();")
              output.append(
                "element(by.xpath(\"//input[@id='" + x + "']\")).sendKeys(\"" + month_split[0] + "\");"
                                                                                                 "element(by.xpath(\"//input[@id='" + x + "']\")).sendKeys(protractor.Key.TAB);"
                                                                                                                                          "element(by.xpath(\"//input[@id='" + x + "']\")).sendKeys(\"" +
                month_split[1] + "\");"
              )

          elif (xy == "email"):
            email_val = split_values[index]
          # if x not in unique_id1:
            unique_id1.append(x)
            unique_id[x] = 2
            if email_val.__contains__("$"):
              output.append(
                "element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + email_val.split("$")[
                  1] + ").clear();")
              output.append(
                "element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + email_val.split("$")[
                  1] + ").sendKeys(\"" + email_val.split("$")[0] + "\");")
            else:
              output.append(
                "element(by.xpath(\"//input[@id='" + x + "']\")).clear();")
              output.append(
                "element(by.xpath(\"//input[@id='" + x + "']\")).sendKeys(\"" + split_values[
                  split_element.index(val)] + "\");")

          elif (xy == "number"):
            number_val = split_values[index]
          # if number_val.__contains__("$"):
            output.append(
              "element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + str(
                number_val.split("$")[1]) + ").sendKeys(" + str(
                number_val.split("$")[0]) + ");")

          elif (xy == "button" or xy == "submit" or xy == "checkbox" or xy == "reset"):
            but_val = split_values[index]
            if but_val.__contains__("$"):
              output.append(
                "element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + but_val.split("$")[1] + ").click();")
            else:
              output.append("element(by.xpath(\"//input[@id='" + x + "']\")).click();")
          elif (xy == "password"):
            pass_val = split_values[index]
            if (str(split_values[split_element.index(val)]).isalnum()):
              if pass_val.__contains__("$"):
                output.append(
                  "element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + str(
                    pass_val.split("$")[1]) + ").sendKeys(\'" + pass_val.split("$")[0] + "\');")
              else:
                output.append(
                  "element(by.xpath(\"//input[@id='" + x + "']\")).sendKeys(\'" + split_values[
                    split_element.index(val)] + "\');")
            elif (isinstance((split_values[split_element.index(val)]), float) or isinstance(
              (split_values[split_element.index(val)]), int)):
              if pass_val.__contains__("$"):
                output.append(
                  "element.all(by.xpath(\"//input[@id='" + x + "']\")).get(" + str(
                    pass_val.split("$")[1]) + ").sendKeys(\'" + str(int(pass_val.split("$")[0])) + "\');")
              else:
                output.append("element(by.xpath(\"//input[@id='" + x + "']\")).sendKeys(" + str(
                  int(split_values[split_element.index(val)])) + ");")
            else:
              if pass_val.__contains__("$"):
                output.append("element(by.xpath(\"//input[@id='" + x + "']\")).get(" + str(
                  pass_val.split("$")[1]) + ").sendKeys(\'" + str(
                  pass_val.split("$")[0]) + "\');")
              else:
                output.append("element(by.xpath(\"//input[@id='" + x + "']\")).sendKeys(\'" + str(
                  (split_values[split_element.index(val)])) + "\');")
          input_count += 1
      elif (str(val).find("<select") >= 0):
        x = str(val.split("id=")[1]).split('"')[1].strip()
        sel_val_id = split_values[index]
        opt_val = str(val).split("<option")
        for mk in opt_val:
          if (mk.find("value=\"" + split_values[sp_val] + "\"") >= 0):
            sp_value = opt_val.index(mk)
        sel_val = opt_val[int(sp_value)].split(">")[1].split("<")[0]
      # if x not in unique_name1:
        unique_id1.append(x)
        unique_id[x] = 2
        if sel_val_id.__contains__("$"):
          output.append(
            "element.all(by.xpath(\"//select[@id=\'" + x + "\']\")).get(" + str(
              sel_val_id.split("$")[1]) + ").sendKeys('" + sel_val + "');")
        else:
          output.append(
            "element(by.xpath(\"//select[@id=\'" + x + "\']\")).sendKeys('" + sel_val + "');")

      elif (str(val).find("<textarea") > 0):
        flag = 1
        x = str(val.split(" id=")[1]).split('"')[1].strip()
        text_val = split_values[index]
        if text_val.__contains__("$"):
          output.append(
            "element.all(by.xpath('//textarea')).get(" + str(
              text_val.split("$")[1]) + ").sendKeys('" + str(
              str(text_val.split("$")[0])).strip() + "');")
        else:
          output.append(
            "element(by.xpath('//textarea[position()=" + str(text_area) + "]')).sendKeys('" + str(
              split_values[split_element.index(val)]).strip() + "');")
        text_area += 1
      elif (str(val).find("<button") > 0):
        x = val.split("id=")[1].strip().split('"')[1].strip('"')
        but_val_un = split_values[index]
      # if x not in unique_id1:
        unique_id1.append(x)
        unique_id[x] = 2
        if but_val_un.__contains__("$"):#click$2
          output.append("element.all(by.xpath(\"//button[@id='" + x + "']\")).get(" + str(
            but_val_un.split("$")[1]) + ").click();")
        else:
          output.append("element(by.xpath(\"//button[@id='" + x + "']\")).click();")

      else:
        click_val = split_values[split_element.index(val)]
        elm = val.strip().split()[0].split("<")[1].split(">")[0]
        x = val.split("id=")[1].strip().split()[0].strip('"')
        if click_val == "click":
          file1 = open(base_location+"\\" + outerFolderName + "\\" + folderName + "\\" + "inputfile.json")
          data_out = json.load(file1)
          find_val = str(val).strip().strip('"') + ">"
          find_val = find_val.replace("wasClicked", "").replace("block", "none").replace(" ", "")
          ind = 0
          # print("element", elm)
          # print(data_out[counter - 1][elm])
          # print("finfval12", find_val)
          for el_ind, el_val in enumerate(data_out[counter - 1][elm]):
            # print("inside el", el_val)
            # print("finfval12", find_val)

            el_fin_val = el_val.replace("block", "none").replace("wasClicked", "").replace(" ", "")
            if el_fin_val == find_val:
              print("inside")
              ind = el_ind + 1
          output.append("element.all(by.xpath(\"//" + elm + "[@id=\'"+str(x)+"\'] \")).click();")

          glob_dict[elm] += 1

        else:
          if click_val.__contains__("$"):

            ind_val_else=click_val.split("$")[1]
            click_else=click_val.split("$")[0]
            x = val.split("id=")[1].strip().split()[0].strip('"')
            if click_else!="click":
              output.append("element.all(by.xpath(\"//" + elm + "[@id=\'"+str(x)+"\']\")).get(" + str(
                ind_val_else) + ").sendKeys(\'" + click_else + "\');")
            elif(click_else=="click"):
              output.append("element.all(by.xpath(\"//" + elm + "[@id=\'"+str(x)+"\']\")).get(" + str(
                ind_val_else) + ").click();")

          # else:
          #   output.append("element.all(by.xpath(\"//" + elm + "\")).get(" + str(
          #     glob_dict[elm]) + ").sendKeys(\'" + glob_dict[elm] + "\');")
          # glob_dict[elm] += 1

          # output.append("element.all(by.xpath(\"//" + elm + "\")).get(" + str(
          #   glob_dict[elm]) + ").sendKeys(\'" + click_val + "\');")
          glob_dict[elm] += 1

    elif (
      str(val).strip().find("<span") >= 0 and str(val).strip().find("</span>") >= 0 and str(val).find(
      "<button") == -1):
      if (str(val).find("class=") >= 0):
        if (str(val).split("class=")[1].find("glyphicon") >= 0):
          file1 = open(base_location+"\\" + outerFolderName + '\\' + folderName +"\\inputfile.json")
          data_out = json.load(file1)
          find_val = str(val).strip().strip('"')
          ind = data_out[counter - 1]["span"].index(find_val)
          span_count += 1
          output.append(
            "element.all(by.xpath(\"//span\")).get(" + str(ind) + ").click();")

      else:
        file1 = open(base_location+"\\" + outerFolderName + '\\' + folderName + "\\inputfile.json")
        data_out = json.load(file1)
        find_val = str(val).strip().strip('"')
        ind = data_out[counter - 1]["span"].index(find_val)
        span_count += 1
        output.append(
          "element.all(by.xpath(\"//span\")).get(" + str(ind) + ").click();")

    elif (str(val).find(" name=") >= 0 and str(val).find(" id=") >= -1
          and (str(val).find("<a") == -1 or str(val).find("</a>") == -1) and str(val).find("<span") == -1 and str(
        val).strip().find("<button") == -1):
      if (str(val).find(" <input") >= 0):
        flag = 1
        inp_pos += 1
        if (len(val.split(" name=")) >= 3):
          x = val.split(" name=")[1].strip().split()[0].strip('"')
        # if x not in unique_name1:
          unique_name1.append(x)
          unique_name[x] = 2
          inp_val = split_values[index]
          if inp_val.__contains__("$"):
            output.append("element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + str(
                inp_val.split("$")[1]) + ").clear();")
            output.append(
              "element.all(by.xpath(\"//input[@name=\'" + x + "\']\")).get(" + str(
                inp_val.split("$")[1]) + ").sendKeys('" + inp_val.split("$")[0] + "');")
          else:
            output.append("element(by.xpath(\"//input[@name='" + x + "']\")).clear();")
            output.append(
              "element(by.xpath(\"//input[@name=\'" + x + "\']\")).sendKeys('" + split_values[
                split_element.index(val)] + "');")

        else:
          x = str(val.split(" name=")[1]).split('"')[1].strip()
          if (val.find("type") == -1):
            xy = "text"
          else:
            xy = str(val.split(" type=")[1]).split('"')[1]
            y = str(val.split(" name=")[0]).split('"')
          if (xy == "radio"):
            # rad = split_values[split_element.index(val)]
            rad = split_values[index]
            if (val.find("value=") >= 0):
              if rad.__contains__("$"):
                rad_ind = rad.split("$")[1]
                output.append('element.all(by.xpath(\'//input[@value="' +
                              str(rad).strip() + '"]\')).get(' + rad_ind + ').click();')
              else:
                output.append('element(by.xpath(\'//input[@value="' +
                              str(rad).strip() + '"]\')).click();')
            elif (val.find("name=") >= 0 and val.find("value=") == -1):
              if rad.__contains__("$"):
                rad_ind = rad.split("$")[1]
                output.append('element.all(by.xpath(\'//input[@id="' +
                              str(val.split("id=\"")[1].split('"')[0]) + '"]\')).get(' + rad_ind + ').click();')
              else:
                output.append('element(by.xpath(\'//input[@name="' +
                              str(rad).strip() + '"]\')).click();')
            input_count += 1
          elif (xy == "text"):
            txt = split_values[index]
          # if x not in unique_name1:
            unique_name1.append(x)
            unique_name[x] = 2
            if txt.__contains__("$"):
              # ind_txt=txt.split("$")[1]
              output.append("element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + str(
                  txt.split("$")[1]) + ").clear();")
              output.append(
                "element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + str(
                  txt.split("$")[1]) + ").sendKeys(\"" + txt.split("$")[0] + "\");")
            else:
              output.append("element(by.xpath(\"//input[@name='" + x + "']\")).clear();")
              output.append(
                "element(by.xpath(\"//input[@name='" + x + "']\")).sendKeys(\"" + split_values[
                  split_element.index(val)] + "\");")

          elif (xy == "time"):
            time = split_values[index]
          # if x not in unique_id1:
            unique_id1.append(x)
            unique_id[x] = 2
            if time.__contains__("$"):
              output.append("element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + str(
                  time.split("$")[1]) + ").clear();")
              output.append(
                "element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + str(
                  time.split("$")[1]) + ").sendKeys(\"" + split_values[
                  index] + "\");")
            else:
              output.append("element(by.xpath(\"//input[@id='" + x + "']\")).clear();")
              output.append(
                "element(by.xpath(\"//input[@id='" + x + "']\")).sendKeys(\"" + split_values[
                  split_element.index(val)] + "\");")

          elif (xy == "tel"):
            tel = split_values[index]
          # if x not in unique_name1:
            unique_name1.append(x)
            unique_name[x] = 2
            if tel.__contains__("$"):
              output.append("element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + str(
                  tel.split("$")[1]) + ").clear();")
              output.append(
                "element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + str(
                  tel.split("$")[1]) + ").sendKeys(\"" + tel.split("$")[0] + "\");")
            else:
              output.append("element(by.xpath(\"//input[@name='" + x + "']\")).clear();")
              output.append(
                "element(by.xpath(\"//input[@name='" + x + "']\")).sendKeys(\"" + split_values[
                  split_element.index(val)] + "\");")

          elif (xy == "week"):
            week = split_values[index]
          # if x not in unique_name1:
            unique_name1.append(x)
            unique_name[x] = 2
            if week.__contains__("$"):
              output.append("element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + str(
                  week.split("$")[1]) + ").clear();")
              output.append(
                "element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + str(
                  week.split("$")[1]) + ").sendKeys(\"" + week.split("$")[0] + "\");")
            else:
              output.append("element(by.xpath(\"//input[@name='" + x + "']\")).clear();")
              output.append(
                "element(by.xpath(\"//input[@name='" + x + "']\")).sendKeys(\"" + split_values[
                  split_element.index(val)] + "\");")

          elif (xy == "range"):
            range_val = split_values[index]
          # if x not in unique_name1:
            unique_name1.append(x)
            unique_name[x] = 2
            if range_val.__contains__("$"):
              output.append("element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + str(
                  range_val.split("$")[1]) + ").clear();")
              output.append(
                "element.all(by.xpath('//input[@name=\"" + x +
                "\"]')).get(" + str(
                  range_val.split("$")[1]) + ").sendKeys(protractor.Key.SHIFT, protractor.Key.HOME);")
              for ran in range(int(range_val.split("$")[0])):
                output.append(
                  "element.all(by.xpath('//input[@name=\"" + x + "\"]')).get(" + range_val.split("$")[
                    1] + ").sendKeys(protractor.Key.UP);")
            else:
              output.append("element(by.xpath(\"//input[@name='" + x + "']\")).clear();")
              output.append(
                "element(by.xpath('//input[@name=\"" + x +
                "\"]')).sendKeys(protractor.Key.SHIFT, protractor.Key.HOME);")
              for ran in range(int(split_values[split_element.index(val)])):
                output.append(
                  "element(by.xpath('//input[@name=\"" + x + "\"]')).sendKeys(protractor.Key.UP);")

          elif (xy == "search"):
            search_val = split_values[index]
          # if x not in unique_name1:
            unique_name1.append(x)
            unique_name[x] = 2
            if search_val.__contains__("$"):
              output.append("element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + search_val.split("$")[
                  1] + ").clear();")
              output.append(
                "element(by.xpath(\"//input[@name='" + x + "']\")).get(" + search_val.split("$")[
                  1] + ").sendKeys(\"" +
                search_val.split("$")[0] + "\");")
            else:
              output.append("element(by.xpath(\"//input[@name='" + x + "']\")).clear();")
              output.append(
                "element(by.xpath(\"//input[@name='" + x + "']\")).sendKeys(\"" + split_values[
                  split_element.index(val)] + "\");")

          elif (xy == "date"):
            date_val = split_values[index]
          # if x not in unique_name1:
            unique_name1.append(x)
            unique_name[x] = 2
            if date_val.__contains__("$"):
              output.append("element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + date_val.split("$")[
                  1] + ").clear();")
              output.append(
                "element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + date_val.split("$")[
                  1] + ").sendKeys(\"" + date_val.split("$")[0] + "\");")
            else:
              output.append("element(by.xpath(\"//input[@name='" + x + "']\")).clear();")
              output.append(
                "element(by.xpath(\"//input[@name='" + x + "']\")).sendKeys(\"" + split_values[
                  split_element.index(val)] + "\");")

          elif (xy == "month"):
            month_val = split_values[index]
          # if x not in unique_name1:
            unique_name1.append(x)
            unique_name[x] = 2
            if month_val.__contains__("$"):
              output.append("element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + str(
                  month_val.split("$")[1]) + ").clear();")
              output.append(
                "element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + str(
                  month_val.split("$")[1]) + ").sendKeys(\"" + month_split[0] + "\");"
                                                                                "element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + str(
                  month_val.split("$")[1]) + ").sendKeys(protractor.Key.TAB);"
                                             "element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + str(
                  month_val.split("$")[1]) + ").sendKeys(\"" + month_split[1] + "\");")
            else:
              output.append("element(by.xpath(\"//input[@name='" + x + "']\")).clear();")
              month_split = (split_values[split_element.index(val)]).split()
              output.append(
                "element(by.xpath(\"//input[@name='" + x + "']\")).sendKeys(\"" + month_split[
                  0] + "\");"
                       "element(by.xpath(\"//input[@name='" + x + "']\")).sendKeys(protractor.Key.TAB);"
                                                                  "element(by.xpath(\"//input[@name='" + x + "']\")).sendKeys(\"" +
                month_split[1] + "\");"
              )

          elif (xy == "email"):
            email_val = split_values[index]
          # if x not in unique_name1:
            unique_name1.append(x)
            unique_name[x] = 2
            if email_val.__contains__("$"):
              output.append(
                "element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + email_val.split("$")[
                  1] + ").sendKeys(\"" + email_val.split("$")[0] + "\");")
            else:
              output.append(
                "element(by.xpath(\"//input[@name='" + x + "']\")).sendKeys(\"" + split_values[
                  split_element.index(val)] + "\");")

          elif (xy == "number"):
            number_val = split_values[index]
            if number_val.__contains__("$"):
              output.append(
                "element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + str(
                  number_val.split("$")[1]) + ").sendKeys(" + str(
                  number_val.split("$")[0]) + ");")
            else:
              output.append(
                "element(by.xpath(\"//input[@name='" + x + "']\")).sendKeys(" + str(
                  split_values[split_element.index(val)]) + ");")
          elif (xy == "button" or xy == "submit" or xy == "checkbox" or xy == "reset"):
            but_val = split_values[index]
            if but_val.__contains__("$"):
              output.append(
                "element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + but_val.split("$")[1] + ").click();")
            else:
              output.append("element(by.xpath(\"//input[@name='" + x + "']\")).click();")
          elif (xy == "password"):
            pass_val = split_values[index]
            if (str(split_values[split_element.index(val)]).isalnum()):
              if pass_val.__contains__("$"):
                output.append(
                  "element.all(by.xpath(\"//input[@name='" + x + "']\")).get(" + str(
                    pass_val.split("$")[1]) + ").sendKeys(\'" + pass_val.split("$")[0] + "\');")
            else:
              output.append(
                "element(by.xpath(\"//input[@name='" + x + "']\")).sendKeys(\'" + str(split_values[
                                                                                        split_element.index(
                                                                                          val)]) + "\');")
        input_count += 1
      elif (str(val).find(" <select") > 0):
        flag = 1
        x = str(val.split("name=")[1]).split('"')[1].strip()
        sel_val_id = split_values[index]
        opt_val = str(val).split("<option")
        for mk in opt_val:
          if (mk.find("value=\"" + split_values[sp_val] + "\"") >= 0):
            sp_value = opt_val.index(mk)
        sel_val = opt_val[int(sp_value)].split(">")[1].split("<")[0]
      # if x not in unique_name1:
        unique_name1.append(x)
        unique_name[x] = 2
        if sel_val_id.__contains__("$"):
          output.append(
            "element.all(by.xpath(\"//select[@name=\'" + x + "\']\")).get(" + str(
              sel_val_id.split("$")[1]) + ").sendKeys('" + sel_val + "');")
        else:
          output.append(
            "element(by.xpath(\"//select[@name=\'" + x + "\']\")).sendKeys('" + sel_val + "');")

      elif (str(val).find("<textarea") > 0):
        flag = 1
        x = str(val.split("name")[1]).split('"')[1].strip()
        text_val = split_values[index]
        if text_val.__contains__("$"):
          output.append(
            "element.all(by.xpath('//textarea')).get(" + str(
              text_val.split("$")[1]) + ").sendKeys('" + str(
              str(text_val.split("$")[0])).strip() + "');")
        else:
          output.append("element(by.xpath('//textarea[position()=" + str(text_area) + "]')).sendKeys('" + str(
            split_values[split_element.index(val)]).strip() + "');")
        text_area += 1
      elif (str(val).find("<button") > 0):
        flag = 1
        x = val.split("name=")[1].strip().split('"')[1].strip('"')
        but_val_un = split_values[index]
      # if x not in unique_name1:
        unique_name1.append(x)
        unique_name[x] = 2
        if but_val_un.__contains__("$"):
          output.append("element.all(by.xpath(\"//button[@name='" + x + "']\")).get(" + str(
            but_val_un.split("$")[1]) + ").click();")
        else:
          output.append("element(by.xpath(\"//button[@name='" + x + "']\")).click();")

      else:
        click_val = split_values[split_element.index(val)]
        elm = val.strip().split()[0].split("<")[1].split(">")[0]
        if click_val == "click":
          output.append("element.all(by.xpath(\"//" + elm + "\")).get(" + str(glob_dict[elm]) + ").click();")
          glob_dict[elm] += 1

        else:
          output.append("element.all(by.xpath(\"//" + elm + "\")).get(" + str(
            glob_dict[elm]) + ").sendKeys(\'" + click_val + "\');")
          glob_dict[elm] += 1




    else:
      if val != '':
        # if val != '' and val.__contains__("<p")==False:
        click_val = split_values[split_element.index(val)]
        elm = ''
        if val.__contains__("<"):
          elm = val.strip().split()[0].split("<")[1]
          if elm.__contains__(">"):
            elm = elm.split(">")[0]
        # elm = val.strip().split()[0].split("<")[1].split(">")[0]
        file1 = open(base_location+"\\" + outerFolderName + "\\" + folderName + "\\" + "inputfile.json")
        data_out = json.load(file1)
        find_val = str(val).strip().strip('"') + ">"
        find_val = find_val.replace("wasClicked", "").replace("block", "none").replace(" ", "")
        ind = ''
        # print("element", elm)
        # print(data_out[counter - 1][elm])
        # print("finfval12", find_val)
        for el_ind, el_val in enumerate(data_out[counter - 1][elm]):
          # print("inside el", el_val)
          # print("finfval12", find_val)
          el_fin_val = el_val.replace("block", "none").replace("wasClicked", "").replace(" ", "")
          if el_fin_val == find_val:
            # print("inside")
            ind = el_ind
        # ind = data_out[counter - 1][elm].index(find_val)
        if click_val == "click":
          # output.append("element.all(by.xpath(\"//" + elm + "\")).get(" + str(ind) + ").click();")
          # glob_dict[elm] += 1
          if elm == "p":
            close_count = 0
            val_count_para = val.count(">")
            if val_count_para > 2:
              for para_line in range(val_count_para):
                para_val = val.split(">", close_count)
                print(para_line)
                if para_val[para_line].__contains__("<"):
                  print("one")
                  req_para_line = para_val[para_line].split(">")[1].split("<")[0]
                  if req_para_line != '':
                    break
                close_count += 1
            else:
              req_para_line = val.split(">")[1].split("<")[0]
            output.append(
              "element.all(by.xpath(\"//" + elm + "[contains(text(),\'" + str(req_para_line) + "\')]\")).click();")
          else:
            output.append("element.all(by.xpath(\"//" + elm + "\")).get(" + str(ind) + ").click();")
        else:
          if click_val.__contains__("$"):
            ind_val_else=click_val.split("$")[1]
            click_else=click_val.split("$")[0]
            if click_else!="click":
              output.append("element.all(by.xpath(\"//" + elm + "\")).get(" + str(
                ind_val_else) + ").sendKeys(\'" + click_else + "\');")
            elif(click_else=="click"):
              output.append("element.all(by.xpath(\"//" + elm + "\")).get(" + str(
                ind_val_else) + ").click();")

          else:
            output.append("element.all(by.xpath(\"//" + elm + "\")).get(" + str(
              glob_dict[elm]) + ").sendKeys(\'" + str(click_val) + "\');")
          glob_dict[elm] += 1

    if (flag == 1):
      flag = 0
    sp_val += 1
  flag_delay = False
  for ind, out in enumerate(output):
    print("index", ind, len(output))
    # if ind == 0:
    if out.__contains__(".clear("):
      final += "browser.wait(until.presenceOf(" + out.split(".clear(")[0
      ] + "), delay, 'Element taking too long to appear in the DOM');"
    if out.__contains__(".sendKeys("):
      final += "browser.wait(until.presenceOf(" + out.split(".sendKeys(")[0
      ] + "), delay, 'Element taking too long to appear in the DOM');"
    if out.__contains__(".click("):
      final += "browser.wait(until.presenceOf(" + out.split(".click(")[
        0] + "), delay, 'Element taking too long to appear in the DOM');"
    if (ind == len(output) - 1):
      flag_delay = True
      print("inside screenshot")
      final += '''browser.takeScreenshot().then(function (png) {
        var dir="'''+base_location.replace("\\","\\\\")+"\\\\" + outerFolderName + '\\\\' + folderName + '''\\\oldFiles\\\\";
  		  var fname="page''' + str(counter) + ".png\"; " + \
               '''var stream1 = fs.createWriteStream(path.join(dir, fname));
    stream1.write(new Buffer(png, 'base64'));
    stream1.end();
    });'''

    # if (ind == len(output)):
    #   final += 'browser.sleep(2000);'


    if flag_delay == True:
      final += out + "browser.sleep("+str(ui_delay)+");"
    else:
      final += out + "browser.sleep("+str(ui_minimumdelay)+");"
  output_val += '''it('Test ''' + str(1) + "\'" + ''', function() {
                        browser.get(\'''' + url + '''\');''' + final + "});"
  fin_out += output_val + '\n'
  output_file.write(fin_out)
  print("output", output)
  logger.info('Test case generated successfully')
  return flow_completed
  # except Exception as e:
  #   logger.error(e)

final_line = ''
whole_script = ''
modal_count = 0

@app.route("/generateTestCase")
def generateTestCase():  # To generate spec file
  global folderName
  global outerFolderName
  # folderName = request.args.get("folderName")
  logger = logging.getLogger('myapp')
  hdlr = logging.FileHandler(
    base_location + '\\' + outerFolderName + '\\' + folderName + '\\oldFiles\\generate_test.log')
  formatter = logging.Formatter('%(asctime)s %(levelname)s %(message)s')
  hdlr.setFormatter(formatter)
  logger.addHandler(hdlr)
  logger.setLevel(logging.INFO)

  flow = loadTestCase()
  length = int(1)
  modal_flag = request.args.get(
    "modalFlag")  # If it contains both id and name or name alone, we extract it using name
  global modal_count

  f_modal = 0
  f = open(
    base_location + "\\" + outerFolderName + '\\' + folderName + "\\uniqure_output_file.txt")  # Writes test case into this file
  f1 = open(base_location + "\\" + outerFolderName + '\\' + folderName + "\\OldFiles\\spec.js",
            "w+")  # Structured test cases
  f2 = open(base_location + "\\" + outerFolderName + '\\' + folderName + "\\OldFiles\\conf.js", "w+")
  conf = '''
  exports.config = {
    framework: 'jasmine',
    seleniumAddress: 'http://localhost:4444/wd/hub',
    specs: ['spec.js'],
    commonCapabilities: {
     'browserName': 'chrome',
     'chromeOptions': {
                                                'args': [ '--disable-gpu','-disable-dev-shm-usage','--no-sandbox','-disable-popup-blocking','--start-maximized','--disable-web-security','--allow-running-insecure-content','--disable-infobars']
                                }
   },
   onPrepare: function() {
      return browser.takeScreenshot();
    },
   getPageTimeout: 20000,
                allScriptsTimeout: 3000000,
                jasmineNodeOpts: {defaultTimeoutInterval: 3000000},
                defaultTimeoutInterval: 3000000
  }

  '''
  f2.write(conf)
  extractScriptValue = '''
async function extractModal() {
	browser.sleep(5000);
	var urlList = [];
	var inputList = [];
	var dropdownList = [];
	var textareaList = [];
	var linkList = [];
	var spanList = [];
	var buttonList = [];
	var divTagIdx = [];
	var inputJson = {};
	var aTagIdx = [];
	var anchorList = [];
	var tot_List=[]
	var tot_anchorList=[]
	var tot_linkList=[]
	var url = '';
	var divCounter=0;
	var glob_count=0;
	var divAccList = [];
	var tagName_1=[];
	var orderOfIpElems = {};
	var complete_List=[];
	var un_handle=["html", "head", "style", "title", "meta", "option","link", "br", "body", "script","input",
		"textarea", "a", "button", "span", "select","app-root","app-dashboard","app-sidebar", "section","div", "nav","td", "tr", "app-router", "app-footer",
		"app-header", "router-outlet", "base", "app-login", "app-layout","form", "footer", "header", "li","dd","dt","dl" ];
	var un_app="app-";
	var un_ng="ng-";
	var un_crst="crst-";
	/*element.all(by.xpath("//div[contains(@style, 'display: block')]//*")).getTagName("outerHTML").then(function(outerHTML){
		//outerHTML.forEach(function (arrayValue) {
		//	glob_count+=1;
		//});
		glob_count = outerHTML.length;
		if(glob_count>1000){
			browser.sleep(15000);
		} else if(500<glob_count<1000){
			browser.sleep(12000);
		}
		else{
			browser.sleep(10000);
		}
	});*/



	element.all(By.xpath("//div[contains(@style, 'display: block')]//a")).getAttribute("outerHTML").then(function(outerhtml){
	console.log("inside div datatoggle modal");
		urlList.push('modal$page');
		inputJson["url"] = urlList;
		outerhtml.forEach(function (arrayItem, idx) {
			const dom = new jsdom.JSDOM(arrayItem);
			tot_List.push(arrayItem);
			if((dom.window.document.querySelector("a").getAttribute("data-toggle")) == "collapse"
				&& (dom.window.document.querySelector("a").getAttribute("aria-expanded")) == "false"){
				aTagIdx.push(idx);
				anchorList.push(idx);
				tot_anchorList.push(arrayItem)
			} else {
				linkList.push(idx);
				tot_linkList.push(arrayItem)
			}
		});
		//inputJson["a"] = anchorList;
		//inputJson["links"] = linkList;
		var remaining = tot_List.filter(item => tot_anchorList.indexOf(item) < 0);
		if(aTagIdx.length!=0){
			aTagIdx.reverse();
			aTagIdx.forEach(function(idx){
				element(by.xpath("//div[contains(@style, 'display: block')]//a["+(idx+1)+"]")).click();
			})
		}
	});


	element.all(by.xpath('//div[contains(@style, "display: block")]//div[contains(@data-toggle,"collapse")]')).getAttribute("outerHTML").then(function(outerHTML){
        outerHTML.forEach(function(arrayItem, idx){

            const dom = new jsdom.JSDOM(arrayItem)
            var div = dom.window.document.querySelector("div");
            divTagIdx.push(idx);

        });
        if(divTagIdx.length!=0)
        {
            console.log("div Tag Idx b4", divTagIdx);
            divTagIdx.reverse();
            console.log("div Tag Idx after", divTagIdx);
            console.log("Click reversed");
            divTagIdx.forEach(function(idx){
                element.all(by.xpath('//div[contains(@data-toggle,"collapse")]')).get(idx).click();
                console.log("click happening", idx);
                browser.sleep(1000);
            });
        }
    });

	element.all(By.xpath("//div[contains(@style, 'display: block')]//input")).getAttribute("outerHTML").then(function(outerhtml){
	console.log("inside input datatoggle modal");
		outerhtml.forEach(function (arrayItem) {
			inputList.push(arrayItem);
		});
		inputJson["linklist"] = tot_List;
		inputJson["input"] = inputList;
	});

	element.all(By.xpath("//div[contains(@style, 'display: block')]//select")).getAttribute("outerHTML").then(function(outerhtml){
	console.log("inside select datatoggle modal");
		outerhtml.forEach(function (arrayItem) {
			dropdownList.push(arrayItem);
		});
		inputJson["select"] = dropdownList;
	});
	element.all(by.xpath("//div[contains(@style, 'display: block')]//span")).getAttribute("outerHTML").then(function(outerHTML){
	console.log("inside span datatoggle modal");
		outerHTML.forEach(function (arrayItem) {
			spanList.push(arrayItem);
		});
		inputJson["span"] = spanList;
	});
	element.all(By.xpath("//div[contains(@style, 'display: block')]//button")).getAttribute("outerHTML").then(function(outerhtml){
	console.log("inside button datatoggle modal");
		outerhtml.forEach(function (arrayItem) {
			buttonList.push(arrayItem);
		});
		inputJson["button"] = buttonList;
	});

	element.all(By.xpath("//div[contains(@style, 'display: block')]//textarea")).getAttribute("outerHTML").then(function(outerhtml){
	console.log("inside textarea datatoggle modal");
		outerhtml.forEach(function (arrayItem) {
			textareaList.push(arrayItem);
		});
		inputJson["textarea"] = textareaList;
	});


	element.all(By.xpath("//div[contains(@style, 'display: block')]//input")).getLocation().then(function(pos){
	console.log("inside input getloc modal");
		orderOfIpElems['input']=[];
		for(var i=0;i<pos.length;i++){
			orderOfIpElems['input'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
		}
	});

	element.all(By.xpath("//div[contains(@style, 'display: block')]//select")).getLocation().then(function(pos){
	console.log("inside select getloc modal");
		orderOfIpElems['select']=[];
		for(var i=0;i<pos.length;i++){
			orderOfIpElems['select'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
		}
	});
	element.all(by.xpath("//div[contains(@style, 'display: block')]//a")).getLocation().then(function(pos){
	console.log("inside anchor getloc modal");
		orderOfIpElems['linklist']=[];
		for(var i=0;i<pos.length;i++){
			orderOfIpElems['linklist'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
		}
	});
	element.all(by.xpath("//div[contains(@style, 'display: block')]//span")).getLocation().then(function(pos){
	console.log("inside span getloc modal");
		orderOfIpElems['span']=[];
		for(var i=0;i<pos.length;i++){
			orderOfIpElems['span'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
		}
	});
	element.all(By.xpath("//div[contains(@style, 'display: block')]//button")).getLocation().then(function(pos){
	console.log("inside button getloc modal");
		orderOfIpElems['button']=[];
		for(var i=0;i<pos.length;i++){
			orderOfIpElems['button'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
		}
	});
	element.all(by.xpath("//div[contains(@style, 'display: block')]//*")).getTagName().then(function(outerHTML){
	console.log("inside all getloc modal");
		outerHTML.forEach(function (arrayValue, idx) {
			if(tagName_1.indexOf(arrayValue)<0 && un_handle.indexOf(arrayValue)<0 && arrayValue.indexOf(un_app)<0 && arrayValue.indexOf(un_crst)<0  && arrayValue.indexOf(un_ng)<0){
				tagName_1.push(arrayValue);
				console.log("tagname",tagName_1);

				element.all(by.xpath("//div[contains(@style, 'display: block')]//"+arrayValue)).getAttribute("outerHTML").then(function(outerHTML_1){
					outerHTML_1.forEach(function (arrayItem) {
							console.log("arrayValue",arrayValue);
							complete_List.push(arrayItem);
					});
					console.log("complete List", complete_List);
					inputJson[arrayValue] = complete_List;
					complete_List=[];
					console.log("ip json", inputJson);
				});


				element.all(by.xpath("//div[contains(@style, 'display: block')]//"+arrayValue)).getLocation().then(function(pos){
					orderOfIpElems[arrayValue]=[];
					for(var i=0;i<pos.length;i++){
						orderOfIpElems[arrayValue].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
					}
					console.log('Elements extracted..')
				});
			}
		});
	});
	element.all(By.xpath("//div[contains(@style, 'display: block')]//textarea")).getLocation().then(function(pos){
	console.log("inside textarea getloc modal");
	orderOfIpElems['textarea']=[];
		for(var i=0;i<pos.length;i++){
			orderOfIpElems['textarea'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
		}
		inputJson['order'] = [orderOfIpElems]
		ipJson.push(inputJson);

		var data = JSON.stringify(ipJson);
		fs.writeFileSync(fileUrl, data.replace(/\\n/g, '').replace(/\\t/g, ''));
		fs.writeFileSync(fileUrl1, data.replace(/\\n/g, '').replace(/\\t/g, ''));
		console.log('Elements extracted..')

		urlList = [];
		inputList = [];
		dropdownList = [];
		buttonList = [];
		textareaList = [];
		linkList = [];
		inputJson = {};
		aTagIdx = [];
		anchorList = [];
		url = '';
		complete_List=[];
		tagName_1=[];
		orderOfIpElems = {};
		browser.waitForAngularEnabled(true)
	});
}
async function extract() {
	browser.sleep(20000);
	var urlList = [];
	var inputList = [];
	var dropdownList = [];
	var textareaList = [];
	var linkList = [];
	var buttonList = [];
	var divTagIdx = [];
	var spanList =[];
	var inputJson = {};
	var aTagIdx = [];
	var anchorList = [];
	var tot_List=[]
	var tot_anchorList=[]
	var tot_linkList=[]
	var url = '';
	var divCounter=0;
	var divAccList = [];
	var glob_count=0;
	var complete_List=[];
	var tagName_1=[];
	var orderOfIpElems = {};
	var un_handle=["html", "head", "style", "title", "meta", "option","link", "br", "body", "script","input",
		"textarea", "a", "button", "span", "select","app-root","app-dashboard","app-sidebar", "section","div", "nav", "td", "tr", "app-router", "app-footer",
		"app-header", "router-outlet", "base", "app-login","app-layout", "form", "footer", "header", "li","dd","dt","dl" ];
	var un_app='app-';
	var un_crst="crst-";
	var un_ng="ng-";
	/*element.all(by.xpath('//*')).getTagName("outerHTML").then(function(outerHTML){
		//outerHTML.forEach(function (arrayValue) {
		//	glob_count+=1;
		//});
		glob_count = outerHTML.length;
		if(glob_count>1000){
			browser.sleep(15000);
		} else if(500<glob_count<1000){
			browser.sleep(12000);
		}
		else{
			browser.sleep(1000);
		}
	});*/

	element.all(by.xpath('//a')).getAttribute("outerHTML").then(function(outerHTML){
	urlList.push('new$page');
	inputJson["url"] = urlList;
	console.log("inside extract");
	console.log("inside anchor");
	outerHTML.forEach(function (arrayItem, idx) {
			tot_List.push(arrayItem);
			const dom = new jsdom.JSDOM(arrayItem);

			if((dom.window.document.querySelector("a").getAttribute("data-toggle")) == "collapse"
				&& (dom.window.document.querySelector("a").getAttribute("aria-expanded")) == "false"){
				aTagIdx.push(idx);
				anchorList.push(idx);
				tot_anchorList.push(arrayItem)
			} else {
				linkList.push(idx);
				tot_linkList.push(arrayItem)
			}
		});
		inputJson["linklist"] = tot_List;

		//inputJson["a"] = anchorList;
		//inputJson["links"] = linkList;
		var remaining = tot_List.filter(item => tot_anchorList.indexOf(item) < 0);
		if(aTagIdx.length!=0){
			aTagIdx.reverse();
			aTagIdx.forEach(function(idx){
				element(by.xpath("//a["+(idx+1)+"]")).click();
			})
		}
		console.log("a completed");
		console.log("length", tot_List.length)
	});
	//console.log("after element");

	element.all(by.xpath('//div[contains(@data-toggle,"collapse")]')).getAttribute("outerHTML").then(function(outerHTML){
		outerHTML.forEach(function(arrayItem, idx){
			const dom = new jsdom.JSDOM(arrayItem)
			var div = dom.window.document.querySelector("div");
			divTagIdx.push(idx);
		});
		if(divTagIdx.length!=0)
		{
			console.log("div Tag Idx b4", divTagIdx);
			divTagIdx.reverse();
			console.log("div Tag Idx after", divTagIdx);
			console.log("Click reversed");
			divTagIdx.forEach(function(idx){
				element.all(by.xpath('//div[contains(@data-toggle,"collapse")]')).get(idx).click();
				console.log("click happening", idx);
				browser.sleep(1000);
			});
		}
	});

	element.all(by.xpath('//input')).getAttribute("outerHTML").then(function(outerHTML){
	console.log("inside input ");
	console.log("inside input");
		outerHTML.forEach(function (arrayItem) {
			inputList.push(arrayItem);
		});

		inputJson["input"] = inputList;
	});
	element.all(by.xpath('//button')).getAttribute("outerHTML").then(function(outerHTML){
	console.log("inside button");
		outerHTML.forEach(function (arrayItem) {
			buttonList.push(arrayItem);
		});
		inputJson["button"] = buttonList;
	});
	element.all(by.xpath('//span')).getAttribute("outerHTML").then(function(outerHTML){
		console.log("inside span");
		outerHTML.forEach(function (arrayItem) {
			spanList.push(arrayItem);
		});
		inputJson["span"] = spanList;
	});
	element.all(by.xpath('//select')).getAttribute("outerHTML").then(function(outerHTML){
	console.log("inside select");
		outerHTML.forEach(function (arrayItem) {
			dropdownList.push(arrayItem);
		});
		inputJson["select"] = dropdownList;
	});

	element.all(by.xpath('//textarea')).getAttribute("outerHTML").then(function(outerHTML){
	console.log("inside textarea");
		outerHTML.forEach(function (arrayItem) {
			textareaList.push(arrayItem);
		});
		inputJson["textarea"] = textareaList;
	});

	element.all(by.xpath('//input')).getLocation().then(function(pos){
	console.log("inside input getLocation");
		orderOfIpElems['input']=[];
		for(var i=0;i<pos.length;i++){
			orderOfIpElems['input'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
		}
	})
	element.all(by.xpath('//button')).getLocation().then(function(pos){
	console.log("inside button getLocation");
		orderOfIpElems['button']=[];
		for(var i=0;i<pos.length;i++){
			orderOfIpElems['button'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
		}
	});
	element.all(by.xpath('//a')).getLocation().then(function(pos){
	console.log("inside anchor getLocation");
		orderOfIpElems['linklist']=[];
		for(var i=0;i<pos.length;i++){
			orderOfIpElems['linklist'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
		}
	});
	element.all(by.xpath('//span')).getLocation().then(function(pos){
	console.log("inside span getLocation");
		orderOfIpElems['span']=[];
		for(var i=0;i<pos.length;i++){
			orderOfIpElems['span'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
		}
	});
	element.all(by.xpath('//select')).getLocation().then(function(pos){
	console.log("inside select getLocation");
		orderOfIpElems['select']=[];
		for(var i=0;i<pos.length;i++){
			orderOfIpElems['select'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
		}
	});
	//console.log("after element before textarea");
	element.all(by.xpath('//*')).getTagName().then(function(outerHTML){
	console.log("inside all getLocation");
		outerHTML.forEach(function (arrayValue, idx) {
			if(tagName_1.indexOf(arrayValue)<0 && un_handle.indexOf(arrayValue)<0 && arrayValue.indexOf(un_app)<0 && arrayValue.indexOf(un_crst)<0 && arrayValue.indexOf(un_ng)<0 ){
				tagName_1.push(arrayValue);
				console.log("tagname",tagName_1);
				element.all(by.xpath('//'+arrayValue)).getAttribute("outerHTML").then(function(outerHTML_1){
					outerHTML_1.forEach(function (arrayItem) {
							console.log("arrayValue",arrayValue);
							complete_List.push(arrayItem);
					});
					console.log("complete List", complete_List);
					inputJson[arrayValue] = complete_List;
					complete_List=[];
					console.log("ip json", inputJson);
				});
				element.all(by.xpath('//'+arrayValue)).getLocation().then(function(pos){
					orderOfIpElems[arrayValue]=[];
					for(var i=0;i<pos.length;i++){
						orderOfIpElems[arrayValue].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
					}
					console.log('Elements extracted..')
				});
			}
		});
	});
	element.all(by.xpath('//textarea')).getLocation().then(function(pos){
		orderOfIpElems['textarea']=[];
		for(var i=0;i<pos.length;i++){
			orderOfIpElems['textarea'].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
		}
		inputJson['order'] = [orderOfIpElems]
		ipJson.push(inputJson);

		var data = JSON.stringify(ipJson);
		fs.writeFileSync(fileUrl, data.replace(/\\n/g, '').replace(/\\t/g, ''));
		fs.writeFileSync(fileUrl1, data.replace(/\\n/g, '').replace(/\\t/g, ''));
		console.log('Elements extracted..')

		urlList = [];
		inputList = [];
		dropdownList = [];
		buttonList = [];
		textareaList = [];
		linkList = [];
		inputJson = {};
		aTagIdx = [];
		anchorList = [];
		url = '';
		tagName_1=[];
		complete_List=[];
		orderOfIpElems = {};
		browser.waitForAngularEnabled(true);
	});
}

	async function extractUnhandled(){
				const jsdom = require("jsdom");
				const fs = require('fs');
				var orderOfIpElems = {};
				var fullElem={};
				var inputJson = {};
				var tagNames=[];
				var tagName_1=[]
				var un_handle=["html", "head", "style", "title", "link", "body", "script","input","textarea", "linklist", "button", "span", "select","app-root","app-dashboard","app-sidebar", "section","div", "nav","dd","dt","dl"]
                var un_app="app-"

				element.all(by.xpath('//*')).getTagName("outerHTML").then(function(outerHTML){
				console.log("inside all ");
                    outerHTML.forEach(function (arrayValue, idx) {
					if(tagName_1.indexOf(arrayValue)<0 && un_handle.indexOf(arrayValue)<0 && arrayValue.indexOf(un_app)<0){
						tagName_1.push(arrayValue);
						console.log("tagname",tagName_1);
						fullElem[arrayValue]=[];


					element.all(by.xpath('//'+arrayValue)).getLocation().then(function(pos){
						orderOfIpElems[arrayValue]=[];
						for(var i=0;i<pos.length;i++){
							orderOfIpElems[arrayValue].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
						}

					fs.writeFileSync("base_location\\requried_3.txt", tagName_1);
					console.log('Elements extracted..')

					});
					}



					});

					});
	}

	async function extractUnhandledModal(){
				const jsdom = require("jsdom");
				const fs = require('fs');
				var orderOfIpElems = {};
				var fullElem={};
				var inputJson = {};
				var tagNames=[];
				var tagName_1=[]
				var un_handle=["html", "head", "style", "title", "link", "body", "script","input","textarea", "linklist", "button", "span", "select","app-root","app-dashboard","app-sidebar", "section","div", "nav","dd","dt","dl" ]
                var un_app="app-";

				element.all(by.xpath("//div[contains(@style, 'display: block')]//*")).getTagName("outerHTML").then(function(outerHTML){
				console.log("inside all");
                    outerHTML.forEach(function (arrayValue, idx) {
					if(tagName_1.indexOf(arrayValue)<0 && un_handle.indexOf(arrayValue)<0 && arrayValue.indexOf(un_app)<0 ){
						tagName_1.push(arrayValue);
						console.log("tagname",tagName_1);
						fullElem[arrayValue]=[];


					element.all(by.xpath("//div[contains(@style, 'display: block')]//"+arrayValue)).getLocation().then(function(pos){
						orderOfIpElems[arrayValue]=[];
						for(var i=0;i<pos.length;i++){
							orderOfIpElems[arrayValue].push("(" + Math.floor(pos[i].x) + "," + Math.floor(pos[i].y) + ")");
						}

					fs.writeFileSync("''' + base_location + '''\\requried_3.txt", tagName_1);
					console.log('Elements extracted..')

					});
					}



					});

					});
	}
  '''
  es = open(base_location + "\\" + outerFolderName + '\\' + folderName + "\\ExtractScript.txt", "w+")
  es.write(extractScriptValue)
  script = open(base_location + "\\" + outerFolderName + '\\' + folderName + "\\ExtractScript.txt",
                "r+")  # File that is already present
  global whole_script
  new_final = ''
  popup = request.args.get("popup")
  name_id = request.args.get("nameId")
  output_string = request.args.get("outputString")
  idVal = request.args.get("idVal")
  flag = 0
  output = []
  f_out = ''
  m = ''
  for lines in f.readlines():  # From uniqure_output_file
    if (lines.strip().startswith("it")):
      flag = 1
    if (flag == 1):
      f_out += lines.strip()
    if (lines.strip().endswith(";});")):
      flag = 0
      output.append(f_out)
      f_out = ''
  global final_line

  for j in range(length):
    for i in range(j, len(output), length):
      url = output[i].split("browser.get(")[1].split("\'")[1]
      if (output[i].startswith("it")):
        m += output[i][output[i].index("browser.wait("):]

  file1 = open(base_location + "\\" + outerFolderName + '\\' + folderName + "\\inputfile.json")
  data_out = json.load(file1)
  cntr = len(data_out) - 1
  keys_data = data_out[cntr]['order'][0].keys()
  tot_count = 0
  for key in keys_data:
    tot_count += len(data_out[cntr]['order'][0][key])
  print('tot_count of elements', tot_count)
  if tot_count < 1000 and tot_count > 0:
    delay = 10000
  elif tot_count < 3000 and tot_count > 1000:
    delay = 12000
  elif tot_count < 5000 and tot_count > 3000:
    delay = 14000
  elif tot_count < 8000 and tot_count > 5000:
    delay = 16000
  elif tot_count < 10000 and tot_count > 8000:
    delay = 18000
  else:
    delay = 20000

  final = '''
const jsdom = require("jsdom");
const fs = require('fs');
const path = require('path');
const fileUrl = \'''' + base_location_front + '/' + outerFolderName + '/' + folderName + '''/inputfile.json';
const fileUrl1 = \'''' + base_location_front + '/' + outerFolderName + '/' + folderName + '''/inputfile1.json';
var ipJson = require(fileUrl);
describe('Protractor Demo App', function() {
\t\tit('Test 1', function() {\n
\t\tbrowser.ignoreSynchronization =true;
\t\tbrowser.driver.manage().window().maximize();

\t\tbrowser.get(\'''' + url + '''\');
\t\tbrowser.waitForAngularEnabled(false);
\t\tbrowser.sleep(10);
\t\tvar until = protractor.ExpectedConditions;
\t\tvar delay = ''' + str(delay) + ''';'''
  # default_var = '''
  #           browser.sleep(3000);
  #           browser.waitForAngularEnabled(false);
  #           browser.sleep(5000);'''
  default_var = ''

  # default_var = '''
  #             browser.sleep(5000);
  #             browser.waitForAngularEnabled(false);
  #             browser.sleep(30000);'''
  if (flow == "yes"):  # If user clicks complete flow in page
    final_line += m[:-3] + default_var
    if whole_script.find("extract()") >= 0:  # If the file is already created
      if modal_flag == "yes":  # If user clicks modal window button
        f_modal = 1
        whole_script = whole_script[:whole_script.index("browser.wait(")] + final_line + default_var + \
                       '''
\t\t//extract();
\t\textractModal();
\t\t//extractUnhandled();
\t\t//extractUnhandledModal();
\t\t});
''' + script.read() + '});'


      else:  # If user doesn't clicks modal window button

        if f_modal == 1:  # If it went inside the modal window
          f_modal = 0
          whole_script = whole_script[:whole_script.index("browser.wait(")] + final_line + \
                         '''''' + default_var + \
                         '''
                         \t\textract();
                         \t\t//extractModal();
                         \t\t//extractUnhandled();
                         \t\t//extractUnhandledModal();
                         \t\t});
                         ''' + script.read() + '});'
        else:  # If it doesn't went inside modal window
          whole_script = whole_script[:whole_script.index("browser.wait(")] + final_line + default_var + \
                         '''
                         \t\textract();
                         \t\t//extractModal();
                         \t\t//extractUnhandled();
                         \t\t//extractUnhandledModal();
                         \t\t});
                         ''' + script.read() + '});'

    else:  # If it is a new file
      if (modal_flag == "yes"):  # If the modal window button is clicked
        f_modal = 1
        final = final + m[:-3] + default_var
        final += '''
\t\t//extract();
\t\textractModal();
\t\t//extractUnhandled();
\t\t//extractUnhandledModal();
\t\t});
''' + script.read() + '});'
        whole_script = final

      else:
        if f_modal == 1:  # If the modal window button is not clicked
          f_modal = 0
          final = final + m[:-3] + default_var
          final += '''
\t\textract();
\t\t//extractModal();
\t\t//extractUnhandled();
\t\t//extractUnhandledModal();
\t\t});
''' + script.read() + '});'
        else:  # If the modal window is open
          final = final + m[:-3] + default_var + '''
\t\textract();
\t\t//extractModal();
\t\t//extractUnhandled();
\t\t//extractUnhandledModal();
\t\t});
''' + script.read() + '});'
        whole_script = final
    f1.write(whole_script)

  else:  # If the user clicks complete user flow
    if f_modal == 1:  # If the modal window is opened
      f_modal = 0
      if modal_count == 0:  # If the modal window is not opened once
        final_line += m[
                      :-3] + default_var
        # + "let list = element.all(by.xpath(\"//span[@class='close']\"));list.get(0).click();browser.sleep(2000);"

        modal_count += 1
      else:
        final_line += m[:-3] + "list.get(" + str(modal_count) + ").click(); browser.sleep(2000);"
        modal_count += 1
    else:
      final_line += m[:-3] + default_var
    if (whole_script.find("extract()") >= 0):
      if modal_flag == "yes":  # If the open modal button is clicked
        if (f_modal == 1):  # If the modal window is  not opened
          if modal_count == 0:  # Number of modal window
            # final_line += "let list = element.all(by.xpath(\"//span[@class='close']\"));list.get(0).click();browser.sleep(2000);"
            final_line += default_var
            modal_count += 1
          else:  # Number of modal window is greater then 1
            final_line += "list.get(" + str(modal_count) + ").click(); "
            modal_count += 1
        whole_script = whole_script[:whole_script.index("browser.wait(")] + final_line + default_var + '''
\t\t//extract();
\t\textractModal();
\t\t//extractUnhandled();
\t\t//extractUnhandledModal();
\t\t});
''' + script.read() + '});'
      else:  # If the modal window is opened
        if (popup != None and popup != "other" and name_id != None and output_string != None and idVal != None):
          if (popup == "browser-pop-up"):
            add_output = '''browser.switchTo().alert().then((alrt) => {
              alrt.getText().then(function(txt)
            {expect(txt.toLowerCase()).toContain("''' + output_string + '''".toLowerCase())});alrt.accept();});'''
          else:
            add_output = "element(by.xpath(\"//*[@" + name_id + "='" + idVal + "']\")).getText().then(function(txt)" \
                                                                               "{expect(txt.toLowerCase()).toContain(\'" + output_string + "\'.toLowerCase()" \
                                                                                                                                           ");});"
        else:
          add_output = ''
          # element(by.xpath("//*[@id='successMsg']")).getText().then(function(txt)
          # {
          #   expect(txt.toLowerCase()).toContain('Success'.toLowerCase());
          # });
        whole_script = whole_script[:whole_script.index("browser.wait(")] + \
                       final_line + default_var + \
                       add_output + '''
\t\t//extract();
\t\t//extractModal();
\t\t//extractUnhandled();
\t\t//extractUnhandledModal();
\t\t});
''' + script.read() + '});'
    else:  # If it is a new file
      if modal_flag == "yes":  # If the user clicks open modal
        # if (popup != None and popup!="other" and name_id != None and output_string != None and idVal != None):
        #   if(popup=="browser-pop-up"):
        #     add_output='''browser.switchTo().alert().then((alrt) = > {
        #       alrt.getText().then(function(txt)
        #     {expect(txt.toLowerCase()).toContain("'''+output_string+".toLowerCase())});});"
        #   else:
        #     add_output = "element(by.xpath(\"*[@" + name_id + "=" + idVal + "]\")).getText().then(function(txt)" \
        #                                                                   "{expect(txt.toLowerCase()).toContain(\'" + output_string + "\'.toLowerCase()" \
        #                                                                                                                               ");});"
        # else:
        #   add_output=''
        final = final + m[:-3] + default_var + '''
\t\t//extract();
\t\textractModal();
\t\t//extractUnhandled();
\t\t//extractUnhandledModal();
\t\t});
''' + script.read() + '});'
        whole_script = final
      else:  # If the user clicks complete user flow
        if (popup != None and popup != "other" and name_id != None and output_string != None and idVal != None):
          if (popup == "browser-pop-up"):
            add_output = '''browser.switchTo().alert().then((alrt) => {
              alrt.getText().then(function(txt)
            {expect(txt.toLowerCase()).toContain("''' + output_string + '''".toLowerCase())});alrt.accept();});'''
          else:
            add_output = "element(by.xpath(\"//*[@" + name_id + "='" + idVal + "']\")).getText().then(function(txt)" \
                                                                               "{expect(txt.toLowerCase()).toContain(\'" + output_string + "\'.toLowerCase()" \
                                                                                                                                           ");});"
        else:
          add_output = ''

        final = final + m[:-3] + default_var + \
                add_output + '''
\t\t//extract();
\t\t//extractModal();
\t\t//extractUnhandled()
\t\t//extractUnhandledModal();
\t\t});
''' + script.read() + '});'
        whole_script = final

    f1.write(whole_script)

  logger.info('Test file generated')

  # logger.error(e)
  format_spec(outerFolderName, folderName)
  testDescription()
  return jsonify({"status": "success"})

final_line_1 = ''
whole_script_1 = ''
modal_count_1 = 0

@app.route("/updategenerateTestCase")
def updategenerateTestCase():
  logger = logging.getLogger('myapp')
  hdlr = logging.FileHandler('base_location\oldFiles\generate_test.log')
  formatter = logging.Formatter('%(asctime)s %(levelname)s %(message)s')
  hdlr.setFormatter(formatter)
  logger.addHandler(hdlr)
  logger.setLevel(logging.INFO)

  flow = loadTestCase()
  length = int(1)
  modal_flag = request.args.get(
    "modalFlag")  # If it contains both id and name or name alone, we extract it using name
  global modal_count_1
  f_modal = 0
  f = open(r"base_location\uniqure_output_file.txt")  # Writes test case into this file
  f1 = open(r"base_location\OldFiles\updatespec.js", "w+")  # Structured test cases
  script = open(r"base_location\ExtractScript.txt", "r+")  # File that is already present
  global whole_script_1
  new_final = ''
  popup = request.args.get("popup")
  name_id = request.args.get("nameId")
  output_string = request.args.get("outputString")
  idVal = request.args.get("idVal")
  flag = 0
  output = []
  f_out = ''
  m = ''
  for lines in f.readlines():  # From uniqure_output_file
    if (lines.strip().startswith("it")):
      flag = 1
    if (flag == 1):
      f_out += lines.strip()
    if (lines.strip().endswith(";});")):
      flag = 0
      output.append(f_out)
      f_out = ''
  global final_line_1

  for j in range(length):
    for i in range(j, len(output), length):
      url = output[i].split("browser.get(")[1].split("\'")[1]
      if (output[i].startswith("it")):
        m += output[i][output[i].index("element"):]

  final = '''const jsdom = require("jsdom");
           const fs = require('fs');
           const path = require('path');
           const fileUrl = \'''' +base_location_front+'/'+ outerFolderName + '/' + folderName + '''/inputfile.json';
           var ipJson = require(fileUrl);

           describe('Protractor Demo App', function() { it('Test 1', function() {
           browser.ignoreSynchronization =true;
           browser.driver.manage().window().maximize();

           browser.get(\'''' + url + '''\');
           browser.waitForAngularEnabled(false);
           browser.sleep(10);'''
  default_var = '''
            browser.sleep(8000);
            browser.waitForAngularEnabled(false);
            browser.sleep(10000);'''

  # default_var = '''
  #             browser.sleep(5000);
  #             browser.waitForAngularEnabled(false);
  #             browser.sleep(30000);'''
  if (flow == "yes"):  # If user clicks complete flow in page
    final_line_1 += m[:-3] + default_var
    if whole_script_1.find("extract()") >= 0:  # If the file is already created
      if modal_flag == "yes":  # If user clicks modal window button
        f_modal = 1
        whole_script_1 = whole_script_1[:whole_script_1.index("element(")] + final_line_1 + default_var + \
                         '''
                  //extract();
                                                            extractModal();

                                                            });
                                                            ''' + script.read() + '});'


      else:  # If user doesn't clicks modal window button

        if f_modal == 1:  # If it went inside the modal window
          f_modal = 0
          whole_script_1 = whole_script_1[:whole_script_1.index("element")] + final_line_1 + \
                           '''''' + default_var + \
                           '''
               extract();
  //extractModal();

  });
  ''' + script.read() + '});'
        else:  # If it doesn't went inside modal window
          whole_script_1 = whole_script_1[:whole_script_1.index("element")] + final_line_1 + default_var + \
                           '''
                                                    extract();
                                                    //extractModal();

                                                    });
                                                    ''' + script.read() + '});'

    else:  # If it is a new file
      if (modal_flag == "yes"):  # If the modal window button is clicked
        f_modal = 1
        final = final + m[:-3] + default_var
        final += '''//extract();
                    extractModal();

                                });
                                ''' + script.read() + '});'
        whole_script_1 = final

      else:
        if f_modal == 1:  # If the modal window button is not clicked
          f_modal = 0
          final = final + m[:-3] + default_var
          final += '''extract();
                        //extractModal();

                                               });
                                               ''' + script.read() + '});'
        else:  # If the modal window is open
          final = final + m[:-3] + default_var + '''
                          extract();
                          //extractModal();

                                                             });
                                                             ''' + script.read() + '});'
        whole_script_1 = final
    f1.write(whole_script_1)

  else:  # If the user clicks complete user flow
    if f_modal == 1:  # If the modal window is opened
      f_modal = 0
      if modal_count_1 == 0:  # If the modal window is not opened once
        final_line_1 += m[
                        :-3] + default_var
        # + "let list = element.all(by.xpath(\"//span[@class='close']\"));list.get(0).click();browser.sleep(2000);"

        modal_count_1 += 1
      else:
        final_line_1 += m[:-3] + "list.get(" + str(modal_count_1) + ").click(); browser.sleep(2000);"
        modal_count_1 += 1
    else:
      final_line_1 += m[:-3] + default_var
    if (whole_script_1.find("extract()") >= 0):
      if modal_flag == "yes":  # If the open modal button is clicked
        if (f_modal == 1):  # If the modal window is  not opened
          if modal_count_1 == 0:  # Number of modal window
            # final_line_1 += "let list = element.all(by.xpath(\"//span[@class='close']\"));list.get(0).click();browser.sleep(2000);"
            final_line_1 += default_var
            modal_count_1 += 1
          else:  # Number of modal window is greater then 1
            final_line_1 += "list.get(" + str(modal_count_1) + ").click(); "
            modal_count_1 += 1
        whole_script_1 = whole_script_1[:whole_script_1.index("element")] + final_line_1 + default_var + '''
                                  //extract();
                                  extractModal();

                                   });
                                   ''' + script.read() + '});'
      else:  # If the modal window is opened
        if (popup != None and popup != "other" and name_id != None and output_string != None and idVal != None):
          if (popup == "browser-pop-up"):
            add_output = '''browser.switchTo().alert().then((alrt) => {
              alrt.getText().then(function(txt)
            {expect(txt.toLowerCase()).toContain("''' + output_string + '''".toLowerCase())});alrt.accept();});'''
          else:
            add_output = "element(by.xpath(\"//*[@" + name_id + "='" + idVal + "']\")).getText().then(function(txt)" \
                                                                               "{expect(txt.toLowerCase()).toContain(\'" + output_string + "\'.toLowerCase()" \
                                                                                                                                           ");});"
        else:
          add_output = ''
          # element(by.xpath("//*[@id='successMsg']")).getText().then(function(txt)
          # {
          #   expect(txt.toLowerCase()).toContain('Success'.toLowerCase());
          # });
        whole_script_1 = whole_script_1[:whole_script_1.index("element")] + \
                         final_line_1 + default_var + \
                         add_output + '''//extract();
                                      //extractModal();
                                      });
                                      ''' + script.read() + '});'
    else:  # If it is a new file
      if modal_flag == "yes":  # If the user clicks open modal
        # if (popup != None and popup!="other" and name_id != None and output_string != None and idVal != None):
        #   if(popup=="browser-pop-up"):
        #     add_output='''browser.switchTo().alert().then((alrt) = > {
        #       alrt.getText().then(function(txt)
        #     {expect(txt.toLowerCase()).toContain("'''+output_string+".toLowerCase())});});"
        #   else:
        #     add_output = "element(by.xpath(\"*[@" + name_id + "=" + idVal + "]\")).getText().then(function(txt)" \
        #                                                                   "{expect(txt.toLowerCase()).toContain(\'" + output_string + "\'.toLowerCase()" \
        #                                                                                                                               ");});"
        # else:
        #   add_output=''
        final = final + m[:-3] + default_var + '''
                         //extract();
                         extractModal();
                        });
                        ''' + script.read() + '});'
        whole_script_1 = final
      else:  # If the user clicks complete user flow
        if (popup != None and popup != "other" and name_id != None and output_string != None and idVal != None):
          if (popup == "browser-pop-up"):
            add_output = '''browser.switchTo().alert().then((alrt) => {
              alrt.getText().then(function(txt)
            {expect(txt.toLowerCase()).toContain("''' + output_string + '''".toLowerCase())});alrt.accept();});'''
          else:
            add_output = "element(by.xpath(\"//*[@" + name_id + "='" + idVal + "']\")).getText().then(function(txt)" \
                                                                               "{expect(txt.toLowerCase()).toContain(\'" + output_string + "\'.toLowerCase()" \
                                                                                                                                           ");});"
        else:
          add_output = ''

        final = final + m[:-3] + default_var + \
                add_output + '''//extract();
               //extractModal();
              });
              ''' + script.read() + '});'
        whole_script_1 = final

    f1.write(whole_script_1)

  logger.info('Test file generated')
  # f.flush()
  # f1 = open(r"base_location\uniqure_output_file.txt", "w+")
  # logger.error(e)
  return jsonify({"status": "success"})  # Not requried

@app.route("/testReportGeneration")
def testReportGeneration():  # Output report
  # folderName = request.args.get("folderName")
  open_file = open(
    base_location+"\\" + outerFolderName + '\\' + folderName + "\oldFiles\multiple_conf.log")
  read_f = (open_file.readlines())
  passStr = '[32m.[0m'
  failStr = '[31mF[0m'
  resultList = []
  header = []
  header1 = []
  data = []
  js = {}
  # passStr=".F."
  for wholeStr in read_f:
    if (wholeStr.__contains__(passStr) or wholeStr.__contains__(failStr)):
      outputLine = [(wholeStr[i:i + 10]) for i in range(0, len(wholeStr), 10)]
      for result in outputLine:
        if result == passStr:
          resultList.append('Pass')
        elif result == failStr:
          resultList.append('Fail')

    # if(r.__contains__(failStr)):

  file_name = base_location+"\\" + outerFolderName + '\\' + folderName + "\\sample1.xlsx"  # Path to your file
  df = pd.read_excel(file_name)
  # df["result"]=resultList[0]
  r, c = df.shape
  df['result'] = ""
  ind = 0
  for row in range(r):
    df["result"][row] = resultList[ind]
    ind += 1
  df.to_excel(base_location+"\\" + outerFolderName + '\\' + folderName + "\\new_excel_sheet.xlsx",
              index=False)
  x = df.head()
  time.sleep(1)
  excel_sheet = pd.read_excel(
    base_location+"\\" + outerFolderName + '\\' + folderName + "\\new_excel_sheet.xlsx")
  js["data"] = excel_sheet.to_dict(orient="records")
  for colName in excel_sheet.head():
    header.append(colName)
  # for h in header:
  js["headers"] = header
  totPass = resultList.count('Pass')
  totFail = resultList.count('Fail')
  testResult = [totPass, totFail]
  js["testResult"] = testResult
  # for inew in x:
  #   header.append(inew)
  #   js1=df[inew].to_json()
  #   js2=dict(js1)
  # for k,v in df[inew].to_json().items():
  # js={"inew",df[inew]}
  # data.append(js)
  return jsonify(js)

@app.route("/testDescription")  # Test case description
def testDescription(): #Based on elements ,  browser.get and etc test description will be generated
  # folderName = request.args.get("folderName")
  # folderName="XYZ_Deposit"
  global outerFolderName
  global folderName
  # outerFolderName="1View_Tool"
  # folderName="Test48"
  open_file = open(base_location+"\\" + outerFolderName + '\\' + folderName + "\oldFiles\spec.js")
  # new_file = open(r"base_location\\" + outerFolderName + '\\' + folderName + "\oldFiles\TestDescription.html",
  #                 "w+")
  lines = open_file.readlines()
  full_line = ''
  flag = 0
  # global sentence
  sentence = ''
  for line in lines:
    if line.__contains__("extractModal();"):
      break
    if line.strip().startswith("browser.get("):
      flag = 1
    if flag == 1:
      full_line += line

  number = 0
  for sline in full_line.split(";"):
    if sline.__contains__("browser.get("):  # Get Url
      number += 1
      url_line=sline.split("browser.get(")[1].split(")")[0]
      if len(url_line)<50:
        sentence += str(number) + ". Open browser and navigate to URL: " + str(
          url_line) + " <br>"
      else:
        sentence += str(number) + ". Open browser and navigate to URL: " + "<br>"+ str(
          url_line)+ " <br>"
    # if sline.__contains__("browser.sleep("):
    #   sentence+="Browser will wait for "+sline.split("browser.sleep(")[1].split(")")[0]+"MilliSeconds<br>"
    if sline.__contains__("until.presenceOf") == False and \
      sline.__contains__("by.xpath(") and \
      sline.strip().endswith("click()"):  # button click
      if sline.__contains__("@id="):
        number += 1
        sentence += str(number) + ". Click on the button with id " + sline.split("@id=")[1].split("]")[0]+"<br>"
      elif sline.__contains__("@name="):
        number += 1
        sentence += str(number) + ". Click on the button with name " + sline.split("@name=")[1].split("]")[0]+"<br>"
      elif sline.__contains__(".get("):
        number += 1
        x = str(sline.split(".get(")[1].split(")")[0]).strip()
        if x.endswith("1"):
          inBetween = str(floor(int(x) / 10)) + "1st"
        elif x.endswith("2"):
          inBetween = str(floor(int(x) / 10)) + "2nd"
        elif x.endswith("3"):

          inBetween = str(floor(int(x) / 10)) + "3rd"
        else:
          inBetween = x + "th"
        sentence += str(number) + ". Click on " + inBetween + " button in the current page<br>"

    elif sline.__contains__("until.presenceOf") == False and \
      sline.__contains__("by.xpath(") and \
      sline.__contains__("click()") == False and \
      sline.__contains__("sendKeys"):  # Sendkey
      if sline.__contains__("@id="):
        number += 1
        if sline.__contains__(".get("):
          x = str(sline.split(".get(")[1].split(")")[0]).strip()
          if x.endswith("1"):
            inBetween = str(floor(int(x) / 10)) + "1st"
          elif x.endswith("2"):
            inBetween = str(floor(int(x) / 10)) + "2nd"
          elif x.endswith("3"):
            inBetween = str(floor(int(x) / 10)) + "3rd"
          else:
            inBetween = x + "th"
            if sline.lower().__contains__("password"):
              sentence += str(number) + ". Enter value " + "******" + " on the" + inBetween + "field with id " + sline.split("@id=")[1].split("]")[0] + "<br>"
            else:
              sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                0] + " on the"+inBetween+"field with id " + sline.split("@id=")[1].split("]")[0] + "<br>"
        else:
          if sline.lower().__contains__("password"):
            sentence += str(number) + ". Enter value " + "******" + " in the field with id " + sline.split("@id=")[1].split("]")[0] + "<br>"
          else:
            sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
              0] + " in the field with id " + sline.split("@id=")[1].split("]")[0] + "<br>"
      elif sline.__contains__("@name="):
        number += 1
        if sline.__contains__(".get("):
          x = str(sline.split(".get(")[1].split(")")[0]).strip()
          if x.endswith("1"):
            inBetween = str(floor(int(x) / 10)) + "1st"
          elif x.endswith("2"):
            inBetween = str(floor(int(x) / 10)) + "2nd"
          elif x.endswith("3"):
            inBetween = str(floor(int(x) / 10)) + "3rd"
          else:
            inBetween = x + "th"
            if sline.lower().__contains__("password"):
              sentence += str(number) + ". Enter value " + "******"+ " in the field with name " + sline.split("@name=")[1].split("]")[0] + "<br>"
            else:
              sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                0] + " in the field with name " + sline.split("@name=")[1].split("]")[0] + "<br>"
        else:
          if sline.lower().__contains__("password"):
            sentence += str(number) + ". Enter value " + "*****" + " in the field with name " + sline.split("@name=")[1].split("]")[0] + "<br>"
          else:
            sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
            0] + " in the field with name " + sline.split("@name=")[1].split("]")[0] + "<br>"
      elif sline.__contains__(".get(") and sline.__contains__("@id=")==False:
        number += 1
        x = str(sline.split(".get(")[1].split(")")[0]).strip()
        if x.endswith("1"):
          inBetween = str(floor(int(x) / 10)) + "1st"
        elif x.endswith("2"):
          inBetween = str(floor(int(x) / 10)) + "2nd"
        elif x.endswith("3"):
          inBetween = str(floor(int(x) / 10)) + "3rd"
        else:
          inBetween = x + "th"
        # sentence += "Enter value " + sline.split("sendKeys(")[1].split(")")[0] + " in the field with id " + \
        #             inBetween+" in the field from the webpage<br>"
        sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
          0] + " on " + inBetween + " field in the current page<br>"
  # new_file.write(sentence)
  # print(sentence)
  connections = db.folderName.find()
  myquery = {"appName": outerFolderName, "scenarioName": folderName}
  newvalues = {"$set": {"test_Description": sentence}}
  db.folderName.update_many(myquery, newvalues)

def callNpm():
  npm_batch='''START "cmd \B cd D:\XYZ\XYZ_Bank" && npm install jsdom'''
  if os.path.isdir(base_location + "\\" + outerFolderName+"\\node_modules")==False:
    trigger_npm = open(base_location + "\\" + outerFolderName+"\\npm_batch.bat","w+")
    trigger_npm.write(npm_batch)
    trigger_npm.close()
    subprocess.call(base_location+'\\' + outerFolderName +"\\npm_batch.bat")

@app.route("/checkFileExist")  # Check whether folder exist inside ProtractorWebdriverTest/outerFolder
def checkFileExist(): #checks the scenerio name folder exists or not
  global folderName
  global outerFolderName
  folderName = request.args.get("folderName")
  outerFolderName = request.args.get("outerFolder")

  newFile = []
  if os.path.isdir(base_location + "\\" + outerFolderName) == False:
    os.mkdir(base_location + "\\" + outerFolderName)
  for file in os.listdir(base_location + "\\" + outerFolderName):
    if not file.lower().__contains__("."):
      newFile.append(file.lower())


  if folderName.lower() in newFile:
    return jsonify("Fail")
  else:

    return jsonify("Success")

@app.route("/checkFolderExist")  # Check folder exist inside ProtractorWebdriverTest
def checkFolderExist(): # checks whether app name exists or not
  # global folderName
  folderName1 = request.args.get("folderName")
  newFile = []
  if os.path.isdir(base_location) == False:
    os.mkdir(base_location)
  for file in os.listdir(base_location):
    if not file.lower().__contains__("."):
      newFile.append(file.lower())

  if folderName1.lower() in newFile:
    return jsonify("Fail")
  else:
    return jsonify("Success")

@app.route("/getSavedTestCases")  # Fetch data from db
def getSavedTestCases(): # fetching app name scenerio name and description from db
  connections = db.folderName.find()
  outDict = {}
  finalList = []
  finalDict = {}
  for collection in connections:
    outDict["scenario_Name"] = collection["scenarioName"]
    outDict["application_Name"] = collection["appName"]
    outDict["test_Description"] = collection["test_Description"]
    finalList.append(outDict)
    outDict = {}
  finalDict["data"] = finalList
  finalDict["headers"] = ["application_Name", "scenario_Name", "test_Description", "Actions"]
  return jsonify(finalDict)

@app.route("/runSavedTestCase")  # Run already exist data
def runSavedTestCase(): # run the test case triggering batch file using subprocess
  application_name = request.args.get("application_name")
  scenario_name = request.args.get("scenario_name")
  connections = db.folderName.find()
  for connection in connections:
    if connection["appName"] == application_name and connection["scenarioName"] == scenario_name:
      appPath = connection["path"]
      scenarioPath = connection["innerpath"]
      if os.path.isdir(appPath):
        if os.path.isdir(scenarioPath):
          if os.path.exists(scenarioPath + "\\Batchfiles\\Test_batch.bat"):
            setPath = scenarioPath + "\\Batchfiles\\Test_batch.bat"
            subprocess.call(setPath)
            conf_file = open(scenarioPath + "\\OldFiles\\conf.log").read()
            return jsonify(conf_file)
          else:
            setPath = scenarioPath + "\\Batchfiles\\batch.bat"
            subprocess.call(setPath)
            conf_file = open(scenarioPath + "\\confApr9.log").read()
            return jsonify(conf_file)
        else:
          return jsonify("scenarioFail")
      else:
        return jsonify("appFail")

@app.route("/download_spec_scenario")  # Download specific scenario
def download_spec_scenario():
  # folderName = request.args.get("folderName")
  application_name = request.args.get("application_name")
  scenario_name = request.args.get("scenario_name")
  connections = db.folderName.find()
  for connection in connections:
    if connection["appName"] == application_name and connection["scenarioName"] == scenario_name:
      appPath = connection["path"]
      scenarioPath = connection["innerpath"]
      if os.path.isdir(appPath):
        if os.path.isdir(scenarioPath):
          spec_file = open(scenarioPath + "\oldFiles\spec.js")
          headers = []
          values = []
          open_spec = spec_file.read()
          json_format = {}
          final_json = {}
          headerList = []
          final_list = []
          split_file = open_spec.split("browser.driver.manage().window().maximize();")[1].split("extract()")[0].split(
            ";")
          for elm in split_file:
            if (elm.find("element") >= 0):
              if (elm.find("sendKeys") >= 0):

                fin_elm = elm.split("xpath")[1].split(".sendKeys")[0]
                headers.append(fin_elm)
                # fin_val = elm.split("sendKeys")[1].split("'")[1]
                if elm.split("sendKeys")[1].__contains__("'"):
                  fin_val = elm.split("sendKeys")[1].split("'")[1]
                elif elm.split("sendKeys")[1].__contains__('"'):
                  fin_val = elm.split("sendKeys")[1].split('"')[1]
                values.append(fin_val)
          workbook = xlwt.Workbook()
          sheet = workbook.add_sheet(scenario_name)

          header_style = xlwt.easyxf('font: bold 1')
          val_style = xlwt.easyxf('font: bold 0')

          for index in range(len(headers)):

            if headers[index] not in headerList:
              headerList.append(headers[index])
            json_format[headers[index]] = values[index]

          final_list.append(json_format)
          final_list_1 = final_list[::-1]

          final_json = {"headers": headerList, "data": final_list_1}
          return jsonify(final_json)
        else:
          return "scenarioFail"
      else:
        return "appFail"

@app.route("/uploadExcelScenario", methods=['GET', 'POST'])  # Upload specific sceanrio
def uploadExcelScenario():
  # folderName = request.args.get("folderName")
  import pandas as pd
  file = request.get_data()
  str_file = file.decode("utf-8")
  # temp_file = list(eval(str_file))
  print(str_file)
  temp_List = []
  if type(list(eval(str_file))[0]) == str:
    return jsonify("failure")
  else:
    temp_file = list(eval(str_file))
    df = pd.DataFrame.from_dict(temp_file)
  application_name = request.args.get("application_name")
  scenario_name = request.args.get("scenario_name")
  connections = db.folderName.find()
  for connection in connections:
    if connection["appName"] == application_name and connection["scenarioName"] == scenario_name:
      appPath = connection["path"]
      scenarioPath = connection["innerpath"]
      if os.path.isdir(appPath):
        if os.path.isdir(scenarioPath):
          if os.path.exists(scenarioPath + "\\oldFiles"):
            df.to_excel(scenarioPath + "\\oldFiles" + '\sample1.xlsx',
                        index=False)
            time.sleep(1)


        else:
          return "scenarioFail"
      else:
        return "appFail"

  returnVal = uploadUI(application_name, scenario_name)
  return returnVal

@app.route("/uploadUI")  # Generate excel sheet while uploading
def uploadUI(application_name, scenario_name):
  replace_spec = ''
  head_list = []
  cell_list = []
  # folderName = request.args.get("folderName")
  connections = db.folderName.find()
  for connection in connections:
    if connection["appName"] == application_name and connection["scenarioName"] == scenario_name:
      appPath = connection["path"]
      scenarioPath = connection["innerpath"]
      if os.path.isdir(appPath):
        if os.path.isdir(scenarioPath):
          if os.path.exists(scenarioPath + "\\oldFiles\\\spec.js"):
            spec_file = open(
              scenarioPath + "\\oldFiles\\\spec.js")
            path = scenarioPath + "\\oldFiles\\\sample1.xlsx"
            new_spec_file = open(
              scenarioPath + "\\oldFiles\\multiple_scenario_file.js",
              "w+")
            new_conf_file = open(scenarioPath + "\\oldFiles\\multiple_conf_file.js", 'w+')
            if os.path.isdir(scenarioPath + "\\Batchfiles") == False:
              os.mkdir(scenarioPath + "\\Batchfiles")
            batch_file = open(scenarioPath + "\\Batchfiles\\multiple_batch_file.bat", "w+")
            wb = xlrd.open_workbook(path)
            sheet = wb.sheet_by_index(0)
            start_row = 1
            end_row = sheet.nrows
            num_of_columns = sheet.ncols
            open_file = spec_file.read()
            split_file = open_file.split("extract")[0]
            if (split_file.endswith("//")):
              split_file = split_file[:-2] + "});"
            sec_test = "it('Test 1', function() {" + \
                       "browser.driver.manage().window().maximize();" + \
                       split_file.split("browser.driver.manage().window().maximize();")[1]

            for r in range(start_row, end_row):
              head_list = []
              cell_list = []
              for c in range(num_of_columns):
                cell_obj = sheet.cell_value(r, c)
                head_obj = sheet.cell_value(0, c)
                head_list.append(head_obj)
                cell_list.append(cell_obj)
              for line in sec_test.split(";"):
                for idx, header in enumerate(head_list):
                  if (
                    line.__contains__(header) and not line.__contains__("browser.wait(until") and not line.__contains__(
                    ".clear(")):
                    if (type(cell_list[idx]) == int or type(cell_list[idx]) == float):
                      line = line.split("sendKeys")[0] + "sendKeys(" + str(int(cell_list[idx])) + ");"
                    else:
                      line = line.split("sendKeys")[0] + "sendKeys('" + str(cell_list[idx]) + "');"
                    break
                replace_spec += line + ";"
            new_spec = '''const jsdom = require("jsdom");
                         const fs = require('fs');
                         const path = require('path');
                         describe('Protractor Demo App', function() { ''' + str(replace_spec) + "});"
            new_conf = '''
                    exports.config = {
     framework: 'jasmine',
     seleniumAddress: 'http://localhost:4444/wd/hub',
     specs: ['multiple_scenario_file.js'],
     commonCapabilities: {
      'browserName': 'chrome',
      'chromeOptions': {
                                                 'args': [ '--disable-gpu','-disable-dev-shm-usage','--no-sandbox','-disable-popup-blocking','--start-maximized','--disable-web-security','--allow-running-insecure-content','--disable-infobars']
                                 }
    },
    onPrepare: function() {
       return browser.takeScreenshot();
     },
    getPageTimeout: 20000,
                 allScriptsTimeout: 3000000,
                 jasmineNodeOpts: {defaultTimeoutInterval: 3000000},
                 defaultTimeoutInterval: 3000000
   }

             '''
            # START /B cmd "cd base_location\" && protractor base_location\oldFiles\multiple_conf.js > base_location\oldFiles\multiple_conf.log
            multipleBatchFile = "START /B cmd " + scenarioPath + " && protractor " + scenarioPath + "\\oldFiles\\multiple_conf_file.js >" + scenarioPath + "\\oldFiles\\multiple_conf_file.log"
            new_conf_file.write(new_conf)
            new_spec_file.write(new_spec)
            new_conf_file.close()
            new_spec_file.close()
            batch_file.write(multipleBatchFile)
            batch_file.close()
            subprocess.call(scenarioPath + '\\Batchfiles\\multiple_batch_file.bat')
            conf_log = open(scenarioPath + "\\oldFiles\\multiple_conf_file.log").read()
            logVal = multipletestReportGeneration(scenarioPath)
            time.sleep(1)

            return jsonify(logVal)

        else:
          return "scenarioFail"
      else:
        return "appFail"

@app.route("/multipletestReportGeneration")  # Run multiple test case
def multipletestReportGeneration(scenarioPath):
  # folderName = request.args.get("folderName")
  open_file = open(scenarioPath + "\\oldFiles\\multiple_conf_file.log")
  log_file = open(scenarioPath + "\\oldFiles\\multiple_conf_file.log").read()
  read_f = (open_file.readlines())
  passStr = '[32m.[0m'
  failStr = '[31mF[0m'
  resultList = []
  header = []
  header1 = []
  data = []
  js = {}
  # passStr=".F."
  for wholeStr in read_f:
    if (wholeStr.__contains__(passStr) or wholeStr.__contains__(failStr)):
      outputLine = [(wholeStr[i:i + 10]) for i in range(0, len(wholeStr), 10)]
      for result in outputLine:
        if result == passStr:
          resultList.append('Pass')
        elif result == failStr:
          resultList.append('Fail')

    # if(r.__contains__(failStr)):

  file_name = scenarioPath + "\\oldFiles\\sample1.xlsx"  # Path to your file
  df = pd.read_excel(file_name)
  # df["result"]=resultList[0]
  r, c = df.shape
  df['result'] = ""
  ind = 0
  for row in range(r):
    df["result"][row] = resultList[ind]
    ind += 1
  df.to_excel(scenarioPath + "\\oldFiles\\new_excel_sheet.xlsx", index=False)
  x = df.head()
  time.sleep(1)
  excel_sheet = pd.read_excel(scenarioPath + "\\oldFiles\\new_excel_sheet.xlsx")
  js["data"] = excel_sheet.to_dict(orient="records")
  for colName in excel_sheet.head():
    header.append(colName)
  # for h in header:
  js["headers"] = header
  totPass = resultList.count('Pass')
  totFail = resultList.count('Fail')
  testResult = [totPass, totFail]
  js["testResult"] = testResult
  js["log"] = log_file
  return js

@app.route("/deleteScenario")  # Delete selected scenario
def deleteScenario():
  application_name = request.args.get("application_name")
  scenario_name = request.args.get("scenario_name")
  my_query = {"appName": application_name, "scenarioName": scenario_name}
  connections = db.folderName.find()
  for connection in connections:
    if connection["appName"] == application_name and connection["scenarioName"] == scenario_name:
      path = connection["path"]
      innerPath = connection["innerpath"]
      shutil.rmtree(innerPath, ignore_errors=True)   #delete the entire directory tree
  db.folderName.delete_many(my_query)# delete  whole folder
  return jsonify("success")

@app.route("/moveImage")  # Move Screenshot from our folder to UI
def moveImage():
  application_name = request.args.get("application_name")
  scenario_name = request.args.get("scenario_name")
  connections = db.folderName.find()
  fileName = []
  finalJson = {}
  destination_folder = ui_location+"\\src\\assets\\images\\" + scenario_name
  if os.path.isdir(destination_folder) == False:
    os.mkdir(destination_folder)
  for connection in connections:
    if connection["appName"] == application_name and connection["scenarioName"] == scenario_name:
      appPath = connection["path"]
      scenarioPath = connection["innerpath"]
      # destination_folder="D:\\TestAutomationTool\\src\\assets\\images"
      print(appPath, 'appPath')
      print(scenarioPath, 'scenarioPath')
      if os.path.isdir(appPath):
        # print(appPath, 'appPath1')
        if os.path.isdir(scenarioPath):
          # print(scenarioPath, 'scenarioPath1')
          for img in os.listdir(scenarioPath + "\\oldFiles"):
            if img.endswith(".png"):
              fileName.append(img)
              # print(img)
              shutil.copy(scenarioPath + "\\oldFiles\\" + img, destination_folder)#coping images using shutil
          finalJson["fileName"] = fileName
          finalJson["status"] = "success"
          return jsonify(finalJson)
        else:
          return "scenarioFail"
      else:
        return "appFail"


@app.route("/getAllScenarios")  # Get all saved scenario
def getAllScenarios():
  collections = db.folderName.find()# fetching all scenerios
  string = []
  for collection in collections:
    string.append(collection['scenarioName'] + "[" + collection["appName"] + "]")#appending app name and scenerio name as string

  return jsonify(string)
  # return jsonify(["Deposit_flow[XYZ_Bank]","Withdraw_flow[XYZ_Bank]", "OpenAccount_flow[XYZ_Bank]", "CreateOrderFlow[CRST]", "Download_dashboard_charts_Flow[LCaaS]", "MasterInvFlow[LCaaS]"])

@app.route("/checkTestRunExist")
def checkTestRunExist():  # Check whether folder exists for specific scenario
  value = request.args.get("testName") # if existed name gives it will fail else its success
  # value="testRun1"
  connections = db.testRun.find()
  flag = 0
  for connection in connections:
    if value == connection["test_name"]:
      # if value == 'testRun1':
      flag = 1
      break

  if flag == 1:
    return jsonify("Fail")
  else:
    return jsonify("Success")

@app.route("/getSavedTestRuns")
def getSavedTestRuns():  # Fetch all saved multiple scenario
  connections = db.testRun.find()
  dictionary = {}
  finalList = []
  finalDict = {}
  for connection in connections:
    dictionary["test_name"] = connection["test_name"]
    dictionary["test_scenarios_used"] = connection["test_scenarios_used"]
    finalList.append(dictionary)#appending all saved test sceneros  in finallist
    dictionary = {}

  finalDict["headers"] = ["test_name", "test_scenarios_used", "Actions"]
  finalDict["data"] = finalList
  return jsonify(finalDict)

@app.route("/testRunCreation")
def testRunCreation(): # Making a test job using selected scenerios
  scenarios = request.args.get("scenarios")
  testRunName = request.args.get("testRunName")
  scenario = scenarios
  # .replace(",", ", ")
  if os.path.isdir(base_location + "\\testRun") == False:
    os.mkdir(base_location + "\\testRun")
  os.mkdir(base_location+"\\testRun\\" + testRunName)
  split_scenario = scenarios.split(",")
  dbData = [{
    "test_name": testRunName,
    "test_scenarios_used": scenario,
    "path": base_location + "\\testRun\\" + testRunName
  }]
  db.testRun.insert_many(dbData)# storing in db with test job name and scenerio names and path
  open(base_location + "\\testRun\\" + testRunName + "\\" + "conf.log", "w+").close()
  log = runSavedTestRun1(testRunName)

  if log == 'Success':
    js = { "status": "success", "log": "Test Job Created.." }
  else:
    js = { "status": "success", "log": log }
  return jsonify(js)

@app.route("/runSavedTestRun")
def runSavedTestRun():   # run the test 
  testRunName = request.args.get("testName")
  testCount = 0
  requiredLines = ''
  fulString = '''
            const jsdom = require("jsdom");
            const fs = require('fs');
            const path = require('path');
            describe('Protractor Demo App', function() {
   '''
  # scenarios="Deposit_Flow[XYZ_BANK], Bank_flow[XYZ_Bank]"
  collections = db.testRun.find()
  for collection in collections:
    if collection["test_name"] == testRunName:
      scenarios = collection["test_scenarios_used"]
  split_scenario = scenarios.split(",")
  if os.path.isdir(base_location + "\\testRun") == False:
    os.mkdir(base_location + "\\testRun")
  if os.path.isdir(base_location + "\\testRun\\" + testRunName) == False:
    os.mkdir(base_location + "\\testRun\\" + testRunName)

  scenario_file = open(base_location + "\\testRun\\" + testRunName + "\\" + testRunName + ".js", "w+")
  conf_file = open(base_location + "\\testRun\\" + testRunName + "\\" + "conf" + ".js", "w+")

  conf_string = '''
     exports.config = {
       framework: 'jasmine',
       seleniumAddress: 'http://localhost:4444/wd/hub',
       specs: ["''' + testRunName + '''.js"],
       commonCapabilities: {
        'browserName': 'chrome',
        'chromeOptions': {
                                                   'args': [ '--disable-gpu','-disable-dev-shm-usage','--no-sandbox','-disable-popup-blocking','--start-maximized','--disable-web-security','--allow-running-insecure-content','--disable-infobars']
                                   }
      },
      onPrepare: function() {
         return browser.takeScreenshot();
       },
      getPageTimeout: 20000,
                   allScriptsTimeout: 3000000,
                   jasmineNodeOpts: {defaultTimeoutInterval: 3000000},
                   defaultTimeoutInterval: 3000000
     }

     '''
  conf_file.write(conf_string)
  conf_file.close()
  batch_file = open(base_location + "\\testRun\\" + testRunName + "\\" + "batch.bat", "w+")
  script1 = '''START /B cmd "cd '''+ base_location+'''\\" && protractor '''+ base_location+\
            '''\\''' + "testRun" + '\\' + testRunName + '''\\conf.js > '''+ base_location+'''\\''' +\
            "testRun" + '\\' + testRunName + '''\\conf.log'''
  batch_file.write(script1)
  batch_file.close()
  for scene_app in split_scenario:
    testCount += 1

    flag = 0
    scenario_name = scene_app.split("[")[0].strip()
    application_name = scene_app.split("[")[1].split("]")[0].strip()
    path_var = base_location + "\\" + application_name + "\\" + scenario_name
    print(path_var)
    if os.path.isdir(path_var) == False:
      return "Fail"
    else:
      specFile = open(path_var + "\\OldFiles\\spec.js").readlines()
      print(specFile)
      for lines in specFile:
        if flag == 1:
          requiredLines += lines

        if lines.__contains__("describe('Protractor Demo App'") and lines.__contains__("it("):
          flag = 1
          requiredLines += "it(" + lines.split("it(")[1].replace("1", str(testCount))

        elif lines.__contains__("it('Test"):
          flag = 1
          requiredLines += "it(" + lines.split("it(")[1].replace("1", str(testCount))

        if lines.__contains__("extract();"):
          requiredLines += "});"
          flag = 0
          break

  fulString += requiredLines + "});"
  scenario_file.write(fulString)
  scenario_file.close()

  subprocess.call(base_location+'\\' + "testRun" + '\\' + testRunName + '\\batch.bat')# trigger batch file using subprocess
  log_file = open(base_location+'\\' + "testRun" + '\\' + testRunName + '\\conf.log').read()
  return jsonify(log_file)

@app.route("/scheduledJobList")
def scheduledJobList(): #fetching scheduled jobs
  cursor=db.testRun.find()
  totList=[]
  for collection in cursor:
    totList.append(collection["test_name"])
  return jsonify(totList)

@app.route("/runSavedTestRun1")
def runSavedTestRun1(testRunName):#run combine scenerio
  # testRunName = request.args.get("testName")
  testCount = 0
  requiredLines = ''
  fulString = '''
            const jsdom = require("jsdom");
            const fs = require('fs');
            const path = require('path');
            describe('Protractor Demo App', function() {
   '''
  # scenarios="Deposit_Flow[XYZ_BANK], Bank_flow[XYZ_Bank]"
  collections = db.testRun.find()
  for collection in collections:
    if collection["test_name"] == testRunName:
      scenarios = collection["test_scenarios_used"]
  split_scenario = scenarios.split(",")
  if os.path.isdir(base_location + "\\testRun") == False:
    os.mkdir(base_location + "\\testRun")
  if os.path.isdir(base_location + "\\testRun\\" + testRunName) == False:
    os.mkdir(base_location + "\\testRun\\" + testRunName)

  scenario_file = open(base_location + "\\testRun\\" + testRunName + "\\" + testRunName + ".js", "w+")
  conf_file = open(base_location + "\\testRun\\" + testRunName + "\\" + "conf" + ".js", "w+")
  conf_string = '''
     exports.config = {
       framework: 'jasmine',
       seleniumAddress: 'http://localhost:4444/wd/hub',
       specs: ["''' + testRunName + '''.js"],
       commonCapabilities: {
        'browserName': 'chrome',
        'chromeOptions': {
                                                   'args': [ '--disable-gpu','-disable-dev-shm-usage','--no-sandbox','-disable-popup-blocking','--start-maximized','--disable-web-security','--allow-running-insecure-content','--disable-infobars']
                                   }
      },
      onPrepare: function() {
         return browser.takeScreenshot();
       },
      getPageTimeout: 20000,
                   allScriptsTimeout: 3000000,
                   jasmineNodeOpts: {defaultTimeoutInterval: 3000000},
                   defaultTimeoutInterval: 3000000
     }

     '''
  conf_file.write(conf_string)
  conf_file.close()
  batch_file = open(base_location + "\\testRun\\" + testRunName + "\\" + "batch.bat", "w+")
  script1 = '''START /B cmd "cd '''+ base_location+'''\\" && protractor '''+ base_location+'''\\''' + "testRun" + '\\' + testRunName + '''\\conf.js > '''+ base_location+'''\\''' + "testRun" + '\\' + testRunName + '''\\conf.log'''
  for scene_app in split_scenario:
    testCount += 1

    flag = 0
    scenario_name = scene_app.split("[")[0].strip()
    application_name = scene_app.split("[")[1].split("]")[0].strip()
    path_var = base_location + "\\" + application_name + "\\" + scenario_name
    if os.path.isdir(path_var) == False:
      return "Fail"
    else:
      specFile = open(path_var + "\\OldFiles\\spec.js").readlines()
      for lines in specFile:
        if flag == 1:
          requiredLines += lines

        if lines.strip().startswith("it("):
          flag = 1
          requiredLines += "it(" + lines.split("it(")[1].replace("1", str(testCount))

        if lines.__contains__("extract();"):
          requiredLines += "});"
          flag = 0
          break

  fulString += requiredLines + "});"
  scenario_file.write(fulString)
  scenario_file.close()
  batch_file.write(script1)
  batch_file.close()
  #subprocess.call(base_location+'\\' + "testRun" + '\\' + testRunName + '\\batch.bat')
  #log_file = open(base_location+'\\' + "testRun" + '\\' + testRunName + '\\conf.log').read()
  # return log_file
  return "Success"

@app.route("/download_run_scenario")
def download_run_scenario():
  #folderName = request.args.get("folderName")
  testRunName = request.args.get("testName")
  flag = 0
  collections = db.testRun.find()
  complete_list = []
  complete_json_format = {}
  for collection in collections:
    if collection["test_name"] == testRunName:
      scenarios = collection["test_scenarios_used"]
      path = collection["path"]
      flag = 1
  if flag == 1:
    split_scenario = scenarios.split(",")
    for app_scenario in split_scenario:
      app_name = app_scenario.split("[")[1].split("]")[0]
      scenario_name = app_scenario.split("[")[0]
      spec_file = open(base_location+"\\" + app_name.strip() + '\\' + scenario_name.strip() + "\oldFiles\spec.js")
      headers = []
      values = []
      open_spec = spec_file.read()
      json_format = {}
      final_json = {}
      headerList = []
      final_list = []
      split_file = open_spec.split("browser.driver.manage().window().maximize();")[1].split("extract()")[0].split(";")
      for elm in split_file:
        if (elm.find("element") >= 0):
          if (elm.find("sendKeys") >= 0):
            fin_elm = elm.split("xpath")[1].split(".sendKeys")[0]
            headers.append(fin_elm)
            # fin_val = elm.split("sendKeys")[1].split("'")[1]
            if elm.split("sendKeys")[1].__contains__("'"):
              fin_val = elm.split("sendKeys")[1].split("'")[1]
            elif elm.split("sendKeys")[1].__contains__('"'):
              fin_val = elm.split("sendKeys")[1].split('"')[1]
            values.append(fin_val)
      workbook = xlwt.Workbook()
      sheet = workbook.add_sheet(testRunName)
      # Specifying style
      header_style = xlwt.easyxf('font: bold 1')
      val_style = xlwt.easyxf('font: bold 0')

      for index in range(len(headers)):
        if headers[index] not in headerList:
          headerList.append(headers[index])
        json_format[headers[index]] = values[index]
      final_list.append(json_format)
      # final_list_1 = final_list[::-1]
      final_json = {"headers": headerList, "data": final_list}
      complete_list.append(final_json)
      final_json = {}
      headerList = []
      final_list_1 = []
      final_list = []

    for index, val in enumerate(complete_list):
      complete_json_format[str(index + 1)] = val
    return jsonify(complete_json_format)
  else:
    return "appFail"


@app.route("/uploadExcelRun", methods=['GET', 'POST'])
def uploadExcelRun():
  # folderName = request.args.get("folderName")
  import pandas as pd
  global mulipleScenarioPath
  file = request.get_data()
  final_val = []
  str_file = file.decode("utf-8")
  temp_List = []
  sheetLength = 0
  print("str_file", str_file)
  json_format = json.loads(str_file)
  print("json_format", json_format)

  # Checking the muliple scenario file exist or not, if exists remove the file
  testRunName = request.args.get("testName")
  connections = db.testRun.find()
  for connection in connections:
    if connection["test_name"] == testRunName:
      mulipleScenarioPath = connection["path"]
      if os.path.exists(mulipleScenarioPath+"\\combined_multi_scenario_file.js"):
        p = os.path.join(mulipleScenarioPath+"\\combined_multi_scenario_file.js")
        os.remove(p)

  for k, v in json_format.items():
    v1 = json.dumps(v['data'])
    if type(list(eval(v1))[0]) == "str":
      return jsonify("failure")
    else:
      temp_file = list(eval(v1))
      df = pd.DataFrame.from_dict(temp_file)
    testRunName = request.args.get("testName")
    connections = db.testRun.find()
    for connection in connections:
      if connection["test_name"] == testRunName:
        mulipleScenarioPath = connection["path"]
        if os.path.isdir(mulipleScenarioPath):
          df.to_excel(mulipleScenarioPath + '\\sample.xlsx',
                      index=False)
          time.sleep(1)
          append_file = open(mulipleScenarioPath + "\\append_spec.txt", "a+")
          returnVal = uploadUIRun(testRunName, sheetLength)
          final_val.append(returnVal)
          sheetLength += 1
        else:
          return "appFail"
  final_js = {}
  testResult = [0, 0]
  f_log = ''
  for idx, f_val in enumerate(final_val):
    final_js[str(idx)] = f_val
    testResult = [sum(i) for i in zip(testResult, f_val["testResult"])]
    for log_val in f_val["log"]:
      f_log += log_val
  final_js["testResult"] = testResult
  final_js["log"] = f_log

  comb_multi_file = open(mulipleScenarioPath+"\\combined_multi_scenario_file.js","a+")
  comb_multi_file.write("});")
  return jsonify(final_js)

# @app.route("/uploadUIRun")
# def uploadUIRun(testRunName, sheetLength):
#   replace_spec = ''
#   head_list = []
#   cell_list = []
#   filePath = ''
#   # folderName = request.args.get("folderName")
#   connections = db.testRun.find()
#   for connection in connections:
#     if connection["test_name"] == testRunName:
#       appPath = connection["path"]
#       scenario = connection["test_scenarios_used"]
#       split_scenario = scenario.split(",")[sheetLength]
#       appNameRun = split_scenario.split("[")[1].split("]")[0]
#       scenarioNameRun = split_scenario.split("[")[0]
#       collections = db.folderName.find()
#       for collection in collections:
#         if collection["appName"] == appNameRun and collection["scenarioName"] == scenarioNameRun:
#           filePath = collection["innerpath"]
#       if os.path.isdir(filePath):
#         if os.path.exists(filePath + "\\OldFiles\\spec.js"):
#           spec_file = open(
#             filePath + "\\OldFiles\\spec.js")
#           path = appPath + "\\sample.xlsx"
#           new_spec_file = open(
#             appPath + "\\multiple_scenario_file.js",
#             "w+")
#
#           combined_multi_scenario_file = open(appPath + "\\combined_multi_scenario_file.js","a+")
#           new_conf_file = open(appPath + "\\multiple_conf_file.js", 'w+')
#           combined_conf_file = open(appPath+ "\\combined_conf_file.js", 'w+')
#           append_file = open(appPath + "\\append_spec", "a+")
#           batch_file = open(appPath + "\\multiple_batch_file.bat", "w+")
#           combined_batch_file = open(appPath + "\\combined_batch_file.bat", "w+")
#
#           wb = xlrd.open_workbook(path)
#           sheet = wb.sheet_by_index(0)
#           start_row = 1
#           end_row = sheet.nrows
#           num_of_columns = sheet.ncols
#           open_file = spec_file.read()
#           split_file = open_file.split("extract")[0]
#           if (split_file.endswith("//")):
#             split_file = split_file[:-2] + "});"
#           sec_test = "it('Test 1', function() {" + \
#                      "browser.driver.manage().window().maximize();" + \
#                      split_file.split("browser.driver.manage().window().maximize();")[1]
#
#           for r in range(start_row, end_row):
#             head_list = []
#             cell_list = []
#             for c in range(num_of_columns):
#               cell_obj = sheet.cell_value(r, c)
#               head_obj = sheet.cell_value(0, c)
#               head_list.append(head_obj)
#               cell_list.append(cell_obj)
#             for line in sec_test.split(";"):
#               for idx, header in enumerate(head_list):
#                 line = str(line)
#                 if (line.__contains__(header) and not line.__contains__("browser.wait(until")
#                   and not line.__contains__(".clear(")):
#                   if (type(cell_list[idx]) == int or type(cell_list[idx]) == float):
#                     line = line.split("sendKeys")[0] + "sendKeys(" + str(int(cell_list[idx])) + ");"
#                   else:
#                     line = line.split("sendKeys")[0] + "sendKeys('" + str(cell_list[idx]) + "');"
#                   break
#               replace_spec += line + ";"
#           new_spec = '''const jsdom = require("jsdom");
#                          const fs = require('fs');
#                          const path = require('path');
#                          const {SpecReporter} = require("jasmine-spec-reporter");
#                         jasmine.getEnv().addReporter(new SpecReporter({spec: {displayStacktrace: true, displayDuration: true}}));
#                          describe('Protractor Demo App', function() { ''' + str(replace_spec) + "});"
#           if os.stat(appPath + "\\combined_multi_scenario_file.js").st_size == 0:
#             new_combined_file = '''const jsdom = require("jsdom");
#                                      const fs = require('fs');
#                                      const path = require('path');
#                                      describe('Protractor Demo App', function() { ''' + str(replace_spec)
#             combined_multi_scenario_file.write(new_combined_file)
#             # print("Combined_spec_file", new_combined_file)
#           else:
#             new_combined_file = str(replace_spec)
#             combined_multi_scenario_file.write(new_combined_file)
#             # print("Combined_spec_file", new_combined_file)
#
#           new_conf = '''
#              exports.config = {
#                framework: 'jasmine',
#                seleniumAddress: 'http://localhost:4444/wd/hub',
#                specs: ['multiple_scenario_file.js'],
#                commonCapabilities: {
#                 'browserName': 'chrome',
#                 'chromeOptions': {
#                                                            'args': [ '--disable-gpu','-disable-dev-shm-usage','--no-sandbox','-disable-popup-blocking','--start-maximized','--disable-web-security','--allow-running-insecure-content','--disable-infobars']
#                                            }
#               },
#               onPrepare: function() {
#                  return browser.takeScreenshot();
#                },
#               getPageTimeout: 20000,
#                            allScriptsTimeout: 3000000,
#                            jasmineNodeOpts: {defaultTimeoutInterval: 3000000},
#                            defaultTimeoutInterval: 3000000
#              }
#
#              '''
#
#           comb_conf = '''
#              exports.config = {
#                framework: 'jasmine',
#                seleniumAddress: 'http://localhost:4444/wd/hub',
#                specs: ['combined_multi_scenario_file.js'],
#                commonCapabilities: {
#                 'browserName': 'chrome',
#                 'chromeOptions': {
#                                                            'args': [ '--disable-gpu','-disable-dev-shm-usage','--no-sandbox','-disable-popup-blocking','--start-maximized','--disable-web-security','--allow-running-insecure-content','--disable-infobars']
#                                            }
#               },
#               onPrepare: function() {
#                  return browser.takeScreenshot();
#                },
#               getPageTimeout: 20000,
#                            allScriptsTimeout: 3000000,
#                            jasmineNodeOpts: {defaultTimeoutInterval: 3000000},
#                            defaultTimeoutInterval: 3000000
#              }
#
#              '''
#           # START /B cmd "cd base_location\" && protractor base_location\oldFiles\multiple_conf.js > D:\ProtractorWebDriverTest\oldFiles\multiple_conf.log
#           multipleBatchFile = "START /B cmd " + appPath + " && protractor " + appPath + "\\multiple_conf_file.js >" + appPath + "\\multiple_conf_file.log"
#           comb_batch_file_data = "START /B cmd " + appPath + " && protractor " + appPath + "\\combined_multi_scenario_file.js >" + appPath + "\\multiple_conf_file.log"
#
#           new_conf_file.write(new_conf)
#           combined_conf_file.write(comb_conf)
#
#           new_spec_file.write(new_spec)
#           append_file.write(new_spec)
#           new_conf_file.close()
#           new_spec_file.close()
#           batch_file.write(multipleBatchFile)
#           combined_batch_file.write(comb_batch_file_data)
#           batch_file.close()
#           combined_batch_file.close()
#           combined_multi_scenario_file.close()
#
#           subprocess.call(appPath + '\\multiple_batch_file.bat')
#           conf_log = open(appPath + "\\multiple_conf_file.log").read()
#           logVal = multipletestReportGenerationRun(appPath)
#           time.sleep(1)
#
#           return logVal
#
#       else:
#         return "appFail"

@app.route("/uploadUIRun")
def uploadUIRun(testRunName, sheetLength):
  replace_spec = ''
  head_list = []
  cell_list = []
  filePath = ''
  # folderName = request.args.get("folderName")
  connections = db.testRun.find()
  for connection in connections:
    if connection["test_name"] == testRunName:
      appPath = connection["path"]
      scenario = connection["test_scenarios_used"]
      split_scenario = scenario.split(",")[sheetLength]
      appNameRun = split_scenario.split("[")[1].split("]")[0]
      scenarioNameRun = split_scenario.split("[")[0]
      print("scenario Name run", scenarioNameRun)
      collections = db.folderName.find()
      scenarioNamedict={}
      for collection in collections:
        if collection["appName"] == appNameRun and collection["scenarioName"] == scenarioNameRun:
          filePath = collection["innerpath"]
      if os.path.isdir(filePath):
        if os.path.exists(filePath + "\\OldFiles\\spec.js"):
          spec_file = open(
            filePath + "\\OldFiles\\spec.js")
          path = appPath + "\\sample.xlsx"
          new_spec_file = open(
            appPath + "\\multiple_scenario_file.js",
            "w+")

          combined_multi_scenario_file = open(appPath + "\\combined_multi_scenario_file.js", "a+")
          new_conf_file = open(appPath + "\\multiple_conf_file.js", 'w+')
          combined_conf_file = open(appPath + "\\combined_conf_file.js", 'w+')
          append_file = open(appPath + "\\append_spec", "a+")
          batch_file = open(appPath + "\\multiple_batch_file.bat", "w+")
          combined_batch_file = open(appPath + "\\combined_batch_file.bat", "w+")

          wb = xlrd.open_workbook(path)
          sheet = wb.sheet_by_index(0)
          start_row = 1
          end_row = sheet.nrows
          num_of_columns = sheet.ncols
          open_file = spec_file.read()
          split_file = open_file.split("extract")[0]
          if (split_file.endswith("//")):
            split_file = split_file[:-2] + "});"

          new_list=[]

          if scenarioNameRun in new_list:
            scenarioNamedict[scenarioNameRun]+=1
          else:
            scenarioNamedict[scenarioNameRun] =0
          new_list.append(scenarioNameRun)
          sec_test = "it(\'"+scenarioNameRun+" "+str(scenarioNamedict[scenarioNameRun])+"\', function() {" + \
                     "browser.driver.manage().window().maximize();" + \
                     split_file.split("browser.driver.manage().window().maximize();")[1]

          for r in range(start_row, end_row):
            head_list = []
            cell_list = []
            for c in range(num_of_columns):
              cell_obj = sheet.cell_value(r, c)
              head_obj = sheet.cell_value(0, c)
              head_list.append(head_obj)
              cell_list.append(cell_obj)
            for line in sec_test.split(";"):
              for idx, header in enumerate(head_list):
                line = str(line)
                if (line.__contains__(header) and not line.__contains__("browser.wait(until")
                  and not line.__contains__(".clear(")):
                  if (type(cell_list[idx]) == int or type(cell_list[idx]) == float):
                    line = line.split("sendKeys")[0] + "sendKeys(" + str(int(cell_list[idx])) + ");"
                  else:
                    line = line.split("sendKeys")[0] + "sendKeys('" + str(cell_list[idx]) + "');"
                  break
              replace_spec += line + ";"
          new_spec = '''const jsdom = require("jsdom");
                         const fs = require('fs');
                         const path = require('path');
                         const {SpecReporter} = require("jasmine-spec-reporter");
                        jasmine.getEnv().addReporter(new SpecReporter({spec: {displayStacktrace: true, displayDuration: true}}));
                         describe('Protractor Demo App', function() { ''' + str(replace_spec) + "});"
          if os.stat(appPath + "\\combined_multi_scenario_file.js").st_size == 0:
            new_combined_file = '''const jsdom = require("jsdom");
                                     const fs = require('fs');
                                     const path = require('path');
                                     describe('Protractor Demo App', function() { ''' + str(replace_spec)
            combined_multi_scenario_file.write(new_combined_file)
            # print("Combined_spec_file", new_combined_file)
          else:
            new_combined_file = str(replace_spec)
            combined_multi_scenario_file.write(new_combined_file)
            # print("Combined_spec_file", new_combined_file)

          new_conf = '''
             exports.config = {
               framework: 'jasmine',
               seleniumAddress: 'http://localhost:4444/wd/hub',
               specs: ['multiple_scenario_file.js'],
               commonCapabilities: {
                'browserName': 'chrome',
                'chromeOptions': {
                                                           'args': [ '--disable-gpu','-disable-dev-shm-usage','--no-sandbox','-disable-popup-blocking','--start-maximized','--disable-web-security','--allow-running-insecure-content','--disable-infobars']
                                           }
              },
              onPrepare: function() {
                 return browser.takeScreenshot();
               },
              getPageTimeout: 20000,
                           allScriptsTimeout: 3000000,
                           jasmineNodeOpts: {defaultTimeoutInterval: 3000000},
                           defaultTimeoutInterval: 3000000
             }

             '''

          comb_conf = '''
             exports.config = {
               framework: 'jasmine',
               seleniumAddress: 'http://localhost:4444/wd/hub',
               specs: ['combined_multi_scenario_file.js'],
               commonCapabilities: {
                'browserName': 'chrome',
                'chromeOptions': {
                                                           'args': [ '--disable-gpu','-disable-dev-shm-usage','--no-sandbox','-disable-popup-blocking','--start-maximized','--disable-web-security','--allow-running-insecure-content','--disable-infobars']
                                           }
              },
              onPrepare: function() {
                 return browser.takeScreenshot();
               },
              getPageTimeout: 20000,
                           allScriptsTimeout: 3000000,
                           jasmineNodeOpts: {defaultTimeoutInterval: 3000000},
                           defaultTimeoutInterval: 3000000
             }

             '''
          # START /B cmd "cd base_location\" && protractor base_location\oldFiles\multiple_conf.js > D:\ProtractorWebDriverTest\oldFiles\multiple_conf.log
          multipleBatchFile = "START /B cmd " + appPath + " && protractor " + appPath + "\\multiple_conf_file.js >" + appPath + "\\multiple_conf_file.log"
          comb_batch_file_data = "START /B cmd " + appPath + " && protractor " + appPath + "\\combined_multi_scenario_file.js >" + appPath + "\\multiple_conf_file.log"

          new_conf_file.write(new_conf)
          combined_conf_file.write(comb_conf)

          new_spec_file.write(new_spec)
          append_file.write(new_spec)
          new_conf_file.close()
          new_spec_file.close()
          batch_file.write(multipleBatchFile)
          combined_batch_file.write(comb_batch_file_data)
          batch_file.close()
          combined_batch_file.close()
          combined_multi_scenario_file.close()

          subprocess.call(appPath + '\\multiple_batch_file.bat')
          conf_log = open(appPath + "\\multiple_conf_file.log").read()
          logVal = multipletestReportGenerationRun(appPath)
          time.sleep(1)

          return logVal

      else:
        return "appFail"

@app.route("/multipletestReportGenerationRun")
def multipletestReportGenerationRun(scenarioPath):
  # folderName = request.args.get("folderName")
  print(scenarioPath)
  open_file = open(scenarioPath + "\\multiple_conf_file.log")
  log_file = open(scenarioPath + "\\multiple_conf_file.log").read()
  read_f = (open_file.readlines())
  passStr = '[32m.[0m'
  failStr = '[31mF[0m'
  resultList = []
  header = []
  header1 = []
  data = []
  js = {}
  # passStr=".F."
  for wholeStr in read_f:
    if (wholeStr.__contains__(passStr) or wholeStr.__contains__(failStr)):
      outputLine = [(wholeStr[i:i + 10]) for i in range(0, len(wholeStr), 10)]
      for result in outputLine:
        if result == passStr:
          resultList.append('Pass')
        elif result == failStr:
          resultList.append('Fail')

    # if(r.__contains__(failStr)):
  print(resultList)
  file_name = scenarioPath + "\\sample.xlsx"  # Path to your file
  df = pd.read_excel(file_name)
  # df["result"]=resultList[0]
  r, c = df.shape
  df['result'] = ""
  ind = 0
  for row in range(r):
    df["result"][row] = resultList[ind]
    ind += 1
  df.to_excel(scenarioPath + "\\new_excel_sheet.xlsx", index=False)
  x = df.head()
  time.sleep(1)
  excel_sheet = pd.read_excel(scenarioPath + "\\new_excel_sheet.xlsx")
  js["data"] = excel_sheet.to_dict(orient="records")
  for colName in excel_sheet.head():
    header.append(colName)
  # for h in header:
  js["headers"] = header
  totPass = resultList.count('Pass')
  totFail = resultList.count('Fail')
  testResult = [totPass, totFail]
  js["testResult"] = testResult;
  js["log"] = log_file
  return js

@app.route("/moveRunImage")
def moveRunImage():
  testRun = request.args.get("testRun")
  collections = db.testRun.find()
  fileName = []
  for collection in collections:
    if collection["test_name"] == testRun:
      appPath = collection["path"]
      scenario = collection["test_scenarios_used"]
      split_scenario = scenario.split(",")
      for app_scene in split_scenario:
        application_name = app_scene.split("[")[1].split("]")[0]
        scenario_name = app_scene.split("[")[0]
        connections = db.folderName.find()
        finalJson = {}
        destination_folder = ui_location+"\\src\\assets\\images\\" + testRun
        if os.path.isdir(destination_folder) == False:
          os.mkdir(destination_folder)
        for connection in connections:
          if connection["appName"] == application_name and connection["scenarioName"] == scenario_name:
            appPath = connection["path"]
            scenarioPath = connection["innerpath"]
            # destination_folder="D:\\TestAutomationTool\\src\\assets\\images"
            print(appPath, 'appPath', os.path.isdir(appPath))
            if os.path.isdir(appPath):
              if os.path.isdir(scenarioPath):
                for img in os.listdir(scenarioPath + "\\oldFiles"):
                  if img.endswith(".png"):
                    fileName.append(img)
                    shutil.copy(scenarioPath + "\\oldFiles\\" + img, destination_folder)
                finalJson["fileName"] = fileName
                finalJson["status"] = "success"
            else:
              return jsonify("appFail")
      break
  print(finalJson)
  return jsonify(finalJson)

@app.route("/deleteRun")
def deleteRun():
  testRun = request.args.get("testName")
  connections = db.testRun.find()
  myQuery = {"test_name": testRun}
  # "D:\\ProtractorWebDriverTest\\testRun\\" + testRun
  for connection in connections:
    if connection["test_name"] == testRun:
      path = connection["path"]
      shutil.rmtree(path, ignore_errors=True)
  db.testRun.delete_one(myQuery)
  return jsonify("success")

@app.route("/getScheduledJobs")  # Fetch data from db
def getScheduledJobs():
  connections = db.scheduledJobs.find()
  outDict = {}
  finalList = []
  finalDict = {}
  for collection in connections:
    outDict["scheduler_name"] = collection["scheduler_name"]
    outDict["test_job_name"] = collection["test_job_name"]
    outDict["scheduled_time"] = collection["scheduled_time"]
    finalList.append(outDict)
    outDict = {}
  finalDict["data"] = finalList
  finalDict["headers"] = ["scheduler_name", "test_job_name", "scheduled_time", "Actions"]
  return jsonify(finalDict)

def daily_job(appPath, schedulerName):
  print("job function")
  subprocess.call(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")

def job(appPath, schedulerName, scheduledDate):
  mon, day, year = scheduledDate.split("/")
  print(mon, day, year)

  datestring = str(year) + "-" + str(mon) + "-" + str(day)
  print(scheduledDate, "job function")
  print(date.today(), "joob function")
  print(datestring, "jobbb function")

  if date.today() == date(int(year), int(mon), int(day)):
    print("job function")
    subprocess.call(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")

def sched():
  while True:
    schedule.run_pending()
    time.sleep(5)
    print("Checking schedule")


@app.route("/ScheduleTestRun")
def ScheduleTestRun():
  thread = threading.Thread(target=sched)
  thread.start()
  with app.app_context():
    ScheduleDateTime = request.args.get("ScheduleDateTime")
    dailyFlag = request.args.get("ScheduleDateTimeFlag")
    testRunName = request.args.get("testRunName")
    email = request.args.get("email")
    schedulerName = request.args.get("schedulerName")
    ScheduleDate = ScheduleDateTime.split("T")[0]
    ScheduleTime = ScheduleDateTime.split("T")[1]
    hours = ScheduleTime.split(":")[0]
    mins = ScheduleTime.split(":")[1]
    meridian = "AM"

    print('hours: ', hours, type(hours), int(hours) < 10)
    print(ScheduleTime, ScheduleDate, "Date & time")
    if int(hours) > 12:
      hours = int(hours) - 12
      if hours < 10:
        hours = '0' + str(hours)
      meridian = "PM"
    elif int(hours) < 10 and int(hours) != 0:
      print(int(hours))
      hours = '0' + str(hours)
      meridian = "AM"
    elif int(hours) == 12:
      meridian = "PM"
    ScheduleTime1 = str(hours) + ":" + str(mins)
    ScheduleDate = ScheduleDate.split("-")[1] + "/" + ScheduleDate.split("-")[2] + "/" + ScheduleDate.split("-")[0]
    print(ScheduleTime, ScheduleDate)

    connections = db.testRun.find()
    for connection in connections:
      if connection["test_name"] == testRunName:
        print("entere")
        appPath = connection["path"]
        appPath = appPath.replace("\\\\", "\\")
        os.mkdir(appPath + "\\" + schedulerName)
        open(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat", "w").close()
        batch_file = open(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat", "w+")

        if dailyFlag == 'false':
          batchScript = f'''
                                    START /B cmd cd {base_location}\\ && protractor {appPath}\\conf.js > {appPath}\\conf.log    
                        '''
          #db_scheduledTime = ScheduleDate + " at " + str(hours) + ":" + str(mins) +" "+ meridian
          if int(hours) == 00:
            db_scheduledTime = ScheduleDate + " at " + str(int(hours)+12) + ":" + str(mins) + " " +meridian
          else:
            db_scheduledTime = ScheduleDate + " at " + str(hours) + ":" + str(mins) + " " + meridian

          dbData = [{"scheduler_name": schedulerName,
                     "test_job_name": testRunName,
                     "scheduled_time": db_scheduledTime,
                     "path": appPath}]

          db.scheduledJobs.insert_many(dbData)

          batch_file.write(batchScript)
          batch_file.close()

          # job(appPath, schedulerName, ScheduleDate)

          month, day, year = (int(i) for i in ScheduleDate.split('/'))
          # print(year, month, day)
          born = datetime(year, month, day)

          my_day = born.strftime("%A")
          print(my_day)

          if my_day == "Monday":
            schedule.every().monday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Tuesday":
            schedule.every().tuesday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Wednesday":
            schedule.every().wednesday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Thursday":
            schedule.every().thursday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Friday":
            schedule.every().friday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Saturday":
            schedule.every().saturday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Sunday":
            schedule.every().sunday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))

        else:
          batchScript = f'''
                                    START /B cmd cd {base_location}\\ && protractor {appPath}\\conf.js > {appPath}\\conf.log    
                        '''
          #db_scheduledTime = "Everyday at " + str(hours) + ":" + str(mins) + " " + meridian
          if int(hours) == 00:
            db_scheduledTime = "Everyday at " + str(int(hours)+12) + ":" + str(mins) + " " +meridian
          else:
            db_scheduledTime = "Everyday at " + str(hours) + ":" + str(mins) + " " + meridian
          dbData = [{"scheduler_name":schedulerName,
                     "test_job_name": testRunName,
                     "scheduled_time": db_scheduledTime,
                     "path": appPath}]
          db.scheduledJobs.insert_many(dbData)
          batch_file.write(batchScript)
          batch_file.close()

          # daily_job(appPath, schedulerName)

          print(ScheduleTime, "ScheduleTime")

          schedule.every().day.at(ScheduleTime).do(lambda: daily_job(appPath, schedulerName))

          # subprocess.call(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")
    send_mail(email, testRunName)
    return jsonify("Success")


@app.route("/deleteBatchFile")
def deleteBatchFile():
  schedulerName = request.args.get("schedulerName")
  testRunName = request.args.get("testRunName")
  ScheduleDateTime=request.args.get("ScheduleDateTime")
  my_query = {"test_job_name": testRunName, "scheduled_time": ScheduleDateTime}
  connections = db.scheduledJobs.find()
  for connection in connections:
    if connection["scheduler_name"]==schedulerName and connection["test_job_name"] == testRunName \
            and connection["scheduled_time"] == ScheduleDateTime:
      path = connection["path"]
      os.remove(path+"\\"+schedulerName+"\\scheduleTestRun.bat")
      break
  db.scheduledJobs.delete_many(my_query)
  return jsonify("Success")

def deleteTestJob(schedulerName, testRunName, ScheduleDateTime):
  my_query = {"scheduler_name": schedulerName, "test_job_name": testRunName, "scheduled_time": ScheduleDateTime}
  connections = db.scheduledJobs.find()
  for connection in connections:
    if connection["scheduler_name"] == schedulerName and connection["test_job_name"] == testRunName and connection["scheduled_time"] == ScheduleDateTime:
      path = connection["path"]
      os.remove(path + "\\" + schedulerName + "\\scheduleTestRun.bat")
      break
  db.scheduledJobs.delete_many(my_query)

@app.route("/editScheduleRun")
def editScheduleRun():
  ScheduleDateTime = request.args.get("ScheduleDateTime")
  oldScheduleDateTime = request.args.get("oldScheduleDateTime")
  testRunName = request.args.get("testRunName")
  dailyFlag = request.args.get("ScheduleDateTimeFlag")
  email = request.args.get("email")
  schedulerName = request.args.get("schedulerName")

  ScheduleDate = ScheduleDateTime.split("T")[0]
  ScheduleTime = ScheduleDateTime.split("T")[1]
  hours = ScheduleTime.split(":")[0]
  mins = ScheduleTime.split(":")[1]
  meridian = "AM"
  if int(hours) > 12:
    hours = int(hours) - 12
    meridian = "PM"
    if hours < 10:
      hours = '0' + str(hours)
  elif int(hours) < 10 and int(hours) != 0:
    hours = '0' + str(hours)
  elif int(hours) == 12:
    meridian = "PM"
  ScheduleTime1 = str(hours) + ":" + str(mins)
  ScheduleDate = ScheduleDate.split("-")[1] + "/" + ScheduleDate.split("-")[2] + "/" + ScheduleDate.split("-")[0]
  print(ScheduleTime, ScheduleDate)
  connections = db.scheduledJobs.find()

  for connection in connections:
    if connection["test_job_name"] == testRunName:
      if connection["scheduled_time"] == oldScheduleDateTime:
        appPath = connection["path"]
        print(appPath)
        appPath = appPath.replace("/", "\\").replace("\\\\", "\\")
        open(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat", "w").close()
        deleteTestJob(schedulerName, testRunName, oldScheduleDateTime)
        batch_file = open(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat", "w+")

        if dailyFlag == "false":
          batchScript = f'''
                        START /B cmd "cd {base_location}\\" && protractor {appPath}\\conf.js > {appPath}\\conf.log                            
                    '''
          # db_scheduledTime = ScheduleDate + " at " + str(hours) + ":" + str(mins) + " " + meridian
          if int(hours) == 00:
            db_scheduledTime = ScheduleDate + " at " + str(int(hours) + 12) + ":" + str(mins) + " " + meridian
          else:
            db_scheduledTime = ScheduleDate + " at " + str(hours) + ":" + str(mins) + " " + meridian

          dbData = [{"scheduler_name": schedulerName,
                     "test_job_name": testRunName,
                     "scheduled_time": db_scheduledTime,
                     "path": appPath}]

          db.scheduledJobs.insert_many(dbData)

          batch_file.write(batchScript)
          batch_file.close()
          month, day, year = ScheduleDate.split('/')
          born = datetime(int(year), int(month), int(day))
          my_day = born.strftime("%A")
          print(my_day, ScheduleTime)

          if my_day == "Monday":
            schedule.every().monday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Tuesday":
            schedule.every().tuesday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Wednesday":
            schedule.every().wednesday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Thursday":
            schedule.every().thursday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Friday":
            print("entered into this friday")

            schedule.every().friday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
            # schedule.every().friday.at("15:55").do(lambda: job('Hello ', 'world!'))
            print("schedule function is called..")
          elif my_day == "Saturday":
            schedule.every().saturday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Sunday":
            schedule.every().sunday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          # subprocess.call(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")
        else:
          batchScript = f'''
                      START /B cmd "cd {base_location}\\" && protractor {appPath}\\conf.js > {appPath}\\conf.log
                            '''
          # db_scheduledTime = "Everyday at " + str(hours) + ":" + str(mins) + " " + meridian
          if int(hours) == 00:
            db_scheduledTime = "Everyday at " + str(int(hours) + 12) + ":" + str(mins) + " " + meridian
          else:
            db_scheduledTime = "Everyday at " + str(hours) + ":" + str(mins) + " " + meridian

          dbData = [{"scheduler_name": schedulerName,
                     "test_job_name": testRunName,
                     "scheduled_time": db_scheduledTime,
                     "path": appPath}]

          db.scheduledJobs.insert_many(dbData)
          batch_file.write(batchScript)
          batch_file.close()

          schedule.every().day.at(ScheduleTime).do(lambda: daily_job(appPath, schedulerName))
          # subprocess.call(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")
  send_mail(email, testRunName)
  return jsonify("Success")

@app.route("/editScheduleRunWithExcel", methods=["GET","POST"])
def editScheduleRunWithExcel():
  ScheduleDateTime = request.args.get("ScheduleDateTime")
  oldScheduleDateTime = request.args.get("oldScheduleDateTime")
  testRunName = request.args.get("testRunName")
  dailyFlag = request.args.get("ScheduleDateTimeFlag")
  email = request.args.get("email")
  schedulerName = request.args.get("schedulerName")
  file = request.get_data()

  catchit = uploadingExcelInSchedule(schedulerName, file, testRunName)

  print(catchit)
  ScheduleDate = ScheduleDateTime.split("T")[0]
  ScheduleTime = ScheduleDateTime.split("T")[1]
  hours = ScheduleTime.split(":")[0]
  mins = ScheduleTime.split(":")[1]
  meridian = "AM"
  if int(hours) > 12:
    hours = int(hours) - 12
    meridian = "PM"
    if hours < 10:
      hours = '0' + str(hours)
  elif int(hours) < 10 and int(hours) != 0:
    hours = '0' + str(hours)
  elif int(hours) == 12:
    meridian = "PM"
  ScheduleTime1 = str(hours) + ":" + str(mins)
  ScheduleDate = ScheduleDate.split("-")[1] + "/" + ScheduleDate.split("-")[2] + "/" + ScheduleDate.split("-")[0]
  print(ScheduleTime, ScheduleDate)
  connections = db.scheduledJobs.find()

  for connection in connections:
    if connection["test_job_name"] == testRunName:
      if connection["scheduled_time"] == oldScheduleDateTime:
        appPath = connection["path"]
        print(appPath)
        appPath = appPath.replace("/", "\\").replace("\\\\", "\\")
        open(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat", "w").close()
        deleteTestJob(schedulerName, testRunName, oldScheduleDateTime)
        batch_file = open(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat", "w+")
        if os.path.isdir(appPath+"\\"+schedulerName+"\\combined_conf_file.log") == False:
          open(appPath+"\\"+schedulerName+"\\combined_conf_file.log", "w").close()
        if dailyFlag == "false":
          batchScript = f'''
                       START /B cmd "cd {base_location}\\" && protractor {appPath}\\{schedulerName}\\combined_conf_file.js > {appPath}\\{schedulerName}\\combined_conf_file.log
                            '''
          # db_scheduledTime = ScheduleDate + " at " + str(hours) + ":" + str(mins) + " " + meridian
          if int(hours) == 00:
            db_scheduledTime = ScheduleDate + " at " + str(int(hours) + 12) + ":" + str(mins) + " " + meridian
          else:
            db_scheduledTime = ScheduleDate + " at " + str(hours) + ":" + str(mins) + " " + meridian

          dbData = [{"scheduler_name": schedulerName,
                     "test_job_name": testRunName,
                     "scheduled_time": db_scheduledTime,
                     "path": appPath}]

          db.scheduledJobs.insert_many(dbData)

          batch_file.write(batchScript)
          batch_file.close()
          month, day, year = ScheduleDate.split('/')
          born = datetime(int(year), int(month), int(day))
          my_day = born.strftime("%A")
          print(my_day, ScheduleTime)

          if my_day == "Monday":
            schedule.every().monday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Tuesday":
            schedule.every().tuesday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Wednesday":
            schedule.every().wednesday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Thursday":
            schedule.every().thursday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Friday":
            print("entered into this friday")

            schedule.every().friday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
            # schedule.every().friday.at("15:55").do(lambda: job('Hello ', 'world!'))
            print("schedule function is called..")
          elif my_day == "Saturday":
            schedule.every().saturday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Sunday":
            schedule.every().sunday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          # subprocess.call(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")
        else:
          batchScript = f'''
                               START /B cmd "cd {base_location}\\" && protractor {appPath}\\{schedulerName}\\combined_conf_file.js > {appPath}\\{schedulerName}\\combined_conf_file.log
                        '''
          # db_scheduledTime = "Everyday at " + str(hours) + ":" + str(mins) + " " + meridian
          if int(hours) == 00:
            db_scheduledTime = "Everyday at " + str(int(hours) + 12) + ":" + str(mins) + " " + meridian
          else:
            db_scheduledTime = "Everyday at " + str(hours) + ":" + str(mins) + " " + meridian

          dbData = [{"scheduler_name": schedulerName,
                     "test_job_name": testRunName,
                     "scheduled_time": db_scheduledTime,
                     "path": appPath}]

          db.scheduledJobs.insert_many(dbData)
          batch_file.write(batchScript)
          batch_file.close()
          schedule.every().day.at(ScheduleTime).do(lambda: daily_job(appPath, schedulerName))
          # subprocess.call(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")
  send_mail(email, testRunName)
  return jsonify("Success")

@app.route("/checkMultiScenarioPath")
def checkMultiScenarioPath():
  testRunName = request.args.get("testRunName")
  schedulerName = request.args.get("schedulerName")
  if os.path.exists(base_location + "\\testRun\\"+ testRunName +"\\"+ schedulerName +"\\combined_multi_scenario_file.js"):
    return jsonify({"status":"success"})
  else:
    return jsonify({"status":"failure"})


@app.route("/ScheduleTestRunWithExcel", methods=['GET', 'POST'])
def ScheduleTestRunWithExcel():
  thread = threading.Thread(target=sched)
  thread.start()

  ScheduleDateTime = request.args.get("ScheduleDateTime")
  dailyFlag = request.args.get("ScheduleDateTimeFlag")
  testRunName = request.args.get("testRunName")
  email = request.args.get("email")
  schedulerName = request.args.get("schedulerName")
  file = request.get_data()
  print('file', file)

  catchit = uploadingExcelInSchedule(schedulerName, file, testRunName)

  print(catchit)
  with app.app_context():

    ScheduleDate = ScheduleDateTime.split("T")[0]
    ScheduleTime = ScheduleDateTime.split("T")[1]
    hours = ScheduleTime.split(":")[0]
    mins = ScheduleTime.split(":")[1]
    meridian = "AM"
    print('hours: ', hours, type(hours), int(hours) < 10)
    print(ScheduleTime, ScheduleDate, "Date & time")
    if int(hours) > 12:
      hours = int(hours) - 12
      if hours < 10:
        hours = '0' + str(hours)
      meridian = "PM"
    elif int(hours) < 10 and int(hours) != 0:
      print(int(hours))
      hours = '0' + str(hours)
      meridian = "AM"
    elif int(hours) == 12:
      meridian = "PM"
    ScheduleTime1 = str(hours) + ":" + str(mins)
    ScheduleDate = ScheduleDate.split("-")[1] + "/" + ScheduleDate.split("-")[2] + "/" + ScheduleDate.split("-")[0]
    print(ScheduleTime, ScheduleDate)

    connections = db.testRun.find()
    for connection in connections:
      if connection["test_name"] == testRunName:
        appPath = connection["path"]
        appPath = appPath.replace("\\\\", "\\")
        if os.path.isdir(appPath + "\\" + schedulerName) == False:
          os.mkdir(appPath + "\\" + schedulerName)
        if os.path.isdir(appPath + "\\" + schedulerName + "\\combined_conf_file.log") == False:
          open(appPath + "\\" + schedulerName + "\\combined_conf_file.log", "w").close()
        open(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat", "w").close()
        batch_file = open(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat", "w+")

        if dailyFlag == 'false':
          batchScript = f'''

                           START /B cmd "cd {base_location}\\" && protractor {appPath}\\{schedulerName}\\combined_conf_file.js > {appPath}\\{schedulerName}\\combined_conf_file.log
                       '''
          # db_scheduledTime = ScheduleDate + " at " + str(hours) + ":" + str(mins) +" "+ meridian
          if int(hours) == 00:
            db_scheduledTime = ScheduleDate + " at " + str(int(hours) + 12) + ":" + str(mins) + " " + meridian
          else:
            db_scheduledTime = ScheduleDate + " at " + str(hours) + ":" + str(mins) + " " + meridian

          dbData = [{"scheduler_name": schedulerName,
                     "test_job_name": testRunName,
                     "scheduled_time": db_scheduledTime,
                     "path": appPath}]

          db.scheduledJobs.insert_many(dbData)

          batch_file.write(batchScript)
          batch_file.close()

          # job(appPath, schedulerName, ScheduleDate)

          month, day, year = ScheduleDate.split('/')
          born = datetime(int(year), int(month), int(day))
          my_day = born.strftime("%A")
          print(my_day, ScheduleTime)

          if my_day == "Monday":
            schedule.every().monday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Tuesday":
            schedule.every().tuesday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Wednesday":
            schedule.every().wednesday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Thursday":
            schedule.every().thursday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Friday":
            print("entered into this friday")

            schedule.every().friday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
            # schedule.every().friday.at("15:55").do(lambda: job('Hello ', 'world!'))
            print("schedule function is called..")
          elif my_day == "Saturday":
            schedule.every().saturday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))
          elif my_day == "Sunday":
            schedule.every().sunday.at(ScheduleTime).do(lambda: job(appPath, schedulerName, ScheduleDate))

          # subprocess.call(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")
        else:

          batchScript = f'''

                                     START /B cmd "cd {base_location}\\" && protractor {appPath}\\{schedulerName}\\combined_conf_file.js > {appPath}\\{schedulerName}\\combined_conf_file.log

                         '''

          # db_scheduledTime = "Everyday at " + str(hours) + ":" + str(mins) + " " + meridian
          if int(hours) == 00:
            db_scheduledTime = "Everyday at " + str(int(hours) + 12) + ":" + str(mins) + " " + meridian
          else:
            db_scheduledTime = "Everyday at " + str(hours) + ":" + str(mins) + " " + meridian
          dbData = [{"scheduler_name": schedulerName,
                     "test_job_name": testRunName,
                     "scheduled_time": db_scheduledTime,
                     "path": appPath}]
          db.scheduledJobs.insert_many(dbData)
          batch_file.write(batchScript)
          batch_file.close()

          # daily_job(appPath, schedulerName)
          schedule.every().day.at(ScheduleTime).do(lambda: daily_job(appPath, schedulerName))
          # subprocess.call(appPath + "\\" + schedulerName + "\\scheduleTestRun.bat")
    send_mail(email, testRunName)
    return jsonify("Success")

def uploadUIWithoutRun(schedulerName, testRunName, sheetLength):
  replace_spec = ''
  head_list = []
  cell_list = []
  filePath = ''
  # folderName = request.args.get("folderName")
  connections = db.testRun.find()
  for connection in connections:
    if connection["test_name"] == testRunName:
      appPath = connection["path"]
      scenario = connection["test_scenarios_used"]
      split_scenario = scenario.split(",")[sheetLength]
      appNameRun = split_scenario.split("[")[1].split("]")[0]
      scenarioNameRun = split_scenario.split("[")[0]
      collections = db.folderName.find()
      for collection in collections:
        if collection["appName"] == appNameRun and collection["scenarioName"] == scenarioNameRun:
          filePath = collection["innerpath"]
      if os.path.isdir(appPath+"\\"+schedulerName) == False:
        os.mkdir(appPath + "\\" + schedulerName)
      if os.path.isdir(filePath):
        if os.path.exists(filePath + "\\OldFiles\\spec.js"):
          spec_file = open(
            filePath + "\\OldFiles\\spec.js")
          path = appPath + "\\"+ schedulerName+ "\\sample.xlsx"
          new_spec_file = open(
            appPath + "\\" + schedulerName + "\\multiple_scenario_file.js",
            "w+")

          combined_multi_scenario_file = open(appPath  +"\\" + schedulerName + "\\combined_multi_scenario_file.js","a+")
          new_conf_file = open(appPath +"\\" + schedulerName + "\\multiple_conf_file.js", 'w+')
          combined_conf_file = open(appPath +"\\" + schedulerName +"\\combined_conf_file.js", 'w+')
          append_file = open(appPath +"\\" + schedulerName + "\\append_spec", "a+")
          batch_file = open(appPath +"\\" + schedulerName + "\\multiple_batch_file.bat", "w+")
          combined_batch_file = open(appPath +"\\" + schedulerName + "\\combined_batch_file.bat", "w+")

          wb = xlrd.open_workbook(path)
          sheet = wb.sheet_by_index(0)
          start_row = 1
          end_row = sheet.nrows
          num_of_columns = sheet.ncols
          open_file = spec_file.read()
          split_file = open_file.split("extract")[0]
          if (split_file.endswith("//")):
            split_file = split_file[:-2] + "});"
          sec_test = "it('Test 1', function() {" + \
                     "browser.driver.manage().window().maximize();" + \
                     split_file.split("browser.driver.manage().window().maximize();")[1]

          for r in range(start_row, end_row):
            head_list = []
            cell_list = []
            for c in range(num_of_columns):
              cell_obj = sheet.cell_value(r, c)
              head_obj = sheet.cell_value(0, c)
              head_list.append(head_obj)
              cell_list.append(cell_obj)
            for line in sec_test.split(";"):
              for idx, header in enumerate(head_list):
                line = str(line)
                if (line.__contains__(header) and not line.__contains__("browser.wait(until")
                  and not line.__contains__(".clear(")):
                  if (type(cell_list[idx]) == int or type(cell_list[idx]) == float):
                    line = line.split("sendKeys")[0] + "sendKeys(" + str(int(cell_list[idx])) + ");"
                  else:
                    line = line.split("sendKeys")[0] + "sendKeys('" + str(cell_list[idx]) + "');"
                  break
              replace_spec += line + ";"
          new_spec = '''const jsdom = require("jsdom");
                         const fs = require('fs');
                         const path = require('path');
                         describe('Protractor Demo App', function() { ''' + str(replace_spec) + "});"

          if os.stat(appPath +"\\" + schedulerName + "\\combined_multi_scenario_file.js").st_size == 0:
            new_combined_file = '''const jsdom = require("jsdom");
                                     const fs = require('fs');
                                     const path = require('path');
                                     describe('Protractor Demo App', function() { ''' + str(replace_spec)
            combined_multi_scenario_file.write(new_combined_file)
            print("Combined_spec_file", combined_multi_scenario_file)
          else:
            new_combined_file = str(replace_spec)
            combined_multi_scenario_file.write(new_combined_file)
            print("Combined_spec_file", combined_multi_scenario_file)

          new_conf = '''
             exports.config = {
               framework: 'jasmine',
               seleniumAddress: 'http://localhost:4444/wd/hub',
               specs: ['multiple_scenario_file.js'],
               commonCapabilities: {
                'browserName': 'chrome',
                'chromeOptions': {
                                                           'args': [ '--disable-gpu','-disable-dev-shm-usage','--no-sandbox','-disable-popup-blocking','--start-maximized','--disable-web-security','--allow-running-insecure-content','--disable-infobars']
                                           }
              },
              onPrepare: function() {
                 return browser.takeScreenshot();
               },
              getPageTimeout: 20000,
                           allScriptsTimeout: 3000000,
                           jasmineNodeOpts: {defaultTimeoutInterval: 3000000},
                           defaultTimeoutInterval: 3000000
             }

             '''

          comb_conf = '''
             exports.config = {
               framework: 'jasmine',
               seleniumAddress: 'http://localhost:4444/wd/hub',
               specs: ['combined_multi_scenario_file.js'],
               commonCapabilities: {
                'browserName': 'chrome',
                'chromeOptions': {
                                                           'args': [ '--disable-gpu','-disable-dev-shm-usage','--no-sandbox','-disable-popup-blocking','--start-maximized','--disable-web-security','--allow-running-insecure-content','--disable-infobars']
                                           }
              },
              onPrepare: function() {
                 return browser.takeScreenshot();
               },
              getPageTimeout: 20000,
                           allScriptsTimeout: 3000000,
                           jasmineNodeOpts: {defaultTimeoutInterval: 3000000},
                           defaultTimeoutInterval: 3000000
             }

             '''
          # START /B cmd "cd base_location\" && protractor base_location\oldFiles\multiple_conf.js > D:\ProtractorWebDriverTest\oldFiles\multiple_conf.log

          multipleBatchFile = "START /B cmd " + appPath + " && protractor " + appPath +"\\" + schedulerName + "\\multiple_conf_file.js >" + appPath +"\\" + schedulerName + "\\multiple_conf_file.log"
          comb_batch_file_data = "START /B cmd " + appPath + " && protractor " + appPath +"\\" + schedulerName + "\\combined_multi_scenario_file.js >" + appPath +"\\" + schedulerName + "\\multiple_conf_file.log"

          new_conf_file.write(new_conf)
          combined_conf_file.write(comb_conf)

          new_spec_file.write(new_spec)
          append_file.write(new_spec)
          new_conf_file.close()
          new_spec_file.close()
          batch_file.write(multipleBatchFile)
          combined_batch_file.write(comb_batch_file_data)
          batch_file.close()
          combined_batch_file.close()
          combined_multi_scenario_file.close()

          # subprocess.call(appPath + '\\multiple_batch_file.bat')
          # conf_log = open(appPath + "\\multiple_conf_file.log").read()
          # logVal = multipletestReportGenerationRun(appPath)
          time.sleep(1)

          return "Success"

      else:
        return "appFail"

def uploadingExcelInSchedule(schedulerName, file, testRunName):
  # folderName = request.args.get("folderName")
  import pandas as pd
  global mulipleScenarioPath

  final_val = []

  str_file = file.decode("utf-8")
  # temp_file = list(eval(str_file))
  temp_List = []
  sheetLength = 0
  json_format = json.loads(str_file)
  print("str_file", str_file)

  # Checking the muliple scenario file exist or not, if exists remove the file
  connections = db.testRun.find()
  for connection in connections:
    if connection["test_name"] == testRunName:
      mulipleScenarioPath = connection["path"]
      if os.path.isdir(mulipleScenarioPath + "\\" + schedulerName) == False:
        os.mkdir(mulipleScenarioPath+"\\"+schedulerName)
      if os.path.exists(mulipleScenarioPath +"\\"+schedulerName+ "\\combined_multi_scenario_file.js"):
        p = os.path.join(mulipleScenarioPath +"\\"+schedulerName+ "\\combined_multi_scenario_file.js")
        os.remove(p)

  for k, v in json_format.items():
    v1 = json.dumps(v['data'])
    if type(list(eval(v1))[0]) == "str":
      return jsonify("failure")
    else:
      temp_file = list(eval(v1))
      df = pd.DataFrame.from_dict(temp_file)

    connections = db.testRun.find()
    for connection in connections:
      if connection["test_name"] == testRunName:
        mulipleScenarioPath = connection["path"]
        if os.path.isdir(mulipleScenarioPath):
          df.to_excel(mulipleScenarioPath + "\\"+ schedulerName + '\\sample.xlsx',
                      index=False)
          time.sleep(1)
          append_file = open(mulipleScenarioPath + "\\" + schedulerName + "\\append_spec.txt", "a+")
          returnVal = uploadUIWithoutRun(schedulerName, testRunName, sheetLength)
          # final_val.append(returnVal)
          sheetLength += 1
        else:
          return "appFail"

  comb_multi_file = open(mulipleScenarioPath+"\\" + schedulerName +"\\combined_multi_scenario_file.js","a+")
  comb_multi_file.write("});")
  return jsonify("Success")

@app.route("/editTestDescription")
def editTestDescription():
  appName = request.args.get("appName")
  scenarioName = request.args.get("scenarioName")
  testDescription = request.args.get("testDescription")

  myquery = { "appName" : appName, "scenarioName": scenarioName  }
  newvalues = { "$set": { "test_Description": testDescription } }
  db.folderName.update_one(myquery, newvalues)
  return jsonify("Success")

def send_mail(MailID, testRunName):
  try:
    pythoncom.CoInitialize()
    outlook = win32.Dispatch('outlook.application')
    mail = outlook.CreateItem(0)
    mail.To = MailID

    mail.Subject = testRunName + ' Test Report'
    mail.Body = 'Message body'
    cont = """\
    <html>
     <head></head>
     <body>
       <p>Dear iTAP User,<br><br>
          Test Job Name : """ + testRunName + """<br><br>
          Date and Time : """ + datetime.now().strftime("%d/%m/%Y %H:%M:%S") + """<br><br>
          Attached the report of above Test Job, Please find attached document.<br><br>

          Note : Please do not reply to this mail, This is an auto-generated e-mail.<br><br>
          Regards,<br>
          iTAP Team
       </p>
     </body>
    </html>
    """
    mail.HTMLBody = cont  # this field is optional
    try:
      file_one = base_location + "\\testRun\\" + testRunName + "\\new_excel_sheet.xlsx"
      file_two = base_location + "\\testRun\\" + testRunName + "\\iTAPTestDataReport-" + testRunName + ".xlsx"
      os.rename(file_one, file_two)

      attachmentPath = file_two
      mail.Attachments.Add(attachmentPath)
      mail.Send()
      os.rename(file_two, file_one)
      return "Success"
    except:
      return "No File"
  except:
    return "Outlook Missing"

@app.route("/checkSchedulerNameExists")
def checkSchedulerNameExists():
  schedulerName = request.args.get("schedulerName")
  collections = db.scheduledJobs.find()
  if collections.count() == 0:
    return jsonify("Success")
  else:
    for collection in collections:
      if collection['scheduler_name'] != schedulerName:
        return jsonify("Success")
      else:
        return jsonify("Failure")

def split_list(a_list):
    half = len(a_list) // 2
    return a_list[:half], a_list[half:]

def json_to_json(given_json):
  given_json = str(given_json)
  # given_json = given_json.replace("'","\\\"")
  # gvn_json = "\"" + given_json + "\""
  given_json = given_json.replace("'", "\"")
  print(type(given_json),' ',given_json)
  json_value = json.loads(given_json)
  headers = ["browser_compatibility_test", "chrome", "firefox"]
  data = []
  for val in json_value.keys():
    x = {}
    x = json_value[val][0]
    if val == "time_taken":
      val = str(val).replace("time_taken", "Time Taken (In seconds)")
    if val == "total_specs":
      val = str(val).replace("total_specs", "Total Script Results")
    if val == "total_pass_results":
      val = str(val).replace("total_pass_results", "Total Scripts Passed")
    if val == "total_fail_results":
      val = str(val).replace("total_fail_results", "Total Scripts Failed")
    x["browser_compatibility_test"] = val
    final_json = x
    data.append(final_json)
  lst = [0] * 2
  lst[1] = headers
  lst[0] = data
  print(lst)
  return lst

# def json_json_json(str_file, jobName):
#   # str_file=open("D:\ProtractorWebDriverTest\output_json.js","r").read()
#   # jobName="something"
#   print(str_file)
#   str_file=str(str_file).replace("'", "\"")
#   json_line=json.loads(str_file)
#   Test_Job=json.loads(str(json_line["data"][2]).replace("'", "\""))
#   num_of_scripts_chrome=len(Test_Job["chrome"].split(","))
#   num_of_scripts_firefox=len(Test_Job["firefox"].split(","))
#   chrome_scripts=str(Test_Job["chrome"]).replace("[","").replace("]","")
#   firefox_script=str(Test_Job["firefox"]).replace("[","").replace("]","")
#   # print("chrome scripts",chrome_scripts, firefox_script)
#   tot_value=max(num_of_scripts_chrome, num_of_scripts_firefox)
#   test_job_list=[]
#   for val in range(1,tot_value+1):
#     test_job_list.append(jobName+"_DataSet_"+str(val))
#
#   # print(test_job_list)
#   result=json.loads(str(json_line["data"][1]).replace("'", "\""))
#   newresult1=json.loads(str(json_line["data"][3]).replace("'", "\""))
#   newresult2=json.loads(str(json_line["data"][4]).replace("'", "\""))
#   total_chrome=(newresult1["chrome"],newresult2["chrome"])
#   total_firefox=(newresult1["firefox"],newresult2["firefox"])
#   # print("total chrome", total_chrome, total_firefox)
#   resultchrome=result["chrome"].append(str(total_chrome))
#   resultfirefox=(result["firefox"]).append(str(total_firefox))
#   tot_time=json.loads(str(json_line["data"][0]).replace("'", "\""))
#   tot_time_chrome=float(tot_time["chrome"])
#   tot_time_firefox=float(tot_time["firefox"])
#   new_tot_time_chrome=str(int(tot_time_chrome//60))+"mins"+str(int(tot_time_chrome%60))+"secs"
#   new_tot_time_firefox=str(int(tot_time_firefox//60))+"mins"+str(int(tot_time_firefox%60))+"secs"
#   print(new_tot_time_chrome, new_tot_time_firefox)
#   # print("result chrome",result["chrome"], result["firefox"])
#   total_time_chrome=chrome_scripts.split(",")
#   total_time_firefox=firefox_script.split(",")
#   total_time_chrome.append(new_tot_time_chrome)
#   total_time_firefox.append(new_tot_time_firefox)
#   print("total time chrome",tot_time_chrome ,tot_time_firefox )
#   json_format={}
#   test_job_list.append("Total")
#   json_format["Test_Job"]=test_job_list
#   json_format["Time_Taken_chrome"]=total_time_chrome
#   json_format["Time_Taken_Firefox"]=total_time_firefox
#   json_format["Result_Chrome"]=result["chrome"]
#   json_format["Result_Firefox"]=result["firefox"]
#   # print(json.dumps(json_format, indent=4))
#
#
#   #{"Test_job:"something_DataSet_0", "Time Taken chrome":"1 min 23 secs","Time Taken Firefox":"38 secs"}
#   new_json_format={}
#   full_list=[]
#   counter=0
#   for index_val in range(len(json_format["Test_Job"])):
#     print(index_val)
#     for key,value in json_format.items():
#         # print(key)
#         new_json_format[key]=value[index_val]
#     full_list.append(new_json_format)
#     new_json_format={}
#
#   print(full_list)
#   return full_list

def json_json_json(str_file, jobName):
  # str_file=open("D:\ProtractorWebDriverTest\output_json.js","r").read()
  # jobName="something"
  print(str_file)
  str_file=str(str_file)
  str_file=str_file.replace("'", "\"")
  print(str_file)
  json_line=json.loads(str_file)
  Test_Job=json.loads(str(json_line["data"][2]).replace("'", "\""))
  num_of_scripts_chrome=len(Test_Job["chrome"].split(","))
  num_of_scripts_firefox=len(Test_Job["firefox"].split(","))
  chrome_scripts=str(Test_Job["chrome"]).replace("[","").replace("]","")
  firefox_script=str(Test_Job["firefox"]).replace("[","").replace("]","")
  # print("chrome scripts",chrome_scripts, firefox_script)
  tot_value=max(num_of_scripts_chrome, num_of_scripts_firefox)
  test_job_list=[]
  ful_count=count_func_name(jobName)
  count_dict={}
  arr_val=[]
  for val in ful_count:
    if val.split()[0] in arr_val:
      count_dict[val.split()[0]]+=1
    else:
      count_dict[val.split()[0]] = 1

    arr_val.append(val.split()[0])
    test_job_list.append(val.split()[0]+"_DataSet_"+str(count_dict[val.split()[0]]))

  # print(test_job_list)
  result=json.loads(str(json_line["data"][1]).replace("'", "\""))
  newresult1=json.loads(str(json_line["data"][3]).replace("'", "\""))
  newresult2=json.loads(str(json_line["data"][4]).replace("'", "\""))
  total_chrome=(newresult1["chrome"],newresult2["chrome"])
  total_firefox=(newresult1["firefox"],newresult2["firefox"])
  # print("total chrome", total_chrome, total_firefox)
  resultchrome=result["chrome"].append(str(total_chrome))
  resultfirefox=(result["firefox"]).append(str(total_firefox))
  tot_time=json.loads(str(json_line["data"][0]).replace("'", "\""))
  tot_time_chrome=float(tot_time["chrome"])
  tot_time_firefox=float(tot_time["firefox"])
  new_tot_time_chrome=str(int(tot_time_chrome//60))+"mins"+str(int(tot_time_chrome%60))+"secs"
  new_tot_time_firefox=str(int(tot_time_firefox//60))+"mins"+str(int(tot_time_firefox%60))+"secs"
  print(new_tot_time_chrome, new_tot_time_firefox)
  # print("result chrome",result["chrome"], result["firefox"])
  total_time_chrome=chrome_scripts.split(",")
  total_time_firefox=firefox_script.split(",")
  total_time_chrome.append(new_tot_time_chrome)
  total_time_firefox.append(new_tot_time_firefox)
  print("total time chrome",tot_time_chrome ,tot_time_firefox )
  json_format={}
  test_job_list.append("Total")
  json_format["Test_Job"]=test_job_list
  json_format["Time_Taken_chrome"]=total_time_chrome
  json_format["Time_Taken_Firefox"]=total_time_firefox
  json_format["Result_Chrome"]=result["chrome"]
  json_format["Result_Firefox"]=result["firefox"]
  print(json.dumps(json_format, indent=4))


  #{"Test_job:"something_DataSet_0", "Time Taken chrome":"1 min 23 secs","Time Taken Firefox":"38 secs"}
  new_json_format={}
  full_list=[]
  counter=0
  for index_val in range(len(json_format["Test_Job"])):
    print(index_val)
    for key,value in json_format.items():
        print("key",key)
        new_json_format[key]=value[index_val]
    full_list.append(new_json_format)
    new_json_format={}

  print(full_list)
  return full_list

def count_func_name(jobName):
  full_array=[]
  collections=db.testRun.find({"test_name":jobName})
  for collection in collections:
    path=collection["path"]
    f=open(path+"\combined_multi_scenario_file.js").readlines()
    for line in f:
      if line.__contains__("it('"):
        # print(line)
        # print(line.split("it('")[1].split("'")[0])
        full_array.append(line.split("it('")[1].split("'")[0])
    print(full_array)
    # full_array=['Deposit_Flow 0', 'Deposit_Flow 0', 'Withdraw_flow 0', 'Withdraw_flow 0']
    return full_array

def getParallelJobDetails(jobName, file):
    with app.app_context():
      final_json = {}
      connections = db.testRun.find()
      for connection in connections:
        if connection["test_name"] == jobName:
          path = connection["path"]
          break
      parallelConf = open(path + "\\"+ file +".log").readlines()
      time_taken_seconds = []
      passStr = '[32m.[0m'
      failStr = '[31mF[0m'
      resultList = []
      for line in parallelConf:
        if line.__contains__("Finished in"):
          time_taken_seconds.append(line.split(" ")[4])
        if line.__contains__(passStr) or line.__contains__(failStr):
          line = line.split(" ")[2]
          outputLine = [(line[i:i + 10]) for i in range(0, len(line), 10)]
          for result in outputLine:
            if result == passStr:
              resultList.append('Pass')
            elif result == failStr:
              resultList.append('Fail')
      final_json["time_taken"] = []
      final_json["time_taken"].append({"chrome": time_taken_seconds[0], "firefox": time_taken_seconds[1]})
      final_json["total_specs"] = []
      chromeList, FirefoxList = split_list(resultList)
      final_json["total_specs"].append({"chrome": chromeList, "firefox": FirefoxList})
      final_json["total_pass_results"] = []
      final_json["total_pass_results"].append(
        {"chrome": chromeList.count('Pass'), "firefox": FirefoxList.count('Pass')})
      final_json["total_fail_results"] = []
      final_json["total_fail_results"].append(
        {"chrome": chromeList.count('Fail'), "firefox": FirefoxList.count('Fail')})
      print(final_json)
      data_header_list = json_to_json(final_json)
      return data_header_list

def elapseTime(jobName, fileName):
  collections=db.testRun.find({"test_name":jobName})
  if collections.count()>0:
    for collection in collections:
      chrome_list=[]
      firefox_list=[]
      json_format={}
      path=collection["path"]
      f=open(path+"\\"+fileName+".log","r")
      for line in f:
        # print(line)
        if line.__contains__("secs)") or line.__contains__("min)") or line.__contains__("mins)"):
          # print(line)
          if line.__contains__("chrome"):
            chrome_list.append(line.split("(")[1].split(")")[0])
          if line.__contains__("firefox"):
            firefox_list.append(line.split("(")[1].split(")")[0])

      json_format["browser_compatibility_test"]="Time Taken Individual"
      chrometimer=''
      print (chrome_list, 'chrome_list')
      for chrome in chrome_list:
        print(chrome, 'chrome')
        chrometimer+= '['+chrome +'],'
      firefoxtimer = ''
      for firefox in firefox_list:
        firefoxtimer += '['+firefox + '],'
      json_format["chrome"]=chrometimer[:-1]
      json_format["firefox"]=firefoxtimer[:-1]
      print(json_format)
      return json_format
  else:
    return "No db found"

@app.route("/browserCompatibilityTest")
def browserCompatibilityTest():
  jobName = request.args.get("test_run")
  parallel_conf = '''
  exports.config={
    framework:"jasmine",
	seleniumAddress:"http://localhost:4444/wd/hub",
    specs:["combined_multi_scenario_file.js"],
	multiCapabilities: [
		{
			"browserName":"chrome",
			"shardTestFiles":true,
			"chromeOptions":{
				"args":["--disable-gpu","-disable-dev-shm-usage","--no-sandbox","-disable-popup-blocking","--start-maximized","--disable-web-security","--allow-running-insecure-content","--disable-infobars"]
			}
		},
		{
			"browserName":"firefox",
			"shardTestFiles":true,
			"moz:firefoxOptions":{
				"args":["--disable-gpu","-disable-dev-shm-usage","--no-sandbox","-disable-popup-blocking","--start-maximized","--disable-web-security","--allow-running-insecure-content","--disable-infobars"]
			},
			'firefoxPath' : 'C:\\Program Files\\Mozilla Firefox\\firefox.exe'
		}
	],
	getPageTimeout:1000000,
	allScriptsTimeout:3000000,
	jasmineNodeOpts:{defaultTimeoutInterval:3000000},
	defaultTimeoutInterval:3000000
};'''
  collections = db.testRun.find()
  for collection in collections:
    if collection["test_name"] == jobName:
      path = collection["path"]
      break
  open(path + "\\parallel_conf.js", "w").close()
  conf_file = open(path+"\\parallel_conf.js",'w+')
  conf_file.write(parallel_conf)
  conf_file.close()
  parallel_batch_file = open(path + "\\parallelTestRun.bat", "w+")
  multipleBatchFile = "START /B cmd " + base_location + " && protractor " + path + "\\parallel_conf.js > " + path + "\\parallel_conf_file.log"
  parallel_batch_file.write(multipleBatchFile)
  parallel_batch_file.close()
  os.chdir(path)
  subprocess.call(path + "\\parallelTestRun.bat")
  conf_data = open(path+"\\parallel_conf_file.log").read()
  # print(conf_data)
  data = getParallelJobDetails(jobName, 'parallel_conf_file')
  json_timer = elapseTime(jobName, 'parallel_conf_file')
  data[0].insert(2, json_timer)
  dataList = json_json_json({"data": data[0], "headers": data[1]},jobName)
  print('T$$T', {"log": conf_data, "data": dataList, "headers":["Test_Job", "Time_Taken_chrome", "Time_Taken_Firefox", "Result_Chrome", "Result_Firefox"]})
  return jsonify({"log":conf_data, "data": dataList, "headers": ["Test_Job", "Time_Taken_chrome", "Time_Taken_Firefox", "Result_Chrome", "Result_Firefox"]})

def uploadUIWithoutRunForBrowserCompatibility(testRunName, sheetLength):
  replace_spec = ''
  head_list = []
  cell_list = []
  filePath = ''
  # folderName = request.args.get("folderName")
  connections = db.testRun.find()
  for connection in connections:
    if connection["test_name"] == testRunName:
      appPath = connection["path"]
      scenario = connection["test_scenarios_used"]
      split_scenario = scenario.split(",")[sheetLength]
      appNameRun = split_scenario.split("[")[1].split("]")[0]
      scenarioNameRun = split_scenario.split("[")[0]
      collections = db.folderName.find()
      for collection in collections:
        if collection["appName"] == appNameRun and collection["scenarioName"] == scenarioNameRun:
          filePath = collection["innerpath"]
      if os.path.isdir(filePath):
        if os.path.exists(filePath + "\\OldFiles\\spec.js"):
          spec_file = open(
            filePath + "\\OldFiles\\spec.js")
          path = appPath +"\\sample.xlsx"

          parallel_multi_scenario_file = open(appPath+"\\"+ "parallel_multi_scenario_file.js","a+")
          parallel_conf_file = open(appPath + "\\parallel_conf_file.js", 'w+')
          parallel_batch_file = open(appPath + "\\parallel_batch_file.bat", "w+")

          wb = xlrd.open_workbook(path)
          sheet = wb.sheet_by_index(0)
          start_row = 1
          end_row = sheet.nrows
          num_of_columns = sheet.ncols
          open_file = spec_file.read()
          split_file = open_file.split("extract")[0]
          if (split_file.endswith("//")):
            split_file = split_file[:-2] + "});"
          new_list = []
          if scenarioNameRun in new_list:
            scenarioNamedict[scenarioNameRun] += 1
          else:
            scenarioNamedict[scenarioNameRun] = 0
          new_list.append(scenarioNameRun)
          sec_test = "it(\'" + scenarioNameRun + " " + str(scenarioNamedict[scenarioNameRun]) + "\', function() {" + \
                     "browser.driver.manage().window().maximize();" + \
                     split_file.split("browser.driver.manage().window().maximize();")[1]

          for r in range(start_row, end_row):
            head_list = []
            cell_list = []
            for c in range(num_of_columns):
              cell_obj = sheet.cell_value(r, c)
              head_obj = sheet.cell_value(0, c)
              head_list.append(head_obj)
              cell_list.append(cell_obj)
            for line in sec_test.split(";"):
              for idx, header in enumerate(head_list):
                line = str(line)
                if (line.__contains__(header) and not line.__contains__("browser.wait(until")
                  and not line.__contains__(".clear(")):
                  if (type(cell_list[idx]) == int or type(cell_list[idx]) == float):
                    line = line.split("sendKeys")[0] + "sendKeys(" + str(int(cell_list[idx])) + ");"
                  else:
                    line = line.split("sendKeys")[0] + "sendKeys('" + str(cell_list[idx]) + "');"
                  break
              replace_spec += line + ";"
          new_spec = '''const jsdom = require("jsdom");
                         const fs = require('fs');
                         const path = require('path');
                         const {SpecReporter} = require("jasmine-spec-reporter");
                          jasmine.getEnv().addReporter(new SpecReporter({spec: {displayStacktrace: true, displayDuration: true}}));
                         describe('Protractor Demo App', function() { ''' + str(replace_spec) + "});"

          if os.stat(appPath + "\\parallel_multi_scenario_file.js").st_size == 0:
            new_parallel_file = '''const jsdom = require("jsdom");
                                     const fs = require('fs');
                                     const path = require('path');
                                     describe('Protractor Demo App', function() { ''' + str(replace_spec)
            parallel_multi_scenario_file.write(new_parallel_file)
            print("parallel_multi_scenario_file", parallel_multi_scenario_file)
          else:
            new_parallel_file = str(replace_spec)
            parallel_multi_scenario_file.write(new_parallel_file)
            print("parallel_multi_scenario_file", parallel_multi_scenario_file)
          parallel_multi_scenario_file.close()

          comb_conf = '''
             exports.config = {
               framework: 'jasmine',
               seleniumAddress: 'http://localhost:4444/wd/hub',
               specs: ['parallel_multi_scenario_file.js'],
               commonCapabilities: {
                'browserName': 'chrome',
                'chromeOptions': {
                                                           'args': [ '--disable-gpu','-disable-dev-shm-usage','--no-sandbox','-disable-popup-blocking','--start-maximized','--disable-web-security','--allow-running-insecure-content','--disable-infobars']
                                           }
              },
              onPrepare: function() {
                 return browser.takeScreenshot();
               },
              getPageTimeout: 20000,
                           allScriptsTimeout: 3000000,
                           jasmineNodeOpts: {defaultTimeoutInterval: 3000000},
                           defaultTimeoutInterval: 3000000
             }

             '''
          parallel_conf_file.write(comb_conf)
          parallel_conf_file.close()

          parallel_batch_file_data = "START /B cmd " + appPath + " && protractor " + appPath + "\\parallel_multi_scenario_file.js >" + appPath  + "\\parallel_conf_file.log"
          parallel_batch_file.write(parallel_batch_file_data)
          parallel_batch_file.close()
          time.sleep(1)
          return "Success"
      else:
        return "appFail"

def uploadingExcelForBrowserCompatibility(file, testRunName):
  import pandas as pd
  global mulipleScenarioPath

  str_file = file.decode("utf-8")

  sheetLength = 0
  json_format = json.loads(str_file)
  print("str_file", str_file)

  # Checking the muliple scenario file exist or not, if exists remove the file
  connections = db.testRun.find()
  for connection in connections:
    if connection["test_name"] == testRunName:
      mulipleScenarioPath = connection["path"]
      if os.path.exists(mulipleScenarioPath + "\\parallel_multi_scenario_file.js"):
        p = os.path.join(mulipleScenarioPath  + "\\parallel_multi_scenario_file.js")
        os.remove(p)

  for k, v in json_format.items():
    v1 = json.dumps(v['data'])
    if type(list(eval(v1))[0]) == "str":
      return jsonify("failure")
    else:
      temp_file = list(eval(v1))
      df = pd.DataFrame.from_dict(temp_file)

    connections = db.testRun.find()
    for connection in connections:
      if connection["test_name"] == testRunName:
        mulipleScenarioPath = connection["path"]
        if os.path.isdir(mulipleScenarioPath):
          df.to_excel(mulipleScenarioPath + "\\" + '\\sample.xlsx',
                      index=False)
          time.sleep(1)
          append_file = open(mulipleScenarioPath  + "\\append_spec.txt", "a+")
          returnVal = uploadUIWithoutRunForBrowserCompatibility(testRunName, sheetLength)
          sheetLength += 1
        else:
          return "appFail"

  comb_multi_file = open(mulipleScenarioPath +"\\parallel_multi_scenario_file.js","a+")
  comb_multi_file.write("});")
  return jsonify("Success")

@app.route("/browserCompatibilityTestWithExcel", methods=["GET","POST"])
def browserCompatibilityTestWithExcel():
  jobName = request.args.get("test_run")
  file = request.get_data()
  print(file)

  catchit = uploadingExcelForBrowserCompatibility(file, jobName)

  print(catchit)

  parallel_conf = '''exports.config = {
    framework: 'jasmine',
    seleniumAddress: 'http://localhost:4444/wd/hub',
    specs: ['parallel_multi_scenario_file.js'],
    multiCapabilities: [
      {
         'browserName': 'chrome',
         'shardTestFiles': true,
         'chromeOptions': {
         'args': [ '--disable-gpu','-disable-dev-shm-usage','--no-sandbox','-disable-popup-blocking','--start-maximized','--disable-web-security','--allow-running-insecure-content','--disable-infobars']
         }
      },
      {
         'browserName': 'firefox',
         'shardTestFiles': true,
         'moz:firefoxOptions': {
         'args': [ '--disable-gpu','-disable-dev-shm-usage','--no-sandbox','-disable-popup-blocking','--start-maximized','--disable-web-security','--allow-running-insecure-content','--disable-infobars']
         },
         'firefoxPath' : 'C:\\Program Files\\Mozilla Firefox\\firefox.exe'
      }
    ],
    onPrepare: function() {
      return browser.takeScreenshot();
   },
   getPageTimeout: 1000000,
   allScriptsTimeout: 3000000,
   jasmineNodeOpts: {defaultTimeoutInterval: 3000000},
   defaultTimeoutInterval: 3000000
  }'''
  collections = db.testRun.find()
  for collection in collections:
    if collection["test_name"] == jobName:
      path = collection["path"]
      break
  open(path + "\\parallel_new_conf.js", "w").close()
  conf_file = open(path+"\\parallel_new_conf.js","w")
  conf_file.write(parallel_conf)
  conf_file.close()
  if not os.path.isfile(path + "\\parallelTestRun1.bat"):
    parallel_batch_file = open(path + "\\parallelTestRun1.bat", "w+")
    multipleBatchFile = "START /B cmd " + path + " && protractor " + path + "\\parallel_new_conf.js >" + path + "\\parallel_new_conf_file.log"
    parallel_batch_file.write(multipleBatchFile)
    parallel_batch_file.close()
  subprocess.call(path + "\\parallelTestRun1.bat")
  conf_data = open(path + "\\parallel_new_conf_file.log").read()
  data = getParallelJobDetails(jobName, 'parallel_new_conf_file')
  json_timer = elapseTime(jobName, 'parallel_new_conf_file')
  data[0].insert(2, json_timer)
  dataList = json_json_json({"data": data[0], "headers": data[1]}, jobName)
  print('T$$T', {"log": conf_data, "data": dataList,
                 "headers": ["Test_Job", "Time_Taken_chrome", "Time_Taken_Firefox", "Result_Chrome", "Result_Firefox"]})
  return jsonify({"log": conf_data, "data": dataList,
          "headers": ["Test_Job", "Time_Taken_chrome", "Time_Taken_Firefox", "Result_Chrome", "Result_Firefox"]})
  # return jsonify({"log": conf_data, "data": data[0], "headers": data[1]})

@app.route("/checkFileExistForParallel")
def checkFileExistForParallel():
  testRunName = request.args.get("testRunName")
  if os.path.exists(base_location + "\\testRun\\"+ testRunName +"\\combined_multi_scenario_file.js"):
    print("success")
    return jsonify({"status":"success"})
  else:
    print("failure")
    return jsonify({"status":"failure"})

def getOverViewDetailsforTestJobDoc(jobName):
    '''
    Get section 1 for generating Test Job Document
    :param jobName: test job name
    :return: json object with section 1 data
    '''
    final_json = {}
    # add job name to json
    final_json["test_job_name"] = jobName
    print(jobName)
    # find all records in job table
    collections = db.testRun.find()
    # loop through cursor
    for collection in collections:
      # find current job record
      if collection["test_name"] == jobName:
        # get scenarios used and path and exit loop
        test_scenarios_used = collection["test_scenarios_used"]
        jobPath = collection["path"]
        break
    # split and find scenarios used and add to final json
    scenarios_used = test_scenarios_used.split(',')
    final_json["test_scenarios_used"] = scenarios_used
    final_json["test_description"] = []
    # loop through scenarios used and do the following
    for scenario_used in scenarios_used:
      test_desc = {}
      # get application name and scenario name for each record
      scenario_name = scenario_used.split('[')[0]
      application_name = scenario_used.split('[')[1][:-1]
      test_desc["scenario_name"] = scenario_name
      test_desc["application_name"] = application_name
      # get all scenarios from scenario table
      scenario_collections = db.folderName.find()
      # for every scenario in job, do the following
      for scenario_collection in scenario_collections:
        # find current loop record
        if scenario_collection["scenarioName"] == scenario_name and \
                scenario_collection["appName"] == application_name:
          # find test description and add to temp json and continue for next record
          test_description = scenario_collection["test_Description"]
          test_desc["test_description"] = test_description
          break
      # add test steps to final json
      final_json["test_description"].append(test_desc)
      test_desc = {}
    # get report location and add to final json
    final_json["report_loc"] = jobPath + "\\new_excel_sheet.xlsx"
    # print('Overview: ', final_json)
    # return final json
    return final_json

def getMetricsDetailsforTestJobDoc(jobName):
  '''
  Get Section 2 to generate Test Job Report Document
  :param jobName: test job name
  :return: json object with dataset
  '''
  final_json = {}
  final_json["test_job_name"] = jobName
  # find all records from test job table
  collections = db.testRun.find()
  # loop through the cursor object, do the following
  for collection in collections:
    # find the corresponding job record
    if collection["test_name"] == jobName:
      # if found, find path
      path = collection["path"]
      # find last modified time of conf.log to find when the batch was run last
      mod_time = time.ctime(os.path.getmtime(path + "\\conf.log"))
      # add the last modified time to final json and exit loop
      final_json["last_run"] = "last modified: %s" % time.ctime(os.path.getmtime(path + "\\conf.log"))
      break
  # read config log generated in path
  confLog = open(path + "\\conf.log").readlines()
  # for every line, do following
  for line in confLog:
    # find Finished time (in sections)
    if line.__contains__("Finished in"):
      time_secs = line.split(" ")[2]
      # add time taken to final json and break loop
      t_min = str(strftime("%H:%M:%S", gmtime(float(line.split(" ")[2])))).split(":")[1]
      t_secs = str(strftime("%H:%M:%S", gmtime(float(line.split(" ")[2])))).split(":")[2]
      final_json["time_elapsed"] = line.split(" ")[2].split(".")[0] + " seconds " + "[" + t_min + "m," + t_secs + "s]"
      break
  final_json["time_stamp"] = []
  time_stamp = {}
  # get end time from last modified time
  print(final_json["last_run"])
  time_stamp["end_time"] = final_json["last_run"].split(": ")[1]
  print(mod_time)
  last_run_time = mod_time.strip().split(" ")[3].strip()
  # to find start time, subtract with time taken
  sub_min = str(strftime("%H:%M:%S", gmtime(float(time_secs)))).split(":")[1]
  sub_secs = str(strftime("%H:%M:%S", gmtime(float(time_secs)))).split(":")[2]
  FMT = '%H:%M:%S'
  print(last_run_time)
  tdelta = datetime.strptime(last_run_time, FMT) - timedelta(minutes=int(sub_min), seconds=int(sub_secs))
  # add start time to temp json
  time_stamp["start_time"] = mod_time.replace(last_run_time, str(tdelta).split(" ")[1])
  # add time stamp to final json
  final_json["time_stamp"].append(time_stamp)
  # initialize pass count and fail count
  passCount = 0
  failCount = 0
  # load workbook to calculate total pass & fail
  wb = load_workbook(path + "\\new_excel_sheet.xlsx")
  # find number of sheets
  tot_sheets = len(wb.worksheets)
  tot_runs = 0
  # for every sheet, do the following
  for sheet in wb.worksheets:
    # find tot rows and tot cols in sheet
    tot_rows = sheet.max_row
    # calculate tot runs with rows count
    tot_runs += tot_rows
    column_count = sheet.max_column
    # for every row, do the following
    for row in sheet.iter_rows(1, tot_rows):
      # find last column value and increase count for pass and fail
      if row[column_count - 1].value == 'Pass':
        passCount += 1
      if row[column_count - 1].value == 'Fail':
        failCount += 1
  # add pass count and fail count to final json
  testResult = {"pass": passCount, "fail": failCount}
  final_json["detailed_result"] = testResult
  # calculate avg
  avg = tot_sheets / (tot_runs - tot_sheets)
  # add avg to final json
  final_json["avgerage_time"] = avg
  # print('Metrics: ',final_json)
  # return final json
  return final_json

def getEnvironmentDetailsforTestJobDoc(jobName):
  '''
  Get Section 3 to populate Test Job Document
  :param jobName: test job name
  :return: final json with data
  '''
  final_json = {}
  final_json["test_job_name"] = jobName
  # get all test job records
  collections = db.testRun.find()
  # loop through cursor object
  for collection in collections:
    # find current job name
    if collection["test_name"] == jobName:
      # get path of the job and exit loop
      path = collection["path"]
      break
  # with path information, read the spec file
  jsFile = open(path + "\\" + jobName + ".js").readlines()
  urls_unedited = []
  # for each line in spec js file, do the following
  for line in jsFile:
    # find all urls used
    if line.__contains__("browser.get("):
      # append the value to list
      urls_unedited.append(line)
  urls = []
  # loop through lines, and extract urls
  for url_unedited in urls_unedited:
    urls.append(url_unedited.strip().replace('\n', '').split('\'')[1].split('\'')[0])
  # append urls to final json
  final_json["urls"] = urls
  # read config file to find browser used
  confFile = open(path + "\\conf.js").readlines()
  final_json["browser"] = []
  # for every line in config file, do the following
  for line in confFile:
    # split the required browser name and save to final json
    if line.__contains__("browserName"):
      final_json["browser"].append(line.split(':')[1].split('\'')[1].split('\'')[0])
  # print('Environment: ',final_json)
  # return final json
  return final_json

def getLogsAndScreenshotsforTestJobDoc(jobName, screenshotFlag):
  '''
  Get section 4 for populating test job document
  :param jobName: Test Job Name, for which information has to be fetched
  :param screenshotFlag: Boolean - True / False to get screenshots on user's details
  :return: json object with data
  '''
  final_json = {}
  final_json["test_job_name"] = jobName
  # Only when user requests for screenshot to be added.
  if screenshotFlag == True or screenshotFlag == 'true':
    final_json["screenshots"] = []
    # Find all records from test job table
    collections = db.testRun.find()
    # loop through the cursor object
    for collection in collections:
      # find current job record
      if collection["test_name"] == jobName:
        # if found, get path and scnearios used
        appPath = collection["path"]
        scenario = collection["test_scenarios_used"]
        split_scenario = scenario.split(",")
        # for ever scenario used, do the following
        for app_scene in split_scenario:
          # get application name and scenario name from the string
          application_name = app_scene.split("[")[1].split("]")[0]
          scenario_name = app_scene.split("[")[0]
          # get all scenarios from scenario table
          scnearios_connections = db.folderName.find()
          # loop through the cursor object
          for scneario_connections in scnearios_connections:
            fileName = []
            # find corresponding record by matching app name and scenario name
            if scneario_connections["appName"] == application_name \
                    and scneario_connections["scenarioName"] == scenario_name:
              # get path and innerpath
              s_appPath = scneario_connections["path"]
              scenarioPath = scneario_connections["innerpath"]
              # check if folder exists
              if os.path.isdir(s_appPath):
                # check if inner folder exists
                if os.path.isdir(scenarioPath):
                  # find all images (.png) files
                  for img in os.listdir(scenarioPath + "\\oldFiles"):
                    if img.endswith(".png"):
                      fileName.append(img)
                  # save image names with path to final json
                  list_names = {s_appPath + '\\' + scenario_name: fileName}
                  final_json["screenshots"].append(list_names)
  else:
    # if user doesn't want screenshots in report
    # find current record in test run table and get its path
    collections = db.testRun.find()
    for collection in collections:
      if collection["test_name"] == jobName:
        appPath = collection["path"]
  # finally get the conf log details and save to final json
  conf_log = open(appPath + "\\conf.log", "r").readlines()
  final_json["log"] = ''.join(conf_log)
  # print(final_json)
  # return final json
  return final_json

def combineReturnJsons(jsonList):
  '''
  Combine list of jsons / dicts to a single json / dict
  :param jsonList: list of json / dict objects
  :return: final json / dict value
  '''
  document_details = {}
  # loop through the list of dict
  for json_dict in jsonList:
    # for ever key in json, do the following
    for k in json_dict:
      # find key's value and append to final json
      document_details[k] = json_dict[k]
  # after processing, return the final json
  return document_details

def valid_xml_char_ordinal(c):
  codepoint = ord(c)
  # conditions ordered by presumed frequency
  return (
          0x20 <= codepoint <= 0xD7FF or
          codepoint in (0x9, 0xA, 0xD) or
          0xE000 <= codepoint <= 0xFFFD or
          0x10000 <= codepoint <= 0x10FFFF
  )

def testDescMapping(outerFolderName, folderName):
  open_file = open(base_location+"\\" + outerFolderName + '\\' + folderName + "\\OldFiles\\spec.js")

  lines = open_file.readlines()
  full_line = ''
  flag = 0
  # global sentence
  sentence = ''
  for line in lines:
    if line.__contains__("extractModal();"):
      break
    if line.strip().startswith("browser.get(") or line.strip().startswith("browser.takeScreenshot("):
      flag = 1

    if flag == 1:
      full_line += line

  print(full_line)
  sentence_list = []
  number = 0
  for item in full_line.split("browser.takeScreenshot("):
    sentence = ''
    for sline in item.split(";"):

      if sline.__contains__("browser.get("):  # Get Url
        number += 1
        url_line=sline.split("browser.get(")[1].split(")")[0]
        if len(url_line) < 50:
          sentence += str(number) + ". Open browser and navigate to URL: " + str(
            url_line) + " <br>"
        else:
          sentence += str(number) + ". Open browser and navigate to URL: " + "<br>"+ str(
            url_line) + " <br>"

      if sline.__contains__("until.presenceOf") == False and \
        sline.__contains__("by.xpath(") and \
        sline.strip().endswith("click()"):  # button click
        if sline.__contains__("@id="):
          number += 1
          sentence += str(number) + ". Click on the button with id " + sline.split("@id=")[1].split("]")[0]+"<br>"
        elif sline.__contains__("@name="):
          number += 1
          sentence += str(number) + ". Click on the button with name " + sline.split("@name=")[1].split("]")[0]+"<br>"
        elif sline.__contains__(".get("):
          number += 1
          x = str(sline.split(".get(")[1].split(")")[0]).strip()
          if x.endswith("1"):
            inBetween = str(floor(int(x) / 10)) + "1st"
          elif x.endswith("2"):
            inBetween = str(floor(int(x) / 10)) + "2nd"
          elif x.endswith("3"):

            inBetween = str(floor(int(x) / 10)) + "3rd"
          else:
            inBetween = x + "th"
          sentence += str(number) + ". Click on " + inBetween + " button in the current page<br>"

      elif sline.__contains__("until.presenceOf") == False and \
        sline.__contains__("by.xpath(") and \
        sline.__contains__("click()") == False and \
        sline.__contains__("sendKeys"):  # Sendkey
        if sline.__contains__("@id="):
          number += 1
          if sline.__contains__(".get("):
            x = str(sline.split(".get(")[1].split(")")[0]).strip()
            if x.endswith("1"):
              inBetween = str(floor(int(x) / 10)) + "1st"
            elif x.endswith("2"):
              inBetween = str(floor(int(x) / 10)) + "2nd"
            elif x.endswith("3"):
              inBetween = str(floor(int(x) / 10)) + "3rd"
            else:
              inBetween = x + "th"
              if sline.lower().__contains__("password"):
                sentence += str(number) + ". Enter value " + "******" + " on the" + inBetween + "field with id " + sline.split("@id=")[1].split("]")[0] + "<br>"
              else:
                sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                  0] + " on the"+inBetween+"field with id " + sline.split("@id=")[1].split("]")[0] + "<br>"
          else:
            if sline.lower().__contains__("password"):
              sentence += str(number) + ". Enter value " + "******" + " in the field with id " + sline.split("@id=")[1].split("]")[0] + "<br>"
            else:
              sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                0] + " in the field with id " + sline.split("@id=")[1].split("]")[0] + "<br>"
        elif sline.__contains__("@name="):
          number += 1
          if sline.__contains__(".get("):
            x = str(sline.split(".get(")[1].split(")")[0]).strip()
            if x.endswith("1"):
              inBetween = str(floor(int(x) / 10)) + "1st"
            elif x.endswith("2"):
              inBetween = str(floor(int(x) / 10)) + "2nd"
            elif x.endswith("3"):
              inBetween = str(floor(int(x) / 10)) + "3rd"
            else:
              inBetween = x + "th"
              if sline.lower().__contains__("password"):
                sentence += str(number) + ". Enter value " + "******"+ " in the field with name " + sline.split("@name=")[1].split("]")[0] + "<br>"
              else:
                sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                  0] + " in the field with name " + sline.split("@name=")[1].split("]")[0] + "<br>"
          else:
            if sline.lower().__contains__("password"):
              sentence += str(number) + ". Enter value " + "*****" + " in the field with name " + sline.split("@name=")[1].split("]")[0] + "<br>"
            else:
              sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
              0] + " in the field with name " + sline.split("@name=")[1].split("]")[0] + "<br>"
        elif sline.__contains__(".get(") and sline.__contains__("@id=")==False:
          number += 1
          x = str(sline.split(".get(")[1].split(")")[0]).strip()
          if x.endswith("1"):
            inBetween = str(floor(int(x) / 10)) + "1st"
          elif x.endswith("2"):
            inBetween = str(floor(int(x) / 10)) + "2nd"
          elif x.endswith("3"):
            inBetween = str(floor(int(x) / 10)) + "3rd"
          else:
            inBetween = x + "th"
          # sentence += "Enter value " + sline.split("sendKeys(")[1].split(")")[0] + " in the field with id " + \
          #             inBetween+" in the field from the webpage<br>"
          sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
            0] + " on " + inBetween + " field in the current page<br>"
    if sentence != "":
      sentence = sentence.replace("<br>", "\n")
      sentence_list.append(sentence.strip())

  myquery = {"appName": outerFolderName, "scenarioName": folderName}
  newvalues = {"$set": {"test_Description": sentence}}
  # db.folderName.update_many(myquery, newvalues)
  print(sentence)
  print(sentence_list)
  print(len(sentence_list))

  final_list = []
  img_no = 0
  for sen in sentence_list:
    sen_dict ={}
    sen_dict["ScenarioName"] = folderName
    sen_dict["TestStep"] = sen
    # sen_dict["ScreenShot"] = "Page" + str(img_no+1) + ".png"
    img_no = img_no+1
    final_list.append(sen_dict)
  print(final_list)
  return final_list

def DocWrite(TestJobDetails):
    Mergeval = TestJobDetails

    if "screenshots" in Mergeval.keys():
      # Define the templates - assumes they are in the same directory as the code
      template_path = base_location + "\\" + "Detailed Test Report - with SS.docx"
      pythoncom.CoInitialize()
      word = win32.gencache.EnsureDispatch('Word.Application')
      pythoncom.CoInitialize()
      doc = word.Documents.Open(template_path)
      word.Visible = False

      try:
        if os.path.exists(Mergeval['report_loc']):
          doc.InlineShapes.AddOLEObject(FileName=Mergeval['report_loc'], DisplayAsIcon=1,
                                        Range=doc.Paragraphs(47).Range, IconLabel="Individual Data Set Report",
                                        IconFileName="Individual Data Set Report",
                                        LinkToFile=False)

        # number = 63
        # for item in Mergeval['screenshots']:
        #     print(item)
        #     for k in item:
        #         print(item[k])
        #         for v in item[k]:
        #             folderName = k.split("\\")[-1]
        #             if os.path.exists(k + "\\OldFiles"+"\\" + v):
        #                 doc.InlineShapes.AddOLEObject(FileName=k + "\\OldFiles" + "\\" + v, DisplayAsIcon=1,
        #                                               Range=doc.Paragraphs(number+3).Range, IconLabel=folderName+"-"+v,
        #                                               IconFileName="Image",
        #                                               LinkToFile=False)

        word.ActiveDocument.SaveAs(
          base_location + "\\" + Mergeval['test_job_name'] + ".docx")
        doc.Close()
      except Exception as e:
        print(e)

    else:
      # Define the templates - assumes they are in the same directory as the code
      template_path = base_location + "\\" + "Detailed Test Report.docx"
      pythoncom.CoInitialize()
      word = win32.gencache.EnsureDispatch('Word.Application')
      pythoncom.CoInitialize()
      doc = word.Documents.Open(template_path)
      word.Visible = False

      try:
        if os.path.exists(Mergeval['report_loc']):
          doc.InlineShapes.AddOLEObject(FileName=Mergeval['report_loc'], DisplayAsIcon=1,
                                        Range=doc.Paragraphs(46).Range, IconLabel="Report",
                                        IconFileName="Individual Data Set Report",
                                        LinkToFile=False)

        word.ActiveDocument.SaveAs(
          base_location + "\\" + Mergeval['test_job_name'] + ".docx")
        doc.Close()
      except Exception as e:
        print(e)

    cleaned_string = ''.join(c for c in Mergeval['log'] if valid_xml_char_ordinal(c))
    test_description = ""

    for item in Mergeval['test_description']:
      print(item['scenario_name'])
      test_description += "\n" + item['scenario_name'] + "\n"
      test_description += item['test_description'].replace("<br>", "\n") + "\n"
    # Show a simple example
    document_1 = MailMerge(base_location + "\\" + Mergeval['test_job_name'] + ".docx")
    print("Fields included in {}: {}".format(template_path,
                                             document_1.get_merge_fields()))

    Separator = "\n"
    document_1.merge(Date='{:%d-%b-%Y}'.format(date.today()),
                     TestJobName=str(Mergeval['test_job_name']),
                     Scenario=str(Separator.join(Mergeval['test_scenarios_used'])),
                     ScenarioDetails=test_description,
                     LastRun=str(Mergeval['last_run']),
                     StartTime=str(Mergeval['time_stamp'][0]['start_time']),
                     EndTime=Mergeval['time_stamp'][0]['end_time'],
                     TimeElapsed=Mergeval['time_elapsed'],
                     TestCasePassed=str(Mergeval['detailed_result']['pass']),
                     TestCaseFailed=str(Mergeval['detailed_result']['fail']),
                     URL=str(Separator.join(Mergeval['urls'])),
                     Browser=str(Separator.join(Mergeval['browser'])),
                     Logs=cleaned_string
                     )
    if "screenshots" in Mergeval.keys():

      test_desc_mapping_list = []

      for item in Mergeval['screenshots']:
        print(item)
        for k in item:
          print(item[k])
          for v in item[k]:
            folderName = k.split("\\")[-1]
            outerFolderName = k.split("\\")[-2]
            if os.path.exists(k + "\\OldFiles" + "\\" + v):
              catch_it = testDescMapping(outerFolderName, folderName)
              print(catch_it)
              print(len(catch_it))
              test_desc_mapping_list.append(catch_it)
              break
      print(test_desc_mapping_list)

      tDesc = []
      for item in test_desc_mapping_list:
        for ite in item:
          tDesc.append(ite)
      print(tDesc)

      document_1.merge_rows('TestStep', tDesc)

      fileName = "Detailed Test Report " + Mergeval['test_job_name'] + '.docx'
      # Get all data in test job database
      connections = db.testRun.find()
      # Loop through cursor object
      for connection in connections:
        # Find current Job record
        if connection["test_name"] == Mergeval['test_job_name']:
          # If found, check if document exists
          document_1.write(connection["path"] + "\\" + fileName)
          doc = Document(connection["path"] + "\\" + fileName)
          # document_1.write(base_location + "\\Detailed Test Report " + Mergeval['test_job_name'] + ".docx")
          # doc = Document(base_location + "\\Detailed Test Report " + Mergeval['test_job_name'] + ".docx")
          tables = doc.tables

          outerFolderName = Mergeval["test_scenarios_used"][0].split("[")[1].strip().replace("]", "")

          iteration_num = len(tDesc)
          print(iteration_num)
          num = 0
          for catch_it in test_desc_mapping_list:
            for i in range(len(catch_it)):
              print(i)
              folderName = catch_it[0]['ScenarioName']
              if num <= iteration_num:
                p = tables[0].rows[num + 1].cells[2].add_paragraph()
                r = p.add_run()

                img_loc = base_location + "\\" + outerFolderName + "\\" + folderName + "\\OldFiles\\" + 'page' + str(
                  i + 1) + '.png'
                print(img_loc)
                if os.path.exists(img_loc):
                  r.add_picture(img_loc, width=Inches(3.0), height=Inches(2.0))
                num += 1
          # doc.save(base_location + "\\Detailed Test Report " + Mergeval['test_job_name'] + ".docx")
          doc.save(connection["path"] + "\\" + fileName)
          break

          # Save the document as example 1
      # document_1.write(base_location + "\\Detailed Test Report " + Mergeval['test_job_name'] + ".docx")
      # doc = Document(base_location + "\\Detailed Test Report " + Mergeval['test_job_name'] + ".docx")
      # tables = doc.tables
      #
      # outerFolderName = Mergeval["test_scenarios_used"][0].split("[")[1].strip().replace("]", "")
      #
      # iteration_num = len(tDesc)
      # print(iteration_num)
      # num = 0
      # for catch_it in test_desc_mapping_list:
      #     for i in range(len(catch_it)):
      #         print(i)
      #         folderName = catch_it[0]['ScenarioName']
      #         if num <= iteration_num:
      #             p = tables[0].rows[num+1].cells[2].add_paragraph()
      #             r = p.add_run()
      #
      #             img_loc = base_location + "\\" + outerFolderName + "\\" + folderName + "\\OldFiles\\" + 'page' + str(
      #                 i + 1) + '.png'
      #             print(img_loc)
      #             if os.path.exists(img_loc):
      #                 r.add_picture(img_loc, width=Inches(3.0), height=Inches(2.0))
      #             num += 1
      # doc.save(base_location + "\\Detailed Test Report " + Mergeval['test_job_name'] + ".docx")
    else:
      fileName = "Detailed Test Report " + Mergeval['test_job_name'] + '.docx'
      # Get all data in test job database
      connections = db.testRun.find()
      # Loop through cursor object
      for connection in connections:
        # Find current Job record
        if connection["test_name"] == Mergeval['test_job_name']:
          # If found, check if document exists
          document_1.write(connection["path"] + "\\" + fileName)
          break
      # document_1.write(base_location + "\\Detailed Test Report " + Mergeval['test_job_name'] + ".docx")

    return "Success"

@app.route("/generateAndMoveDoc")
def generateAndMoveDoc():
  jobName = request.args.get("testRunName")
  screenshotFlag = request.args.get("screenshotFlag")
  print(screenshotFlag, 'screenshotFlag')
  group_jsons = []
  overviewDetailsJson = getOverViewDetailsforTestJobDoc(jobName)
  group_jsons.append(overviewDetailsJson)
  metricDetailJson = getMetricsDetailsforTestJobDoc(jobName)
  group_jsons.append(metricDetailJson)
  environmentDetailJson = getEnvironmentDetailsforTestJobDoc(jobName)
  group_jsons.append(environmentDetailJson)
  screenshotAndLogDetailJson = getLogsAndScreenshotsforTestJobDoc(jobName, screenshotFlag)
  group_jsons.append(screenshotAndLogDetailJson)
  document_detials = combineReturnJsons(group_jsons)
  print('Document Generation Data: ', document_detials)

  # call document generation method to generate
  DocWrite(document_detials)

  destination_folder = ui_location + "\\src\\assets\\doc\\"
  # Frame the string of file name
  fileName = "Detailed Test Report " + jobName + '.docx'
  # Get all data in test job database
  connections = db.testRun.find()
  # Loop through cursor object
  for connection in connections:
    # Find current Job record
    if connection["test_name"] == jobName:
      # If found, check if document exists
      if os.path.exists(connection["path"] + "\\" + fileName) == True:
        # If exists, Copy file to UI folder & Return Success
        shutil.copy(connection["path"] + "\\" + fileName, destination_folder)
        return jsonify("Success")
      else:
        # If doesn't exist, Return Failure
        return jsonify("File Not Generated")

@app.route("/exportScenarioAsZip")
def exportScenarioAsZip():
  app_name = request.args.get("appName")
  scenario_name = request.args.get("scenarioName")
  if os.path.isdir(base_location + "\\zip_files\\" + scenario_name == False):
    os.mkdir(base_location + "\\zip_files\\" + scenario_name)
  output_filename = base_location + "\\zip_files\\" + scenario_name
  dir_name = base_location + "\\" + app_name + "\\" + scenario_name
  archive_format = "zip"
  collections = db["folderName"].find({"$and":[{"appName":app_name},{"scenarioName":scenario_name}]})
  js_struct = {}
  js_file=open(base_location+"\\"+app_name+"\\"+scenario_name+"\\"+'json_file.json',"w+")
  for collection in collections:
    print("collection",collection)
    js_struct["appName"]=app_name
    js_struct["scenarioName"]=scenario_name
    js_struct["test_Description"]=collection["test_Description"]
    js_struct["path"]=collection['path']
    js_struct["innerpath"]=collection["innerpath"]
    print(js_struct)
    js_file.write(json.dumps(js_struct))
    js_file.close()
    break
  shutil.make_archive(output_filename, archive_format, dir_name)
  destination_folder = ui_location + "\\src\\assets\\zip_files\\"

  shutil.copy(output_filename+'.zip', destination_folder)
  time.sleep(2)
  return jsonify("Success")

def populateDb(tmpLoc):
  with open(tmpLoc+"\\json_file.json", "r") as f:
    data = json.load(f)
  print(data)

  appName = data["appName"]
  scenarioName = data["scenarioName"]

  connections = db.folderName.find()
  print(connections)

  added_updated_count = 0
  for collection in connections:
    if appName == collection["appName"]:
      if scenarioName == collection["scenarioName"]:
        print("this falling under this category")
        myquery = {"appName": data["appName"], "scenarioName": data["scenarioName"]}
        newvalues = {"$set": {"test_Description": data["test_Description"], "path": data["path"], "innerpath": data["innerpath"]}}
        db.folderName.update_many(myquery, newvalues)
        added_updated_count += 1
        break
      else:
        dbData = [{"path": data["path"],
                   "innerpath": data["innerpath"],
                   "appName": data["appName"],
                   "scenarioName": data["scenarioName"],
                   "test_Description": data["test_Description"]}]
        db.folderName.insert_many(dbData)
        added_updated_count += 1
        break
    else:
      continue
  if added_updated_count == 0:
    dbData = [{"path": data["path"],
               "innerpath": data["innerpath"],
               "appName": data["appName"],
               "scenarioName": data["scenarioName"],
               "test_Description": data["test_Description"]}]
    db.folderName.insert_many(dbData)
  shutil.copytree(tmpLoc, data["innerpath"])
  return "Success"

@app.route('/uploadRecordedZip', methods=["GET","POST"])
def uploadRecordedZip():
  file = request.get_data()
  fname = request.args.get("fname")
  z = zipfile.ZipFile(io.BytesIO(file))
  if os.path.isdir(base_location + "\\TmpZip") == False:
    os.mkdir(base_location + "\\TmpZip")
  z.extractall(base_location + "\\TmpZip")
  populateDb(base_location + "\\TmpZip")
  return jsonify("Success")

@app.route('/uploadFile', methods=["GET","POST"])
def uploadFile():
  from PIL import Image
  from io import BytesIO, StringIO
  from docxtpl import DocxTemplate
  from docx import Document
  import xlsxwriter

  file = request.get_data()
  fname = request.args.get("fname")
  fileSavePath = base_location + "\\uploader"
  if os.path.isdir(fileSavePath) == False:
    os.path.mkdir(fileSavePath)
  fileExtension = fname.split(".")[1]
  print('file: ',file)
  # for image
  if fileExtension.lower() == 'jpg' or fileExtension.lower() == 'jpeg' or \
          fileExtension.lower() == 'png':
    trial = file.split(b"Content-Type: ")[1].split(b"------WebKitForm")[0]
    if fileExtension.lower() == 'jpg':
      if str(trial).__contains__("image/jpg"):
        trial = trial.split(b"image/jpg")[1]
      if str(trial).__contains__("image/jpeg"):
        trial = trial.split(b"image/jpeg")[1]
    if fileExtension.lower() == 'jpeg':
      if str(trial).__contains__("image/jpg"):
        trial = trial.split(b"image/jpg")[1]
      if str(trial).__contains__("image/jpeg"):
        trial = trial.split(b"image/jpeg")[1]
    if fileExtension.lower() == 'png':
      trial = trial.split(b"image/png")[1]
    while 1:
      if trial.startswith(b"\n") or trial.startswith(b"\r"):
        trial=trial[2:]
      elif trial.endswith(b"\n") or trial.endswith(b"\r"):
        trial=trial[:-2]
      else:
        break
    stream = io.BytesIO(trial)
    stream.seek(0)
    image = Image.open(stream)
    if os.path.exists(fileSavePath + '\\' + fname):
      p = os.path.join(fileSavePath + '\\' + fname)
      os.remove(p)
    image.save(fileSavePath + "\\" +fname)
  elif fileExtension.lower() == 'pdf' or fileExtension.lower() == 'txt':
    if fileExtension.lower() == 'txt':
      trial = file.split(b"Content-Type: ")[1].split(b"------WebKitForm")[0]
      trial = trial.split(b"text/plain")[1]
      while 1:
        if trial.startswith(b"\n") or trial.startswith(b"\r"):
          trial = trial[2:]
        elif trial.endswith(b"\n") or trial.endswith(b"\r"):
          trial = trial[:-2]
        else:
          break
        open(fileSavePath + "\\" + fname, 'wb').close()
        f = open(fileSavePath + "\\" + fname, 'wb')
        f.write(trial)
        f.close()
    else:
      open(fileSavePath+"\\"+fname, 'wb').close()
      f = open(fileSavePath+"\\"+fname, 'wb')
      f.write(file)
      f.close()
  elif fileExtension.lower() == 'doc' or fileExtension.lower() == 'docx':
    print(file)
    trial = file.split(b"Content-Type: ")[1].split(b"------WebKitForm")[0]
    trial = trial.split(b"application/vnd.openxmlformats-officedocument.wordprocessingml.document")[1]
    while 1:
      if trial.startswith(b"\n") or trial.startswith(b"\r"):
        trial=trial[2:]
      elif trial.endswith(b"\n") or trial.endswith(b"\r"):
        trial=trial[:-2]
      else:
        break
    print(trial)
    f = open(fileSavePath + "\\" + fname, "wb")
    f.write(trial)
    f.close()
    print('docx saved')
  elif fileExtension.lower() == 'xlsx' or fileExtension.lower()=='xls':
    str_file = file.decode("utf-8")
    json_format = json.loads(str_file)
    if os.path.exists(fileSavePath + '\\' + fname):
      p = os.path.join(fileSavePath + '\\' + fname)
      os.remove(p)
    writer = pd.ExcelWriter(fileSavePath + '\\' + fname)
    i=0
    for k, v in json_format.items():
      i+=1
      v1 = json.dumps(v['data'])
      if type(list(eval(v1))[0]) == "str":
        return jsonify("failure")
      else:
        temp_file = list(eval(v1))
        df = pd.DataFrame.from_dict(temp_file)
        df.to_excel(writer, 'Sheet '+ str(i), index=False)
    writer.save()
  else:
    open(fileSavePath + "\\" + fname, 'wb').close()
    f = open(fileSavePath + "\\" + fname, 'wb')
    f.write(file)
    f.close()
  return jsonify("Success")

@app.route("/browsercheck")
def browsercheck():
  browsers=["chrome","firefox"]
  for b in browsers:
    flag=0
    for app in winapps.search_installed(b):
      if flag==0:
        print(app)
        flag=1
    else:
      if flag==0:
        print("browser", b, "is not installed")
        return jsonify({"status": "failure", "browser": b})
  return jsonify({"status": "success"})

@app.route("/combineFiles")#combined scenarios
def combine_files():# removing the common part of scenerios and combining the uncommon part
  scenarios = request.args.get("scenarios")
  outerFolderName = request.args.get("appName")
  folderName = request.args.get("testRunName")
  # outerFolderName="combineFolder"
  print("scenarios",scenarios)
  if os.path.isdir(base_location + "\\" + outerFolderName ) == False:
    os.mkdir(base_location + "\\" + outerFolderName)
  # scenarios = "Deposit_Flow[XYZ_Bank],OpenAccountFlow[XYZ_Bank],Withdraw_flow[XYZ_Bank]"
  # testRunName = request.args.get("testRunName")
  # split_scenario=scenarios.split(",")
  split_scenario = scenarios.split(",")
  array = []
  scenario_array=[]
  dictionary = {}
  for uniq in split_scenario:
    application_name = uniq.split("[")[1].split("]")[0].strip()
    array.append(str(application_name))
    scenario_name=uniq.split("[")[0].strip()
    scenario_array.append(scenario_name)

  if len(set(array)) == 1:
    # folderName = array[0]
    for scene_app in split_scenario:
      scenario_name = scene_app.split("[")[0].strip()
      application_name = scene_app.split("[")[1].split("]")[0].strip()
      dictionary[scenario_name] = []
      print(scenario_name)
      print(application_name)
      file1 = open(base_location + "\\" + application_name + "\\" + scenario_name + "\\OldFiles\spec.js").read()
      f_line = file1.split("protractor.ExpectedConditions;")[1].split("extract();")[0]
      if f_line.strip().endswith("//"):
        f_line = f_line[:-2]
      split_line = f_line.split(";")
      for l1 in split_line:
        if l1.__contains__("sendKeys") or l1.__contains__("click") or l1.__contains__("clear")\
                or l1.__contains__("switchTo") or l1.__contains__("browser.takScreenshot") :
          dictionary[scenario_name].append(l1.strip().strip("\n").strip("\t"))
    # print(len(dictionary),dictionary.values())
    full_array = dictionary.values()
    full_array = list(full_array)
    print(full_array)
    # i = 0
    # line = 0

    # for l in range(len(full_array)):
    i = 0
    f = 0
    out = ''
    while (1):
      # print("i value",i)
      for l in range(len(full_array) - 1):
        # print("full array",len(full_array))
        if i < len(full_array[l]) and i < len(full_array[l + 1]) and full_array[l][i] == full_array[l + 1][i] \
                and full_array[l][i].__contains__("browser.takeScreenshot") == False \
                and full_array[l + 1][i].__contains__("browser.takScreenshot") == False:
          # print(l,full_array[l][i])
          pass
        else:
          f = 1
          break
      if f == 1:
        break
      i += 1
    var = 2
    tot_string = ''

    # for lis in full_array[1:]:
    #   # print("lists",lis[i:])
    #   out = ''
    #   tot_string+="\n//nextDB\n"
    #   for n in lis[i:]:
    #     out += n + ";\nbrowser.sleep(2000);\n"
    #   print(out)
    #   tot_string += out
    #   var += 1
    for lis in full_array[1:]:
      # print("lists",lis[i:])
      tot_string += "\n//nextDB\n"
      out = ''
      # var_x=''
      page_num = 0
      for index, n in enumerate(lis[i:]):
        if n.strip().startswith('//'):
          continue
        if index % 3 == 0:
          page_num += 1
          out += '''
           browser.takeScreenshot().then(function (png) {var dir="''' + base_location.replace('\\','\\\\') + "\\\\" + outerFolderName + "\\\\" + folderName + "\\\\"'''OldFiles\\\\\";
           var fname="page''' + str(page_num) + '''.png";
           var stream1 = fs.createWriteStream(path.join(dir, fname));
           stream1.write(new Buffer(png, 'base64'));
           stream1.end();
          });
           '''
        if n.__contains__("sendKeys("):
          var_x = n.split(".sendKeys(")[0]
        elif n.__contains__("click("):
          var_x = n.split(".click(")[0]
        else:
          var_x = n
        if str(var_x).__contains__("alert()") == False:
          out += '''\n\t\t\tbrowser.wait(until.presenceOf(''' + var_x + '''),delay, 'Element taking too long to appear in the DOM');\n'''
        out += "\t\t\t" + n + ";\n\t\t\tbrowser.sleep(2000);\n"
      print(out)
      tot_string += out
      var += 1
    print("split", split_scenario[0])
    scene_app = split_scenario[0]
    scenario_name = scene_app.split("[")[0].strip()
    application_name = scene_app.split("[")[1].split("]")[0].strip()
    # f4 = open(base_location + "\\" + application_name + "\\" + scenario_name + "\\OldFiles\\new_spec1.js", "w+")
    if os.path.isdir(base_location + "\\" + outerFolderName + "\\" + folderName)==False:
      os.mkdir(base_location + "\\" + outerFolderName + "\\" + folderName)
    if os.path.isdir(base_location + "\\" + outerFolderName + "\\" + folderName + "\\OldFiles")==False:
      os.mkdir(base_location + "\\" + outerFolderName + "\\" + folderName + "\\OldFiles")
    new_file = open(base_location + "\\" +  application_name + "\\" + scenario_name + "\\OldFiles\spec.js").readlines()
    new_spec_file = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\\OldFiles\spec.js",
                         "w+")
    # new_spec_file.write(new_file)
    f3 = new_file
    f7 = ''
    for ind, l in enumerate(f3):
      # print(l)
      if l.__contains__("extract()"):
        # print("inside smthng")
        # print(l,ind, f3[ind+1])
        if str(f3[ind + 1]).strip().__contains__("extractModal()"):
          # print("inside somthing")
          f3[ind-1] = tot_string + ";\nbrowser.sleep(2000);\n"
    # new_spec_file.write(f3)
    print("f3", f3)
    for f in f3:
      f7 += f
    new_spec_file.write(f7)

    conf = '''

      exports.config = {
        framework: 'jasmine',
        seleniumAddress: 'http://localhost:4444/wd/hub',
        specs: ['spec.js'],
        commonCapabilities: {
         'browserName': 'chrome',
         'chromeOptions': {
                                                    'args': [ '--disable-gpu','-disable-dev-shm-usage','--no-sandbox','-disable-popup-blocking','--start-maximized','--disable-web-security','--allow-running-insecure-content','--disable-infobars']
                                    }
       },
       onPrepare: function() {
          return browser.takeScreenshot();
        },
       getPageTimeout: 20000,
                    allScriptsTimeout: 3000000,
                    jasmineNodeOpts: {defaultTimeoutInterval: 3000000},
                    defaultTimeoutInterval: 3000000
      }


      '''
    scene_app = split_scenario[0]
    f5 = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\\OldFiles\\conf.js", "w+")
    f5.write(conf)
    f5.close()
    cbn_batch = '''
      START /B cmd "cd ''' + base_location + "\\\" && protractor " + base_location + "\\" + outerFolderName + "\\" + folderName + "\\OldFiles\\conf.js >" + base_location + "\\" + outerFolderName + "\\" + folderName + '''\\OldFiles\\conf.log
      '''
    if os.path.isdir(base_location + "\\" + outerFolderName + "\\" + folderName + "\\Batchfiles") == False:
      os.mkdir(base_location + "\\" + outerFolderName + "\\" + folderName + "\\Batchfiles")
    f6 = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\\Batchfiles\Test_batch.bat", "w+")
    f6.write(cbn_batch)
    f6.close()
    dbData = [{"path": base_location + "\\" + outerFolderName,
               "innerpath": base_location + "\\" + outerFolderName + "\\" + folderName,
               "appName": outerFolderName,
               "scenarioName": folderName,
               "test_Description": ''}]
    db.folderName.insert_many(dbData)
    combineTestDescription(outerFolderName, folderName,scenario_array)
    spec_screenshot(outerFolderName, folderName)
    format_spec(outerFolderName, folderName)
    num_img(outerFolderName, folderName)
    #subprocess.call(base_location + "\\" + outerFolderName + "\\" + folderName + "\\Batchfiles\combineBatch.bat")
    return jsonify("Success")
  else:
    return jsonify("Fail")
  # f1=open(r"D:\ProtractorWebDriverTest\Student_Portal\Delete_Student\OldFiles\spec.js").read()
  # f2=open(r"D:\ProtractorWebDriverTest\Student_Login\Add_User\OldFiles\spec.js").read()

  # for scene_app in split_scenario:
  #   scenario_name = scene_app.split("[")[0].strip()
  #   application_name = scene_app.split("[")[1].split("]")[0].strip()

def combineTestDescription(outerFolderName,folderName,scenario_array): # test description for combind scenerio
  full_line = ''
  for arr in scenario_array:
    full_line+=arr+"<br>"
  sentence = "Combined Scenarios of :<br>"+full_line[:-4]
  print("combined sentence",sentence)
  myquery = {"appName": outerFolderName, "scenarioName": folderName}
  newvalues = {"$set": {"test_Description": sentence}}
  db.folderName.update_many(myquery, newvalues)
  return "done"

@app.route("/getAllAppNames")
def getAllAppNames():
  appName=[]
  collections = db.folderName.find()
  for collection in collections:
    appName.append(collection["appName"])
  appName = list(dict.fromkeys(appName))
  return jsonify(appName)

@app.route("/getAllScenariosWithAppName")
def getAllScenariosWithAppName():
  scenariosList = []
  appName = request.args.get("appName")
  collections = db.folderName.find()
  for collection in collections:
    if collection["appName"] == appName:
      scenariosList.append(collection["scenarioName"] + '[' + collection["appName"] + ']')
  return jsonify(scenariosList)

def format_spec(outerFolderName, folderName):
  print("insid eformat spec")
  file_one=open(base_location+"\\"+outerFolderName+"\\"+folderName+"\OldFiles\spec.js").readlines()
  file_wirte=open(base_location+"\\"+outerFolderName+"\\"+folderName+"\OldFiles\spec_write.js","w+")
  string=""
  flag=0
  for f in file_one:
    if f.__contains__("browser.ignoreSynchronization"):
      print("inside")
      f=f.replace(";",";\n\t\t\t")
      # f=f
      flag=1
    if flag==1:
      # print(flag, "inside")
      f =  f.replace(";",";\n\t\t\t")
      f = "\t" + f
    if f.__contains__("extract()"):
      # print(f)
      flag=0
    if flag==0:
      f=f.replace(";",";\n")
    if f.strip()!='':
      string+=f

  # string=string.replace("\n","")
  whole_string=''''''
  for s in string.split("\n"):
    if s.strip()!="":
      whole_string+=s+"\n"
  file_wirte.write(whole_string)
  spec_file=open(base_location+"\\"+outerFolderName+"\\"+folderName+"\OldFiles\spec.js","w+")
  new_spec = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\OldFiles\spec_write.js").read()
  n_line=''
  for s1 in new_spec:
    n_line+=s1
  spec_file.write(n_line)

def spec_screenshot(outerFolderName, folderName):
  spec_file=open(base_location+"\\"+outerFolderName+"\\"+folderName+"\\"+"OldFiles\\spec.js").readlines()

  whole_str=''
  for lin in spec_file:
    # print(lin)

    if lin.__contains__("browser.takeScreenshot()") and (lin.__contains__(outerFolderName)==False or\
      lin.__contains__(folderName)==False):
      print(lin)
      path="D:\\\\ProtractorWebDriverTest\\\\"+outerFolderName+"\\\\"+folderName+"\\\\"+"OldFiles\\"
      full_lin=lin.split("dir=")[0]+"dir=\""+path+"\\\";\n"+lin.split(";",1)[1]
      print(full_lin)
      whole_str+=full_lin
    else:
      whole_str += lin
  spec_file = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\\" + "OldFiles\\spec.js", "w+")
  spec_file.write(whole_str)

def num_img(outerFolderName, folderName):
  spec_file = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\\" + "OldFiles\\spec.js").readlines()
  num_line=0
  whole_line=''
  for lin in spec_file:
    if  lin.__contains__(".png"):
      num_line+=1
      full_line=lin.split("page")[0]+"page"+str(num_line)+"."+lin.split(".")[1]
      print(full_line)
      print(lin)
      whole_line+=full_line
    else:
      whole_line+=lin
  if whole_line.__contains__("async function extractModal"):
    whole_line = whole_line[:whole_line.index("async function extractModal")] + "});"
  spec_file = open(base_location + "\\" + outerFolderName + "\\" + folderName + "\\" + "OldFiles\\spec.js","w+")
  spec_file.write(whole_line)

@app.route("/getCombinedTestDescription")
def getCombinedTestDescription():
    appName = request.args.get("appName")
    scenarioName = request.args.get("scenarioName")
    value = modifiedtestDescription(appName, scenarioName)
    if value != '':
        json_format = {
            "status": "success",
            "test_Description": value
        }
    else:
        json_format = {
            "status": "fail",
            "test_Description": value
        }
    return jsonify(json_format)
    # collections=db.folderName.find({"$and":[{"appName":appName},{"scenarioName":scenarioName}]})
    # if collections.count_documents>0:
    #   for collection in collections:
    #     path=collection["innerpath"]
    #     spec_file=open(path+"\\OldFiles\\spec.js")

def modifiedtestDescription(outerFolderName, folderName):
    # folderName = request.args.get("folderName")
    # folderName="XYZ_Deposit"
    # global outerFolderName
    # global folderName
    # outerFolderName="1View_Tool"
    # folderName="Test48"
    if os.path.exists(base_location + "\\" + outerFolderName + '\\' + folderName + "\oldFiles\spec.js"):
        open_file = open(base_location + "\\" + outerFolderName + '\\' + folderName + "\oldFiles\spec.js")
        # new_file = open(r"base_location\\" + outerFolderName + '\\' + folderName + "\oldFiles\TestDescription.html",
        #                 "w+")
        lines = open_file.readlines()
        full_line = ''
        flag = 0
        # global sentence
        sentence = 'Combined Scenarios Test Steps :<br>'
        for line in lines:
            if line.__contains__("extractModal();"):
                break
            if line.strip().startswith("browser.get("):
                flag = 1
            if flag == 1:
                full_line += line

        number = 0
        for sline in full_line.split(";"):
            if sline.__contains__("browser.get("):  # Get Url
                number += 1
                url_line = sline.split("browser.get(")[1].split(")")[0]
                if len(url_line) < 50:
                    sentence += str(number) + ". Open browser and navigate to URL: " + str(
                        url_line) + " <br>"
                else:
                    sentence += str(number) + ". Open browser and navigate to URL: " + "<br>" + str(
                        url_line) + " <br>"
            # if sline.__contains__("browser.sleep("):
            #   sentence+="Browser will wait for "+sline.split("browser.sleep(")[1].split(")")[0]+"MilliSeconds<br>"
            if sline.__contains__("until.presenceOf") == False and \
                    sline.__contains__("by.xpath(") and \
                    sline.strip().endswith("click()"):  # button click
                if sline.__contains__("@id="):
                    number += 1
                    sentence += str(number) + ". Click on the button with id " + sline.split("@id=")[1].split("]")[
                        0] + "<br>"
                elif sline.__contains__("@name="):
                    number += 1
                    sentence += str(number) + ". Click on the button with name " + sline.split("@name=")[1].split("]")[
                        0] + "<br>"
                elif sline.__contains__(".get("):
                    number += 1
                    x = str(sline.split(".get(")[1].split(")")[0]).strip()
                    if x.endswith("1"):
                        inBetween = str(floor(int(x) / 10)) + "1st"
                    elif x.endswith("2"):
                        inBetween = str(floor(int(x) / 10)) + "2nd"
                    elif x.endswith("3"):

                        inBetween = str(floor(int(x) / 10)) + "3rd"
                    else:
                        inBetween = x + "th"
                    sentence += str(number) + ". Click on " + inBetween + " button in the current page<br>"

            elif sline.__contains__("until.presenceOf") == False and \
                    sline.__contains__("by.xpath(") and \
                    sline.__contains__("click()") == False and \
                    sline.__contains__("sendKeys"):  # Sendkey
                if sline.__contains__("@id="):
                    number += 1
                    if sline.__contains__(".get("):
                        x = str(sline.split(".get(")[1].split(")")[0]).strip()
                        if x.endswith("1"):
                            inBetween = str(floor(int(x) / 10)) + "1st"
                        elif x.endswith("2"):
                            inBetween = str(floor(int(x) / 10)) + "2nd"
                        elif x.endswith("3"):
                            inBetween = str(floor(int(x) / 10)) + "3rd"
                        else:
                            inBetween = x + "th"
                            if sline.lower().__contains__("password"):
                                sentence += str(
                                    number) + ". Enter value " + "******" + " on the" + inBetween + "field with id " + \
                                            sline.split("@id=")[1].split("]")[0] + "<br>"
                            else:
                                sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                                    0] + " on the" + inBetween + "field with id " + sline.split("@id=")[1].split("]")[
                                                0] + "<br>"
                    else:
                        if sline.lower().__contains__("password"):
                            sentence += str(number) + ". Enter value " + "******" + " in the field with id " + \
                                        sline.split("@id=")[1].split("]")[0] + "<br>"
                        else:
                            sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                                0] + " in the field with id " + sline.split("@id=")[1].split("]")[0] + "<br>"
                elif sline.__contains__("@name="):
                    number += 1
                    if sline.__contains__(".get("):
                        x = str(sline.split(".get(")[1].split(")")[0]).strip()
                        if x.endswith("1"):
                            inBetween = str(floor(int(x) / 10)) + "1st"
                        elif x.endswith("2"):
                            inBetween = str(floor(int(x) / 10)) + "2nd"
                        elif x.endswith("3"):
                            inBetween = str(floor(int(x) / 10)) + "3rd"
                        else:
                            inBetween = x + "th"
                            if sline.lower().__contains__("password"):
                                sentence += str(number) + ". Enter value " + "******" + " in the field with name " + \
                                            sline.split("@name=")[1].split("]")[0] + "<br>"
                            else:
                                sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                                    0] + " in the field with name " + sline.split("@name=")[1].split("]")[0] + "<br>"
                    else:
                        if sline.lower().__contains__("password"):
                            sentence += str(number) + ". Enter value " + "*****" + " in the field with name " + \
                                        sline.split("@name=")[1].split("]")[0] + "<br>"
                        else:
                            sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                                0] + " in the field with name " + sline.split("@name=")[1].split("]")[0] + "<br>"
                elif sline.__contains__(".get(") and sline.__contains__("@id=") == False:
                    number += 1
                    x = str(sline.split(".get(")[1].split(")")[0]).strip()
                    if x.endswith("1"):
                        inBetween = str(floor(int(x) / 10)) + "1st"
                    elif x.endswith("2"):
                        inBetween = str(floor(int(x) / 10)) + "2nd"
                    elif x.endswith("3"):
                        inBetween = str(floor(int(x) / 10)) + "3rd"
                    else:
                        inBetween = x + "th"
                    # sentence += "Enter value " + sline.split("sendKeys(")[1].split(")")[0] + " in the field with id " + \
                    #             inBetween+" in the field from the webpage<br>"
                    sentence += str(number) + ". Enter value " + sline.split("sendKeys(")[1].split(")")[
                        0] + " on " + inBetween + " field in the current page<br>"

    else:
        sentence = ''
    sentence = sentence.replace('<br>', '\n')
    return sentence
    # new_file.write(sentence)
    # print(sentence)
    # connections = db.folderName.find()
    # myquery = {"appName": outerFolderName, "scenarioName": folderName}
    # newvalues = {"$set": {"test_Description": sentence}}
    # db.folderName.update_many(myquery, newvalues)

if __name__ == "__main__":
  app.run(host='0.0.0.0', port=5010, debug=True)
